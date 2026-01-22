#!/usr/bin/env python3
"""
Извлечение полного содержимого шаблона
"""

from docx import Document

template_path = "/Users/vladog/Desktop/отчет_petr_2_s1_e2_2024_09_12_rus/отчет_petr_2_s1_e2_2024_09_12_rus.docx"

doc = Document(template_path)

print("=" * 80)
print("ПОЛНОЕ СОДЕРЖИМОЕ ДОКУМЕНТА")
print("=" * 80)
print()

# Все параграфы
for i, para in enumerate(doc.paragraphs, 1):
    text = para.text.strip()
    if text:
        print(f"[Para {i}]")
        print(text)
        print()

print("\n" + "=" * 80)
print("СОДЕРЖИМОЕ ТАБЛИЦЫ (первые 10 строк)")
print("=" * 80)
print()

if doc.tables:
    table = doc.tables[0]
    for i, row in enumerate(table.rows[:10]):
        cells_text = [cell.text.strip() for cell in row.cells]
        print(f"Строка {i}: {cells_text}")
        print()

