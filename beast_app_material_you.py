"""
Beast Auto Reporter - Compact Material You Style Version
Компактный интерфейс в стиле Material You с поддержкой перетаскивания файлов
"""

import sys
import os
import logging
import shutil
import subprocess
from pathlib import Path
from datetime import datetime
import re
import io
import importlib.util

# Добавляем корневую директорию в путь (нужно для локальных модулей src при запуске из любого cwd)
ROOT_DIR = Path(__file__).resolve().parent
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))


def _maybe_add_user_site() -> None:
    try:
        import site
        user_site = site.getusersitepackages()
        if user_site and user_site not in sys.path:
            site.addsitedir(user_site)
        # Если модуль все еще не находится, пробуем подключить user-site других версий Python
        if importlib.util.find_spec("PyQt5") is None:
            user_base = Path.home() / "Library" / "Python"
            if user_base.exists():
                for ver_dir in sorted(user_base.iterdir()):
                    if not ver_dir.is_dir():
                        continue
                    candidate = ver_dir / "lib" / "python" / "site-packages"
                    if candidate.exists() and (candidate / "PyQt5").exists():
                        if str(candidate) not in sys.path:
                            site.addsitedir(str(candidate))
    except Exception:
        pass


_maybe_add_user_site()


def _ensure_pyqt5_sip_compat() -> None:
    try:
        if importlib.util.find_spec("PyQt5.sip") is None:
            if importlib.util.find_spec("PyQt5_sip") is not None:
                import PyQt5_sip
                sys.modules.setdefault("PyQt5.sip", PyQt5_sip)
    except Exception:
        pass


_ensure_pyqt5_sip_compat()


def _can_import_pyqt5(candidate: str) -> bool:
    try:
        result = subprocess.run(
            [candidate, "-c", "import PyQt5.sip"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=10
        )
        return result.returncode == 0
    except Exception:
        return False


def _maybe_reexec_with_pyqt5() -> None:
    if os.environ.get("BEAST_PYQT5_BOOTSTRAP") == "1":
        return
    if importlib.util.find_spec("PyQt5.sip") is not None:
        return

    candidates = []
    candidates.append(sys.executable)
    candidates.append("/Library/Developer/CommandLineTools/usr/bin/python3")
    candidates.append("/usr/bin/python3")

    which_py = shutil.which("python3")
    if which_py:
        candidates.append(which_py)

    repo_root = Path(__file__).resolve().parent
    venv_py = repo_root / "venv" / "bin" / "python"
    if venv_py.exists():
        candidates.append(str(venv_py))

    # Deduplicate while preserving order
    seen = set()
    deduped = []
    for c in candidates:
        if not c:
            continue
        key = str(Path(c).resolve())
        if key in seen:
            continue
        seen.add(key)
        deduped.append(c)

    current = str(Path(sys.executable).resolve())
    for candidate in deduped:
        try:
            if str(Path(candidate).resolve()) == current:
                continue
        except Exception:
            pass
        if _can_import_pyqt5(candidate):
            os.environ["BEAST_PYQT5_BOOTSTRAP"] = "1"
            os.execv(candidate, [candidate] + sys.argv)


_maybe_reexec_with_pyqt5()


def _require_module(module_name: str, install_hint: str) -> None:
    if importlib.util.find_spec(module_name) is None:
        raise SystemExit(
            f"Модуль '{module_name}' не установлен.\n"
            f"{install_hint}\n"
            "Или запустите приложение из собранного .app"
        )


# Проверяем наличие зависимостей до импорта модулей, которые их используют
_require_module("docx", "Установите зависимости: pip install -r requirements_pinned.txt")
_require_module("PyQt5", "Установите зависимости: pip install -r requirements_pinned.txt")
_require_module("fitz", "Установите зависимости: pip install -r requirements_pinned.txt")

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QRadioButton, QTextEdit, QProgressBar,
    QCheckBox, QListWidget, QListWidgetItem, QFrame, QMessageBox, QFileDialog,
    QGroupBox, QScrollArea, QSizePolicy
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QMimeData
from PyQt5.QtGui import QFont, QDragEnterEvent, QDropEvent, QPalette, QColor, QIcon

import fitz

from src.exact_report_generator import ExactReportGenerator
from src.technical_info_extractor import TechnicalInfoExtractor
from src.csv_importer import CSVImporter
from src.pdf_extractor import PDFExtractor
from src.conclusion_generator import ConclusionGenerator
from src.parameter_exporter import ParameterExporter

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


_CHANNEL_51_RE = re.compile(
    r'(^|[^0-9])(?:5\.1|5_1|5-1|5\.0|51)([^0-9]|$)|(^|[^a-z0-9])(?:surround|6\s*ch|6ch|6\s*channel)([^a-z0-9]|$)',
    re.IGNORECASE
)
_CHANNEL_20_RE = re.compile(
    r'(^|[^0-9])(?:2\.0|2_0|2-0|2\.1|20)([^0-9]|$)|(^|[^a-z0-9])(?:stereo|2\s*ch|2ch|2\s*channel|lr|l\s*r)([^a-z0-9]|$)',
    re.IGNORECASE
)
_UNCENS_RE = re.compile(r'(^|[^a-z0-9])(?:uncens|uncensored|uc)([^a-z0-9]|$)', re.IGNORECASE)
_CENS_RE = re.compile(r'(^|[^a-z0-9])(?:cens|censored|c)([^a-z0-9]|$)', re.IGNORECASE)
_INCORRECT_AUDIO_TYPE_PATTERNS = [
    r'_50_(cens|uncens|mix|c|uc)',
    r'_50$',
    r'_21_(cens|uncens|mix|c|uc)',
    r'_21$',
    r'_5\.0_(cens|uncens|mix)',
    r'_2\.1_(cens|uncens|mix)',
]


def detect_channel_from_name(name: str):
    name_lower = name.lower()
    if _CHANNEL_51_RE.search(name_lower):
        return "51"
    if _CHANNEL_20_RE.search(name_lower):
        return "20"
    return None


def detect_cens_state(name: str):
    name_lower = name.lower()
    if _UNCENS_RE.search(name_lower):
        return "uc"
    if _CENS_RE.search(name_lower):
        return "c"
    return None


def has_incorrect_audio_marker(name: str) -> bool:
    name_lower = name.lower()
    for pattern in _INCORRECT_AUDIO_TYPE_PATTERNS:
        if re.search(pattern, name_lower, re.IGNORECASE):
            return True
    return False


def order_pdfs_by_channel(pdf_paths: list) -> list:
    if not pdf_paths:
        return pdf_paths
    scored = []
    for idx, p in enumerate(pdf_paths):
        name = Path(p).stem
        channel = detect_channel_from_name(name) or ""
        cens_state = detect_cens_state(name) or ""
        # channel order: 20 -> 51 -> unknown
        ch_score = 2
        if channel == "20":
            ch_score = 0
        elif channel == "51":
            ch_score = 1
        # cens order: cens -> uncens -> unknown
        cens_score = 2
        if cens_state == "c":
            cens_score = 0
        elif cens_state == "uc":
            cens_score = 1
        scored.append((ch_score, cens_score, idx, p))
    scored.sort(key=lambda t: (t[0], t[1], t[2]))
    return [p for _, __, ___, p in scored]


def sanitize_base_name(name: str) -> str:
    base = Path(name).stem
    # Удаляем служебные маркеры каналов/цензуры/типа
    patterns = [
        r'(?i)(^|[ _\.-])(5\.?1|5_1|5-1|5\.0|51)(?=$|[ _\.-])',
        r'(?i)(^|[ _\.-])(2\.?0|2_0|2-0|2\.1|20|stereo)(?=$|[ _\.-])',
        r'(?i)(^|[ _\.-])(uncens|uncensored|uc)(?=$|[ _\.-])',
        r'(?i)(^|[ _\.-])(cens|censored|c)(?=$|[ _\.-])',
        r'(?i)(^|[ _\.-])(audio|video)(?=$|[ _\.-])',
    ]
    for pat in patterns:
        base = re.sub(pat, "_", base)
    base = re.sub(r'[_\.\- ]+', '_', base).strip('_')
    return base or "отчет"


def normalize_stem(name: str) -> str:
    stem = Path(name).stem.lower()
    stem = re.sub(r'[^a-z0-9]+', '_', stem)
    stem = re.sub(r'_+', '_', stem).strip('_')
    return stem


class DropZone(QFrame):
    """Compact Drag & Drop зона для файлов"""
    
    files_dropped = pyqtSignal(list)  # Сигнал при добавлении файлов
    
    def __init__(self):
        super().__init__()
        self.setAcceptDrops(True)
        self.init_ui()
    
    def init_ui(self):
        """Инициализация компактного UI зоны"""
        self.setMinimumHeight(120)
        self.setMaximumHeight(160)
        self.setFrameStyle(QFrame.NoFrame)
        
        # Material You стили
        self.default_style = """
            QFrame {
                background-color: #F8F9FA;
                border: 2px dashed #DEE2E6;
                border-radius: 16px;
                padding: 16px;
            }
            QFrame:hover {
                background-color: #E9ECEF;
                border: 2px dashed #ADB5BD;
            }
        """
        
        # Стиль при наведении
        self.hover_style = """
            QFrame {
                background-color: #E8F5E9;
                border: 2px dashed #4CAF50;
                border-radius: 16px;
                padding: 16px;
            }
        """
        
        self.setStyleSheet(self.default_style)
        
        # Layout
        layout = QVBoxLayout()
        layout.setAlignment(Qt.AlignCenter)
        
        # Иконка
        icon_label = QLabel("📥")
        icon_label.setFont(QFont("SF Pro Display", 28))
        icon_label.setAlignment(Qt.AlignCenter)
        icon_label.setStyleSheet("background: transparent; border: none; color: #6C757D;")
        layout.addWidget(icon_label)
        
        # Текст
        text_label = QLabel("Перетащите файлы сюда")
        text_label.setFont(QFont("SF Pro Display", 12, QFont.Medium))
        text_label.setAlignment(Qt.AlignCenter)
        text_label.setStyleSheet("background: transparent; border: none; color: #495057;")
        layout.addWidget(text_label)
        
        # Подсказка
        hint_label = QLabel("аудио, видео, CSV, PDF")
        hint_label.setFont(QFont("SF Pro Text", 9))
        hint_label.setAlignment(Qt.AlignCenter)
        hint_label.setStyleSheet("background: transparent; border: none; color: #868E96;")
        layout.addWidget(hint_label)
        
        self.setLayout(layout)
    
    def dragEnterEvent(self, event: QDragEnterEvent):
        """Обработка входа файла в зону"""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            self.setStyleSheet(self.hover_style)
    
    def dragLeaveEvent(self, event):
        """Обработка выхода файла из зоны"""
        self.setStyleSheet(self.default_style)
    
    def dropEvent(self, event: QDropEvent):
        """Обработка сброса файлов"""
        self.setStyleSheet(self.default_style)
        
        files = []
        for url in event.mimeData().urls():
            file_path = url.toLocalFile()
            if os.path.isfile(file_path):
                files.append(file_path)
        
        if files:
            self.files_dropped.emit(files)


class PdfOnlyDocxBuilder:
    """Создает простой DOCX из первых страниц PDF (по одному PDF на страницу)."""

    @staticmethod
    def build(output_path: Path, pdf_paths: list) -> None:
        doc = Document()

        # A3 landscape как в основном отчете
        section = doc.sections[0]
        section.page_width = Cm(42.02)
        section.page_height = Cm(29.70)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.orientation = 1  # Landscape

        for idx, pdf_path in enumerate(pdf_paths):
            if idx > 0:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_break(WD_BREAK.PAGE)

            pdf_doc = fitz.open(str(pdf_path))
            page = pdf_doc[0]
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            img_stream = io.BytesIO(pix.tobytes("png"))
            doc.add_picture(img_stream, width=Cm(38.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            pdf_doc.close()

        doc.save(str(output_path))


def convert_dragdrop_to_files(files_data):
    """Преобразует drag & drop файлы в формат v5.11"""
    files = {
        'audio_20_c': None,
        'audio_20_uc': None,
        'audio_51_c': None,
        'audio_51_uc': None,
        'video': None,
        'csv': None,
        'pdf_20': None,
        'pdf_51': None,
        'pdf_20_c': None,
        'pdf_20_uc': None,
        'pdf_51_c': None,
        'pdf_51_uc': None,
        'params': None,
        'all_files': []
    }
    
    # CSV
    if files_data.get('csv'):
        files['csv'] = Path(files_data['csv'][0])
        files['all_files'].append(files['csv'])
    
    # Video
    if files_data.get('video'):
        files['video'] = Path(files_data['video'][0])
        files['all_files'].append(files['video'])
    
    # Обработка аудио файлов (логика из v5.11)
    for audio_file in files_data.get('audio', []):
        file_path = Path(audio_file)
        files['all_files'].append(file_path)
        name_lower = file_path.name.lower()
        channel = detect_channel_from_name(name_lower)
        cens_state = detect_cens_state(name_lower)
        
        # Проверяем 5.1 ПЕРЕД 2.0 (как в v5.11)
        if channel == "51":
            if cens_state == "uc":
                files['audio_51_uc'] = file_path
            else:
                files['audio_51_c'] = file_path
        
        # Проверяем 2.0
        elif channel == "20":
            if cens_state == "uc":
                files['audio_20_uc'] = file_path
            else:
                files['audio_20_c'] = file_path
    
    # Обработка PDF файлов (логика из v5.11)
    for pdf_file in files_data.get('pdf', []):
        file_path = Path(pdf_file)
        files['all_files'].append(file_path)
        name_lower = file_path.name.lower()
        channel = detect_channel_from_name(name_lower)
        cens_state = detect_cens_state(name_lower)
        
        # Проверяем 5.1 ПЕРЕД 2.0
        if channel == "51":
            if cens_state == "uc":
                files['pdf_51_uc'] = file_path
            elif cens_state == "c":
                files['pdf_51_c'] = file_path
            else:
                files['pdf_51'] = file_path
        
        # Проверяем 2.0
        elif channel == "20":
            if cens_state == "uc":
                files['pdf_20_uc'] = file_path
            elif cens_state == "c":
                files['pdf_20_c'] = file_path
            else:
                files['pdf_20'] = file_path
    
    return files


class ProcessingThread(QThread):
    """Поток для обработки файлов и генерации отчета"""
    
    status_update = pyqtSignal(str)
    progress_update = pyqtSignal(int)
    finished = pyqtSignal(bool, str)
    
    def __init__(self, app, files_data, report_type, output_folder, pyloudnorm_enabled=False):
        super().__init__()
        self.app = app
        self.files_data = files_data
        self.report_type = report_type
        self.output_folder = output_folder
        self.pyloudnorm_enabled = pyloudnorm_enabled
    
    def run(self):
        """Запуск обработки"""
        logger.debug("=== ProcessingThread.run() STARTED ===")
        try:
            import shutil
            from pathlib import Path
            
            logger.info(f"=== НАЧАЛО ОБРАБОТКИ ===")
            logger.info(f"Тип отчета: {self.report_type}")
            logger.info(f"PyLoudNorm включен: {self.pyloudnorm_enabled}")
            logger.info(f"Файлов: {sum(len(v) for v in self.files_data.values())}")
            logger.info(f"Output folder: {self.output_folder}")
            logger.debug(f"files_data keys = {list(self.files_data.keys())}")
            
            # Определяем имя файла из аудио или видео
            audio_files = self.files_data.get('audio', [])
            video_files = self.files_data.get('video', [])
            logger.debug(f"audio_files = {audio_files}")
            logger.debug(f"video_files = {video_files}")
            
            # Также проверяем специфичные ключи для 5.1
            audio_51_c = self.files_data.get('audio_51_c', [])
            audio_51_uc = self.files_data.get('audio_51_uc', [])
            audio_20_c = self.files_data.get('audio_20_c', [])
            audio_20_uc = self.files_data.get('audio_20_uc', [])
            logger.debug(f"audio_51_c = {audio_51_c}")
            logger.debug(f"audio_51_uc = {audio_51_uc}")
            csv_files = self.files_data.get('csv', [])
            pdf_files = self.files_data.get('pdf', [])
            
            logger.info(f"Audio files found: {len(audio_files)}")
            logger.info(f"Video files found: {len(video_files)}")
            logger.info(f"CSV files found: {len(csv_files)}")
            logger.info(f"PDF files found: {len(pdf_files)}")
            
            # ДИАГНОСТИКА: ЛОГИРУЕМ PDF ФАЙЛЫ
            if pdf_files:
                logger.info("=== PDF FILES DETAILS ===")
                for i, pdf in enumerate(pdf_files):
                    logger.info(f"  [{i}] {pdf}")
                logger.info("=== END PDF FILES ===")
            else:
                logger.warning("⚠️ NO PDF FILES LOADED!")
            
            # Логируем report_type для диагностики
            logger.info(f"Report type: {self.report_type}")
            
            # Объединяем все аудио файлы для pyloudnorm анализа
            all_audio_files = list(audio_files)
            all_audio_files.extend(audio_51_c)
            all_audio_files.extend(audio_51_uc)
            all_audio_files.extend(audio_20_c)
            all_audio_files.extend(audio_20_uc)
            logger.debug(f"all_audio_files for analysis = {all_audio_files}")
            
            base_name = "отчет"
            if audio_files:
                base_name = sanitize_base_name(audio_files[0])
            elif video_files:
                base_name = sanitize_base_name(video_files[0])
            
            # Папка уже создана в start_processing через create_output_folder()
            output_dir = Path(self.output_folder)
            logger.info(f"Выходная папка: {output_dir}")

            # === PDF-only режим: вставка PDF в DOCX в самом начале ===
            if pdf_files and not audio_files and not video_files and not csv_files:
                self.status_update.emit("📄 Вставка PDF в DOCX...")
                self.progress_update.emit(10)
                
                copied_pdfs = []
                for pdf_file in pdf_files:
                    dest = output_dir / Path(pdf_file).name
                    shutil.copy2(pdf_file, dest)
                    copied_pdfs.append(dest)
                
                simple_docx_path = output_dir / f"pdf_{base_name}.docx"
                ordered_pdfs = order_pdfs_by_channel(copied_pdfs)
                PdfOnlyDocxBuilder.build(simple_docx_path, ordered_pdfs)
                
                self.progress_update.emit(100)
                self.status_update.emit("✅ Готово!")
                
                files_in_output = list(output_dir.glob('*'))
                success_msg = (
                    f"✅ PDF вставлены в DOCX!\n"
                    f"📁 Папка: {output_dir.name}\n"
                    f"📄 Файлов: {len(files_in_output)}"
                )
                self.finished.emit(True, success_msg)
                logger.info(f"✅ PDF-only DOCX создан: {simple_docx_path}")
                return
            
            # Копируем PDF файлы в выходную папку и разделяем по типам
            self.status_update.emit("📋 Копирование файлов...")
            self.progress_update.emit(10)
            
            # Словари для хранения путей к разным типам PDF
            pdf_paths = {
                '20_c': None,
                '20_uc': None,
                '20': None,
                '51_c': None,
                '51_uc': None,
                '51': None
            }
            copied_pdfs = []
            pdf_by_stem = {}
            
            for pdf_file in pdf_files:
                dest = output_dir / Path(pdf_file).name
                shutil.copy2(pdf_file, dest)
                copied_pdfs.append(dest)
                logger.info(f"✅ Скопирован PDF: {Path(pdf_file).name}")
                stem_key = normalize_stem(pdf_file)
                pdf_by_stem.setdefault(stem_key, []).append(dest)
                
                filename = Path(pdf_file).stem.lower()
                
                # Определяем тип PDF с приоритетом (regex, чтобы ловить cens_20 / cens51 и т.п.)
                channel = detect_channel_from_name(filename)
                is_20 = channel == "20"
                is_51 = channel == "51"
                cens_state = detect_cens_state(filename)
                
                if is_20 and cens_state == "c":
                    pdf_paths['20_c'] = str(dest)
                    logger.info(f"📄 PDF 2.0 CENS: {Path(pdf_file).name}")
                elif is_20 and cens_state == "uc":
                    pdf_paths['20_uc'] = str(dest)
                    logger.info(f"📄 PDF 2.0 UNCENS: {Path(pdf_file).name}")
                elif is_20:
                    pdf_paths['20'] = str(dest)
                    logger.info(f"📄 PDF 2.0: {Path(pdf_file).name}")
                elif is_51 and cens_state == "c":
                    pdf_paths['51_c'] = str(dest)
                    logger.info(f"📄 PDF 5.1 CENS: {Path(pdf_file).name}")
                elif is_51 and cens_state == "uc":
                    pdf_paths['51_uc'] = str(dest)
                    logger.info(f"📄 PDF 5.1 UNCENS: {Path(pdf_file).name}")
                elif is_51:
                    pdf_paths['51'] = str(dest)
                    logger.info(f"📄 PDF 5.1: {Path(pdf_file).name}")
            
            # Выбираем PDF для вставки в отчет (приоритет: cens > uncens > общий)
            copied_pdf_20 = pdf_paths['20_c'] or pdf_paths['20_uc'] or pdf_paths['20']
            copied_pdf_51 = pdf_paths['51_c'] or pdf_paths['51_uc'] or pdf_paths['51']

            # Фолбэк: если тип не определился, пробуем по имени среди всех PDF
            if not copied_pdf_20 or not copied_pdf_51:
                for p in copied_pdfs:
                    name = p.stem.lower()
                    if not copied_pdf_20 and detect_channel_from_name(name) == "20":
                        copied_pdf_20 = str(p)
                    if not copied_pdf_51 and detect_channel_from_name(name) == "51":
                        copied_pdf_51 = str(p)

            # Последний фолбэк: берем первые два PDF в списке
            if copied_pdfs:
                if not copied_pdf_20:
                    copied_pdf_20 = str(copied_pdfs[0])
                if not copied_pdf_51 and len(copied_pdfs) > 1:
                    copied_pdf_51 = str(copied_pdfs[1])
            
            logger.info(f"Выбран PDF 2.0 для отчета: {copied_pdf_20}")
            logger.info(f"Выбран PDF 5.1 для отчета: {copied_pdf_51}")
            
            # Копируем CSV если есть
            if csv_files:
                dest = output_dir / Path(csv_files[0]).name
                shutil.copy2(csv_files[0], dest)
                logger.info(f"✅ Скопирован CSV: {Path(csv_files[0]).name}")
            
            # Извлечение технической информации
            self.status_update.emit("🔍 Анализ файлов...")
            self.progress_update.emit(20)
            
            tech_extractor = TechnicalInfoExtractor()
            tech_info = {}
            audio_key_by_stem = {}
            
            # Обработка аудио файлов
            for audio_file in audio_files:
                logger.info(f"Обработка аудио: {audio_file}")
                filename = Path(audio_file).stem.lower()
                
                # Извлекаем метаданные для определения количества каналов
                audio_info = tech_extractor.extract_audio_info(audio_file)
                if not audio_info:
                    logger.warning(f"⚠️  Не удалось извлечь информацию из аудио: {filename}")
                    continue
                
                channels = audio_info.get('channels', 0)
                try:
                    channels = int(float(channels)) if channels is not None else 0
                except Exception:
                    channels = 0
                is_51 = channels >= 6  # 6 каналов = 5.1
                is_20 = channels == 2  # 2 канала = 2.0
                
                # Определяем тип по имени файла (гибкая проверка)
                channel_hint = detect_channel_from_name(filename)
                channel_meta = "51" if is_51 else ("20" if is_20 else None)
                name_mismatch = bool(channel_meta and channel_hint and channel_meta != channel_hint)
                has_bad_marker = has_incorrect_audio_marker(filename)
                
                cens_state = detect_cens_state(filename)
                suffix = "uc" if cens_state == "uc" else "c"
                
                # Выбираем канал по метаданным, если есть; иначе по имени
                channel_final = channel_meta or channel_hint
                
                if channel_final == "51":
                    key = f'audio_51_{suffix}'
                elif channel_final == "20":
                    key = f'audio_20_{suffix}'
                else:
                    logger.warning(f"⚠️  Не удалось определить тип аудио: {filename} (channels={channels})")
                    continue
                
                name_score = 0
                if channel_hint and channel_final and channel_hint == channel_final:
                    name_score += 2
                if not has_bad_marker:
                    name_score += 1

                if name_mismatch or has_bad_marker:
                    audio_info['name_channel_mismatch'] = name_mismatch
                    audio_info['name_incorrect_type'] = has_bad_marker
                    if name_mismatch:
                        logger.warning(f"⚠️  Несовпадение типа в имени и метаданных: name={channel_hint}, meta={channel_meta}")
                    if has_bad_marker:
                        logger.warning(f"⚠️  Ошибка в названии формата (5.0/2.1): {filename}")

                audio_info['_name_score'] = name_score
                if key in tech_info and isinstance(tech_info[key], dict):
                    existing_score = tech_info[key].get('_name_score', 0)
                    if name_score <= existing_score:
                        logger.warning(f"⚠️  Найден дубль для {key}, оставляем файл с более корректным именем")
                        continue
                
                tech_info[key] = audio_info
                audio_key_by_stem[normalize_stem(audio_info.get('file_name', audio_file))] = key
                logger.info(f"✅ {key}: {audio_info.get('file_name')} (channels={channels})")

            self.progress_update.emit(40)
            
            # Обработка видео файла
            if video_files:
                video_file = video_files[0]
                logger.info(f"📹 Обработка видео: {video_file}")
                video_info = tech_extractor.extract_video_info(video_file)
                if video_info:
                    tech_info['video'] = video_info
                    logger.info(f"✅ Видео: {video_info.get('file_name')}, fps={video_info.get('fps')}, format={video_info.get('format')}")
                else:
                    logger.warning("⚠️  Не удалось извлечь информацию о видео")
            
            self.progress_update.emit(50)
            
            # === ШАГ: Извлекаем данные из PDF и сохраняем в JSON ===
            # Данные из PDF сначала накапливаются в словаре, затем сохраняются в JSON
            self.status_update.emit("📊 Извлечение данных из PDF...")
            
            logger.info("=== PDF FILES PROCESSING START ===")
            logger.info(f"pdf_files count: {len(pdf_files)}")
            logger.info(f"pdf_files: {pdf_files}")
            
            # Словарь для накопления PDF данных
            pdf_data_all = {}
            pending_pdfs = []
            
            for pdf_file in pdf_files:
                logger.info(f"\n--- Обработка PDF: {pdf_file} ---")
                filename = Path(pdf_file).stem.lower()
                logger.info(f"  filename (lower): {filename}")

                # Гибкая проверка разных форматов: _51_, _51_, - 51 -, _cens_51, stereo, и т.д.
                channel_hint = detect_channel_from_name(filename)
                has_51_marker = channel_hint == "51"
                has_20_marker = channel_hint == "20"

                cens_state = detect_cens_state(filename)
                is_cens = cens_state == "c"
                is_uncens = cens_state == "uc"

                # 1) Если имя PDF совпадает с аудио-файлом — берем тип из аудио метаданных
                matched_audio_key = audio_key_by_stem.get(normalize_stem(pdf_file))
                if matched_audio_key:
                    key = matched_audio_key.replace('audio_', 'pdf_')
                    logger.info(f"  -> Matched PDF to audio by name: {matched_audio_key} -> {key}")
                else:
                    # 2) Определяем тип PDF по имени
                    if has_51_marker and is_cens:
                        key = 'pdf_51_c'
                        logger.info(f"  -> Detected PDF 5.1 CENS, key = '{key}'")
                    elif has_51_marker and is_uncens:
                        key = 'pdf_51_uc'
                        logger.info(f"  -> Detected PDF 5.1 UNCENS, key = '{key}'")
                    elif has_51_marker:
                        key = 'pdf_51'
                        logger.info(f"  -> Detected PDF 5.1 (generic), key = '{key}'")
                    elif has_20_marker and is_cens:
                        key = 'pdf_20_c'
                        logger.info(f"  -> Detected PDF 2.0 CENS, key = '{key}'")
                    elif has_20_marker and is_uncens:
                        key = 'pdf_20_uc'
                        logger.info(f"  -> Detected PDF 2.0 UNCENS, key = '{key}'")
                    elif has_20_marker:
                        key = 'pdf_20'
                        logger.info(f"  -> Detected PDF 2.0 (generic), key = '{key}'")
                    else:
                        key = None

                # Извлекаем данные из PDF
                pdf_data = self.app.pdf_extractor.extract_technical_info(pdf_file)
                if pdf_data is None:
                    pdf_data = {}
                pdf_data['source_pdf'] = Path(pdf_file).name

                # 3) Если ключ не определен — пробуем по контенту PDF (каналы)
                if key is None:
                    logger.info(f"  -> No explicit marker, extracting from PDF content...")
                    if pdf_data:
                        logger.info(f"  -> PDF extracted: lufs={pdf_data.get('lufs')}, channels={pdf_data.get('channels')}")
                    else:
                        logger.warning(f"  -> PDF extraction returned EMPTY for {filename}")
                    
                    if pdf_data and pdf_data.get('channels'):
                        channels_str = str(pdf_data.get('channels', '')).lower()
                        if '5.1' in channels_str or '5.0' in channels_str or '6' in channels_str or 'surround' in channels_str:
                            key = 'pdf_51_c' if is_cens else ('pdf_51_uc' if is_uncens else 'pdf_51')
                            logger.info(f"  -> Detected from channels: 5.1, key = '{key}'")
                        elif '2.0' in channels_str or '2.1' in channels_str or '2' in channels_str or 'stereo' in channels_str:
                            key = 'pdf_20_c' if is_cens else ('pdf_20_uc' if is_uncens else 'pdf_20')
                            logger.info(f"  -> Detected from channels: 2.0, key = '{key}'")
                        else:
                            logger.warning(f"  -> Unknown channel format: {channels_str}")
                    else:
                        logger.warning(f"  -> No channels info in PDF")

                # 4) Если ключ всё ещё не определен — сохраняем в pending
                if key is None:
                    pending_pdfs.append({
                        'path': pdf_file,
                        'data': pdf_data,
                        'cens_state': cens_state
                    })
                    logger.warning(f"  -> PDF key unresolved, added to pending: {Path(pdf_file).name}")
                    continue

                # Сохраняем в словарь для JSON (а не напрямую в tech_info)
                if key in pdf_data_all:
                    # если ключ занят, кладем в общий без цензуры
                    key = 'pdf_20' if '20' in key else 'pdf_51'
                pdf_data_all[key] = pdf_data
                
                if pdf_data:
                    logger.info(f"  -> ADDED to pdf_data_all['{key}']: LUFS={pdf_data.get('lufs')}, Peak={pdf_data.get('true_peak')}, LRA={pdf_data.get('lra')}")
                else:
                    logger.warning(f"  -> PDF extraction returned EMPTY for {filename}")

            # Попытка доопределить pending PDF по совпадению с аудио без маркеров
            if pending_pdfs:
                missing_keys = []
                for prefix in ("20", "51"):
                    for suffix in ("_c", "_uc", ""):
                        k = f"pdf_{prefix}{suffix}"
                        audio_k = f"audio_{prefix}{suffix}"
                        if audio_k in tech_info and (k not in pdf_data_all or pdf_data_all.get(k) is None):
                            missing_keys.append((k, audio_k))

                for item in pending_pdfs:
                    pdf_path = item['path']
                    pdf_data = item['data']
                    best_key = None
                    pdf_norm = normalize_stem(pdf_path)
                    for k, audio_k in missing_keys:
                        audio_norm = normalize_stem(tech_info[audio_k].get('file_name', ''))
                        if audio_norm and (audio_norm in pdf_norm or pdf_norm in audio_norm):
                            best_key = k
                            break
                    if best_key is None and len(missing_keys) == 1:
                        best_key = missing_keys[0][0]
                    
                    if best_key:
                        pdf_data_all[best_key] = pdf_data
                        logger.info(f"  -> Pending PDF matched to {best_key}: {Path(pdf_path).name}")
                    else:
                        logger.warning(f"  -> Pending PDF could not be matched: {Path(pdf_path).name}")
            
            logger.info(f"=== PDF FILES PROCESSING END ===")
            logger.info(f"pdf_data_all keys: {list(pdf_data_all.keys())}")
            logger.info(f"pdf_data_all is empty: {len(pdf_data_all) == 0}")
            
            # === ШАГ: Сохраняем все PDF данные в JSON файл ===
            json_path = output_dir / "pdf_data.json"
            
            logger.info(f"=== JSON SAVE CHECK ===")
            logger.info(f"pdf_data_all is empty: {len(pdf_data_all) == 0}")
            
            if pdf_data_all:
                try:
                    self.app.pdf_extractor.save_tech_info_to_json(pdf_data_all, str(json_path))
                    logger.info(f"✅ PDF данные сохранены в JSON: {json_path}")
                    
                    # Логируем сохраненные данные
                    logger.info("=== PDF данные в JSON ===")
                    for k, v in pdf_data_all.items():
                        if isinstance(v, dict):
                            logger.info(f"  {k}: lufs={v.get('lufs')}, peak={v.get('true_peak')}, lra={v.get('lra')}")
                    logger.info("=== END PDF данные ===")
                    
                    # ДОБАВЛЯЕМ данные из pdf_data_all напрямую в tech_info (для гарантии)
                    # Это обеспечивает PDF → JSON → таблица поток
                    for key, value in pdf_data_all.items():
                        if isinstance(key, str) and key.startswith('pdf_'):
                            tech_info[key] = value
                            logger.info(f"  -> tech_info['{key}'] updated from pdf_data_all")
                    
                except Exception as e:
                    logger.error(f"❌ Ошибка сохранения JSON: {e}")
            else:
                logger.warning("⚠️ Нет PDF данных для сохранения")
            
            self.progress_update.emit(52)
            
            # === ШАГ: Загружаем PDF данные из JSON ===
            # Параметры берутся из JSON файла, в который были сохранены данные из PDF
            self.status_update.emit("💾 Загрузка PDF данных из JSON...")
            
            # Если JSON файл существует, загружаем данные из него
            # ДАННЫЕ УЖЕ ДОБАВЛЕНЫ В tech_info ПРИ СОХРАНЕНИИ (см. выше)
            # Эта загрузка служит как резерв/проверка
            if json_path.exists():
                logger.info(f"📄 JSON файл найден: {json_path}")
                try:
                    # Загружаем техническую информацию из JSON
                    json_data = self.app.pdf_extractor.load_tech_info_from_json(str(json_path))
                    
                    logger.info(f"📄 Загружено из JSON: {list(json_data.keys()) if json_data else 'ПУСТО'}")
                    
                    # Логируем ВСЕ ключи из json_data для диагностики
                    logger.info("=== json_data contents ===")
                    for k, v in json_data.items():
                        if isinstance(v, dict):
                            logger.info(f"  {k}: lufs={v.get('lufs')}, peak={v.get('true_peak')}, lra={v.get('lra')}")
                        else:
                            logger.info(f"  {k}: {v}")
                    logger.info("=== end json_data ===")
                    
                    if json_data:
                        # Обновляем tech_info данными из JSON (дублируем для надежности)
                        # Это гарантирует, что параметры в таблицу вставляются из JSON
                        updated_count = 0
                        for key, value in json_data.items():
                            if isinstance(key, str) and key.startswith('pdf_'):
                                tech_info[key] = value
                                updated_count += 1
                                logger.info(f"  -> tech_info['{key}'] = {value}")
                        logger.info(f"  📊 Обновлено {updated_count} ключей из JSON")
                    else:
                        logger.warning("⚠️ JSON файл пуст или поврежден")
                except Exception as e:
                    logger.error(f"❌ Ошибка чтения JSON: {e}")
            else:
                logger.warning(f"⚠️ JSON файл НЕ найден: {json_path}")
            
            # Диагностика: что в tech_info перед генерацией
            logger.info(f"📊 tech_info перед генерацией: {list(tech_info.keys())}")
            for k in ['pdf_20', 'pdf_51', 'pdf_20_c', 'pdf_51_c', 'pdf_20_uc', 'pdf_51_uc']:
                if k in tech_info:
                    v = tech_info[k]
                    if isinstance(v, dict):
                        logger.info(f"  -> {k}: lufs={v.get('lufs')}, peak={v.get('true_peak')}, lra={v.get('lra')}")
                    else:
                        logger.info(f"  -> {k}: {v}")
            
            # Логируем PDF данные для диагностики
            logger.info("=== PDF данные для таблицы ===")
            for k in tech_info:
                if isinstance(k, str) and k.startswith('pdf_') and isinstance(tech_info[k], dict):
                    v = tech_info[k]
                    logger.info(f"  {k}: lufs={v.get('lufs')}, true_peak={v.get('true_peak')}, lra={v.get('lra')}")
            logger.info("=== END ===")

            # === Принудительно сопоставляем 4 PDF по именам (cens -> uncens) ===
            def remap_pdf_key(target_key: str, pdf_path: str):
                if not pdf_path:
                    return
                target_name = Path(pdf_path).name
                for k, v in tech_info.items():
                    if isinstance(k, str) and k.startswith('pdf_') and isinstance(v, dict):
                        if v.get('source_pdf') == target_name:
                            tech_info[target_key] = v
                            logger.info(f"✅ PDF remap: {target_key} <- {target_name}")
                            return

            remap_pdf_key('pdf_20_c', pdf_paths.get('20_c'))
            remap_pdf_key('pdf_20_uc', pdf_paths.get('20_uc'))
            remap_pdf_key('pdf_51_c', pdf_paths.get('51_c'))
            remap_pdf_key('pdf_51_uc', pdf_paths.get('51_uc'))
            
            # === Уточняем выбор PDF 2.0 / 5.1 по метаданным аудио ===
            def pick_pdf_for_audio(prefix: str):
                for audio_key in (f"audio_{prefix}_c", f"audio_{prefix}_uc", f"audio_{prefix}"):
                    audio_data = tech_info.get(audio_key)
                    if isinstance(audio_data, dict):
                        audio_name = audio_data.get('file_name', '')
                        if audio_name:
                            norm = normalize_stem(audio_name)
                            if norm in pdf_by_stem and pdf_by_stem[norm]:
                                return str(pdf_by_stem[norm][0])
                return None

            picked_20 = pick_pdf_for_audio("20")
            picked_51 = pick_pdf_for_audio("51")
            if picked_20:
                copied_pdf_20 = picked_20
                logger.info(f"✅ PDF 2.0 уточнен по аудио-метаданным: {copied_pdf_20}")
            if picked_51:
                copied_pdf_51 = picked_51
                logger.info(f"✅ PDF 5.1 уточнен по аудио-метаданным: {copied_pdf_51}")

            # Защита от дубликатов (когда выбран один и тот же PDF для 2.0 и 5.1)
            if copied_pdf_20 and copied_pdf_51 and copied_pdf_20 == copied_pdf_51 and len(copied_pdfs) > 1:
                for p in copied_pdfs:
                    if str(p) != copied_pdf_20:
                        copied_pdf_51 = str(p)
                        logger.info(f"⚠️ PDF 5.1 был дубликатом, заменен на {copied_pdf_51}")
                        break
            
            logger.info("✅ Параметры из JSON готовы для вставки в таблицу")
            
            self.progress_update.emit(55)
            
            # === ШАГ 6б: Аудио анализ и экспорт в CSV/HTML ===
            if self.pyloudnorm_enabled:
                self.status_update.emit("📊 PyLoudNorm анализ...")
                
                logger.debug(f"audio_files = {audio_files}")
                
                try:
                    logger.debug(f"Entering try block")
                    if all_audio_files:
                        logger.debug(f"all_audio_files is not empty, length = {len(all_audio_files)}")
                        logger.info(f"🎵 Анализ {len(all_audio_files)} аудиофайлов...")
                        
                        # Создаем экспортер
                        exporter = ParameterExporter({
                            'audio': {
                                'target_lufs': -23.0,
                                'lufs_tolerance': 0.5,
                                'true_peak': -2.0,
                                'lra_max': 18.0,
                                'sample_rate': 48000
                            },
                            'export': {
                                'csv': {
                                    'delimiter': ',',
                                    'encoding': 'utf-8'
                                },
                                'html': {
                                    'include_css': True,
                                    'include_timestamp': True
                                }
                            }
                        })
                        
                        logger.debug(f"output_dir = {output_dir}")
                        
                        # Экспортируем в CSV, HTML и TXT
                        export_results = exporter.analyze_and_export(
                            audio_files=all_audio_files,
                            output_dir=str(output_dir),
                            formats=['csv', 'html', 'txt'],
                            report_name="audio_analysis"
                        )
                        
                        logger.debug(f"export_results = {export_results}")
                        logger.info(f"✅ Аудио анализ экспортирован: {export_results}")
                        
                        # Сохраняем пути к файлам
                        tech_info['audio_analysis_csv'] = export_results.get('csv')
                        tech_info['audio_analysis_html'] = export_results.get('html')
                        tech_info['audio_analysis_txt'] = export_results.get('txt')
                    else:
                        logger.debug(f"all_audio_files is EMPTY")
                        logger.warning("⚠️ Аудиофайлы не найдены для анализа")
                        # Попробуем добавить видео файлы для анализа (они содержат аудио дорожки)
                        if video_files:
                            logger.debug(f"Trying video files for audio analysis: {video_files}")
                            logger.info(f"🎵 Видео файлы будут использованы для аудио анализа: {len(video_files)}")
                            # Добавляем видео файлы в анализ
                            all_audio_files = video_files
                        
                except Exception as e:
                    import traceback
                    logger.debug(f"Exception occurred: {e}")
                    logger.debug(f"Traceback: {traceback.format_exc()}")
                    logger.error(f"❌ Ошибка аудио анализа: {e}")
                    logger.error(f"❌ Traceback: {traceback.format_exc()}")
                
                logger.debug(f"Finished audio analysis step")
            else:
                logger.debug("PyLoudNorm analysis SKIPPED (checkbox not checked)")
                logger.info("PyLoudNorm анализ пропущен (галочка не установлена)")
            
            # Импорт проблем из CSV
            self.status_update.emit("📊 Импорт проблем из CSV...")
            
            issues = []
            if csv_files:
                csv_file = csv_files[0]
                logger.info(f"Импорт CSV: {csv_file}")
                importer = CSVImporter()
                issues = importer.import_issues(csv_file)
                logger.info(f"✅ Импортировано проблем: {len(issues)}")
            
            self.progress_update.emit(70)
            
            # Генерация заключений
            self.status_update.emit("📝 Генерация заключений...")
            
            if self.report_type == "me_ours":
                technical_conclusion = "По технической оценке нареканий не выявлено."
                subjective_conclusion = "Ниже предоставлен список внесенных изменений:"
                logger.info("M&E наши работы: идеальные заключения")
            elif self.report_type == "tifflo":
                technical_conclusion = "По технической оценке нареканий не выявлено."
                subjective_conclusion = "По субъективной оценке нареканий не обнаружено."
                logger.info("TIFFLO: идеальный отчет")
            else:
                params = tech_info.get('params', {})
                technical_conclusion = self.app.conclusion_gen.generate_technical_conclusion(tech_info, params, self.report_type)
                subjective_conclusion = self.app.conclusion_gen.generate_subjective_conclusion(issues)
                logger.info("Стандартная генерация заключений")
            
            self.progress_update.emit(85)
            
            # Генерация отчета
            self.status_update.emit("📄 Создание отчета...")
            
            report_path = output_dir / f"отчет_{base_name}_rus.docx"
            
            logger.info(f"Генерация отчета: {report_path}")
            logger.info(f"PDF 2.0: {copied_pdf_20}")
            logger.info(f"PDF 5.1: {copied_pdf_51}")
            
            # Логируем tech_info для диагностики
            logger.info("=== TECH_INFO CONTENTS ===")
            if tech_info:
                for key, value in tech_info.items():
                    if isinstance(value, dict):
                        logger.info(f"  {key}: dict with keys = {list(value.keys())}")
                        if 'lufs' in value:
                            logger.info(f"    -> LUFS: {value['lufs']}, Peak: {value.get('true_peak')}, LRA: {value.get('lra')}")
                    else:
                        logger.info(f"  {key}: {type(value).__name__} = {value}")
            else:
                logger.info("  tech_info is EMPTY!")
            logger.info("=== END TECH_INFO ===")
            
            # Генерируем отчет через единый генератор
            prepared_by = getattr(self.app, 'prepared_by', '')
            self.app.report_gen.create_exact_report(
                issues=issues,
                output_path=str(report_path),
                tech_info=tech_info,
                pdf_20_path=copied_pdf_20,
                pdf_51_path=copied_pdf_51,
                conclusion_technical=technical_conclusion,
                conclusion_subjective=subjective_conclusion,
                report_type=self.report_type,
                prepared_by=prepared_by
            )
            
            self.progress_update.emit(100)
            self.status_update.emit("✅ Готово!")
            
            logger.info(f"✅ Отчет успешно создан: {report_path}")
            
            # Показываем информацию о созданных файлах
            files_in_output = list(output_dir.glob('*'))
            logger.info(f"Файлы в папке отчета ({len(files_in_output)}):")
            for f in files_in_output:
                logger.info(f"  - {f.name}")
            
            success_msg = f"✅ Отчет создан!\n📁 Папка: {output_dir.name}\n📄 Файлов: {len(files_in_output)}"
            self.finished.emit(True, success_msg)
            
            logger.info(f"=== ЗАВЕРШЕНО УСПЕШНО ===")
            
        except Exception as e:
            logger.error(f"Ошибка при обработке: {e}", exc_info=True)
            self.finished.emit(False, f"Ошибка: {str(e)}")


class BeastApp(QMainWindow):
    """Главное окно приложения с компактным Material You интерфейсом"""
    
    def __init__(self):
        super().__init__()
        
        # Инициализация компонентов (как в v5.11)
        self.report_gen = ExactReportGenerator()
        self.pdf_extractor = PDFExtractor()
        self.conclusion_gen = ConclusionGenerator(use_llm=False)
        self.csv_importer = CSVImporter()
        self.tech_extractor = TechnicalInfoExtractor()
        
        # Хранилище файлов
        self.files_data = {
            'audio': [],
            'video': [],
            'csv': [],
            'pdf': []
        }
        
        self.init_ui()
    
    def init_ui(self):
        """Инициализация компактного Material You интерфейса"""
        self.setWindowTitle("Beast Auto Reporter - Material You")
        self.setGeometry(100, 100, 600, 650)
        
        # Установка Material You темы
        self.apply_material_you_theme()
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Основной layout с компактными отступами
        layout = QVBoxLayout()
        layout.setSpacing(12)
        layout.setContentsMargins(16, 16, 16, 16)
        
        # Заголовок
        title_label = QLabel("🎵 Beast Auto Reporter")
        title_label.setFont(QFont("SF Pro Display", 18, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("color: #1C1B1F; margin-bottom: 8px;")
        layout.addWidget(title_label)
        
        # Scroll area для всех элементов
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setSpacing(12)
        scroll_layout.setContentsMargins(0, 0, 0, 0)
        
        # Карточка выбора типа отчета
        report_card = self.create_compact_card("📋 Тип отчета", self.create_report_type_section())
        scroll_layout.addWidget(report_card)
        
        # Карточка AI и PyLoudNorm опций
        options_card = self.create_compact_card("⚙️ Опции", self.create_options_section())
        scroll_layout.addWidget(options_card)
        
        # Карточка имени подготовившего
        name_card = self.create_compact_card("👤 Подготовил", self.create_name_section())
        scroll_layout.addWidget(name_card)
        
        # Карточка выбора папки
        folder_card = self.create_compact_card("📁 Папка отчета", self.create_output_folder_section())
        scroll_layout.addWidget(folder_card)
        
        # Компактная зона перетаскивания
        drop_card = self.create_compact_card("📥 Файлы", self.create_drop_section())
        scroll_layout.addWidget(drop_card)
        
        # Список файлов
        files_card = self.create_compact_card("📂 Добавленные", self.create_files_list_section())
        scroll_layout.addWidget(files_card)
        
        # Кнопки управления
        buttons_widget = self.create_buttons_section()
        scroll_layout.addWidget(buttons_widget)
        
        # Прогресс бар
        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: none;
                border-radius: 6px;
                background-color: #F1F3F4;
                height: 8px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #6200EE;
                border-radius: 6px;
            }
        """)
        self.progress_bar.setVisible(False)
        scroll_layout.addWidget(self.progress_bar)
        
        # Статус
        self.status_label = QLabel("")
        self.status_label.setFont(QFont("SF Pro Text", 10))
        self.status_label.setAlignment(Qt.AlignCenter)
        self.status_label.setStyleSheet("color: #666666;")
        scroll_layout.addWidget(self.status_label)
        
        scroll_area.setWidget(scroll_content)
        layout.addWidget(scroll_area)
        
        central_widget.setLayout(layout)
    
    def apply_material_you_theme(self):
        """Применение Material You темы"""
        # Палитра Material You
        palette = QPalette()
        palette.setColor(QPalette.Window, QColor(255, 255, 255))  # Surface
        palette.setColor(QPalette.WindowText, QColor(28, 27, 31))  # On Surface
        palette.setColor(QPalette.Base, QColor(255, 255, 255))  # Surface
        palette.setColor(QPalette.AlternateBase, QColor(242, 242, 242))  # Surface Variant
        palette.setColor(QPalette.ToolTipBase, QColor(28, 27, 31))  # On Surface
        palette.setColor(QPalette.ToolTipText, QColor(28, 27, 31))  # On Surface
        palette.setColor(QPalette.Text, QColor(28, 27, 31))  # On Surface
        palette.setColor(QPalette.Button, QColor(255, 255, 255))  # Primary Container
        palette.setColor(QPalette.ButtonText, QColor(28, 27, 31))  # On Primary Container
        palette.setColor(QPalette.BrightText, QColor(255, 255, 255))  # On Primary
        palette.setColor(QPalette.Highlight, QColor(98, 0, 238))  # Primary
        palette.setColor(QPalette.HighlightedText, QColor(255, 255, 255))  # On Primary
        self.setPalette(palette)
    
    def create_compact_card(self, title, content_widget):
        """Создание компактной карточки Material You"""
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: #FFFFFF;
                border-radius: 16px;
                border: 1px solid #E0E0E0;
                padding: 0px;
            }
        """)
        card_layout = QVBoxLayout()
        card_layout.setContentsMargins(16, 12, 16, 12)
        card_layout.setSpacing(8)
        
        if title:
            card_title = QLabel(title)
            card_title.setFont(QFont("SF Pro Display", 12, QFont.Medium))
            card_title.setStyleSheet("color: #1C1B1F; background: transparent; border: none;")
            card_layout.addWidget(card_title)
        
        card_layout.addWidget(content_widget)
        card.setLayout(card_layout)
        return card
    
    def create_report_type_section(self):
        """Создание компактной секции выбора типа отчета"""
        widget = QWidget()
        layout = QHBoxLayout()
        layout.setSpacing(8)
        
        self.report_type_main = QRadioButton("📋 Осн.")
        self.report_type_main.setFont(QFont("SF Pro Text", 10))
        self.report_type_main.setChecked(True)
        self.report_type_main.setStyleSheet("background: transparent; border: none;")
        layout.addWidget(self.report_type_main)
        
        self.report_type_me = QRadioButton("🎧 M&E")
        self.report_type_me.setFont(QFont("SF Pro Text", 10))
        self.report_type_me.setStyleSheet("background: transparent; border: none;")
        layout.addWidget(self.report_type_me)
        
        self.report_type_me_ours = QRadioButton("✔️ Наш")
        self.report_type_me_ours.setFont(QFont("SF Pro Text", 10))
        self.report_type_me_ours.setStyleSheet("background: transparent; border: none; color: #1B5E20;")
        layout.addWidget(self.report_type_me_ours)
        
        self.report_type_tifflo = QRadioButton("🎥 TIFFLO")
        self.report_type_tifflo.setFont(QFont("SF Pro Text", 10))
        self.report_type_tifflo.setStyleSheet("background: transparent; border: none; color: #1B5E20;")
        layout.addWidget(self.report_type_tifflo)
        
        layout.addStretch()
        widget.setLayout(layout)
        return widget
    
    def create_options_section(self):
        """Создание компактной секции опций"""
        widget = QWidget()
        layout = QHBoxLayout()
        layout.setSpacing(12)
        
        self.ai_enabled_checkbox = QCheckBox("🧠 AI")
        self.ai_enabled_checkbox.setFont(QFont("SF Pro Text", 10))
        self.ai_enabled_checkbox.setStyleSheet("background: transparent; border: none;")
        self.ai_enabled_checkbox.setChecked(False)
        self.ai_enabled_checkbox.stateChanged.connect(self.toggle_ai_generation)
        layout.addWidget(self.ai_enabled_checkbox)
        
        self.pyloudnorm_checkbox = QCheckBox("🔊 LoudNorm")
        self.pyloudnorm_checkbox.setFont(QFont("SF Pro Text", 10))
        self.pyloudnorm_checkbox.setStyleSheet("background: transparent; border: none;")
        self.pyloudnorm_checkbox.setChecked(False)
        layout.addWidget(self.pyloudnorm_checkbox)
        
        layout.addStretch()
        widget.setLayout(layout)
        return widget
    
    def create_name_section(self):
        """Создание компактной секции для ввода имени"""
        from PyQt5.QtWidgets import QLineEdit
        
        widget = QWidget()
        layout = QHBoxLayout()
        layout.setSpacing(8)
        
        label = QLabel("Имя:")
        label.setFont(QFont("SF Pro Text", 10, QFont.Medium))
        label.setStyleSheet("background: transparent; border: none; color: #1C1B1F;")
        layout.addWidget(label)
        
        self.name_input = QLineEdit()
        self.name_input.setPlaceholderText("Ваше имя")
        self.name_input.setFont(QFont("SF Pro Text", 10))
        self.name_input.setStyleSheet("""
            QLineEdit {
                background: #F8F9FA;
                border: 1px solid #DEE2E6;
                border-radius: 8px;
                padding: 6px 12px;
                color: #1C1B1F;
            }
            QLineEdit:focus {
                border: 1px solid #6200EE;
            }
        """)
        self.name_input.setMaximumWidth(200)
        
        # Загружаем сохраненное имя
        self.load_saved_name()
        self.name_input.textChanged.connect(self.save_name)
        self.output_folder_path = None
        
        layout.addWidget(self.name_input)
        layout.addStretch()
        widget.setLayout(layout)
        return widget
    
    def create_output_folder_section(self):
        """Компактная секция выбора папки сохранения"""
        widget = QWidget()
        layout = QHBoxLayout()
        layout.setSpacing(8)
        
        label = QLabel("Папка:")
        label.setFont(QFont("SF Pro Text", 10))
        layout.addWidget(label)
        
        self.output_folder_label = QLabel("Рабочий стол")
        self.output_folder_label.setFont(QFont("SF Pro Text", 9))
        self.output_folder_label.setStyleSheet("color: #666666;")
        layout.addWidget(self.output_folder_label, 1)
        
        pick_btn = QPushButton("...")
        pick_btn.setFont(QFont("SF Pro Text", 9))
        pick_btn.setFixedSize(40, 28)
        pick_btn.setStyleSheet("""
            QPushButton {
                background-color: #F8F9FA;
                border: 1px solid #DEE2E6;
                border-radius: 8px;
                padding: 4px 8px;
            }
            QPushButton:hover { background-color: #E9ECEF; }
        """)
        pick_btn.clicked.connect(self.select_output_folder)
        layout.addWidget(pick_btn)
        
        clear_btn = QPushButton("×")
        clear_btn.setFont(QFont("SF Pro Text", 9))
        clear_btn.setFixedSize(30, 28)
        clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #FFFFFF;
                border: 1px solid #DEE2E6;
                border-radius: 8px;
                padding: 4px 8px;
            }
            QPushButton:hover { background-color: #F8F9FA; }
        """)
        clear_btn.clicked.connect(self.clear_output_folder)
        layout.addWidget(clear_btn)
        
        widget.setLayout(layout)
        return widget
    
    def create_drop_section(self):
        """Создание компактной зоны перетаскивания"""
        widget = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        
        self.drop_zone = DropZone()
        self.drop_zone.files_dropped.connect(self.handle_dropped_files)
        layout.addWidget(self.drop_zone)
        
        widget.setLayout(layout)
        return widget
    
    def create_files_list_section(self):
        """Создание компактного списка файлов"""
        widget = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        
        self.files_list = QListWidget()
        self.files_list.setMaximumHeight(100)
        self.files_list.setStyleSheet("""
            QListWidget {
                background-color: #F8F9FA;
                border: 1px solid #DEE2E6;
                border-radius: 8px;
                padding: 6px;
            }
            QListWidget::item {
                padding: 4px;
                border-bottom: 1px solid #E9ECEF;
            }
        """)
        layout.addWidget(self.files_list)
        
        widget.setLayout(layout)
        return widget
    
    def create_buttons_section(self):
        """Создание компактной секции кнопок"""
        widget = QWidget()
        layout = QHBoxLayout()
        layout.setSpacing(8)
        
        clear_btn = QPushButton("🗑️")
        clear_btn.setFont(QFont("SF Pro Text", 10))
        clear_btn.setFixedSize(60, 36)
        clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #F8F9FA;
                border: none;
                border-radius: 12px;
                padding: 8px 12px;
                color: #495057;
            }
            QPushButton:hover {
                background-color: #E9ECEF;
            }
        """)
        clear_btn.clicked.connect(self.clear_files)
        layout.addWidget(clear_btn)
        
        layout.addStretch()
        
        self.generate_btn = QPushButton("✨ Создать")
        self.generate_btn.setFont(QFont("SF Pro Text", 11, QFont.Medium))
        self.generate_btn.setFixedHeight(44)
        self.generate_btn.setStyleSheet("""
            QPushButton {
                background-color: #6200EE;
                border: none;
                border-radius: 12px;
                padding: 12px 24px;
                color: white;
            }
            QPushButton:hover {
                background-color: #3700B3;
            }
            QPushButton:disabled {
                background-color: #E9ECEF;
                color: #ADB5BD;
            }
        """)
        self.generate_btn.clicked.connect(self.start_processing)
        self.generate_btn.setEnabled(False)
        layout.addWidget(self.generate_btn)
        
        widget.setLayout(layout)
        return widget
    
    def load_saved_name(self):
        """Загрузка сохраненного имени"""
        config_file = Path.home() / ".beast_auto_reporter_config.txt"
        if config_file.exists():
            try:
                name = config_file.read_text().strip()
                self.name_input.setText(name)
            except:
                pass
    
    def save_name(self):
        """Сохранение имени"""
        config_file = Path.home() / ".beast_auto_reporter_config.txt"
        try:
            config_file.write_text(self.name_input.text())
        except Exception:
            pass
    
    def select_output_folder(self):
        """Выбор папки сохранения отчета"""
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для отчета")
        if folder:
            self.output_folder_path = Path(folder)
            self.output_folder_label.setText(str(self.output_folder_path.name))
            logger.info(f"Output folder selected: {self.output_folder_path}")
    
    def clear_output_folder(self):
        """Сброс выбора папки"""
        self.output_folder_path = None
        self.output_folder_label.setText("Рабочий стол")
        logger.info("Output folder reset to Desktop")
    
    def handle_dropped_files(self, files):
        """Обработка перетащенных файлов"""
        logger.info(f"Получено файлов: {len(files)}")
        
        for file_path in files:
            file_name = os.path.basename(file_path)
            file_ext = Path(file_path).suffix.lower()
            filename_lower = Path(file_path).stem.lower()
            
            # Определяем тип файла
            if file_ext in ['.wav', '.mp3', '.flac', '.aac', '.mxf', '.caf', '.aiff', '.aif']:
                self.files_data['audio'].append(file_path)
                # Определяем тип аудио для отображения
                if '20' in filename_lower and 'cens' in filename_lower and 'uncens' not in filename_lower:
                    icon = "🔈 [2.0C]"
                elif '20' in filename_lower and 'uncens' in filename_lower:
                    icon = "🔈 [2.0U]"
                elif '51' in filename_lower and 'cens' in filename_lower and 'uncens' not in filename_lower:
                    icon = "🔈 [5.1C]"
                elif '51' in filename_lower and 'uncens' in filename_lower:
                    icon = "🔈 [5.1U]"
                else:
                    icon = "🔈"
            elif file_ext in ['.mp4', '.mov', '.avi', '.mkv', '.m4v', '.webm']:
                self.files_data['video'].append(file_path)
                icon = "🎬"
            elif file_ext == '.csv':
                self.files_data['csv'].append(file_path)
                icon = "📈"
            elif file_ext == '.pdf':
                self.files_data['pdf'].append(file_path)
                # Определяем тип PDF для отображения
                if '20' in filename_lower and 'cens' in filename_lower and 'uncens' not in filename_lower:
                    icon = "📑 [2.0C]"
                elif '20' in filename_lower and 'uncens' in filename_lower:
                    icon = "📑 [2.0U]"
                elif '51' in filename_lower and 'cens' in filename_lower and 'uncens' not in filename_lower:
                    icon = "📑 [5.1C]"
                elif '51' in filename_lower and 'uncens' in filename_lower:
                    icon = "📑 [5.1U]"
                else:
                    icon = "📑"
            else:
                continue
            
            # Добавляем в список
            item = QListWidgetItem(f"{icon} {file_name}")
            item.setToolTip(file_path)
            self.files_list.addItem(item)
        
        # Обновляем доступность кнопки
        total_files = sum(len(v) for v in self.files_data.values())
        self.generate_btn.setEnabled(total_files > 0)
        
        # Показываем детальную статистику
        stats = []
        if self.files_data['audio']:
            stats.append(f"🔈 {len(self.files_data['audio'])}")
        if self.files_data['video']:
            stats.append(f"🎬 {len(self.files_data['video'])}")
        if self.files_data['csv']:
            stats.append(f"📈 {len(self.files_data['csv'])}")
        if self.files_data['pdf']:
            stats.append(f"📑 {len(self.files_data['pdf'])}")
        
        self.status_label.setText(" | ".join(stats) if stats else "Нет файлов")
    
    def clear_files(self):
        """Очистка списка файлов"""
        self.files_data = {
            'audio': [],
            'video': [],
            'csv': [],
            'pdf': []
        }
        self.files_list.clear()
        self.generate_btn.setEnabled(False)
        self.status_label.setText("")
    
    def toggle_ai_generation(self, state):
        """Переключение AI генерации"""
        enabled = state == Qt.Checked
        self.conclusion_gen.use_llm = enabled
        logger.info(f"AI генерация: {'ВКЛЮЧЕНА' if enabled else 'ВЫКЛЮЧЕНА'}")
    
    def get_report_type(self):
        """Получение выбранного типа отчета"""
        if self.report_type_me.isChecked():
            return "me"
        elif self.report_type_me_ours.isChecked():
            return "me_ours"
        elif self.report_type_tifflo.isChecked():
            return "tifflo"
        else:
            return "standard"
    
    def start_processing(self):
        """Запуск обработки файлов"""
        logger.info("=== ЗАПУСК ОБРАБОТКИ ===")
        
        # Получаем базовое имя для папки
        audio_files = self.files_data.get('audio', [])
        video_files = self.files_data.get('video', [])
        csv_files = self.files_data.get('csv', [])
        
        base_name = "отчет"
        if audio_files:
            base_name = sanitize_base_name(audio_files[0])
        elif video_files:
            base_name = sanitize_base_name(video_files[0])
        elif csv_files:
            base_name = sanitize_base_name(csv_files[0])
        
        # Папка вывода: выбранная пользователем или Рабочий стол
        if self.output_folder_path:
            output_folder = self.create_output_folder(base_name, base_dir=Path(self.output_folder_path))
            logger.info(f"Папка отчета: {output_folder}")
        else:
            output_folder = self.create_output_folder(base_name)
            logger.info(f"Папка отчета на Desktop: {output_folder}")
        
        # Отключаем кнопки
        self.generate_btn.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        # Сохраняем имя
        self.prepared_by = self.name_input.text().strip() or "Не указано"
        
        # Запускаем поток с Desktop папкой
        report_type = self.get_report_type()
        pyloudnorm_enabled = self.pyloudnorm_checkbox.isChecked()
        logger.info(f"PyLoudNorm: {pyloudnorm_enabled}, Report type: {report_type}")
        
        self.thread = ProcessingThread(self, self.files_data, report_type, str(output_folder), pyloudnorm_enabled)
        self.thread.status_update.connect(self.status_label.setText)
        self.thread.progress_update.connect(self.progress_bar.setValue)
        self.thread.finished.connect(self.processing_finished)
        self.thread.start()
        logger.info("Thread started")
    
    def processing_finished(self, success, message):
        """Завершение обработки"""
        self.progress_bar.setVisible(False)
        self.generate_btn.setEnabled(True)
        
        if success:
            self.status_label.setText("✅ Готово!")
            logger.info(f"=== SUCCESS ===")
            logger.info(f"{message}")
            QMessageBox.information(self, "Готово!", "Отчет успешно создан!")
            # Очищаем файлы после успешной генерации
            self.clear_files()
        else:
            error_text = message if isinstance(message, str) else "Неизвестная ошибка"
            self.status_label.setText(f"❌ Ошибка: {error_text[:50]}...")
            logger.error(f"=== ERROR ===")
            logger.error(f"{error_text}")
            QMessageBox.critical(self, "Ошибка", error_text)
    
    def extract_base_name(self, filename: str) -> str:
        """Извлечение базового названия из имени файла (из v5.11)"""
        return sanitize_base_name(filename)
    
    def find_files_in_folder(self, folder: Path) -> dict:
        """Поиск файлов в папке (полная логика из v5.11)"""
        import re
        import subprocess
        
        def find_ffprobe():
            """Поиск ffprobe"""
            import shutil
            ffprobe_path = shutil.which('ffprobe')
            if ffprobe_path:
                return ffprobe_path
            return 'ffprobe'
        
        files = {
            'audio_20_c': None,
            'audio_20_uc': None,
            'audio_51_c': None,
            'audio_51_uc': None,
            'video': None,
            'csv': None,
            'pdf_20': None,
            'pdf_51': None,
            'pdf_20_c': None,
            'pdf_20_uc': None,
            'pdf_51_c': None,
            'pdf_51_uc': None,
            'params': None,
            'all_files': []
        }
        
        for file_path in folder.iterdir():
            if file_path.is_file():
                files['all_files'].append(file_path)
                
                name_lower = file_path.name.lower()
                
                # CSV файл
                if file_path.suffix.lower() == '.csv':
                    files['csv'] = file_path
                
                # Параметры.txt
                elif 'параметры' in name_lower or 'parametry' in name_lower:
                    files['params'] = file_path
                
                # Аудио файлы (ВАЖНО: проверяем 5.1 ПЕРЕД 2.0!)
                elif file_path.suffix.lower() in ['.wav', '.mp3', '.aiff', '.flac']:
                    channel_hint = detect_channel_from_name(name_lower)
                    channels = None
                    try:
                        ffprobe_cmd = find_ffprobe()
                        result = subprocess.run(
                            [ffprobe_cmd, '-v', 'error', '-select_streams', 'a:0', 
                             '-show_entries', 'stream=channels', '-of', 
                             'default=noprint_wrappers=1:nokey=1', str(file_path)],
                            capture_output=True, text=True, timeout=5
                        )
                        if result.returncode == 0 and result.stdout.strip():
                            channels = int(result.stdout.strip())
                    except Exception:
                        channels = None
                    
                    channel_meta = "51" if channels and channels >= 6 else ("20" if channels == 2 else None)
                    channel_final = channel_meta or channel_hint
                    cens_state = detect_cens_state(name_lower)
                    suffix = "uc" if cens_state == "uc" else "c"
                    
                    if channel_final == "51":
                        files[f'audio_51_{suffix}'] = file_path
                    elif channel_final == "20":
                        files[f'audio_20_{suffix}'] = file_path
                
                # Видео
                elif file_path.suffix.lower() in ['.mp4', '.mov', '.mkv', '.avi', '.mxf', '.m4v', '.webm', '.flv']:
                    files['video'] = file_path
                
                # PDF файлы
                elif file_path.suffix.lower() == '.pdf':
                    # Проверяем на 5.1
                    if ('5.1' in file_path.name or 
                        '_51_' in name_lower or '_51.' in name_lower or
                        name_lower.endswith('_51') or
                        '_50_' in name_lower or '5.0' in file_path.name):
                        
                        if 'uncens' in name_lower or '_uc' in name_lower:
                            files['pdf_51_uc'] = file_path
                        else:
                            files['pdf_51_c'] = file_path
                        if not files['pdf_51']:
                            files['pdf_51'] = file_path
                    
                    # Проверяем на 2.0
                    elif ('2.0' in file_path.name or 
                          '_20_' in name_lower or '_20.' in name_lower or
                          name_lower.endswith('_20') or
                          '_21_' in name_lower or '2.1' in file_path.name):
                        
                        if 'uncens' in name_lower or '_uc' in name_lower:
                            files['pdf_20_uc'] = file_path
                        else:
                            files['pdf_20_c'] = file_path
                        if not files['pdf_20']:
                            files['pdf_20'] = file_path
                    
                    # Автоопределение
                    else:
                        is_uncens = 'uncens' in name_lower or '_uc' in name_lower
                        if not files['pdf_20_c'] and not files['pdf_20_uc']:
                            files['pdf_20_uc' if is_uncens else 'pdf_20_c'] = file_path
                            if not files['pdf_20']:
                                files['pdf_20'] = file_path
                        elif not files['pdf_51_c'] and not files['pdf_51_uc']:
                            files['pdf_51_uc' if is_uncens else 'pdf_51_c'] = file_path
                            if not files['pdf_51']:
                                files['pdf_51'] = file_path
        
        return files
    
    def create_output_folder(self, base_name: str, base_dir: Path = None) -> Path:
        """Создание выходной папки (из v5.11)"""
        folder_name = f"отчет_{base_name}"
        
        desktop = Path.home() / "Desktop"
        root_dir = base_dir or desktop
        output_folder = root_dir / folder_name
        
        counter = 1
        while output_folder.exists():
            output_folder = root_dir / f"{folder_name}_{counter}"
            counter += 1
        
        output_folder.mkdir(parents=True, exist_ok=True)
        
        return output_folder


def main():
    app = QApplication(sys.argv)
    
    # Устанавливаем стиль приложения
    app.setStyle("Fusion")
    
    window = BeastApp()
    window.show()
    
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
