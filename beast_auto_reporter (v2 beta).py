"""
Beast Auto Reporter v2 beta - macOS Style
Нативный macOS-стиль интерфейс с поддержкой перетаскивания файлов
"""

import sys
import os
import logging
import shutil
import subprocess
from pathlib import Path
from datetime import datetime
import re
import io
import csv
import json
import tempfile
import importlib.util
import math
import html
import yaml

# Добавляем корневую директорию в путь (нужно для локальных модулей src при запуске из любого cwd)
ROOT_DIR = Path(__file__).resolve().parent
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

APP_VERSION = "2.1.0"


def load_app_config() -> dict:
    """Загружает конфиг приложения из config/settings.yaml."""
    config_path = ROOT_DIR / "config" / "settings.yaml"
    try:
        if config_path.exists():
            with config_path.open("r", encoding="utf-8") as fh:
                return yaml.safe_load(fh) or {}
    except Exception as exc:
        logging.getLogger(__name__).warning(f"Не удалось загрузить config/settings.yaml: {exc}")
    return {}


def _maybe_add_user_site() -> None:
    # В замороженном приложении все зависимости уже внутри бандла
    if getattr(sys, 'frozen', False):
        return
    try:
        import site
        user_site = site.getusersitepackages()
        if user_site and user_site not in sys.path:
            site.addsitedir(user_site)
        # Если модуль все еще не находится, пробуем подключить user-site других версий Python
        if importlib.util.find_spec("PyQt5") is None:
            user_base = Path.home() / "Library" / "Python"
            if user_base.exists():
                for ver_dir in sorted(user_base.iterdir()):
                    if not ver_dir.is_dir():
                        continue
                    candidate = ver_dir / "lib" / "python" / "site-packages"
                    if candidate.exists() and (candidate / "PyQt5").exists():
                        if str(candidate) not in sys.path:
                            site.addsitedir(str(candidate))
    except Exception:
        pass


_maybe_add_user_site()


def _ensure_pyqt5_sip_compat() -> None:
    if getattr(sys, 'frozen', False):
        return
    try:
        if importlib.util.find_spec("PyQt5.sip") is None:
            if importlib.util.find_spec("PyQt5_sip") is not None:
                import PyQt5_sip
                sys.modules.setdefault("PyQt5.sip", PyQt5_sip)
    except Exception:
        pass


_ensure_pyqt5_sip_compat()


def _can_import_pyqt5(candidate: str) -> bool:
    try:
        result = subprocess.run(
            [candidate, "-c", "import PyQt5.sip"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=10
        )
        return result.returncode == 0
    except Exception:
        return False


def _maybe_reexec_with_pyqt5() -> None:
    # В замороженном приложении нельзя перезапускаться через другой Python
    if getattr(sys, 'frozen', False):
        return
    if os.environ.get("BEAST_PYQT5_BOOTSTRAP") == "1":
        return
    if importlib.util.find_spec("PyQt5.sip") is not None:
        return

    candidates = []
    candidates.append(sys.executable)
    candidates.append("/Library/Developer/CommandLineTools/usr/bin/python3")
    candidates.append("/usr/bin/python3")

    which_py = shutil.which("python3")
    if which_py:
        candidates.append(which_py)

    repo_root = Path(__file__).resolve().parent
    venv_py = repo_root / "venv" / "bin" / "python"
    if venv_py.exists():
        candidates.append(str(venv_py))

    # Deduplicate while preserving order
    seen = set()
    deduped = []
    for c in candidates:
        if not c:
            continue
        key = str(Path(c).resolve())
        if key in seen:
            continue
        seen.add(key)
        deduped.append(c)

    current = str(Path(sys.executable).resolve())
    for candidate in deduped:
        try:
            if str(Path(candidate).resolve()) == current:
                continue
        except Exception:
            pass
        if _can_import_pyqt5(candidate):
            os.environ["BEAST_PYQT5_BOOTSTRAP"] = "1"
            os.execv(candidate, [candidate] + sys.argv)


_maybe_reexec_with_pyqt5()


def _require_module(module_name: str, install_hint: str) -> None:
    if importlib.util.find_spec(module_name) is None:
        raise SystemExit(
            f"Модуль '{module_name}' не установлен.\n"
            f"{install_hint}\n"
            "Или запустите приложение из собранного .app"
        )


# Проверяем наличие зависимостей до импорта модулей, которые их используют
# В замороженном приложении все зависимости уже внутри бандла
if not getattr(sys, 'frozen', False):
    _require_module("docx", "Установите зависимости: pip install -r requirements_pinned.txt")
    _require_module("PyQt5", "Установите зависимости: pip install -r requirements_pinned.txt")
    _require_module("fitz", "Установите зависимости: pip install -r requirements_pinned.txt")

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QRadioButton, QTextEdit, QTextBrowser, QProgressBar,
    QCheckBox, QListWidget, QListWidgetItem, QFrame, QMessageBox, QFileDialog,
    QGroupBox, QScrollArea, QSizePolicy, QGraphicsDropShadowEffect,
    QDialog, QDialogButtonBox, QFormLayout, QPlainTextEdit, QComboBox,
    QTableWidget, QTableWidgetItem, QProgressDialog, QLineEdit,
    QTreeWidget, QTreeWidgetItem, QInputDialog, QCompleter, QAbstractButton
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QMimeData, QPoint, QEventLoop, QMetaObject, Q_ARG, QTimer, QRectF, QSize, QObject, QEvent, QFileSystemWatcher
from PyQt5.QtGui import QFont, QDragEnterEvent, QDropEvent, QPalette, QColor, QIcon, QPainter, QPainterPath, QBrush, QPen, QPixmap, QRadialGradient, QImage

import fitz

from src.exact_report_generator import ExactReportGenerator
from src.technical_info_extractor import TechnicalInfoExtractor, format_fps
from src.csv_importer import CSVImporter
from src.pdf_extractor import PDFExtractor
from src.conclusion_generator import ConclusionGenerator
from src.parameter_exporter import ParameterExporter
from src.audio_metrics import parse_ebur128_output, run_ffmpeg_ebur128, measure_true_peak_precise
from src.update_checker import (
    check_for_update,
    download_update_asset,
    load_skipped_version,
    save_skipped_version,
)
from src.icons import make_icon, make_icon_pixmap

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


REPORT_TYPE_SIGNAL_EXTENSIONS = {
    ".wav", ".mp3", ".flac", ".aac", ".mxf", ".caf", ".aiff", ".aif",
    ".mp4", ".mov", ".avi", ".mkv", ".m4v", ".webm", ".pdf", ".csv",
}

REPORT_TYPE_SCORE_PATTERNS = {
    "dcp": (
        (re.compile(r"(?:^|\s)dcp(?:\s|$)", re.IGNORECASE), 10, "dcp"),
        (re.compile(r"(?:^|\s)dcp\s*omatic(?:\s|$)", re.IGNORECASE), 8, "dcp-omatic"),
    ),
    "me": (
        (re.compile(r"(?:^|\s)(?:me|mne|m\s*e|m\s*n\s*e|m\s+and\s+e)(?:\s|$)", re.IGNORECASE), 8, "me/mne"),
        (re.compile(r"(?:^|\s)music\s+and\s+effects?(?:\s|$)", re.IGNORECASE), 10, "music and effects"),
    ),
    "standard": (
        (re.compile(r"(?:^|\s)(?:main|standard|std)(?:\s|$)", re.IGNORECASE), 6, "main/standard/std"),
        (re.compile(r"(?:^|\s)ad(?:\s|$)", re.IGNORECASE), 8, "ad"),
        (re.compile(r"(?:^|\s)audio\s+description(?:\s|$)", re.IGNORECASE), 10, "audio description"),
        (re.compile(r"основ", re.IGNORECASE), 6, "основ"),
        (re.compile(r"(?:^|\s)(?:full\s*mix|fullmix|mix|master)(?:\s|$)", re.IGNORECASE), 3, "mix/master"),
        (re.compile(r"(?:^|\s)tifflo(?:\s|$)", re.IGNORECASE), 10, "tifflo"),
        (re.compile(r"(?:^|\s)идеал(?:ьн\w*)?(?:\s|$)", re.IGNORECASE), 4, "идеал"),
        (re.compile(r"(?:^|\s)ideal(?:ly)?(?:\s|$)", re.IGNORECASE), 4, "ideal"),
    ),
}

REPORT_TYPE_OURS_PATTERNS = (
    (re.compile(r"(?:^|\s)(?:ours|our|internal|inhouse)(?:\s|$)", re.IGNORECASE), 4, "ours/internal"),
    (re.compile(r"наши|внутрен", re.IGNORECASE), 4, "наши/внутренний"),
)

REPORT_TYPE_PRIORITY = ("me_ours", "dcp", "me", "standard")


CAT_STATE_ASSETS = {
    "sleeping": ROOT_DIR / "assets" / "cat-states" / "cat-sleeping.png",
    "hover": ROOT_DIR / "assets" / "cat-states" / "cat-hover.png",
    "ready": ROOT_DIR / "assets" / "cat-states" / "cat-ready.png",
    "working": ROOT_DIR / "assets" / "cat-states" / "cat-working.png",
    "done": ROOT_DIR / "assets" / "cat-states" / "cat-done.png",
}

_CHANNEL_51_RE = re.compile(
    r'(^|[^0-9])(?:5\.1|5_1|5-1|5\.0|51)([^0-9]|$)|(^|[^a-z0-9])(?:surround|6\s*ch|6ch|6\s*channel)([^a-z0-9]|$)',
    re.IGNORECASE
)
_CHANNEL_20_RE = re.compile(
    r'(^|[^0-9])(?:2\.0|2_0|2-0|2\.1|20)([^0-9]|$)|(^|[^a-z0-9])(?:stereo|2\s*ch|2ch|2\s*channel|lr|l\s*r)([^a-z0-9]|$)',
    re.IGNORECASE
)
_UNCENS_RE = re.compile(r'(^|[^a-z0-9])(?:uncens|uncensored|uc)([^a-z0-9]|$)', re.IGNORECASE)
_CENS_RE = re.compile(r'(^|[^a-z0-9])(?:cens|censored|c)([^a-z0-9]|$)', re.IGNORECASE)
_INCORRECT_AUDIO_TYPE_PATTERNS = [
    r'_50_(cens|uncens|mix|c|uc)',
    r'_50$',
    r'_21_(cens|uncens|mix|c|uc)',
    r'_21$',
    r'_5\.0_(cens|uncens|mix)',
    r'_2\.1_(cens|uncens|mix)',
]


def detect_channel_from_name(name: str):
    name_lower = name.lower()
    if _CHANNEL_51_RE.search(name_lower):
        return "51"
    if _CHANNEL_20_RE.search(name_lower):
        return "20"
    return None


def detect_cens_state(name: str):
    name_lower = name.lower()
    if _UNCENS_RE.search(name_lower):
        return "uc"
    if _CENS_RE.search(name_lower):
        return "c"
    return None


def has_incorrect_audio_marker(name: str) -> bool:
    name_lower = name.lower()
    for pattern in _INCORRECT_AUDIO_TYPE_PATTERNS:
        if re.search(pattern, name_lower, re.IGNORECASE):
            return True
    return False


def order_pdfs_by_channel(pdf_paths: list) -> list:
    if not pdf_paths:
        return pdf_paths
    scored = []
    for idx, p in enumerate(pdf_paths):
        name = Path(p).stem
        channel = detect_channel_from_name(name) or ""
        cens_state = detect_cens_state(name) or ""
        # channel order: 20 -> 51 -> unknown
        ch_score = 2
        if channel == "20":
            ch_score = 0
        elif channel == "51":
            ch_score = 1
        # cens order: cens -> uncens -> unknown
        cens_score = 2
        if cens_state == "c":
            cens_score = 0
        elif cens_state == "uc":
            cens_score = 1
        scored.append((ch_score, cens_score, idx, p))
    scored.sort(key=lambda t: (t[0], t[1], t[2]))
    return [p for _, __, ___, p in scored]


def sanitize_base_name(name: str) -> str:
    base = Path(name).stem
    # Защищаем даты (YYYYMMDD, YYYY_MM_DD, DD_MM_YYYY и т.п.) от удаления
    _date_placeholder = {}
    def _protect_date(m):
        key = f"\x00DATE{len(_date_placeholder)}\x00"
        _date_placeholder[key] = m.group(0)
        return key
    # Сначала компактные даты YYYYMMDD (8 цифр подряд с правдоподобными месяцем/днём)
    base = re.sub(r'(?:19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])', _protect_date, base)
    base = re.sub(r'(?:0[1-9]|[12]\d|3[01])(?:0[1-9]|1[0-2])(?:19|20)\d{2}', _protect_date, base)
    # Даты с разделителями: YYYY_MM_DD и DD_MM_YYYY (строгие значения месяцев/дней)
    base = re.sub(
        r'(?:19|20)\d{2}[_.\- ](?:0[1-9]|1[0-2])[_.\- ](?:0[1-9]|[12]\d|3[01])'
        r'|(?:0[1-9]|[12]\d|3[01])[_.\- ](?:0[1-9]|1[0-2])[_.\- ](?:19|20)\d{2}',
        _protect_date, base
    )
    # Удаляем только служебные маркеры каналов/цензуры
    patterns = [
        r'(?i)[ _\.-]+(5\.?1|5_1|5-1|5\.0|51)(?=$|[ _\.-])',
        r'(?i)[ _\.-]+(2\.?0|2_0|2-0|2\.1|20|stereo)(?=$|[ _\.-])',
        r'(?i)[ _\.-]+(uncens|uncensored|uc)(?=$|[ _\.-])',
        r'(?i)[ _\.-]+(cens|censored|c)(?=$|[ _\.-])',
    ]
    for pat in patterns:
        base = re.sub(pat, '', base)
    # Восстанавливаем даты
    for key, val in _date_placeholder.items():
        base = base.replace(key, val)
    # Убираем только висячие разделители по краям
    base = re.sub(r'^[ _.\-]+|[ _.\-]+$', '', base)
    return base or "отчет"


def normalize_stem(name: str) -> str:
    stem = Path(name).stem.lower()
    stem = re.sub(r'[^a-z0-9]+', '_', stem)
    stem = re.sub(r'_+', '_', stem).strip('_')
    return stem


def _matching_base_name(name: str) -> str:
    """Базовое имя для сравнения файлов (без расширения, маркеров каналов/цензуры/типа).

    В отличие от sanitize_base_name возвращает нормализованную (lowercase, _) форму,
    подходящую для сопоставления audio / pdf / csv между собой.
    """
    base = sanitize_base_name(name).lower()
    base = re.sub(r'(?i)[ _\.-]+(audio|video|pdf|csv|markers?|маркер[ыа]?)(?=$|[ _\.-])', '', base)
    base = re.sub(r'[^a-z0-9а-яё]+', '_', base, flags=re.IGNORECASE)
    base = re.sub(r'_+', '_', base).strip('_')
    return base


from src.app_paths import CONFIG_DIR, ensure_parent_dir, migrate_legacy_config_file  # noqa: E402
from src import secret_store  # noqa: E402
from src.file_matching import _levenshtein, _bases_match  # noqa: E402
from src.report_filename import parse_report_filename  # noqa: E402
from src.report_uploader import (  # noqa: E402
    save_queue_state, load_queue_state, REPORTS_ROOT,
    fallback_series_key, remember_series_alias, remember_uploaded_report, resolve_manual_pick_target,
)
from src.yandex_ui.helpers import (  # noqa: E402
    _stop_thread, _quick_look_preview, _format_disk_modified_date, _send_system_notification,
    _play_sound,
)
from src.yandex_ui.threads import (  # noqa: E402
    _FallbackFolderFindThread, _IntegrityCheckThread, _MkdirThread,
    YandexDiskFindVersionsThread, YandexDiskFolderVersionsThread,
    YandexDiskCompareThread, YandexDiskUploadThread,
    YandexDiskTokenCheckThread,
)
from src.yandex_ui.edit_sync import YandexEditSyncController  # noqa: E402
from src.yandex_ui.dialogs import (  # noqa: E402
    YandexUploadDiffDialog, YandexVersionPickerDialog, YandexFolderPickerDialog,
    YandexDiskBrowserDialog, YandexUploadQueueDialog, SeriesAliasesDialog,
)
from src.yandex_ui.queue_manager import YandexUploadQueueManager  # noqa: E402


def validate_file_consistency(files_data: dict) -> list:
    """Проверяет, что audio / pdf / csv файлы относятся к одному и тому же материалу.

    Базовые имена сравниваются не только на точное равенство: небольшая
    опечатка в достаточно длинном имени (расстояние Левенштейна <= 2, длина
    >= 10 символов) не считается несоответствием — см. `_bases_match()`.

    Возвращает список строк-предупреждений. Пустой список = всё в порядке.
    """
    items: list = []

    def _add(path_str: str, category: str):
        base = _matching_base_name(str(path_str))
        if not base:
            return
        items.append((base, category, Path(str(path_str)).name))

    for f in files_data.get('audio', []) or []:
        _add(f, 'audio')
    for f in files_data.get('pdf', []) or []:
        _add(f, 'pdf')
    for f in files_data.get('csv', []) or []:
        _add(f, 'csv')

    if len(items) <= 1:
        return []

    # Кластеризуем по похожести базового имени, а не по точному совпадению
    clusters: list = []
    for base, category, name in items:
        target_cluster = next((c for c in clusters if _bases_match(base, c["rep"])), None)
        if target_cluster is not None:
            target_cluster["items"].append((category, name))
        else:
            clusters.append({"rep": base, "items": [(category, name)]})

    if len(clusters) <= 1:
        return []

    lines = [f"Файлы относятся к разным материалам ({len(clusters)} групп(ы)):"]
    for cluster in clusters:
        lines.append(f"\n• «{cluster['rep']}»")
        for category, name in cluster["items"]:
            lines.append(f"    [{category}] {name}")
    return lines


def normalize_report_type_name(name: str) -> str:
    stem = Path(name).stem.lower().replace("&", " and ")
    stem = re.sub(r'[^0-9a-zа-яё]+', ' ', stem, flags=re.IGNORECASE)
    stem = re.sub(r'\s+', ' ', stem).strip()
    return f" {stem} " if stem else ""


def score_report_type_name(name: str) -> dict:
    normalized = normalize_report_type_name(name)
    scores = {report_type: 0 for report_type in REPORT_TYPE_PRIORITY}
    reasons = {report_type: [] for report_type in REPORT_TYPE_PRIORITY}

    if not normalized:
        return {
            "normalized": normalized,
            "scores": scores,
            "reasons": reasons,
        }

    for report_type, rules in REPORT_TYPE_SCORE_PATTERNS.items():
        for pattern, points, reason in rules:
            if pattern.search(normalized):
                scores[report_type] += points
                reasons[report_type].append(reason)

    ours_score = 0
    ours_reasons = []
    for pattern, points, reason in REPORT_TYPE_OURS_PATTERNS:
        if pattern.search(normalized):
            ours_score += points
            ours_reasons.append(reason)

    if scores["me"] > 0 and ours_score > 0:
        scores["me_ours"] = scores["me"] + ours_score + 2
        reasons["me_ours"] = reasons["me"] + ours_reasons

    return {
        "normalized": normalized,
        "scores": scores,
        "reasons": reasons,
    }


def detect_report_type_from_files(files: list):
    relevant_files = []
    aggregate_scores = {report_type: 0 for report_type in REPORT_TYPE_PRIORITY}
    per_file = []

    for file_path in files:
        suffix = Path(file_path).suffix.lower()
        if suffix not in REPORT_TYPE_SIGNAL_EXTENSIONS:
            continue

        relevant_files.append(file_path)
        analysis = score_report_type_name(file_path)
        per_file.append((file_path, analysis))
        for report_type, score in analysis["scores"].items():
            aggregate_scores[report_type] += score

    if not relevant_files:
        return None, {
            "reason": "no_relevant_files",
            "scores": aggregate_scores,
            "per_file": per_file,
        }

    special_scores = {
        report_type: aggregate_scores[report_type]
        for report_type in ("me_ours", "dcp", "me")
    }
    best_special_type = max(special_scores, key=special_scores.get)
    best_special_score = special_scores[best_special_type]

    if best_special_score <= 0:
        return "standard", {
            "reason": "fallback_standard",
            "scores": aggregate_scores,
            "per_file": per_file,
        }

    ranked = sorted(
        aggregate_scores.items(),
        key=lambda item: (-item[1], REPORT_TYPE_PRIORITY.index(item[0])),
    )
    best_type, best_score = ranked[0]
    second_type, second_score = ranked[1]

    if best_score <= 0:
        return "standard", {
            "reason": "fallback_standard",
            "scores": aggregate_scores,
            "per_file": per_file,
        }

    if best_score == second_score:
        return None, {
            "reason": "ambiguous",
            "scores": aggregate_scores,
            "per_file": per_file,
            "leaders": (best_type, second_type),
        }

    return best_type, {
        "reason": "scored_match",
        "scores": aggregate_scores,
        "per_file": per_file,
    }


def generate_loudness_report_isolated(audio_file: str, pdf_path: str, timeout_sec: int = 900):
    """
    Генерирует loudness PDF в отдельном Python-процессе.
    Это защищает основной GUI процесс от падения C-библиотек.
    """
    runner_code = """
import json
import os
import sys
import traceback
from pathlib import Path

audio_file = sys.argv[1]
pdf_path = sys.argv[2]
project_root = Path(sys.argv[3])
os.chdir(str(project_root))
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))

from src.audio_metrics import generate_loudness_report

try:
    metrics = generate_loudness_report(audio_file, pdf_path)
    print(json.dumps({"ok": True, "metrics": metrics}, ensure_ascii=False))
except Exception as exc:
    print(json.dumps({
        "ok": False,
        "error": str(exc),
        "traceback": traceback.format_exc(),
    }, ensure_ascii=False))
    raise
"""
    proc = subprocess.run(
        [sys.executable, "-c", runner_code, str(audio_file), str(pdf_path), str(ROOT_DIR)],
        capture_output=True,
        text=True,
        timeout=timeout_sec,
        check=False,
    )

    parsed = None
    stdout_lines = [line.strip() for line in (proc.stdout or "").splitlines() if line.strip()]
    for line in reversed(stdout_lines):
        try:
            parsed = json.loads(line)
            break
        except Exception:
            continue

    if proc.returncode == 0 and isinstance(parsed, dict) and parsed.get("ok") is True:
        return parsed.get("metrics"), None

    error_parts = [f"exit_code={proc.returncode}"]
    if isinstance(parsed, dict):
        if parsed.get("error"):
            error_parts.append(str(parsed.get("error")))
        if parsed.get("traceback"):
            error_parts.append(str(parsed.get("traceback")))
    stderr_text = (proc.stderr or "").strip()
    if stderr_text:
        error_parts.append(stderr_text)
    return None, " | ".join(error_parts)


def measure_sample_peak_fast(audio_file: str, timeout_sec: int = 180):
    """
    Быстрое измерение SAMPLE PEAK через ffmpeg ebur128 без полного pyloudnorm-анализа.
    Возвращает dBFS (sample peak) либо fallback true peak, если sample peak недоступен.
    """
    try:
        stderr_text = run_ffmpeg_ebur128(audio_file, timeout_sec=timeout_sec)
        parsed = parse_ebur128_output(stderr_text or "")
        sample_peak = parsed.get("sample_peak_dbfs")
        if sample_peak is None:
            sample_peak = parsed.get("true_peak_dbtp")
        return sample_peak, None
    except Exception as exc:
        return None, str(exc)


class CatMascotWidget(QWidget):
    """Animated mascot renderer for the drop zone."""

    _STYLE = {
        "sleeping": {
            "size": 78,
            "stage_inner": QColor(255, 255, 255, 0),
            "stage_outer": QColor(255, 255, 255, 0),
            "floor": QColor(29, 29, 31, 28),
            "halo": None,
        },
        "hover": {
            "size": 80,
            "stage_inner": QColor(255, 255, 255, 0),
            "stage_outer": QColor(255, 255, 255, 0),
            "floor": QColor(0, 122, 255, 54),
            "halo": QColor(0, 122, 255, 42),
        },
        "ready": {
            "size": 78,
            "stage_inner": QColor(255, 255, 255, 0),
            "stage_outer": QColor(255, 255, 255, 0),
            "floor": QColor(29, 29, 31, 42),
            "halo": None,
        },
        "working": {
            "size": 82,
            "stage_inner": QColor(255, 255, 255, 0),
            "stage_outer": QColor(255, 255, 255, 0),
            "floor": QColor(0, 122, 255, 58),
            "halo": QColor(0, 122, 255, 40),
        },
        "done": {
            "size": 84,
            "stage_inner": QColor(255, 255, 255, 0),
            "stage_outer": QColor(255, 255, 255, 0),
            "floor": QColor(52, 199, 89, 58),
            "halo": QColor(52, 199, 89, 40),
        },
    }

    _FRAME_SIZE = {
        "sleeping": QSize(118, 108),
        "hover": QSize(120, 110),
        "ready": QSize(108, 100),
        "working": QSize(108, 102),
        "done": QSize(110, 104),
    }

    _VERTICAL_BIAS = {
        "sleeping": 0.38,
        "hover": 0.34,
        "ready": 0.24,
        "working": 0.23,
        "done": 0.21,
    }

    def __init__(self, parent=None):
        super().__init__(parent)
        self._state = "sleeping"
        self._tick = 0
        self._state_tick = 0
        self._pixmaps = {}
        self._source_rects = {}
        self._load_pixmaps()
        self._apply_frame_size("sleeping")

        self._timer = QTimer(self)
        self._timer.setInterval(45)
        self._timer.timeout.connect(self._advance_frame)
        self._timer.start()

    def _load_pixmaps(self):
        fallback_icon = ROOT_DIR / "app_icon_new.png"
        for state, path in CAT_STATE_ASSETS.items():
            pixmap = QPixmap(str(path))
            if pixmap.isNull() and fallback_icon.exists():
                pixmap = QPixmap(str(fallback_icon))
            cleaned = self._strip_embedded_background(pixmap)
            self._pixmaps[state] = cleaned
            self._source_rects[state] = self._compute_content_rect(cleaned)

    def _strip_embedded_background(self, pixmap: QPixmap) -> QPixmap:
        if pixmap.isNull():
            return pixmap

        image = pixmap.toImage().convertToFormat(QImage.Format_ARGB32)
        width = image.width()
        height = image.height()

        for y in range(height):
            for x in range(width):
                color = image.pixelColor(x, y)
                if color.alpha() == 0:
                    continue

                r, g, b = color.red(), color.green(), color.blue()
                brightness = (r + g + b) / 3.0
                chroma = max(abs(r - g), abs(g - b), abs(r - b))

                # Remove the near-white rounded card baked into mascot PNGs.
                if brightness >= 246 and chroma <= 10:
                    color.setAlpha(0)
                    image.setPixelColor(x, y, color)
                elif brightness >= 236 and chroma <= 12:
                    fade = max(0.0, min(1.0, (246.0 - brightness) / 10.0))
                    color.setAlpha(int(color.alpha() * fade))
                    image.setPixelColor(x, y, color)

        return QPixmap.fromImage(image)

    def _compute_content_rect(self, pixmap: QPixmap) -> QRectF:
        if pixmap.isNull():
            return QRectF()

        image = pixmap.toImage().convertToFormat(QImage.Format_ARGB32)
        width = image.width()
        height = image.height()
        min_x, min_y = width, height
        max_x, max_y = -1, -1

        for y in range(height):
            for x in range(width):
                if image.pixelColor(x, y).alpha() > 12:
                    min_x = min(min_x, x)
                    min_y = min(min_y, y)
                    max_x = max(max_x, x)
                    max_y = max(max_y, y)

        if max_x < min_x or max_y < min_y:
            return QRectF(0, 0, width, height)

        pad_x = max(2.0, (max_x - min_x + 1) * 0.06)
        pad_y = max(2.0, (max_y - min_y + 1) * 0.06)

        left = max(0.0, min_x - pad_x)
        top = max(0.0, min_y - pad_y)
        right = min(float(width), max_x + 1 + pad_x)
        bottom = min(float(height), max_y + 1 + pad_y)
        return QRectF(left, top, right - left, bottom - top)

    def set_state(self, state: str):
        if state not in self._STYLE:
            state = "sleeping"
        if self._state != state:
            self._state = state
            self._state_tick = 0
            self._apply_frame_size(state)
            self.update()
        else:
            self._apply_frame_size(state)

    def _apply_frame_size(self, state: str):
        size = self._FRAME_SIZE.get(state, self._FRAME_SIZE["sleeping"])
        self.setFixedSize(size)

    def _advance_frame(self):
        self._tick += 1
        self._state_tick += 1
        self.update()

    def _draw_star(self, painter: QPainter, center_x: float, center_y: float, size: float, color: QColor):
        path = QPainterPath()
        path.moveTo(center_x, center_y - size)
        path.lineTo(center_x + size * 0.32, center_y - size * 0.32)
        path.lineTo(center_x + size, center_y)
        path.lineTo(center_x + size * 0.32, center_y + size * 0.32)
        path.lineTo(center_x, center_y + size)
        path.lineTo(center_x - size * 0.32, center_y + size * 0.32)
        path.lineTo(center_x - size, center_y)
        path.lineTo(center_x - size * 0.32, center_y - size * 0.32)
        path.closeSubpath()
        painter.fillPath(path, color)

    def paintEvent(self, event):
        del event
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.setRenderHint(QPainter.SmoothPixmapTransform)

        style = self._STYLE.get(self._state, self._STYLE["sleeping"])
        phase = self._tick / 9.0
        drift = math.sin(phase)
        pulse = 0.5 + 0.5 * math.sin(phase * 0.75)

        stage_rect = QRectF(10, 6, self.width() - 20, self.height() - 22)
        gradient = QRadialGradient(stage_rect.center(), stage_rect.width() * (0.42 + pulse * 0.03))
        gradient.setColorAt(0.0, style["stage_inner"])
        gradient.setColorAt(0.72, QColor(style["stage_inner"].red(), style["stage_inner"].green(), style["stage_inner"].blue(), 24))
        gradient.setColorAt(1.0, style["stage_outer"])
        painter.setPen(Qt.NoPen)
        painter.setBrush(QBrush(gradient))
        painter.drawEllipse(stage_rect)

        if style["halo"] is not None:
            halo_rect = QRectF(
                stage_rect.center().x() - stage_rect.width() * 0.28,
                stage_rect.center().y() - stage_rect.height() * 0.25 - pulse * 2,
                stage_rect.width() * 0.56,
                stage_rect.height() * 0.50,
            )
            halo_grad = QRadialGradient(halo_rect.center(), halo_rect.width() * 0.58)
            halo_grad.setColorAt(0.0, style["halo"])
            halo_grad.setColorAt(1.0, QColor(style["halo"].red(), style["halo"].green(), style["halo"].blue(), 0))
            painter.setBrush(QBrush(halo_grad))
            painter.drawEllipse(halo_rect)

        shadow_rect = QRectF(
            self.width() * 0.5 - 28 - pulse * 2,
            self.height() - 26,
            56 + pulse * 4,
            10 + pulse * 1.5,
        )
        painter.setBrush(style["floor"])
        painter.drawEllipse(shadow_rect)

        pixmap = self._pixmaps.get(self._state)
        if pixmap is None or pixmap.isNull():
            painter.setBrush(QColor("#D2D2D7"))
            painter.drawEllipse(QRectF(self.width() * 0.5 - 28, 22, 56, 56))
            return

        size = style["size"]
        scale = 1.0
        offset_y = 0.0
        rotation = 0.0

        if self._state == "sleeping":
            scale = 0.985 + 0.015 * math.sin(phase * 0.7)
            offset_y = 2.0 * math.sin(phase * 0.7)
        elif self._state == "hover":
            scale = 1.0 + 0.035 * abs(math.sin(phase * 1.15))
            offset_y = -4.5 * abs(math.sin(phase * 1.15))
            rotation = -2.25 * math.sin(phase * 1.15)
        elif self._state == "ready":
            scale = 1.0 + 0.015 * math.sin(phase * 0.65)
            offset_y = -2.5 * abs(math.sin(phase * 0.65))
            rotation = 0.7 * math.sin(phase * 0.65)
        elif self._state == "working":
            scale = 1.0 + 0.012 * math.sin(phase * 1.25)
            offset_y = -2.0 * abs(math.sin(phase * 1.25))
            rotation = 1.9 * math.sin(phase * 1.25)
        elif self._state == "done":
            pop_t = min(self._state_tick / 11.0, 1.0)
            pop = 0.82 + 0.22 * math.sin(pop_t * math.pi * 0.5)
            scale = pop + 0.02 * math.sin(phase * 0.8)
            offset_y = -2.5 - 1.5 * math.sin(phase * 0.8)

        render_width = size * scale
        source_rect = self._source_rects.get(self._state) or QRectF(0, 0, pixmap.width(), pixmap.height())
        source_aspect = source_rect.height() / max(source_rect.width(), 1.0)
        render_height = render_width * source_aspect

        available_height = self.height() - 6
        if render_height > available_height:
            fit_scale = available_height / max(render_height, 1.0)
            render_width *= fit_scale
            render_height *= fit_scale

        vertical_bias = self._VERTICAL_BIAS.get(self._state, 0.42)
        target_rect = QRectF(
            self.width() * 0.5 - render_width * 0.5,
            max(2.0, (self.height() - render_height) * vertical_bias + offset_y),
            render_width,
            render_height,
        )

        painter.save()
        painter.translate(target_rect.center())
        painter.rotate(rotation)
        painter.translate(-target_rect.center())
        clip_path = QPainterPath()
        corner_radius = min(target_rect.width(), target_rect.height()) * 0.26
        clip_path.addRoundedRect(target_rect, corner_radius, corner_radius)
        painter.setClipPath(clip_path)
        painter.drawPixmap(target_rect, pixmap, source_rect)
        painter.restore()

        if self._state == "sleeping":
            painter.setPen(QColor(168, 168, 188, 220))
            painter.setFont(QFont("Georgia", 10, QFont.Bold))
            rise = (self._tick % 54) / 54.0
            for idx, (letter, dx, dy, size_add) in enumerate((("z", 18, 16, 0), ("z", 28, 4, 2), ("Z", 40, -10, 4))):
                alpha = int(max(0.0, 1.0 - ((rise + idx * 0.2) % 1.0)) * 220)
                painter.setPen(QColor(168 - idx * 6, 168 - idx * 6, 188 - idx * 6, alpha))
                font = QFont("Georgia", 10 + size_add, QFont.Bold)
                painter.setFont(font)
                painter.drawText(
                    QPoint(
                        int(target_rect.right()) + dx,
                        int(target_rect.top()) + dy - int(((rise + idx * 0.2) % 1.0) * 22),
                    ),
                    letter,
                )

        if self._state == "working":
            report_rect = QRectF(target_rect.right() - 18, target_rect.top() + 16, 24, 32)
            painter.setPen(Qt.NoPen)
            painter.setBrush(QColor(255, 255, 255, 60))
            painter.drawRoundedRect(report_rect, 8, 8)
            painter.setBrush(QColor(0, 122, 255, 36))
            scan = (self._tick % 26) / 26.0
            for idx, color in enumerate((QColor(0, 122, 255, 180), QColor(29, 29, 31, 70), QColor(52, 199, 89, 120))):
                line_y = report_rect.top() + 8 + idx * 7
                line_x = report_rect.left() + 4 + scan * 8
                painter.setBrush(color)
                painter.drawRoundedRect(QRectF(line_x - 5, line_y, 12, 2), 1, 1)
            painter.setBrush(QColor(0, 122, 255, 210 if scan > 0.5 else 110))
            painter.drawEllipse(QRectF(report_rect.right() - 8, report_rect.top() + 4, 5, 5))
            painter.setPen(QPen(QColor(255, 91, 87, 220), 2.0, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin))
            path = QPainterPath()
            path.moveTo(report_rect.left() + 3, report_rect.top() + 21)
            path.cubicTo(
                report_rect.left() + 8,
                report_rect.top() + 15 + math.sin(phase) * 1.5,
                report_rect.left() + 13,
                report_rect.top() + 24,
                report_rect.left() + 19,
                report_rect.top() + 17,
            )
            painter.drawPath(path)

        if self._state == "done":
            sparkle_phase = (self._tick % 36) / 36.0
            sparkles = (
                (target_rect.left() - 8, target_rect.top() + 10, 7, QColor(160, 210, 255, int(130 + 80 * math.sin(phase)))),
                (target_rect.center().x() + 22, target_rect.top() + 2, 8, QColor(195, 175, 255, int(110 + 110 * math.sin(phase * 1.2)))),
                (target_rect.right() - 3, target_rect.top() + 18, 6, QColor(155, 235, 248, int(130 + 90 * math.sin(phase * 0.9)))),
            )
            for x, y, size_hint, color in sparkles:
                shimmer = size_hint * (0.88 + 0.22 * abs(math.sin((sparkle_phase + x * 0.01) * math.pi * 2)))
                self._draw_star(painter, x, y, shimmer, color)


class StatusBubbleWidget(QWidget):
    """Comic-style speech bubble used for the mascot status text."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self._fill = QColor(255, 255, 255, 230)
        self._border = QColor(229, 229, 234, 230)
        self._tail_fill = QColor(self._fill)
        self._tail_border = QColor(self._border)
        self._tail_side = "left"
        self._tail_extent = 16
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self.setMaximumWidth(198)

        self._content_layout = QVBoxLayout(self)
        self._content_layout.setSpacing(1)
        self._update_content_margins()

        self.title_label = QLabel()
        self.title_label.setAlignment(Qt.AlignCenter)
        self.title_label.setFont(QFont(".AppleSystemUIFont", 11, QFont.DemiBold))
        self.title_label.setStyleSheet("background: transparent; border: none;")
        self.title_label.setWordWrap(True)
        self._content_layout.addWidget(self.title_label)

        self.hint_label = QLabel()
        self.hint_label.setAlignment(Qt.AlignCenter)
        self.hint_label.setFont(QFont(".AppleSystemUIFont", 9))
        self.hint_label.setStyleSheet("background: transparent; border: none;")
        self.hint_label.setWordWrap(True)
        self._content_layout.addWidget(self.hint_label)

    def _update_content_margins(self):
        left_margin = 12 + (self._tail_extent if self._tail_side == "left" else 0)
        right_margin = 12 + (self._tail_extent if self._tail_side == "right" else 0)
        self._content_layout.setContentsMargins(left_margin, 10, right_margin, 8)

    def set_content(self, title: str, hint: str):
        self.title_label.setText(title)
        self.hint_label.setText(hint)

    def set_colors(self, title_color: str, hint_color: str, fill: QColor, border: QColor):
        self._fill = QColor(fill)
        self._border = QColor(border)
        self._tail_fill = QColor(fill)
        self._tail_border = QColor(border)
        self.title_label.setStyleSheet(f"background: transparent; border: none; color: {title_color};")
        self.hint_label.setStyleSheet(f"background: transparent; border: none; color: {hint_color};")
        self.update()

    def set_tail_side(self, side: str):
        self._tail_side = side if side in {"left", "right", "none"} else "left"
        self._update_content_margins()
        self.update()

    def paintEvent(self, event):
        del event
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)

        left_inset = self._tail_extent if self._tail_side == "left" else 0
        right_inset = self._tail_extent if self._tail_side == "right" else 0
        rect = QRectF(self.rect().adjusted(1 + left_inset, 1, -1 - right_inset, -1))
        radius = 10.0

        bubble_path = QPainterPath()
        bubble_path.addRoundedRect(rect, radius, radius)

        painter.setPen(QPen(self._border, 1.0))
        painter.setBrush(QBrush(self._fill))
        painter.drawPath(bubble_path)

        if self._tail_side != "none":
            tail_path = QPainterPath()
            tail_center_y = rect.center().y() + 3
            if self._tail_side == "right":
                tail_path.moveTo(rect.right(), tail_center_y - 11)
                tail_path.quadTo(rect.right() + 8, tail_center_y - 8, rect.right() + 12, tail_center_y - 3)
                tail_path.lineTo(rect.right() + self._tail_extent, tail_center_y + 1)
                tail_path.quadTo(rect.right() + 9, tail_center_y + 4, rect.right(), tail_center_y + 11)
            else:
                tail_path.moveTo(rect.left(), tail_center_y - 11)
                tail_path.quadTo(rect.left() - 8, tail_center_y - 8, rect.left() - 12, tail_center_y - 3)
                tail_path.lineTo(rect.left() - self._tail_extent, tail_center_y + 1)
                tail_path.quadTo(rect.left() - 9, tail_center_y + 4, rect.left(), tail_center_y + 11)
            tail_path.closeSubpath()

            painter.setPen(QPen(self._tail_border, 1.0))
            painter.setBrush(QBrush(self._tail_fill))
            painter.drawPath(tail_path)


class DropZone(QFrame):
    """Stateful design-system drag and drop zone with mascot."""

    files_dropped = pyqtSignal(list)
    folder_dropped = pyqtSignal(str)  # путь папки готового отчёта — для отправки на Яндекс.Диск

    _STATE_COPY = {
        "sleeping": ("Перетащите файлы сюда", "аудио, видео, CSV, PDF", "#3D3D3D", "#86868B", "rgba(0, 0, 0, 0.015)", "1.5px dashed rgba(0, 0, 0, 0.13)"),
        "hover": ("Отпустите файлы", "кот уже проснулся и ждёт загрузку", "#007AFF", "#6B7280", "rgba(0, 122, 255, 0.05)", "2px dashed rgba(0, 122, 255, 0.5)"),
        "ready": ("Готов к работе", "Нажмите «Создать» для генерации отчёта", "#1D1D1F", "#86868B", "rgba(0, 122, 255, 0.02)", "1.5px dashed rgba(0, 122, 255, 0.24)"),
        "working": ("Обрабатываю...", "Проверяю файлы и собираю отчёт", "#007AFF", "#6B7280", "rgba(0, 122, 255, 0.035)", "1.5px dashed rgba(0, 122, 255, 0.32)"),
        "done": ("Готово!", "Отчёт создан успешно", "#34C759", "#6B7280", "rgba(52, 199, 89, 0.05)", "1.5px dashed rgba(52, 199, 89, 0.42)"),
    }

    _STATE_HEIGHTS = {
        "sleeping": (138, 154),
        "hover": (142, 158),
        "ready": (112, 124),
        "working": (114, 126),
        "done": (114, 126),
    }

    _BUBBLE_STYLE = {
        "sleeping": (QColor(255, 255, 255, 226), QColor(226, 228, 234, 236)),
        "hover": (QColor(250, 253, 255, 234), QColor(171, 212, 255, 245)),
        "ready": (QColor(255, 255, 255, 232), QColor(215, 225, 238, 240)),
        "working": (QColor(247, 251, 255, 236), QColor(166, 208, 255, 246)),
        "done": (QColor(250, 255, 251, 238), QColor(165, 228, 184, 245)),
    }

    def __init__(self):
        super().__init__()
        self.setAcceptDrops(True)
        self._has_files = False
        self._is_processing = False
        self._is_done = False
        self._drag_active = False
        self._state = "sleeping"
        self.init_ui()

    def init_ui(self):
        self.setMinimumHeight(138)
        self.setMaximumHeight(154)
        self.setFrameStyle(QFrame.NoFrame)
        self.setObjectName("dropZone")

        layout = QVBoxLayout()
        layout.setAlignment(Qt.AlignCenter)
        layout.setSpacing(0)
        layout.setContentsMargins(8, 8, 8, 6)

        row = QHBoxLayout()
        row.setContentsMargins(0, 0, 0, 0)
        row.setSpacing(1)
        row.setAlignment(Qt.AlignCenter)

        self.mascot = CatMascotWidget(self)
        row.addWidget(self.mascot, 0, Qt.AlignVCenter)

        self.status_bubble = StatusBubbleWidget()
        self.status_bubble.set_tail_side("left")
        row.addWidget(self.status_bubble, 0, Qt.AlignVCenter)

        layout.addLayout(row)

        self.setLayout(layout)
        self._apply_state()

    def _sync_state(self):
        if self._is_done:
            state = "done"
        elif self._is_processing:
            state = "working"
        elif self._drag_active:
            state = "hover"
        elif self._has_files:
            state = "ready"
        else:
            state = "sleeping"
        self._set_state(state)

    def _set_state(self, state: str):
        self._state = state
        self._apply_state()

    def _apply_state(self):
        title, hint, title_color, hint_color, bg_color, border = self._STATE_COPY[self._state]
        min_height, max_height = self._STATE_HEIGHTS.get(self._state, self._STATE_HEIGHTS["sleeping"])
        bubble_fill, bubble_border = self._BUBBLE_STYLE.get(self._state, self._BUBBLE_STYLE["sleeping"])
        self.setMinimumHeight(min_height)
        self.setMaximumHeight(max_height)
        self.status_bubble.set_content(title, hint)
        self.status_bubble.set_colors(title_color, hint_color, bubble_fill, bubble_border)
        self.status_bubble.set_tail_side("none" if self._state == "sleeping" else "left")
        self.setStyleSheet(
            f"""
            QFrame#dropZone {{
                background-color: {bg_color};
                border: {border};
                border-radius: 10px;
                padding: 12px;
            }}
            QFrame#dropZone QLabel, QFrame#dropZone StatusBubbleWidget {{
                border: none;
                background: transparent;
            }}
            """
        )
        self.setCursor(Qt.ArrowCursor if (self._is_processing or self._is_done) else Qt.PointingHandCursor)
        self.mascot.set_state(self._state)

    def set_files_loaded(self, has_files: bool):
        self._has_files = has_files
        if not has_files:
            self._is_done = False
        self._sync_state()

    def set_processing_state(self, is_processing: bool):
        self._is_processing = is_processing
        if is_processing:
            self._is_done = False
        self._sync_state()

    def set_completed_state(self, is_done: bool):
        self._is_done = is_done
        if is_done:
            self._is_processing = False
        self._sync_state()

    def dragEnterEvent(self, event: QDragEnterEvent):
        if self._is_processing:
            event.ignore()
            return
        if event.mimeData().hasUrls():
            self._drag_active = True
            event.acceptProposedAction()
            self._sync_state()

    def dragLeaveEvent(self, event):
        del event
        self._drag_active = False
        self._sync_state()

    def dropEvent(self, event: QDropEvent):
        self._drag_active = False
        self._sync_state()

        files = []
        folders = []
        for url in event.mimeData().urls():
            file_path = url.toLocalFile()
            if os.path.isfile(file_path):
                files.append(file_path)
            elif os.path.isdir(file_path):
                folders.append(file_path)

        if not files and len(folders) == 1:
            # Перетащили ровно одну папку без файлов — трактуем как папку
            # уже готового отчёта для отправки на Диск, а не источники для
            # генерации нового отчёта (для этого — files_dropped ниже).
            self.folder_dropped.emit(folders[0])
            return

        if files:
            self.files_dropped.emit(files)


class PdfOnlyDocxBuilder:
    """Создает простой DOCX из первых страниц PDF (по одному PDF на страницу)."""

    @staticmethod
    def build(output_path: Path, pdf_paths: list) -> None:
        doc = Document()

        # A3 landscape как в основном отчете
        section = doc.sections[0]
        section.page_width = Cm(42.02)
        section.page_height = Cm(29.70)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.orientation = 1  # Landscape

        for idx, pdf_path in enumerate(pdf_paths):
            if idx > 0:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_break(WD_BREAK.PAGE)

            pdf_doc = fitz.open(str(pdf_path))
            page = pdf_doc[0]
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            img_stream = io.BytesIO(pix.tobytes("png"))
            doc.add_picture(img_stream, width=Cm(38.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            pdf_doc.close()

        doc.save(str(output_path))


def convert_dragdrop_to_files(files_data):
    """Преобразует drag & drop файлы в формат v5.11"""
    files = {
        'audio_20_c': None,
        'audio_20_uc': None,
        'audio_51_c': None,
        'audio_51_uc': None,
        'video': None,
        'csv': None,
        'pdf_20': None,
        'pdf_51': None,
        'pdf_20_c': None,
        'pdf_20_uc': None,
        'pdf_51_c': None,
        'pdf_51_uc': None,
        'params': None,
        'all_files': []
    }
    
    # CSV
    if files_data.get('csv'):
        files['csv'] = Path(files_data['csv'][0])
        files['all_files'].append(files['csv'])
    
    # Video
    if files_data.get('video'):
        files['video'] = Path(files_data['video'][0])
        files['all_files'].append(files['video'])
    
    # Обработка аудио файлов (логика из v5.11)
    for audio_file in files_data.get('audio', []):
        file_path = Path(audio_file)
        files['all_files'].append(file_path)
        name_lower = file_path.name.lower()
        channel = detect_channel_from_name(name_lower)
        cens_state = detect_cens_state(name_lower)
        
        # Проверяем 5.1 ПЕРЕД 2.0 (как в v5.11)
        if channel == "51":
            if cens_state == "uc":
                files['audio_51_uc'] = file_path
            else:
                files['audio_51_c'] = file_path
        
        # Проверяем 2.0
        elif channel == "20":
            if cens_state == "uc":
                files['audio_20_uc'] = file_path
            else:
                files['audio_20_c'] = file_path
    
    # Обработка PDF файлов (логика из v5.11)
    for pdf_file in files_data.get('pdf', []):
        file_path = Path(pdf_file)
        files['all_files'].append(file_path)
        name_lower = file_path.name.lower()
        channel = detect_channel_from_name(name_lower)
        cens_state = detect_cens_state(name_lower)
        
        # Проверяем 5.1 ПЕРЕД 2.0
        if channel == "51":
            if cens_state == "uc":
                files['pdf_51_uc'] = file_path
            elif cens_state == "c":
                files['pdf_51_c'] = file_path
            else:
                files['pdf_51'] = file_path
        
        # Проверяем 2.0
        elif channel == "20":
            if cens_state == "uc":
                files['pdf_20_uc'] = file_path
            elif cens_state == "c":
                files['pdf_20_c'] = file_path
            else:
                files['pdf_20'] = file_path
    
    return files


REPORT_TYPE_LABELS = {
    "standard": "Основной",
    "me": "ME",
    "me_ours": "ME (наши)",
    "dcp": "DCP",
}

PREVIEW_SLOT_LABELS = {
    "video": "Видео",
    "csv": "CSV",
    "params": "Параметры",
    "audio_20_c": "Аудио 2.0 C",
    "audio_20_uc": "Аудио 2.0 U",
    "audio_51_c": "Аудио 5.1 C",
    "audio_51_uc": "Аудио 5.1 U",
    "pdf_20_c": "PDF 2.0 C",
    "pdf_20_uc": "PDF 2.0 U",
    "pdf_20": "PDF 2.0",
    "pdf_51_c": "PDF 5.1 C",
    "pdf_51_uc": "PDF 5.1 U",
    "pdf_51": "PDF 5.1",
}

PREVIEW_SLOT_ORDER = {
    "video": 10,
    "csv": 20,
    "params": 30,
    "audio_20_c": 40,
    "audio_20_uc": 41,
    "audio_51_c": 42,
    "audio_51_uc": 43,
    "pdf_20_c": 50,
    "pdf_20_uc": 51,
    "pdf_20": 52,
    "pdf_51_c": 53,
    "pdf_51_uc": 54,
    "pdf_51": 55,
}


def preview_slot_label(slot_name: str, report_type: str) -> str:
    if report_type in {"me", "me_ours"}:
        me_labels = {
            "audio_20_c": "Аудио 2.0",
            "audio_20_uc": "Аудио 2.0",
            "audio_51_c": "Аудио 5.1",
            "audio_51_uc": "Аудио 5.1",
            "pdf_20_c": "PDF 2.0",
            "pdf_20_uc": "PDF 2.0",
            "pdf_51_c": "PDF 5.1",
            "pdf_51_uc": "PDF 5.1",
        }
        if slot_name in me_labels:
            return me_labels[slot_name]
    return PREVIEW_SLOT_LABELS.get(slot_name, slot_name)


def filter_preview_warnings(messages: list, report_type: str) -> list:
    if report_type not in {"me", "me_ours"}:
        return list(messages or [])

    filtered = []
    for message in messages or []:
        lowered = message.lower()
        if "cens/uncens" in lowered or "cens-слот" in lowered:
            continue
        filtered.append(message)
    return filtered


def preview_display_filename(file_name: str, report_type: str, kind: str) -> str:
    if report_type not in {"me", "me_ours"} or kind not in {"audio", "pdf"}:
        return file_name

    path_obj = Path(file_name)
    stem = path_obj.stem
    stem = re.sub(r'(?i)([_\-. ]+)(uncens|uncensored|uc)(?=$|[_\-. ])', '', stem)
    stem = re.sub(r'(?i)([_\-. ]+)(cens|censored|c)(?=$|[_\-. ])', '', stem)
    stem = re.sub(r'[_\-. ]{2,}', '_', stem).strip(' _-.')
    return f"{stem or path_obj.stem}{path_obj.suffix}"


def preview_targets(params_metrics: dict, report_type: str) -> dict:
    params_metrics = params_metrics or {}
    return {
        "target_lufs": params_metrics.get("target_lufs", -23.0),
        "true_peak": 0.0 if report_type == "dcp" else params_metrics.get("true_peak", -2.0),
        "lra_max": params_metrics.get("lra_max", 18.0),
        "lufs_tolerance": 0.5,
        "check_lufs_lra": report_type not in {"me", "dcp", "me_ours"},
    }


def preview_metric_issues(pdf_metrics: dict, params_metrics: dict, report_type: str) -> list:
    """Проверка отклонений по тем же правилам, что используются в итоговом отчете."""
    if not pdf_metrics:
        return []

    targets = preview_targets(params_metrics, report_type)
    issues = []

    lufs = pdf_metrics.get("lufs")
    true_peak = pdf_metrics.get("true_peak")
    lra = pdf_metrics.get("lra")

    if targets.get("check_lufs_lra") and lufs is not None:
        if abs(float(lufs) - float(targets["target_lufs"])) > float(targets["lufs_tolerance"]):
            issues.append(
                f"LUFS не соответствует норме: {_preview_numeric(lufs)} вместо {_preview_numeric(targets['target_lufs'])}."
            )

    if true_peak is not None:
        if float(true_peak) > float(targets["true_peak"]):
            issues.append(
                f"True Peak превышает лимит: {_preview_numeric(true_peak)} dBTP при норме {_preview_numeric(targets['true_peak'])} dBTP."
            )

    if targets.get("check_lufs_lra") and lra is not None:
        if float(lra) > float(targets["lra_max"]):
            issues.append(
                f"LRA превышает лимит: {_preview_numeric(lra)} при норме {_preview_numeric(targets['lra_max'])}."
            )

    return issues


def preview_format_issues(audio_metrics: dict, slot_name: str) -> list:
    """Проверка формата файла по тем же правилам, что и в итоговом отчете."""
    if not audio_metrics or not slot_name or "audio" not in slot_name:
        return []

    issues = []
    sample_rate = audio_metrics.get("sample_rate")
    bit_depth = audio_metrics.get("bit_depth")
    channel_order = str(audio_metrics.get("channel_order") or "").strip()

    try:
        sr_khz = int(sample_rate) // 1000 if sample_rate is not None else 0
    except Exception:
        sr_khz = 0

    bd_text = str(bit_depth or "").replace("PCM_", "").replace("bit", "").strip()
    try:
        bd_int = int(bd_text)
    except Exception:
        bd_int = 0

    if sr_khz != 48 or bd_int != 24:
        issues.append(
            f"Формат файла не соответствует норме: {sr_khz or '—'} kHz / {bd_int or '—'} bit вместо 48 kHz / 24 bit."
        )

    if "51" in slot_name:
        standard_51_order = "L R C LFE Ls Rs"
        if not channel_order or "channels" in channel_order.lower() or channel_order.lower() == "unknown":
            issues.append("Не удалось подтвердить порядок каналов 5.1 из метаданных файла.")
        elif channel_order != standard_51_order:
            issues.append(
                f"Порядок каналов 5.1 отличается от нормы: {channel_order} вместо {standard_51_order}."
            )

    return issues


def format_preview_warning(message: str, report_type: str) -> str:
    formatted = str(message or "")
    for slot_name in PREVIEW_SLOT_LABELS:
        formatted = formatted.replace(slot_name, preview_slot_label(slot_name, report_type))
    return formatted


def _preview_numeric(value, digits: int = 1) -> str:
    if value in (None, ""):
        return "—"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    return f"{number:.{digits}f}"


def _preview_detail_pairs(details: list) -> str:
    return ", ".join(part for part in details if part)


def _preview_duration_ms(duration_seconds):
    try:
        duration = float(duration_seconds)
    except (TypeError, ValueError):
        return None
    if duration <= 0:
        return None
    return int(duration * 1000)


def _preview_duration_text(duration_seconds) -> str:
    total_ms = _preview_duration_ms(duration_seconds)
    if total_ms is None:
        return "—"

    hours = total_ms // 3600000
    mins = (total_ms % 3600000) // 60000
    secs = (total_ms % 60000) // 1000
    millis = total_ms % 1000
    return f"{hours}:{mins:02d}:{secs:02d}.{millis:03d}"


def _preview_is_frame_aligned(duration_seconds, fps: float = 25.0) -> bool:
    try:
        duration = float(duration_seconds)
    except (TypeError, ValueError):
        return False
    if duration <= 0:
        return True
    fps = fps or 25
    ms = duration * 1000
    frame_ms = 1000.0 / fps
    nearest_frame = round(ms / frame_ms)
    return abs(ms - nearest_frame * frame_ms) < 0.5


def preview_duration_issues(record: dict, duration_context: dict) -> list:
    """Проверка хронометража по тем же правилам, что и в итоговом отчете."""
    kind = record.get("kind")
    if kind not in {"audio", "video"}:
        return []

    metrics = record.get("metrics") or {}
    duration = metrics.get("duration")
    duration_ms = _preview_duration_ms(duration)
    if duration_ms is None:
        return []

    fps = float(duration_context.get("fps") or 25.0)
    video_duration_ms = duration_context.get("video_duration_ms")
    audio_durations_ms = duration_context.get("audio_durations_ms") or []
    audio_durations_ms_set = duration_context.get("audio_durations_ms_set") or set()
    all_audio_match = bool(duration_context.get("all_audio_match", True))
    issues = []

    if kind == "audio":
        if video_duration_ms is not None:
            if duration_ms != video_duration_ms:
                issues.append("Хронометраж не совпадает с видео.")
            if not _preview_is_frame_aligned(duration, fps):
                issues.append(f"Хронометраж не кратен кадру ({format_fps(fps)} fps).")
        elif not all_audio_match:
            issues.append("Хронометраж отличается от других аудио.")
    elif kind == "video":
        if audio_durations_ms:
            if duration_ms not in audio_durations_ms_set:
                issues.append("Хронометраж не совпадает с аудио.")
        elif not _preview_is_frame_aligned(duration, fps):
            issues.append(f"Хронометраж не кратен кадру ({format_fps(fps)} fps).")

    return issues


def analyze_files_for_preview(files_data: dict, report_type: str = "standard") -> dict:
    """Собирает предпросмотр распознавания файлов и параметров до генерации отчета."""
    snapshot = {
        'audio': list(files_data.get('audio', [])),
        'video': list(files_data.get('video', [])),
        'csv': list(files_data.get('csv', [])),
        'pdf': list(files_data.get('pdf', [])),
        'params': list(files_data.get('params', [])),
    }
    recognized = []
    warnings = []

    detected_report_type, report_analysis = detect_report_type_from_files(
        [item for group in snapshot.values() for item in group]
    )
    if report_analysis.get("reason") == "ambiguous":
        warnings.append("Автоопределение типа отчета конфликтует: проверьте выбранный тип вручную.")

    tech_extractor = TechnicalInfoExtractor()
    pdf_extractor = PDFExtractor()

    selected_audio_by_slot = {}
    audio_key_by_stem = {}
    used_pdf_slots = {}

    def add_record(
        kind: str,
        path: str,
        slot: str = None,
        details: list = None,
        item_warnings: list = None,
        extra: dict = None,
    ):
        record = {
            "kind": kind,
            "path": str(path),
            "name": Path(path).name,
            "slot": slot,
            "details": list(details or []),
            "warnings": list(item_warnings or []),
        }
        if extra:
            record.update(extra)
        recognized.append(record)
        return record

    def collect_warning(message: str):
        if message and message not in warnings:
            warnings.append(message)

    for idx, video_file in enumerate(snapshot['video']):
        item_warnings = []
        slot = "video" if idx == 0 else None
        if idx > 0:
            item_warnings.append("Будет проигнорирован: в отчет идет только первый видеофайл.")
        video_info = tech_extractor.extract_video_info(video_file) if idx == 0 else {}
        details = _preview_detail_pairs([
            f"{format_fps(video_info.get('fps'))} fps" if video_info.get('fps') else "",
            f"{video_info.get('sample_rate')} Hz" if video_info.get('sample_rate') else "",
            f"{video_info.get('channels')} ch" if video_info.get('channels') else "",
            video_info.get('format') or "",
        ])
        record = add_record(
            "video",
            video_file,
            slot=slot,
            details=[details] if details else [],
            item_warnings=item_warnings,
            extra={
                "metrics": {
                    "duration": video_info.get("duration"),
                    "fps": video_info.get("fps"),
                    "sample_rate": video_info.get("sample_rate"),
                    "channels": video_info.get("channels"),
                    "format": video_info.get("format"),
                }
            } if idx == 0 else None,
        )
        for message in record["warnings"]:
            collect_warning(f"{record['name']}: {message}")

    for idx, csv_file in enumerate(snapshot['csv']):
        item_warnings = []
        slot = "csv" if idx == 0 else None
        if idx > 0:
            item_warnings.append("Будет проигнорирован: используется только первый CSV.")
        record = add_record("csv", csv_file, slot=slot, item_warnings=item_warnings)
        for message in record["warnings"]:
            collect_warning(f"{record['name']}: {message}")

    params_preview = None
    for idx, params_file in enumerate(snapshot['params']):
        item_warnings = []
        slot = "params" if idx == 0 else None
        if idx > 0:
            item_warnings.append("Будет проигнорирован: используется только первый файл параметров.")
        details = []
        if idx == 0:
            params_preview = tech_extractor.read_params_file(params_file) or {}
            details = [_preview_detail_pairs([
                f"target {_preview_numeric(params_preview.get('target_lufs'))} LUFS" if params_preview.get('target_lufs') is not None else "",
                f"TP {_preview_numeric(params_preview.get('true_peak'))} dBTP" if params_preview.get('true_peak') is not None else "",
                f"LRA ≤ {_preview_numeric(params_preview.get('lra_max'))}" if params_preview.get('lra_max') is not None else "",
                f"{params_preview.get('sample_rate')} Hz" if params_preview.get('sample_rate') else "",
                f"{params_preview.get('bit_depth')} bit" if params_preview.get('bit_depth') else "",
            ])]
        record = add_record(
            "params",
            params_file,
            slot=slot,
            details=[part for part in details if part],
            item_warnings=item_warnings,
            extra={
                "metrics": {
                    "target_lufs": params_preview.get("target_lufs"),
                    "true_peak": params_preview.get("true_peak"),
                    "lra_max": params_preview.get("lra_max"),
                }
            } if idx == 0 else None,
        )
        for message in record["warnings"]:
            collect_warning(f"{record['name']}: {message}")

    for audio_file in snapshot['audio']:
        filename = Path(audio_file).stem.lower()
        audio_info = tech_extractor.extract_audio_info(audio_file) or {}
        item_warnings = []

        channels = audio_info.get('channels', 0)
        try:
            channels = int(float(channels)) if channels is not None else 0
        except Exception:
            channels = 0

        channel_hint = detect_channel_from_name(filename)
        channel_meta = "51" if channels >= 6 else ("20" if channels == 2 else None)
        cens_state = detect_cens_state(filename)
        suffix = "uc" if cens_state == "uc" else "c"
        channel_final = channel_meta or channel_hint
        slot = f"audio_{channel_final}_{suffix}" if channel_final in {"20", "51"} else None

        if cens_state is None:
            item_warnings.append("В имени нет CENS/UNCENS, будет использован cens-слот по умолчанию.")
        if channel_hint and channel_meta and channel_hint != channel_meta:
            item_warnings.append(f"Имя говорит {channel_hint}, а метаданные файла говорят {channel_meta}.")
        if has_incorrect_audio_marker(filename):
            item_warnings.append("В имени похожий на ошибку маркер канальности (например 5.0 или 2.1).")
        if slot is None:
            item_warnings.append("Не удалось определить слот для этого аудио.")

        name_score = 0
        if channel_hint and channel_final and channel_hint == channel_final:
            name_score += 2
        if not has_incorrect_audio_marker(filename):
            name_score += 1

        selected_for_slot = False
        if slot:
            existing = selected_audio_by_slot.get(slot)
            if existing and name_score <= existing["name_score"]:
                item_warnings.append(
                    f"Слот {slot} уже занят более подходящим файлом {existing['name']}."
                )
            elif existing:
                existing["record"]["warnings"].append(
                    f"Слот {slot} будет перезаписан файлом {Path(audio_file).name}."
                )
                collect_warning(f"{existing['record']['name']}: {existing['record']['warnings'][-1]}")
                selected_audio_by_slot[slot] = {
                    "name": Path(audio_file).name,
                    "name_score": name_score,
                    "record": None,
                }
                selected_for_slot = True
            else:
                selected_audio_by_slot[slot] = {
                    "name": Path(audio_file).name,
                    "name_score": name_score,
                    "record": None,
                }
                selected_for_slot = True

        details = [_preview_detail_pairs([
            f"{channels} ch" if channels else "",
            f"{audio_info.get('sample_rate')} Hz" if audio_info.get('sample_rate') else "",
            str(audio_info.get('bit_depth') or ""),
            str(audio_info.get('channel_order') or ""),
        ])]
        item_warnings.extend(preview_format_issues(audio_info, slot))
        record = add_record(
            "audio",
            audio_file,
            slot=slot,
            details=[part for part in details if part],
            item_warnings=item_warnings,
            extra={
                "metrics": {
                    "duration": audio_info.get("duration"),
                    "sample_rate": audio_info.get("sample_rate"),
                    "bit_depth": audio_info.get("bit_depth"),
                    "channel_order": audio_info.get("channel_order"),
                    "channels": channels,
                }
            },
        )

        if slot and selected_for_slot and selected_audio_by_slot.get(slot):
            selected_audio_by_slot[slot]["record"] = record
            audio_key_by_stem[normalize_stem(audio_info.get('file_name', audio_file))] = slot

        for message in record["warnings"]:
            collect_warning(f"{record['name']}: {message}")

    for pdf_file in snapshot['pdf']:
        filename = Path(pdf_file).stem.lower()
        item_warnings = []
        pdf_data = pdf_extractor.extract_technical_info(pdf_file) or {}
        pdf_data['source_pdf'] = Path(pdf_file).name

        matched_audio_key = audio_key_by_stem.get(normalize_stem(pdf_file))
        if matched_audio_key:
            slot = matched_audio_key.replace('audio_', 'pdf_')
        else:
            channel_hint = detect_channel_from_name(filename)
            cens_state = detect_cens_state(filename)
            if channel_hint == "51" and cens_state == "c":
                slot = "pdf_51_c"
            elif channel_hint == "51" and cens_state == "uc":
                slot = "pdf_51_uc"
            elif channel_hint == "51":
                slot = "pdf_51"
            elif channel_hint == "20" and cens_state == "c":
                slot = "pdf_20_c"
            elif channel_hint == "20" and cens_state == "uc":
                slot = "pdf_20_uc"
            elif channel_hint == "20":
                slot = "pdf_20"
            else:
                slot = None

        if slot is None and pdf_data.get('channels'):
            channels_str = str(pdf_data.get('channels', '')).lower()
            cens_state = detect_cens_state(filename)
            if '5.1' in channels_str or '5.0' in channels_str or '6' in channels_str or 'surround' in channels_str:
                slot = 'pdf_51_c' if cens_state == "c" else ('pdf_51_uc' if cens_state == "uc" else 'pdf_51')
            elif '2.0' in channels_str or '2.1' in channels_str or '2' in channels_str or 'stereo' in channels_str:
                slot = 'pdf_20_c' if cens_state == "c" else ('pdf_20_uc' if cens_state == "uc" else 'pdf_20')

        if slot and slot in used_pdf_slots:
            generic_slot = 'pdf_20' if '20' in slot else 'pdf_51'
            item_warnings.append(f"Слот {slot} уже занят, этот PDF уйдет в {generic_slot}.")
            if generic_slot not in used_pdf_slots:
                slot = generic_slot
            else:
                item_warnings.append(f"Слот {generic_slot} тоже уже занят, нужен ручной выбор.")
                slot = None
        if slot:
            used_pdf_slots[slot] = Path(pdf_file).name
        else:
            item_warnings.append("Не удалось определить слот PDF по имени или содержимому.")

        missing_pdf_values = [
            field_name
            for field_name in ("lufs", "true_peak", "lra")
            if pdf_data.get(field_name) is None
        ]
        if missing_pdf_values:
            item_warnings.append(f"В PDF не найдены поля: {', '.join(missing_pdf_values)}.")

        item_warnings.extend(preview_metric_issues(pdf_data, params_preview or {}, report_type))

        details = [_preview_detail_pairs([
            f"LUFS {_preview_numeric(pdf_data.get('lufs'))}" if pdf_data.get('lufs') is not None else "",
            f"TP {_preview_numeric(pdf_data.get('true_peak'))} dBTP" if pdf_data.get('true_peak') is not None else "",
            f"LRA {_preview_numeric(pdf_data.get('lra'))}" if pdf_data.get('lra') is not None else "",
            f"{pdf_data.get('channels')}" if pdf_data.get('channels') else "",
            f"{pdf_data.get('sample_rate')} Hz" if pdf_data.get('sample_rate') else "",
            f"{pdf_data.get('bit_depth')} bit" if pdf_data.get('bit_depth') else "",
        ])]
        record = add_record(
            "pdf",
            pdf_file,
            slot=slot,
            details=[part for part in details if part],
            item_warnings=item_warnings,
            extra={
                "metrics": {
                    "lufs": pdf_data.get("lufs"),
                    "true_peak": pdf_data.get("true_peak"),
                    "lra": pdf_data.get("lra"),
                }
            },
        )
        for message in record["warnings"]:
            collect_warning(f"{record['name']}: {message}")

    selected_audio_records = [
        data.get("record")
        for data in selected_audio_by_slot.values()
        if data.get("record")
    ]
    selected_audio_durations_ms = [
        _preview_duration_ms((record.get("metrics") or {}).get("duration"))
        for record in selected_audio_records
    ]
    selected_audio_durations_ms = [value for value in selected_audio_durations_ms if value is not None]
    all_audio_match = len(selected_audio_durations_ms) <= 1 or len(set(selected_audio_durations_ms)) == 1

    video_record = next(
        (
            item for item in recognized
            if item.get("kind") == "video" and item.get("slot") == "video"
        ),
        None,
    )
    video_duration_ms = None
    fps = 25.0
    if video_record:
        video_metrics = video_record.get("metrics") or {}
        video_duration_ms = _preview_duration_ms(video_metrics.get("duration"))
        fps = float(video_metrics.get("fps") or 25.0)

    duration_context = {
        "fps": fps,
        "video_duration_ms": video_duration_ms,
        "audio_durations_ms": selected_audio_durations_ms,
        "audio_durations_ms_set": set(selected_audio_durations_ms),
        "all_audio_match": all_audio_match,
    }
    duration_records = list(selected_audio_records)
    if video_record:
        duration_records.append(video_record)

    for record in duration_records:
        for issue in preview_duration_issues(record, duration_context):
            if issue not in record["warnings"]:
                record["warnings"].append(issue)
                collect_warning(f"{record['name']}: {issue}")

    slot_summary = {key: None for key in (
        'video', 'csv', 'params',
        'audio_20_c', 'audio_20_uc', 'audio_51_c', 'audio_51_uc',
        'pdf_20', 'pdf_51', 'pdf_20_c', 'pdf_20_uc', 'pdf_51_c', 'pdf_51_uc',
    )}
    for item in recognized:
        if item.get("slot") in {"video", "csv", "params"} and not slot_summary.get(item["slot"]):
            slot_summary[item["slot"]] = item["name"]
    for slot_name, data in selected_audio_by_slot.items():
        slot_summary[slot_name] = data["name"]
    for slot_name, file_name in used_pdf_slots.items():
        slot_summary[slot_name] = file_name

    recognized.sort(key=lambda item: (
        PREVIEW_SLOT_ORDER.get(item.get("slot"), 999),
        item["name"].lower(),
    ))

    return {
        "recognized": recognized,
        "warnings": warnings,
        "report_type": {
            "detected": detected_report_type,
            "reason": report_analysis.get("reason"),
            "scores": report_analysis.get("scores", {}),
        },
        "slot_summary": slot_summary,
        "counts": {key: len(value) for key, value in snapshot.items()},
        "params": params_preview or {},
    }


class TruePeakResultsDialog(QDialog):
    """Диалог с результатами точного измерения True Peak (ITU-R BS.1770-4, 4x oversampling)"""

    def __init__(self, measurement_results: list, parent=None):
        """
        Args:
            measurement_results: список dict с ключами:
                'key' — ключ в tech_info (напр. 'pdf_20_c')
                'label' — человекочитаемое имя (напр. '2.0 Cens')
                'youlean_value' — значение из Youlean (1 decimal)
                'precise_value' — точное измерение (2 decimals)
                'per_channel' — list True Peak по каналам
                'audio_file' — путь к аудиофайлу
        """
        super().__init__(parent)
        self.setWindowTitle("True Peak — результаты измерения")
        self.setMinimumWidth(520)
        self._accepted = False
        self.setStyleSheet("""
            QDialog {
                background: #FFFFFF;
            }
            QLabel {
                background: transparent;
                border: none;
            }
        """)

        main_layout = QVBoxLayout()
        main_layout.setSpacing(12)
        main_layout.setContentsMargins(20, 20, 20, 20)

        title_label = QLabel("True Peak — измерение")
        title_label.setFont(QFont(".AppleSystemUIFont", 13, QFont.DemiBold))
        title_label.setStyleSheet("color: #1D1D1F;")
        main_layout.addWidget(title_label)

        info_label = QLabel(
            "Измерение True Peak (ITU-R BS.1770-4, 4x oversampling)."
        )
        info_label.setWordWrap(True)
        info_label.setFont(QFont(".AppleSystemUIFont", 11))
        info_label.setStyleSheet("color: #86868B;")
        main_layout.addWidget(info_label)

        # Таблица результатов
        for item in measurement_results:
            row_widget = QFrame()
            row_widget.setStyleSheet("""
                QFrame {
                    background-color: #F5F5F7;
                    border-radius: 8px;
                    border: 1px solid #E5E5EA;
                }
            """)
            row_layout = QVBoxLayout()
            row_layout.setContentsMargins(12, 8, 12, 8)
            row_layout.setSpacing(4)

            # Заголовок: название файла/версии
            header = QLabel(f"{item['label']}")
            header.setFont(QFont(".AppleSystemUIFont", 12, QFont.DemiBold))
            header.setStyleSheet("background: transparent; border: none; color: #1D1D1F;")
            row_layout.addWidget(header)

            # Файл
            if item.get('audio_file'):
                file_label = QLabel(f"Файл: {Path(item['audio_file']).name}")
                file_label.setFont(QFont(".AppleSystemUIFont", 10))
                file_label.setStyleSheet("background: transparent; border: none; color: #86868B;")
                row_layout.addWidget(file_label)

            # Значения
            youlean_str = f"{item['youlean_value']:.1f}" if item['youlean_value'] is not None else "N/A"
            precise = item['precise_value']
            precise_str = f"{precise:.2f}" if precise is not None else "N/A"

            # Определяем цвет: зеленый если ≤ -2.0, красный если > -2.0
            threshold = -2.0
            if precise is not None and precise <= threshold:
                color = "#34C759"  # зеленый
                verdict = "PASS"
            elif precise is not None:
                color = "#FF3B30"  # красный
                verdict = "FAIL"
            else:
                color = "#86868B"
                verdict = "—"

            values_layout = QHBoxLayout()
            values_layout.setSpacing(16)

            yl_label = QLabel(f"Youlean: {youlean_str} dBTP")
            yl_label.setFont(QFont(".AppleSystemUIFont", 11))
            yl_label.setStyleSheet("background: transparent; border: none; color: #86868B;")
            values_layout.addWidget(yl_label)

            arrow_label = QLabel("→")
            arrow_label.setFont(QFont(".AppleSystemUIFont", 13))
            arrow_label.setStyleSheet("background: transparent; border: none; color: #86868B;")
            values_layout.addWidget(arrow_label)

            precise_label = QLabel(f"Точно: {precise_str} dBTP")
            precise_label.setFont(QFont(".AppleSystemUIFont", 12, QFont.Bold))
            precise_label.setStyleSheet(f"background: transparent; border: none; color: {color};")
            values_layout.addWidget(precise_label)

            verdict_label = QLabel(verdict)
            verdict_label.setFont(QFont(".AppleSystemUIFont", 11, QFont.Bold))
            verdict_label.setStyleSheet(f"background: transparent; border: none; color: {color};")
            values_layout.addWidget(verdict_label)

            values_layout.addStretch()
            row_layout.addLayout(values_layout)

            # Per-channel info
            per_ch = item.get('per_channel', [])
            if per_ch and len(per_ch) > 1:
                ch_names = ["L", "R", "C", "LFE", "Ls", "Rs", "Lrs", "Rrs"]
                ch_parts = []
                for i, v in enumerate(per_ch):
                    name = ch_names[i] if i < len(ch_names) else f"Ch{i+1}"
                    ch_parts.append(f"{name}: {v:.2f}")
                ch_label = QLabel("По каналам: " + "  |  ".join(ch_parts))
                ch_label.setFont(QFont(".AppleSystemUIFont", 9))
                ch_label.setStyleSheet("background: transparent; border: none; color: #AEAEB2;")
                ch_label.setWordWrap(True)
                row_layout.addWidget(ch_label)

            row_widget.setLayout(row_layout)
            main_layout.addWidget(row_widget)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.button(QDialogButtonBox.Ok).setText("Применить точные значения")
        buttons.button(QDialogButtonBox.Cancel).setText("Оставить Youlean")
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        cancel_btn = buttons.button(QDialogButtonBox.Cancel)
        ok_btn = buttons.button(QDialogButtonBox.Ok)
        cancel_btn.setStyleSheet("""
            QPushButton {
                background: #FFFFFF;
                color: #1D1D1F;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                padding: 6px 16px;
                font-family: ".AppleSystemUIFont";
                font-size: 12px;
                min-width: 120px;
            }
            QPushButton:hover { background: #F5F5F7; }
        """)
        ok_btn.setStyleSheet("""
            QPushButton {
                background-color: #007AFF;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 6px 16px;
                font-family: ".AppleSystemUIFont";
                font-size: 12px;
                font-weight: 600;
                min-width: 140px;
            }
            QPushButton:hover { background-color: #0063D1; }
        """)
        main_layout.addWidget(buttons)

        self.setLayout(main_layout)
        self._measurement_results = measurement_results

    def get_results(self) -> dict:
        """Возвращает dict {key: precise_true_peak_value}."""
        return {
            item['key']: item['precise_value']
            for item in self._measurement_results
            if item['precise_value'] is not None
        }


class SettingsDialog(QDialog):
    """Диалог настроек приложения"""

    CONFIG_FILE = CONFIG_DIR / "settings.json"
    _LEGACY_CONFIG_FILE = Path.home() / ".beast_auto_reporter_settings.json"

    @classmethod
    def load_settings(cls) -> dict:
        """Загрузка настроек из JSON-файла (с совместимостью со старым .txt)"""
        import json
        migrate_legacy_config_file(cls._LEGACY_CONFIG_FILE, cls.CONFIG_FILE)
        defaults = {
            "name": "",
            "delete_sources_after_copy": False,
            "auto_detect_report_type": False,
            "check_file_consistency": True,
            "auto_reset_after_done": True,
            "extended_analysis_enabled": False,
            "yandex_disk_token": "",
            "yandex_auto_upload": False,
            "yandex_disk_roots": [REPORTS_ROOT],
        }

        if cls.CONFIG_FILE.exists():
            try:
                data = json.loads(cls.CONFIG_FILE.read_text(encoding="utf-8"))
                defaults.update(data)
                if "yandex_disk_roots" not in data and data.get("yandex_disk_root"):
                    # Миграция с предыдущей версии (один корень, до введения
                    # поддержки нескольких) — переносим как единственный
                    # элемент нового списка.
                    defaults["yandex_disk_roots"] = [data["yandex_disk_root"]]
                return defaults
            except Exception as e:
                # Битый файл настроек: сохраняем копию для ручного
                # восстановления — раньше он молча заменялся дефолтами и
                # затирался при следующем сохранении (пользователь терял
                # токен и настройки без единого сообщения).
                backup = cls.CONFIG_FILE.with_name(cls.CONFIG_FILE.name + ".bak")
                try:
                    shutil.copy2(cls.CONFIG_FILE, backup)
                    logger.warning(
                        f"Файл настроек повреждён ({e}) — копия сохранена в {backup}, "
                        f"используются настройки по умолчанию"
                    )
                except OSError as backup_error:
                    logger.warning(
                        f"Файл настроек повреждён ({e}), сделать резервную копию "
                        f"не удалось: {backup_error}"
                    )

        # Обратная совместимость: читаем старый .txt с именем
        old_file = Path.home() / ".beast_auto_reporter_config.txt"
        if old_file.exists():
            try:
                defaults["name"] = old_file.read_text(encoding="utf-8").strip()
            except Exception:
                pass
        return defaults

    @classmethod
    def save_settings(cls, data: dict):
        """Сохранение настроек в JSON-файл"""
        import json
        try:
            ensure_parent_dir(cls.CONFIG_FILE)
            cls.CONFIG_FILE.write_text(
                json.dumps(data, ensure_ascii=False, indent=2),
                encoding="utf-8"
            )
        except Exception as e:
            logger.warning(f"Не удалось сохранить настройки: {e}")

    @classmethod
    def get_yandex_token(cls) -> str:
        """OAuth-токен Яндекс.Диска (пустая строка, если не задан).

        Основное хранилище — Связка ключей macOS (см. src/secret_store.py);
        токен, сохранённый прошлой версией в settings.json открытым текстом,
        при первом чтении прозрачно переносится в Связку и вычищается из
        файла. Если Связка недоступна (ошибка утилиты security) — прежнее
        поведение, токен читается/остаётся в settings.json.
        """
        token = secret_store.load_token()
        if token:
            return token
        legacy = cls.load_settings().get("yandex_disk_token", "").strip()
        if legacy and token == "":
            # Связка доступна, но записи в ней нет — мигрируем токен из
            # файла. Из settings.json вычищаем только после успешной записи.
            if secret_store.save_token(legacy):
                settings = cls.load_settings()
                settings["yandex_disk_token"] = ""
                cls.save_settings(settings)
        return legacy

    @classmethod
    def get_yandex_roots(cls) -> list:
        """Список корневых папок на Диске для отчётов (минимум одна — /отчеты по умолчанию).

        Приложение ищет и создаёт папки серий сразу во всех перечисленных
        корнях — см. find_series_folder/resolve_target_path в report_uploader.py.
        """
        raw = cls.load_settings().get("yandex_disk_roots") or []
        cleaned = []
        for r in raw:
            r = str(r).strip().rstrip("/")
            if not r:
                continue
            if not r.startswith("/"):
                r = f"/{r}"
            if r not in cleaned:
                cleaned.append(r)
        return cleaned or [REPORTS_ROOT]

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Настройки")
        self.setModal(True)
        self.setMinimumWidth(340)
        self.setStyleSheet("""
            QDialog {
                background: #FFFFFF;
            }
            QLabel {
                background: transparent;
                border: none;
            }
        """)

        self._settings = self.load_settings()

        layout = QVBoxLayout()
        layout.setSpacing(12)
        layout.setContentsMargins(20, 16, 20, 16)

        # Заголовок секции
        section_label = QLabel("Настройки")
        section_label.setFont(QFont(".AppleSystemUIFont", 13, QFont.DemiBold))
        section_label.setStyleSheet("color: #1D1D1F;")
        layout.addWidget(section_label)

        # Чекбокс удаления исходников
        checkbox_style = """
            QCheckBox {
                background: transparent; border: none; color: #1D1D1F;
                spacing: 8px; font-size: 12px;
            }
            QCheckBox::indicator {
                width: 16px; height: 16px; border-radius: 4px;
            }
            QCheckBox::indicator:unchecked {
                background-color: transparent; border: 1.5px solid #86868B;
            }
            QCheckBox::indicator:checked {
                background-color: #007AFF; border: 1.5px solid #007AFF;
            }
        """
        self.delete_sources_cb = QCheckBox("Удалять исходники после копирования")
        self.delete_sources_cb.setFont(QFont(".AppleSystemUIFont", 12))
        self.delete_sources_cb.setStyleSheet(checkbox_style)
        self.delete_sources_cb.setIcon(make_icon("trash", "#86868B", 14))
        self.delete_sources_cb.setIconSize(QSize(14, 14))
        self.delete_sources_cb.setChecked(self._settings.get("delete_sources_after_copy", False))
        self.delete_sources_cb.setToolTip(
            "PDF, CSV и файлы параметров будут удалены из исходной папки\n"
            "после копирования в папку отчёта."
        )
        layout.addWidget(self.delete_sources_cb)

        hint = QLabel("PDF, CSV и файлы параметров будут удалены\nиз исходной папки после копирования.")
        hint.setFont(QFont(".AppleSystemUIFont", 10))
        hint.setStyleSheet("color: #86868B; margin-left: 24px;")
        layout.addWidget(hint)

        layout.addSpacing(8)

        # Чекбокс проверки соответствия файлов
        self.check_files_cb = QCheckBox("Проверять соответствие файлов")
        self.check_files_cb.setFont(QFont(".AppleSystemUIFont", 12))
        self.check_files_cb.setStyleSheet(checkbox_style)
        self.check_files_cb.setIcon(make_icon("search", "#86868B", 14))
        self.check_files_cb.setIconSize(QSize(14, 14))
        self.check_files_cb.setChecked(self._settings.get("check_file_consistency", True))
        self.check_files_cb.setToolTip(
            "Перед генерацией отчёта проверять, что аудио, PDF и CSV\n"
            "относятся к одному и тому же материалу (по имени файла)."
        )
        layout.addWidget(self.check_files_cb)

        check_hint = QLabel("Если имена аудио, PDF или CSV различаются,\nперед генерацией появится предупреждение.")
        check_hint.setFont(QFont(".AppleSystemUIFont", 10))
        check_hint.setStyleSheet("color: #86868B; margin-left: 24px;")
        layout.addWidget(check_hint)

        layout.addSpacing(8)

        # Чекбокс автоготовности после завершения отчёта
        self.auto_reset_cb = QCheckBox("Готовность к новому отчёту через 5 сек")
        self.auto_reset_cb.setFont(QFont(".AppleSystemUIFont", 12))
        self.auto_reset_cb.setStyleSheet(checkbox_style)
        self.auto_reset_cb.setIcon(make_icon("copy", "#86868B", 14))
        self.auto_reset_cb.setIconSize(QSize(14, 14))
        self.auto_reset_cb.setChecked(self._settings.get("auto_reset_after_done", True))
        self.auto_reset_cb.setToolTip(
            "После завершения отчёта приложение само очистит список\n"
            "файлов через 5 секунд, не дожидаясь нажатия «Очистить»."
        )
        layout.addWidget(self.auto_reset_cb)

        auto_reset_hint = QLabel("Файлы очищаются автоматически, чтобы можно\nбыло сразу закинуть следующие.")
        auto_reset_hint.setFont(QFont(".AppleSystemUIFont", 10))
        auto_reset_hint.setStyleSheet("color: #86868B; margin-left: 24px;")
        layout.addWidget(auto_reset_hint)

        layout.addSpacing(8)

        # Чекбокс расширенного анализа (PyLoudNorm)
        self.extended_analysis_cb = QCheckBox("Расширенный анализ")
        self.extended_analysis_cb.setFont(QFont(".AppleSystemUIFont", 12))
        self.extended_analysis_cb.setStyleSheet(checkbox_style)
        self.extended_analysis_cb.setIcon(make_icon("gear", "#86868B", 14))
        self.extended_analysis_cb.setIconSize(QSize(14, 14))
        self.extended_analysis_cb.setChecked(self._settings.get("extended_analysis_enabled", False))
        self.extended_analysis_cb.setToolTip(
            "Полный анализ громкости через PyLoudNorm (ITU-R BS.1770)\n"
            "с экспортом дополнительных CSV/HTML данных.\n"
            "Увеличивает время обработки отчёта."
        )
        layout.addWidget(self.extended_analysis_cb)

        extended_analysis_hint = QLabel("Дополнительные данные громкости,\nно отчёт формируется дольше.")
        extended_analysis_hint.setFont(QFont(".AppleSystemUIFont", 10))
        extended_analysis_hint.setStyleSheet("color: #86868B; margin-left: 24px;")
        layout.addWidget(extended_analysis_hint)

        layout.addSpacing(8)

        # Токен Яндекс.Диска
        yandex_label = QLabel("Токен Яндекс.Диска")
        yandex_label.setFont(QFont(".AppleSystemUIFont", 12))
        yandex_label.setStyleSheet("color: #1D1D1F;")
        layout.addWidget(yandex_label)

        self.yandex_token_edit = QLineEdit(self.get_yandex_token())
        self.yandex_token_edit.setEchoMode(QLineEdit.Password)
        self.yandex_token_edit.setPlaceholderText("OAuth-токен для отправки отчётов на Диск")
        self.yandex_token_edit.setFont(QFont(".AppleSystemUIFont", 12))
        self.yandex_token_edit.setStyleSheet("""
            QLineEdit {
                background: #F5F5F7;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                padding: 6px 10px;
                color: #1D1D1F;
            }
            QLineEdit:focus { border: 1px solid #007AFF; }
        """)
        self.yandex_token_edit.setToolTip(
            "OAuth-токен приложения на oauth.yandex.ru с правами disk:write/disk:read.\n"
            "Нужен для кнопки «Отправить на Диск» в главном окне."
        )
        layout.addWidget(self.yandex_token_edit)

        yandex_hint = QLabel("Используется для отправки отчёта на Диск\nи сравнения с предыдущей версией серии.")
        yandex_hint.setFont(QFont(".AppleSystemUIFont", 10))
        yandex_hint.setStyleSheet("color: #86868B;")
        layout.addWidget(yandex_hint)

        yandex_check_row = QHBoxLayout()
        yandex_check_row.setSpacing(8)
        self.yandex_check_status_label = QLabel("")
        self.yandex_check_status_label.setFont(QFont(".AppleSystemUIFont", 10))
        self.yandex_check_status_label.setWordWrap(True)
        yandex_check_row.addWidget(self.yandex_check_status_label, 1)
        self.yandex_check_token_btn = QPushButton("Проверить")
        self.yandex_check_token_btn.setFont(QFont(".AppleSystemUIFont", 11))
        self.yandex_check_token_btn.setStyleSheet("""
            QPushButton {
                background: #FFFFFF;
                color: #007AFF;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                padding: 5px 12px;
            }
            QPushButton:hover { background: #F5F5F7; }
        """)
        self.yandex_check_token_btn.clicked.connect(self._on_check_yandex_token_clicked)
        yandex_check_row.addWidget(self.yandex_check_token_btn)
        layout.addLayout(yandex_check_row)

        self._yandex_check_thread = None

        layout.addSpacing(8)

        # Корневые папки на Диске для отчётов
        yandex_roots_label = QLabel("Папки на Диске для отчётов")
        yandex_roots_label.setFont(QFont(".AppleSystemUIFont", 12))
        yandex_roots_label.setStyleSheet("color: #1D1D1F;")
        layout.addWidget(yandex_roots_label)

        self.yandex_roots_list = QListWidget()
        self.yandex_roots_list.setMaximumHeight(84)
        self.yandex_roots_list.setFont(QFont(".AppleSystemUIFont", 12))
        self.yandex_roots_list.setStyleSheet("""
            QListWidget {
                background: #F5F5F7;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                color: #1D1D1F;
            }
        """)
        self.yandex_roots_list.setToolTip(
            "Папки на Яндекс.Диске, внутри которых ищутся/создаются папки\n"
            "серий и куда отправляются отчёты. Поиск идёт сразу по всем."
        )
        for root in (self._settings.get("yandex_disk_roots") or [REPORTS_ROOT]):
            self.yandex_roots_list.addItem(root)
        layout.addWidget(self.yandex_roots_list)

        yandex_roots_row = QHBoxLayout()
        yandex_roots_row.setSpacing(8)
        self.yandex_new_root_edit = QLineEdit()
        self.yandex_new_root_edit.setPlaceholderText("/новая папка")
        self.yandex_new_root_edit.setFont(QFont(".AppleSystemUIFont", 12))
        self.yandex_new_root_edit.setStyleSheet("""
            QLineEdit {
                background: #F5F5F7;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                padding: 6px 10px;
                color: #1D1D1F;
            }
            QLineEdit:focus { border: 1px solid #007AFF; }
        """)
        self.yandex_new_root_edit.returnPressed.connect(self._on_add_yandex_root)
        yandex_roots_row.addWidget(self.yandex_new_root_edit, 1)

        yandex_add_root_btn = QPushButton("Добавить")
        yandex_add_root_btn.setFont(QFont(".AppleSystemUIFont", 11))
        yandex_add_root_btn.setStyleSheet("""
            QPushButton {
                background: #FFFFFF;
                color: #007AFF;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                padding: 5px 12px;
            }
            QPushButton:hover { background: #F5F5F7; }
        """)
        yandex_add_root_btn.clicked.connect(self._on_add_yandex_root)
        yandex_roots_row.addWidget(yandex_add_root_btn)
        layout.addLayout(yandex_roots_row)

        yandex_remove_root_btn = QPushButton("Удалить выбранную")
        yandex_remove_root_btn.setFont(QFont(".AppleSystemUIFont", 11))
        yandex_remove_root_btn.setStyleSheet("""
            QPushButton {
                background: #FFFFFF;
                color: #FF3B30;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                padding: 5px 12px;
            }
            QPushButton:hover { background: #FFF1F0; }
        """)
        yandex_remove_root_btn.clicked.connect(self._on_remove_yandex_root)
        layout.addWidget(yandex_remove_root_btn)

        yandex_roots_hint = QLabel("Приложение ищет и создаёт папки отчётов\nсразу во всех перечисленных папках.")
        yandex_roots_hint.setFont(QFont(".AppleSystemUIFont", 10))
        yandex_roots_hint.setStyleSheet("color: #86868B;")
        layout.addWidget(yandex_roots_hint)

        layout.addSpacing(8)

        # Автоматическая отправка на Диск после генерации
        self.yandex_auto_upload_cb = QCheckBox("Автоматически отправлять на Диск после генерации")
        self.yandex_auto_upload_cb.setFont(QFont(".AppleSystemUIFont", 12))
        self.yandex_auto_upload_cb.setStyleSheet(checkbox_style)
        self.yandex_auto_upload_cb.setIcon(make_icon("folder_open", "#86868B", 14))
        self.yandex_auto_upload_cb.setIconSize(QSize(14, 14))
        self.yandex_auto_upload_cb.setChecked(self._settings.get("yandex_auto_upload", False))
        self.yandex_auto_upload_cb.setToolTip(
            "Готовый отчёт сам уходит в очередь на Диск, без нажатия «Отправить».\n"
            "Если папка сериала на Диске ещё не найдена — отчёт останется\n"
            "в очереди, дожидаясь ручного решения (диалог не всплывает сам)."
        )
        layout.addWidget(self.yandex_auto_upload_cb)

        yandex_auto_hint = QLabel("Если папка сериала не найдена — отчёт\nостанется в очереди для ручного действия.")
        yandex_auto_hint.setFont(QFont(".AppleSystemUIFont", 10))
        yandex_auto_hint.setStyleSheet("color: #86868B; margin-left: 24px;")
        layout.addWidget(yandex_auto_hint)

        layout.addSpacing(8)

        # Версия приложения + ручная проверка обновлений
        update_row = QHBoxLayout()
        update_row.setSpacing(8)
        version_label = QLabel(f"Версия {APP_VERSION}")
        version_label.setFont(QFont(".AppleSystemUIFont", 11))
        version_label.setStyleSheet("color: #86868B;")
        update_row.addWidget(version_label)
        update_row.addStretch()
        check_updates_btn = QPushButton("Проверить обновления")
        check_updates_btn.setFont(QFont(".AppleSystemUIFont", 11))
        check_updates_btn.setStyleSheet("""
            QPushButton {
                background: #FFFFFF;
                color: #007AFF;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                padding: 5px 12px;
            }
            QPushButton:hover { background: #F5F5F7; }
        """)
        check_updates_btn.clicked.connect(self._on_check_updates_clicked)
        update_row.addWidget(check_updates_btn)
        layout.addLayout(update_row)

        layout.addSpacing(8)

        # Кнопки OK / Отмена
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.button(QDialogButtonBox.Ok).setText("Сохранить")
        buttons.button(QDialogButtonBox.Cancel).setText("Отмена")
        buttons.accepted.connect(self._on_accept)
        buttons.rejected.connect(self.reject)
        buttons.button(QDialogButtonBox.Cancel).setStyleSheet("""
            QPushButton {
                background: #FFFFFF;
                color: #1D1D1F;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                padding: 6px 16px;
                font-family: ".AppleSystemUIFont";
                font-size: 12px;
            }
            QPushButton:hover { background: #F5F5F7; }
        """)
        buttons.button(QDialogButtonBox.Ok).setStyleSheet("""
            QPushButton {
                background-color: #007AFF;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 6px 16px;
                font-family: ".AppleSystemUIFont";
                font-size: 12px;
                font-weight: 600;
            }
            QPushButton:hover { background-color: #0063D1; }
        """)
        layout.addWidget(buttons)

        self.setLayout(layout)

    def _on_accept(self):
        self._settings["delete_sources_after_copy"] = self.delete_sources_cb.isChecked()
        self._settings["check_file_consistency"] = self.check_files_cb.isChecked()
        self._settings["auto_reset_after_done"] = self.auto_reset_cb.isChecked()
        self._settings["extended_analysis_enabled"] = self.extended_analysis_cb.isChecked()
        token = self.yandex_token_edit.text().strip()
        if secret_store.save_token(token):
            # Токен ушёл в Связку ключей — в файле настроек его не храним.
            self._settings["yandex_disk_token"] = ""
        else:
            # Связка недоступна — прежнее поведение (открытым текстом в файле).
            self._settings["yandex_disk_token"] = token
        self._settings["yandex_auto_upload"] = self.yandex_auto_upload_cb.isChecked()
        roots = [self.yandex_roots_list.item(i).text() for i in range(self.yandex_roots_list.count())]
        self._settings["yandex_disk_roots"] = roots or [REPORTS_ROOT]
        self._settings.pop("yandex_disk_root", None)  # устаревший одиночный ключ прошлой версии
        self.save_settings(self._settings)
        self.accept()

    def _on_add_yandex_root(self):
        text = self.yandex_new_root_edit.text().strip().rstrip("/")
        if not text:
            return
        if not text.startswith("/"):
            text = f"/{text}"
        existing = [self.yandex_roots_list.item(i).text() for i in range(self.yandex_roots_list.count())]
        if text not in existing:
            self.yandex_roots_list.addItem(text)
        self.yandex_new_root_edit.clear()

    def _on_remove_yandex_root(self):
        row = self.yandex_roots_list.currentRow()
        if row < 0:
            QMessageBox.information(self, "Не выбрано", "Выберите папку в списке, чтобы удалить.")
            return
        if self.yandex_roots_list.count() <= 1:
            QMessageBox.warning(self, "Нельзя удалить", "Должна остаться хотя бы одна папка.")
            return
        self.yandex_roots_list.takeItem(row)

    def done(self, r):
        # accept()/reject() (Save/Cancel/titlebar-крестик — все три пути идут
        # через done()) — останавливаем фоновую проверку токена, иначе при
        # закрытии диалога до её завершения PyQt валит процесс
        # "QThread: Destroyed while thread is still running".
        _stop_thread(getattr(self, "_yandex_check_thread", None))
        super().done(r)

    def _on_check_updates_clicked(self):
        parent = self.parent()
        if parent is not None:
            parent._check_for_updates(silent=False)

    def _on_check_yandex_token_clicked(self):
        token = self.yandex_token_edit.text().strip()
        if not token:
            self.yandex_check_status_label.setStyleSheet("color: #FF3B30;")
            self.yandex_check_status_label.setText("Введите токен перед проверкой.")
            return

        self.yandex_check_token_btn.setEnabled(False)
        self.yandex_check_status_label.setStyleSheet("color: #86868B;")
        self.yandex_check_status_label.setText("Проверяем...")

        self._yandex_check_thread = YandexDiskTokenCheckThread(token)
        self._yandex_check_thread.finished_check.connect(self._on_yandex_token_check_finished)
        self._yandex_check_thread.start()

    def _on_yandex_token_check_finished(self, success: bool, message: str):
        self.yandex_check_token_btn.setEnabled(True)
        if success:
            self.yandex_check_status_label.setStyleSheet("color: #34C759;")
        else:
            self.yandex_check_status_label.setStyleSheet("color: #FF3B30;")
        self.yandex_check_status_label.setText(message)


class PreviewAnalysisThread(QThread):
    """Фоновый анализ файлов для предпросмотра перед генерацией отчета."""

    preview_ready = pyqtSignal(object, int)
    preview_failed = pyqtSignal(str, int)

    def __init__(self, files_data: dict, epoch: int, report_type: str):
        super().__init__()
        self.files_data = {
            'audio': list(files_data.get('audio', [])),
            'video': list(files_data.get('video', [])),
            'csv': list(files_data.get('csv', [])),
            'pdf': list(files_data.get('pdf', [])),
            'params': list(files_data.get('params', [])),
        }
        self.epoch = epoch
        self.report_type = report_type

    def run(self):
        try:
            preview = analyze_files_for_preview(self.files_data, self.report_type)
            self.preview_ready.emit(preview, self.epoch)
        except Exception as e:
            logger.error("Ошибка предпросмотра: %s", e, exc_info=True)
            self.preview_failed.emit(str(e), self.epoch)


class ProcessingThread(QThread):
    """Поток для обработки файлов и генерации отчета"""

    status_update = pyqtSignal(str)
    progress_update = pyqtSignal(int)
    finished = pyqtSignal(bool, str)
    true_peak_results_ready = pyqtSignal(list)  # результаты измерений для диалога

    def __init__(self, app, files_data, report_type, output_folder, pyloudnorm_enabled=False,
                 tp_verify_enabled=False, delete_sources=False):
        super().__init__()
        self.app = app
        self.files_data = files_data
        self.report_type = report_type
        self.output_folder = output_folder
        self.pyloudnorm_enabled = pyloudnorm_enabled
        self.tp_verify_enabled = tp_verify_enabled
        self.delete_sources = delete_sources
        self._tp_verify_loop = None
        self._tp_verify_results = None
    
    def run(self):
        """Запуск обработки"""
        logger.debug("=== ProcessingThread.run() STARTED ===")
        try:
            import shutil
            from pathlib import Path
            
            logger.info(f"=== НАЧАЛО ОБРАБОТКИ ===")
            logger.info(f"Тип отчета: {self.report_type}")
            logger.info(f"PyLoudNorm включен: {self.pyloudnorm_enabled}")
            logger.info(f"Файлов: {sum(len(v) for v in self.files_data.values())}")
            logger.info(f"Output folder: {self.output_folder}")
            logger.debug(f"files_data keys = {list(self.files_data.keys())}")
            
            # Определяем имя файла из аудио или видео
            audio_files = self.files_data.get('audio', [])
            video_files = self.files_data.get('video', [])
            logger.debug(f"audio_files = {audio_files}")
            logger.debug(f"video_files = {video_files}")
            
            # Также проверяем специфичные ключи для 5.1
            audio_51_c = self.files_data.get('audio_51_c', [])
            audio_51_uc = self.files_data.get('audio_51_uc', [])
            audio_20_c = self.files_data.get('audio_20_c', [])
            audio_20_uc = self.files_data.get('audio_20_uc', [])
            logger.debug(f"audio_51_c = {audio_51_c}")
            logger.debug(f"audio_51_uc = {audio_51_uc}")
            csv_files = self.files_data.get('csv', [])
            pdf_files = self.files_data.get('pdf', [])
            params_files = self.files_data.get('params', [])
            
            logger.info(f"Audio files found: {len(audio_files)}")
            logger.info(f"Video files found: {len(video_files)}")
            logger.info(f"CSV files found: {len(csv_files)}")
            logger.info(f"PDF files found: {len(pdf_files)}")
            logger.info(f"Params files found: {len(params_files)}")
            
            # ДИАГНОСТИКА: ЛОГИРУЕМ PDF ФАЙЛЫ
            if pdf_files:
                logger.info("=== PDF FILES DETAILS ===")
                for i, pdf in enumerate(pdf_files):
                    logger.info(f"  [{i}] {pdf}")
                logger.info("=== END PDF FILES ===")
            else:
                logger.warning("⚠️ NO PDF FILES LOADED!")
            
            # Логируем report_type для диагностики
            logger.info(f"Report type: {self.report_type}")
            
            # Объединяем все аудио файлы для pyloudnorm анализа
            all_audio_files = list(audio_files)
            all_audio_files.extend(audio_51_c)
            all_audio_files.extend(audio_51_uc)
            all_audio_files.extend(audio_20_c)
            all_audio_files.extend(audio_20_uc)
            logger.debug(f"all_audio_files for analysis = {all_audio_files}")
            
            base_name = "отчет"
            if audio_files:
                base_name = sanitize_base_name(audio_files[0])
            elif video_files:
                base_name = sanitize_base_name(video_files[0])
            
            # Папка уже создана в start_processing через create_output_folder()
            output_dir = Path(self.output_folder)
            logger.info(f"Выходная папка: {output_dir}")

            # === PDF-only режим: вставка PDF в DOCX в самом начале ===
            if pdf_files and not audio_files and not video_files and not csv_files:
                self.status_update.emit("📄 Вставка PDF в DOCX...")
                self.progress_update.emit(10)
                
                copied_pdfs = []
                for pdf_file in pdf_files:
                    dest = output_dir / Path(pdf_file).name
                    shutil.copy2(pdf_file, dest)
                    copied_pdfs.append(dest)

                simple_docx_path = output_dir / f"pdf_{base_name}.docx"
                ordered_pdfs = order_pdfs_by_channel(copied_pdfs)
                PdfOnlyDocxBuilder.build(simple_docx_path, ordered_pdfs)

                # Удаляем исходники после завершения работы с ними
                if self.delete_sources:
                    for pdf_file in pdf_files:
                        try:
                            Path(pdf_file).unlink()
                            logger.info(f"🗑 Удалён исходник: {Path(pdf_file).name}")
                        except Exception as e:
                            logger.warning(f"Не удалось удалить {pdf_file}: {e}")

                self.progress_update.emit(100)
                self.status_update.emit("✅ Готово!")

                files_in_output = list(output_dir.glob('*'))
                success_msg = (
                    f"PDF вставлены в DOCX!\n"
                    f"Папка: {output_dir.name}\n"
                    f"Файлов: {len(files_in_output)}"
                )
                self.finished.emit(True, success_msg)
                logger.info(f"✅ PDF-only DOCX создан: {simple_docx_path}")
                return
            
            # Копируем PDF файлы в выходную папку и разделяем по типам
            self.status_update.emit("📋 Копирование файлов...")
            self.progress_update.emit(10)
            
            # Словари для хранения путей к разным типам PDF
            pdf_paths = {
                '20_c': None,
                '20_uc': None,
                '20': None,
                '51_c': None,
                '51_uc': None,
                '51': None
            }
            copied_pdfs = []
            pdf_by_stem = {}
            
            for pdf_file in pdf_files:
                dest = output_dir / Path(pdf_file).name
                shutil.copy2(pdf_file, dest)
                copied_pdfs.append(dest)
                logger.info(f"✅ Скопирован PDF: {Path(pdf_file).name}")
                stem_key = normalize_stem(pdf_file)
                pdf_by_stem.setdefault(stem_key, []).append(dest)
                
                filename = Path(pdf_file).stem.lower()
                
                # Определяем тип PDF с приоритетом (regex, чтобы ловить cens_20 / cens51 и т.п.)
                channel = detect_channel_from_name(filename)
                is_20 = channel == "20"
                is_51 = channel == "51"
                cens_state = detect_cens_state(filename)
                
                if is_20 and cens_state == "c":
                    pdf_paths['20_c'] = str(dest)
                    logger.info(f"📄 PDF 2.0 CENS: {Path(pdf_file).name}")
                elif is_20 and cens_state == "uc":
                    pdf_paths['20_uc'] = str(dest)
                    logger.info(f"📄 PDF 2.0 UNCENS: {Path(pdf_file).name}")
                elif is_20:
                    pdf_paths['20'] = str(dest)
                    logger.info(f"📄 PDF 2.0: {Path(pdf_file).name}")
                elif is_51 and cens_state == "c":
                    pdf_paths['51_c'] = str(dest)
                    logger.info(f"📄 PDF 5.1 CENS: {Path(pdf_file).name}")
                elif is_51 and cens_state == "uc":
                    pdf_paths['51_uc'] = str(dest)
                    logger.info(f"📄 PDF 5.1 UNCENS: {Path(pdf_file).name}")
                elif is_51:
                    pdf_paths['51'] = str(dest)
                    logger.info(f"📄 PDF 5.1: {Path(pdf_file).name}")
            
            # Выбираем PDF для вставки в отчет (приоритет: cens > uncens > общий)
            copied_pdf_20 = pdf_paths['20_c'] or pdf_paths['20_uc'] or pdf_paths['20']
            copied_pdf_51 = pdf_paths['51_c'] or pdf_paths['51_uc'] or pdf_paths['51']

            # Фолбэк: если тип не определился, пробуем по имени среди всех PDF
            if not copied_pdf_20 or not copied_pdf_51:
                for p in copied_pdfs:
                    name = p.stem.lower()
                    if not copied_pdf_20 and detect_channel_from_name(name) == "20":
                        copied_pdf_20 = str(p)
                    if not copied_pdf_51 and detect_channel_from_name(name) == "51":
                        copied_pdf_51 = str(p)

            # Последний фолбэк: берем первые два PDF в списке
            if copied_pdfs:
                if not copied_pdf_20:
                    copied_pdf_20 = str(copied_pdfs[0])
                if not copied_pdf_51 and len(copied_pdfs) > 1:
                    copied_pdf_51 = str(copied_pdfs[1])
            
            logger.info(f"Выбран PDF 2.0 для отчета: {copied_pdf_20}")
            logger.info(f"Выбран PDF 5.1 для отчета: {copied_pdf_51}")
            
            # Копируем CSV если есть, переименовываем по имени папки (без "отчет_")
            if csv_files:
                folder_name = output_dir.name
                # Убираем префикс "отчет_" и возможный суффикс "_1", "_2" и т.д.
                import re
                csv_base = re.sub(r'^отчет_', '', folder_name)
                csv_base = re.sub(r'_\d+$', '', csv_base)
                csv_ext = Path(csv_files[0]).suffix  # .csv
                target_name = f"{csv_base}{csv_ext}"

                # Если CSV уже называется правильно — оставляем как есть
                if Path(csv_files[0]).name == target_name:
                    dest = output_dir / target_name
                else:
                    dest = output_dir / target_name
                shutil.copy2(csv_files[0], dest)
                logger.info(f"✅ Скопирован CSV: {Path(csv_files[0]).name} → {target_name}")

            # Копируем Параметры.txt если есть
            if params_files:
                dest = output_dir / Path(params_files[0]).name
                shutil.copy2(params_files[0], dest)
                logger.info(f"✅ Скопирован файл параметров: {Path(params_files[0]).name}")
            
            # Извлечение технической информации
            self.status_update.emit("🔍 Анализ файлов...")
            self.progress_update.emit(20)
            
            tech_extractor = TechnicalInfoExtractor()
            tech_info = {}
            audio_key_by_stem = {}
            audio_path_by_key = {}
            dcp_audio_metrics_by_key = {}

            # Параметры из файла (если есть)
            if params_files:
                try:
                    tech_info['params'] = tech_extractor.read_params_file(params_files[0])
                    logger.info("✅ Параметры загружены из файла")
                except Exception as e:
                    logger.error(f"❌ Ошибка чтения файла параметров: {e}")
            
            # Обработка аудио файлов
            for audio_file in audio_files:
                logger.info(f"Обработка аудио: {audio_file}")
                filename = Path(audio_file).stem.lower()
                
                # Извлекаем метаданные для определения количества каналов
                audio_info = tech_extractor.extract_audio_info(audio_file)
                if not audio_info:
                    logger.warning(f"⚠️  Не удалось извлечь информацию из аудио: {filename}")
                    continue
                
                channels = audio_info.get('channels', 0)
                try:
                    channels = int(float(channels)) if channels is not None else 0
                except Exception:
                    channels = 0
                is_51 = channels >= 6  # 6 каналов = 5.1
                is_20 = channels == 2  # 2 канала = 2.0
                
                # Определяем тип по имени файла (гибкая проверка)
                channel_hint = detect_channel_from_name(filename)
                channel_meta = "51" if is_51 else ("20" if is_20 else None)
                name_mismatch = bool(channel_meta and channel_hint and channel_meta != channel_hint)
                has_bad_marker = has_incorrect_audio_marker(filename)
                
                cens_state = detect_cens_state(filename)
                suffix = "uc" if cens_state == "uc" else "c"
                
                # Выбираем канал по метаданным, если есть; иначе по имени
                channel_final = channel_meta or channel_hint
                
                if channel_final == "51":
                    key = f'audio_51_{suffix}'
                elif channel_final == "20":
                    key = f'audio_20_{suffix}'
                else:
                    logger.warning(f"⚠️  Не удалось определить тип аудио: {filename} (channels={channels})")
                    continue
                
                name_score = 0
                if channel_hint and channel_final and channel_hint == channel_final:
                    name_score += 2
                if not has_bad_marker:
                    name_score += 1

                if name_mismatch or has_bad_marker:
                    audio_info['name_channel_mismatch'] = name_mismatch
                    audio_info['name_incorrect_type'] = has_bad_marker
                    if name_mismatch:
                        logger.warning(f"⚠️  Несовпадение типа в имени и метаданных: name={channel_hint}, meta={channel_meta}")
                    if has_bad_marker:
                        logger.warning(f"⚠️  Ошибка в названии формата (5.0/2.1): {filename}")

                audio_info['_name_score'] = name_score
                if key in tech_info and isinstance(tech_info[key], dict):
                    existing_score = tech_info[key].get('_name_score', 0)
                    if name_score <= existing_score:
                        logger.warning(f"⚠️  Найден дубль для {key}, оставляем файл с более корректным именем")
                        continue
                
                audio_info['file_path'] = audio_file
                tech_info[key] = audio_info
                audio_key_by_stem[normalize_stem(audio_info.get('file_name', audio_file))] = key
                audio_path_by_key[key] = audio_file
                logger.info(f"✅ {key}: {audio_info.get('file_name')} (channels={channels})")

            if self.report_type == "dcp" and audio_path_by_key:
                self.status_update.emit("🎛️ Расчет SAMPLE PEAK...")
                logger.info("=== DCP SAMPLE PEAK ANALYSIS START ===")
                total_dcp_files = len(audio_path_by_key)
                for idx_dcp, (audio_key, audio_path) in enumerate(audio_path_by_key.items(), start=1):
                    try:
                        self.status_update.emit(
                            f"🎛️ SAMPLE PEAK {idx_dcp}/{total_dcp_files}: {Path(audio_path).name}"
                        )
                        sample_peak, sample_peak_error = measure_sample_peak_fast(audio_path, timeout_sec=180)
                        if sample_peak_error:
                            logger.warning(
                                f"⚠️ DCP sample peak error for {audio_key} ({audio_path}): {sample_peak_error}"
                            )
                        dcp_audio_metrics_by_key[audio_key] = {
                            'sample_peak': sample_peak,
                        }
                        logger.info(
                            f"DCP metrics for {audio_key}: sample_peak={sample_peak}"
                        )
                        # Плавный прогресс внутри этапа DCP, чтобы не выглядело как зависание.
                        dcp_progress = 20 + int((idx_dcp / max(total_dcp_files, 1)) * 18)
                        self.progress_update.emit(min(dcp_progress, 39))
                    except Exception as dcp_error:
                        logger.warning(f"⚠️ DCP metrics failed for {audio_key} ({audio_path}): {dcp_error}")
                logger.info("=== DCP SAMPLE PEAK ANALYSIS END ===")

            self.progress_update.emit(40)
            
            # Обработка видео файла
            if video_files:
                video_file = video_files[0]
                logger.info(f"📹 Обработка видео: {video_file}")
                video_info = tech_extractor.extract_video_info(video_file)
                if video_info:
                    tech_info['video'] = video_info
                    logger.info(f"✅ Видео: {video_info.get('file_name')}, fps={video_info.get('fps')}, format={video_info.get('format')}")
                else:
                    logger.warning("⚠️  Не удалось извлечь информацию о видео")
            
            self.progress_update.emit(50)
            
            # === ШАГ: Извлекаем данные из PDF и сохраняем в JSON ===
            # Данные из PDF сначала накапливаются в словаре, затем сохраняются в JSON
            self.status_update.emit("📊 Извлечение данных из PDF...")
            
            logger.info("=== PDF FILES PROCESSING START ===")
            logger.info(f"pdf_files count: {len(pdf_files)}")
            logger.info(f"pdf_files: {pdf_files}")
            
            # Словарь для накопления PDF данных
            pdf_data_all = {}
            pending_pdfs = []
            
            for pdf_file in pdf_files:
                logger.info(f"\n--- Обработка PDF: {pdf_file} ---")
                filename = Path(pdf_file).stem.lower()
                logger.info(f"  filename (lower): {filename}")

                # Гибкая проверка разных форматов: _51_, _51_, - 51 -, _cens_51, stereo, и т.д.
                channel_hint = detect_channel_from_name(filename)
                has_51_marker = channel_hint == "51"
                has_20_marker = channel_hint == "20"

                cens_state = detect_cens_state(filename)
                is_cens = cens_state == "c"
                is_uncens = cens_state == "uc"

                # 1) Если имя PDF совпадает с аудио-файлом — берем тип из аудио метаданных
                matched_audio_key = audio_key_by_stem.get(normalize_stem(pdf_file))
                if matched_audio_key:
                    key = matched_audio_key.replace('audio_', 'pdf_')
                    logger.info(f"  -> Matched PDF to audio by name: {matched_audio_key} -> {key}")
                else:
                    # 2) Определяем тип PDF по имени
                    if has_51_marker and is_cens:
                        key = 'pdf_51_c'
                        logger.info(f"  -> Detected PDF 5.1 CENS, key = '{key}'")
                    elif has_51_marker and is_uncens:
                        key = 'pdf_51_uc'
                        logger.info(f"  -> Detected PDF 5.1 UNCENS, key = '{key}'")
                    elif has_51_marker:
                        key = 'pdf_51'
                        logger.info(f"  -> Detected PDF 5.1 (generic), key = '{key}'")
                    elif has_20_marker and is_cens:
                        key = 'pdf_20_c'
                        logger.info(f"  -> Detected PDF 2.0 CENS, key = '{key}'")
                    elif has_20_marker and is_uncens:
                        key = 'pdf_20_uc'
                        logger.info(f"  -> Detected PDF 2.0 UNCENS, key = '{key}'")
                    elif has_20_marker:
                        key = 'pdf_20'
                        logger.info(f"  -> Detected PDF 2.0 (generic), key = '{key}'")
                    else:
                        key = None

                # Извлекаем данные из PDF
                pdf_data = self.app.pdf_extractor.extract_technical_info(pdf_file)
                if pdf_data is None:
                    pdf_data = {}
                pdf_data['source_pdf'] = Path(pdf_file).name

                # 3) Если ключ не определен — пробуем по контенту PDF (каналы)
                if key is None:
                    logger.info(f"  -> No explicit marker, extracting from PDF content...")
                    if pdf_data:
                        logger.info(f"  -> PDF extracted: lufs={pdf_data.get('lufs')}, channels={pdf_data.get('channels')}")
                    else:
                        logger.warning(f"  -> PDF extraction returned EMPTY for {filename}")
                    
                    if pdf_data and pdf_data.get('channels'):
                        channels_str = str(pdf_data.get('channels', '')).lower()
                        if '5.1' in channels_str or '5.0' in channels_str or '6' in channels_str or 'surround' in channels_str:
                            key = 'pdf_51_c' if is_cens else ('pdf_51_uc' if is_uncens else 'pdf_51')
                            logger.info(f"  -> Detected from channels: 5.1, key = '{key}'")
                        elif '2.0' in channels_str or '2.1' in channels_str or '2' in channels_str or 'stereo' in channels_str:
                            key = 'pdf_20_c' if is_cens else ('pdf_20_uc' if is_uncens else 'pdf_20')
                            logger.info(f"  -> Detected from channels: 2.0, key = '{key}'")
                        else:
                            logger.warning(f"  -> Unknown channel format: {channels_str}")
                    else:
                        logger.warning(f"  -> No channels info in PDF")

                # 4) Если ключ всё ещё не определен — сохраняем в pending
                if key is None:
                    pending_pdfs.append({
                        'path': pdf_file,
                        'data': pdf_data,
                        'cens_state': cens_state
                    })
                    logger.warning(f"  -> PDF key unresolved, added to pending: {Path(pdf_file).name}")
                    continue

                # Сохраняем в словарь для JSON (а не напрямую в tech_info)
                if key in pdf_data_all:
                    # если ключ занят, кладем в общий без цензуры
                    key = 'pdf_20' if '20' in key else 'pdf_51'
                pdf_data_all[key] = pdf_data
                
                if pdf_data:
                    logger.info(f"  -> ADDED to pdf_data_all['{key}']: LUFS={pdf_data.get('lufs')}, Peak={pdf_data.get('true_peak')}, LRA={pdf_data.get('lra')}")
                else:
                    logger.warning(f"  -> PDF extraction returned EMPTY for {filename}")

            # Попытка доопределить pending PDF по совпадению с аудио без маркеров
            if pending_pdfs:
                missing_keys = []
                for prefix in ("20", "51"):
                    for suffix in ("_c", "_uc", ""):
                        k = f"pdf_{prefix}{suffix}"
                        audio_k = f"audio_{prefix}{suffix}"
                        if audio_k in tech_info and (k not in pdf_data_all or pdf_data_all.get(k) is None):
                            missing_keys.append((k, audio_k))

                for item in pending_pdfs:
                    pdf_path = item['path']
                    pdf_data = item['data']
                    best_key = None
                    pdf_norm = normalize_stem(pdf_path)
                    for k, audio_k in missing_keys:
                        audio_norm = normalize_stem(tech_info[audio_k].get('file_name', ''))
                        if audio_norm and (audio_norm in pdf_norm or pdf_norm in audio_norm):
                            best_key = k
                            break
                    if best_key is None and len(missing_keys) == 1:
                        best_key = missing_keys[0][0]
                    
                    if best_key:
                        pdf_data_all[best_key] = pdf_data
                        logger.info(f"  -> Pending PDF matched to {best_key}: {Path(pdf_path).name}")
                    else:
                        logger.warning(f"  -> Pending PDF could not be matched: {Path(pdf_path).name}")
            
            logger.info(f"=== PDF FILES PROCESSING END ===")
            logger.info(f"pdf_data_all keys: {list(pdf_data_all.keys())}")
            logger.info(f"pdf_data_all is empty: {len(pdf_data_all) == 0}")
            
            # === ШАГ: Сохраняем все PDF данные в JSON файл ===
            json_path = output_dir / "pdf_data.json"
            
            logger.info(f"=== JSON SAVE CHECK ===")
            logger.info(f"pdf_data_all is empty: {len(pdf_data_all) == 0}")
            
            if pdf_data_all:
                try:
                    self.app.pdf_extractor.save_tech_info_to_json(pdf_data_all, str(json_path))
                    logger.info(f"✅ PDF данные сохранены в JSON: {json_path}")
                    
                    # Логируем сохраненные данные
                    logger.info("=== PDF данные в JSON ===")
                    for k, v in pdf_data_all.items():
                        if isinstance(v, dict):
                            logger.info(f"  {k}: lufs={v.get('lufs')}, peak={v.get('true_peak')}, lra={v.get('lra')}")
                    logger.info("=== END PDF данные ===")
                    
                    # ДОБАВЛЯЕМ данные из pdf_data_all напрямую в tech_info (для гарантии)
                    # Это обеспечивает PDF → JSON → таблица поток
                    for key, value in pdf_data_all.items():
                        if isinstance(key, str) and key.startswith('pdf_'):
                            tech_info[key] = value
                            logger.info(f"  -> tech_info['{key}'] updated from pdf_data_all")
                    
                except Exception as e:
                    logger.error(f"❌ Ошибка сохранения JSON: {e}")
            else:
                logger.warning("⚠️ Нет PDF данных для сохранения")
            
            self.progress_update.emit(52)
            
            # === ШАГ: Загружаем PDF данные из JSON ===
            # Параметры берутся из JSON файла, в который были сохранены данные из PDF
            self.status_update.emit("💾 Загрузка PDF данных из JSON...")
            
            # Если JSON файл существует, загружаем данные из него
            # ДАННЫЕ УЖЕ ДОБАВЛЕНЫ В tech_info ПРИ СОХРАНЕНИИ (см. выше)
            # Эта загрузка служит как резерв/проверка
            if json_path.exists():
                logger.info(f"📄 JSON файл найден: {json_path}")
                try:
                    # Загружаем техническую информацию из JSON
                    json_data = self.app.pdf_extractor.load_tech_info_from_json(str(json_path))
                    
                    logger.info(f"📄 Загружено из JSON: {list(json_data.keys()) if json_data else 'ПУСТО'}")
                    
                    # Логируем ВСЕ ключи из json_data для диагностики
                    logger.info("=== json_data contents ===")
                    for k, v in json_data.items():
                        if isinstance(v, dict):
                            logger.info(f"  {k}: lufs={v.get('lufs')}, peak={v.get('true_peak')}, lra={v.get('lra')}")
                        else:
                            logger.info(f"  {k}: {v}")
                    logger.info("=== end json_data ===")
                    
                    if json_data:
                        # Обновляем tech_info данными из JSON (дублируем для надежности)
                        # Это гарантирует, что параметры в таблицу вставляются из JSON
                        updated_count = 0
                        for key, value in json_data.items():
                            if isinstance(key, str) and key.startswith('pdf_'):
                                tech_info[key] = value
                                updated_count += 1
                                logger.info(f"  -> tech_info['{key}'] = {value}")
                        logger.info(f"  📊 Обновлено {updated_count} ключей из JSON")
                    else:
                        logger.warning("⚠️ JSON файл пуст или поврежден")
                except Exception as e:
                    logger.error(f"❌ Ошибка чтения JSON: {e}")
            else:
                logger.warning(f"⚠️ JSON файл НЕ найден: {json_path}")
            
            # Диагностика: что в tech_info перед генерацией
            logger.info(f"📊 tech_info перед генерацией: {list(tech_info.keys())}")
            for k in ['pdf_20', 'pdf_51', 'pdf_20_c', 'pdf_51_c', 'pdf_20_uc', 'pdf_51_uc']:
                if k in tech_info:
                    v = tech_info[k]
                    if isinstance(v, dict):
                        logger.info(f"  -> {k}: lufs={v.get('lufs')}, peak={v.get('true_peak')}, lra={v.get('lra')}")
                    else:
                        logger.info(f"  -> {k}: {v}")
            
            # Логируем PDF данные для диагностики
            logger.info("=== PDF данные для таблицы ===")
            for k in tech_info:
                if isinstance(k, str) and k.startswith('pdf_') and isinstance(tech_info[k], dict):
                    v = tech_info[k]
                    logger.info(f"  {k}: lufs={v.get('lufs')}, true_peak={v.get('true_peak')}, lra={v.get('lra')}")
            logger.info("=== END ===")

            # === Принудительно сопоставляем 4 PDF по именам (cens -> uncens) ===
            def remap_pdf_key(target_key: str, pdf_path: str):
                if not pdf_path:
                    return
                target_name = Path(pdf_path).name
                for k, v in tech_info.items():
                    if isinstance(k, str) and k.startswith('pdf_') and isinstance(v, dict):
                        if v.get('source_pdf') == target_name:
                            tech_info[target_key] = v
                            logger.info(f"✅ PDF remap: {target_key} <- {target_name}")
                            return

            remap_pdf_key('pdf_20_c', pdf_paths.get('20_c'))
            remap_pdf_key('pdf_20_uc', pdf_paths.get('20_uc'))
            remap_pdf_key('pdf_51_c', pdf_paths.get('51_c'))
            remap_pdf_key('pdf_51_uc', pdf_paths.get('51_uc'))
            
            # === Уточняем выбор PDF 2.0 / 5.1 по метаданным аудио ===
            def pick_pdf_for_audio(prefix: str):
                for audio_key in (f"audio_{prefix}_c", f"audio_{prefix}_uc", f"audio_{prefix}"):
                    audio_data = tech_info.get(audio_key)
                    if isinstance(audio_data, dict):
                        audio_name = audio_data.get('file_name', '')
                        if audio_name:
                            norm = normalize_stem(audio_name)
                            if norm in pdf_by_stem and pdf_by_stem[norm]:
                                return str(pdf_by_stem[norm][0])
                return None

            picked_20 = pick_pdf_for_audio("20")
            picked_51 = pick_pdf_for_audio("51")
            if picked_20:
                copied_pdf_20 = picked_20
                logger.info(f"✅ PDF 2.0 уточнен по аудио-метаданным: {copied_pdf_20}")
            if picked_51:
                copied_pdf_51 = picked_51
                logger.info(f"✅ PDF 5.1 уточнен по аудио-метаданным: {copied_pdf_51}")

            # Защита от дубликатов (когда выбран один и тот же PDF для 2.0 и 5.1)
            if copied_pdf_20 and copied_pdf_51 and copied_pdf_20 == copied_pdf_51:
                if len(copied_pdfs) > 1:
                    for p in copied_pdfs:
                        if str(p) != copied_pdf_20:
                            copied_pdf_51 = str(p)
                            logger.info(f"⚠️ PDF 5.1 был дубликатом, заменен на {copied_pdf_51}")
                            break
                else:
                    # Только 1 PDF — убираем дубликат
                    copied_pdf_20 = None
                    logger.info(f"⚠️ Только 1 PDF, убираем дубликат copied_pdf_20")

            ordered_report_pdfs = [str(p) for p in order_pdfs_by_channel(copied_pdfs)]
            logger.info(f"PDF для вставки в отчет ({len(ordered_report_pdfs)}): {ordered_report_pdfs}")

            if self.report_type == "dcp":
                logger.info("=== DCP SAMPLE PEAK MERGE INTO pdf_* START ===")
                audio_to_pdf = {
                    'audio_20_c': 'pdf_20_c',
                    'audio_20_uc': 'pdf_20_uc',
                    'audio_51_c': 'pdf_51_c',
                    'audio_51_uc': 'pdf_51_uc',
                }
                merged_count = 0
                for audio_key, pdf_key in audio_to_pdf.items():
                    dcp_metrics = dcp_audio_metrics_by_key.get(audio_key)
                    if not dcp_metrics:
                        continue
                    pdf_data = tech_info.get(pdf_key)
                    if not isinstance(pdf_data, dict):
                        channel_prefix = "20" if "20" in audio_key else "51"
                        generic_pdf = tech_info.get(f'pdf_{channel_prefix}')
                        if isinstance(generic_pdf, dict):
                            pdf_data = dict(generic_pdf)
                        else:
                            pdf_data = {'source_pdf': None}

                    sample_peak = dcp_metrics.get('sample_peak')
                    if sample_peak is None:
                        sample_peak = pdf_data.get('true_peak')
                    if sample_peak is not None:
                        pdf_data['sample_peak'] = sample_peak
                        if dcp_metrics.get('sample_peak') is not None:
                            pdf_data['sample_peak_source'] = 'app_meter'
                        else:
                            pdf_data['sample_peak_source'] = 'pdf_true_peak_fallback'

                    tech_info[pdf_key] = pdf_data
                    merged_count += 1
                    logger.info(
                        f"DCP merge {audio_key} -> {pdf_key}: "
                        f"sample_peak={pdf_data.get('sample_peak')}"
                    )

                if merged_count == 0:
                    logger.warning("⚠️ DCP merge: нет данных для подстановки SAMPLE PEAK в pdf_*")
                logger.info("=== DCP SAMPLE PEAK MERGE INTO pdf_* END ===")
            
            logger.info("✅ Параметры из JSON готовы для вставки в таблицу")
            
            self.progress_update.emit(55)
            
            # === ШАГ 6б: Аудио анализ и экспорт в CSV/HTML ===
            if self.pyloudnorm_enabled:
                self.status_update.emit("📊 PyLoudNorm анализ...")
                
                logger.debug(f"audio_files = {audio_files}")
                
                try:
                    logger.debug(f"Entering try block")
                    if all_audio_files:
                        logger.debug(f"all_audio_files is not empty, length = {len(all_audio_files)}")
                        logger.info(f"🎵 Анализ {len(all_audio_files)} аудиофайлов...")

                        # Все артефакты анализа храним в отдельной подпапке отчета.
                        analysis_output_dir = output_dir / "pyloudnorm_analysis"
                        analysis_output_dir.mkdir(parents=True, exist_ok=True)
                        logger.info(f"📁 Папка PyLoudNorm анализа: {analysis_output_dir}")
                        
                        # Создаем экспортер
                        exporter = ParameterExporter({
                            'audio': {
                                'target_lufs': -23.0,
                                'lufs_tolerance': 0.5,
                                'true_peak': -2.0,
                                'lra_max': 18.0,
                                'sample_rate': 48000
                            },
                            'export': {
                                'csv': {
                                    'delimiter': ',',
                                    'encoding': 'utf-8'
                                },
                                'html': {
                                    'include_css': True,
                                    'include_timestamp': True
                                }
                            }
                        })
                        
                        logger.debug(f"output_dir = {output_dir}")
                        
                        # Экспортируем в CSV, HTML и TXT
                        export_results = exporter.analyze_and_export(
                            audio_files=all_audio_files,
                            output_dir=str(analysis_output_dir),
                            formats=['csv', 'html', 'txt'],
                            report_name="audio_analysis"
                        )
                        
                        logger.debug(f"export_results = {export_results}")
                        logger.info(f"✅ Аудио анализ экспортирован: {export_results}")
                        
                        # Сохраняем пути к файлам
                        tech_info['audio_analysis_csv'] = export_results.get('csv')
                        tech_info['audio_analysis_html'] = export_results.get('html')
                        tech_info['audio_analysis_txt'] = export_results.get('txt')
                        tech_info['audio_analysis_dir'] = str(analysis_output_dir)

                        # Генерируем PDF loudness-графики (Youlean-style) для каждого аудиофайла.
                        generated_pdf_reports = []
                        total_audio = len(all_audio_files)
                        for idx_audio, audio_path in enumerate(all_audio_files, 1):
                            audio_name = Path(audio_path).name
                            self.status_update.emit(f"📈 Loudness PDF ({idx_audio}/{total_audio}): {audio_name}")
                            try:
                                audio_stem = normalize_stem(audio_path) or Path(audio_path).stem
                                pdf_report_path = analysis_output_dir / f"{audio_stem}_loudness_report.pdf"
                                metrics, generation_error = generate_loudness_report_isolated(
                                    audio_file=str(audio_path),
                                    pdf_path=str(pdf_report_path),
                                )
                                if metrics is None:
                                    raise RuntimeError(
                                        generation_error or "unknown loudness report generation error"
                                    )
                                generated_pdf_reports.append({
                                    'audio_file': str(audio_path),
                                    'pdf_report': str(pdf_report_path),
                                    'metrics': metrics,
                                })
                                logger.info(f"✅ Loudness PDF создан: {pdf_report_path.name}")
                            except Exception as pdf_error:
                                import traceback
                                self.status_update.emit(f"⚠️ PDF ошибка: {audio_name}: {pdf_error}")
                                logger.warning(
                                    f"⚠️ Не удалось создать loudness PDF для {audio_name}: {pdf_error}"
                                )
                                logger.warning("Loudness PDF traceback:\n%s", traceback.format_exc())

                        if generated_pdf_reports:
                            tech_info['audio_analysis_pdf_reports'] = generated_pdf_reports
                            logger.info(f"✅ Loudness PDF создано: {len(generated_pdf_reports)}")
                        else:
                            logger.warning("⚠️ Loudness PDF не были сгенерированы")

                        # Если PDF отсутствуют, заполняем pdf_* из PyLoudNorm CSV
                        if (not pdf_files) or (not any(k.startswith('pdf_') for k in tech_info.keys())):
                            csv_path = tech_info.get('audio_analysis_csv')
                            if csv_path and Path(csv_path).exists():
                                logger.info("📌 PDF не найден — используем PyLoudNorm как источник параметров для таблицы")

                                def _parse_float(value):
                                    if value is None:
                                        return None
                                    if isinstance(value, (int, float)):
                                        return float(value)
                                    value = str(value).strip()
                                    if value == "":
                                        return None
                                    try:
                                        return float(value)
                                    except ValueError:
                                        return None

                                # Маппинг audio_* -> pdf_* для подстановки
                                audio_to_pdf = {
                                    'audio_20_c': 'pdf_20_c',
                                    'audio_20_uc': 'pdf_20_uc',
                                    'audio_51_c': 'pdf_51_c',
                                    'audio_51_uc': 'pdf_51_uc'
                                }

                                # Строим индекс по нормализованному имени файла
                                audio_key_by_norm = {}
                                for audio_key, pdf_key in audio_to_pdf.items():
                                    audio_data = tech_info.get(audio_key)
                                    if isinstance(audio_data, dict):
                                        name = audio_data.get('file_name', '')
                                        if name:
                                            audio_key_by_norm[normalize_stem(name)] = pdf_key

                                injected = 0
                                with open(csv_path, 'r', encoding='utf-8') as f:
                                    reader = csv.DictReader(f)
                                    for row in reader:
                                        file_name = row.get('file_name', '')
                                        norm = normalize_stem(file_name) if file_name else ''
                                        pdf_key = audio_key_by_norm.get(norm)
                                        if not pdf_key:
                                            continue
                                        if pdf_key in tech_info and tech_info.get(pdf_key):
                                            continue

                                        pdf_like = {
                                            'lufs': _parse_float(row.get('integrated_lufs')),
                                            'true_peak': _parse_float(row.get('true_peak_dbtp')),
                                            'lra': _parse_float(row.get('lra')),
                                            'source_pdf': None
                                        }
                                        tech_info[pdf_key] = pdf_like
                                        injected += 1

                                logger.info(f"✅ PyLoudNorm -> pdf_* заполнено: {injected} строк")
                            else:
                                logger.warning("⚠️ PyLoudNorm CSV не найден, заполнение таблицы невозможно")
                    else:
                        logger.debug(f"all_audio_files is EMPTY")
                        logger.warning("⚠️ Аудиофайлы не найдены для анализа")
                        # Попробуем добавить видео файлы для анализа (они содержат аудио дорожки)
                        if video_files:
                            logger.debug(f"Trying video files for audio analysis: {video_files}")
                            logger.info(f"🎵 Видео файлы будут использованы для аудио анализа: {len(video_files)}")
                            # Добавляем видео файлы в анализ
                            all_audio_files = video_files
                        
                except Exception as e:
                    import traceback
                    self.status_update.emit(f"⚠️ Ошибка аудио анализа: {e}")
                    logger.error(f"❌ Ошибка аудио анализа: {e}")
                    logger.error(f"❌ Traceback: {traceback.format_exc()}")
                
                logger.debug(f"Finished audio analysis step")
            else:
                logger.debug("PyLoudNorm analysis SKIPPED (checkbox not checked)")
                logger.info("PyLoudNorm анализ пропущен (галочка не установлена)")
            
            # Импорт проблем из CSV
            self.status_update.emit("📊 Импорт проблем из CSV...")
            
            issues = []
            if csv_files:
                csv_file = csv_files[0]
                logger.info(f"Импорт CSV: {csv_file}")
                importer = CSVImporter()
                issues = importer.import_issues(csv_file)
                logger.info(f"✅ Импортировано проблем: {len(issues)}")
            
            self.progress_update.emit(70)

            # === True Peak verification for borderline values ===
            if self.tp_verify_enabled:
                logger.info("=== TP VERIFY: НАЧАЛО ===")
                self.status_update.emit("🔍 TP verify: анализ True Peak значений...")

                # Диагностика: показываем все pdf_* и audio_* ключи
                pdf_keys_present = [k for k in tech_info if isinstance(k, str) and k.startswith('pdf_') and isinstance(tech_info[k], dict)]
                audio_keys_present = list(audio_path_by_key.keys())
                logger.info(f"  PDF ключи в tech_info: {pdf_keys_present}")
                logger.info(f"  Audio ключи в audio_path_by_key: {audio_keys_present}")
                for pk in pdf_keys_present:
                    pd = tech_info[pk]
                    logger.info(f"    {pk}: true_peak={pd.get('true_peak')!r}, lufs={pd.get('lufs')!r}")

                borderline_items = self._find_tp_verify_items(tech_info, audio_path_by_key)
                if borderline_items:
                    logger.info(f"🔍 Пограничных True Peak: {len(borderline_items)}, запуск точного измерения...")
                    measurement_results = []
                    measurement_errors = []
                    for idx, item in enumerate(borderline_items, 1):
                        audio_path = item.get('audio_file')
                        if not audio_path or not Path(audio_path).exists():
                            msg = f"⚠️ Аудиофайл не найден для {item['key']}: {audio_path}"
                            logger.warning(msg)
                            self.status_update.emit(msg)
                            continue

                        file_size_mb = Path(audio_path).stat().st_size / (1024 * 1024)
                        self.status_update.emit(
                            f"🎯 True Peak {idx}/{len(borderline_items)}: "
                            f"{Path(audio_path).name} ({file_size_mb:.0f} MB)..."
                        )
                        logger.info(f"  Измерение {idx}/{len(borderline_items)}: {Path(audio_path).name} ({file_size_mb:.0f} MB)")

                        try:
                            import time as _time
                            t0 = _time.time()
                            result = measure_true_peak_precise(str(audio_path))
                            elapsed = _time.time() - t0
                            precise_val = result.get('true_peak_dbtp')
                            per_ch = result.get('true_peak_per_channel', [])

                            logger.info(
                                f"  ✅ {item['label']}: Youlean={item['current_value']:.1f} → "
                                f"Precise={precise_val:.2f} dBTP ({elapsed:.1f}s), "
                                f"per_channel={[f'{v:.2f}' for v in per_ch]}"
                            )
                            self.status_update.emit(
                                f"✅ {item['label']}: {item['current_value']:.1f} → {precise_val:.2f} dBTP ({elapsed:.0f}s)"
                            )

                            measurement_results.append({
                                'key': item['key'],
                                'label': item['label'],
                                'youlean_value': item['current_value'],
                                'precise_value': precise_val,
                                'per_channel': per_ch,
                                'audio_file': str(audio_path),
                            })
                        except Exception as e:
                            import traceback
                            tb = traceback.format_exc()
                            logger.error(f"❌ Ошибка измерения True Peak для {audio_path}: {e}\n{tb}")
                            self.status_update.emit(f"❌ Ошибка TP: {Path(audio_path).name}: {e}")
                            measurement_errors.append(f"{item['label']}: {e}")

                    if measurement_errors:
                        logger.warning(f"TP verify: {len(measurement_errors)} ошибок измерений")

                    if measurement_results:
                        logger.info(f"✅ TP verify: {len(measurement_results)} результатов, показываем диалог")
                        # Показываем результаты пользователю
                        self.status_update.emit("🎯 Ожидание подтверждения True Peak...")
                        self._tp_verify_loop = QEventLoop()
                        self.true_peak_results_ready.emit(measurement_results)
                        self._tp_verify_loop.exec_()

                        if self._tp_verify_results:
                            updated_keys = []
                            for key, new_value in self._tp_verify_results.items():
                                pdf_data = tech_info.get(key)
                                if isinstance(pdf_data, dict):
                                    old_val = pdf_data.get('true_peak')
                                    pdf_data['true_peak'] = new_value
                                    pdf_data['true_peak_source'] = 'precise_4x_oversampling'
                                    updated_keys.append(key)
                                    logger.info(
                                        f"✅ True Peak обновлен для {key}: "
                                        f"{old_val} → {new_value:.2f} dBTP (4x oversampling)"
                                    )
                                else:
                                    logger.warning(f"⚠️ tech_info[{key}] не найден или не dict!")
                            self.status_update.emit(f"✅ True Peak обновлено: {', '.join(updated_keys)}")
                        else:
                            logger.info("True Peak: пользователь оставил значения Youlean")
                            self.status_update.emit("ℹ️ True Peak: оставлены значения Youlean")
                    else:
                        logger.warning("⚠️ TP verify: нет успешных измерений")
                        self.status_update.emit("⚠️ TP verify: измерения не удались")
                else:
                    logger.info("✅ Нет пограничных True Peak значений")

                logger.info("=== TP VERIFY: КОНЕЦ ===")

            # Генерация заключений
            self.status_update.emit("📝 Генерация заключений...")
            
            if self.report_type == "me_ours":
                technical_conclusion = "По технической оценке нареканий не выявлено."
                subjective_conclusion = "Ниже предоставлен список внесенных изменений:"
                logger.info("M&E наши работы: идеальные заключения")
            else:
                params = tech_info.get('params', {})
                technical_conclusion = self.app.conclusion_gen.generate_technical_conclusion(
                    tech_info, params, self.report_type
                )
                subjective_conclusion = self.app.conclusion_gen.generate_subjective_conclusion(issues, self.report_type)
                logger.info("Стандартная генерация заключений")
            
            self.progress_update.emit(85)
            
            # Генерация отчета
            self.status_update.emit("📄 Создание отчета...")
            
            report_path = output_dir / f"отчет_{base_name}.docx"
            
            logger.info(f"Генерация отчета: {report_path}")
            logger.info(f"PDF 2.0: {copied_pdf_20}")
            logger.info(f"PDF 5.1: {copied_pdf_51}")
            logger.info(f"PDF app list: {ordered_report_pdfs}")
            
            # Логируем tech_info для диагностики
            logger.info("=== TECH_INFO CONTENTS ===")
            if tech_info:
                for key, value in tech_info.items():
                    if isinstance(value, dict):
                        logger.info(f"  {key}: dict with keys = {list(value.keys())}")
                        if 'lufs' in value:
                            logger.info(f"    -> LUFS: {value['lufs']}, Peak: {value.get('true_peak')}, LRA: {value.get('lra')}")
                    else:
                        logger.info(f"  {key}: {type(value).__name__} = {value}")
            else:
                logger.info("  tech_info is EMPTY!")
            logger.info("=== END TECH_INFO ===")
            
            # Генерируем отчет через единый генератор
            prepared_by = getattr(self.app, 'prepared_by', '')
            self.app.report_gen.create_exact_report(
                issues=issues,
                output_path=str(report_path),
                tech_info=tech_info,
                pdf_20_path=copied_pdf_20,
                pdf_51_path=copied_pdf_51,
                pdf_paths=ordered_report_pdfs,
                conclusion_technical=technical_conclusion,
                conclusion_subjective=subjective_conclusion,
                report_type=self.report_type,
                prepared_by=prepared_by,
            )
            
            self.progress_update.emit(100)
            self.status_update.emit("✅ Готово!")
            
            logger.info(f"✅ Отчет успешно создан: {report_path}")

            # Удаляем временный JSON файл с PDF данными
            if json_path.exists():
                try:
                    json_path.unlink()
                    logger.info(f"🗑️ Временный JSON удален: {json_path}")
                except Exception as e:
                    logger.warning(f"⚠️ Не удалось удалить JSON: {e}")

            # Показываем информацию о созданных файлах
            files_in_output = list(output_dir.glob('*'))
            logger.info(f"Файлы в папке отчета ({len(files_in_output)}):")
            for f in files_in_output:
                logger.info(f"  - {f.name}")
            
            success_msg = f"Отчет создан!\nПапка: {output_dir.name}\nФайлов: {len(files_in_output)}"

            # Удаляем исходные файлы после того, как всё обработано
            if self.delete_sources:
                for src in list(pdf_files) + list(csv_files) + list(params_files):
                    try:
                        p = Path(src)
                        if p.exists():
                            p.unlink()
                            logger.info(f"🗑 Удалён исходник: {p.name}")
                    except Exception as e:
                        logger.warning(f"Не удалось удалить {src}: {e}")

            self.finished.emit(True, success_msg)

            logger.info(f"=== ЗАВЕРШЕНО УСПЕШНО ===")
            
        except Exception as e:
            logger.error(f"Ошибка при обработке: {e}", exc_info=True)
            self.finished.emit(False, f"Ошибка: {str(e)}")

    # --- True Peak borderline helpers ---

    _PDF_KEY_LABELS = {
        'pdf_20_c': '2.0 Cens',
        'pdf_20_uc': '2.0 Uncens',
        'pdf_51_c': '5.1 Cens',
        'pdf_51_uc': '5.1 Uncens',
        'pdf_20': '2.0',
        'pdf_51': '5.1',
    }

    # Маппинг pdf_* -> audio_* для поиска аудиофайлов
    _PDF_TO_AUDIO_KEYS = {
        'pdf_20_c': ['audio_20_c', 'audio_20_uc'],
        'pdf_20_uc': ['audio_20_uc', 'audio_20_c'],
        'pdf_51_c': ['audio_51_c', 'audio_51_uc'],
        'pdf_51_uc': ['audio_51_uc', 'audio_51_c'],
        'pdf_20': ['audio_20_c', 'audio_20_uc'],
        'pdf_51': ['audio_51_c', 'audio_51_uc'],
    }

    @staticmethod
    def _resolve_tp_audio_path(audio_key: str, tech_info: dict, audio_path_by_key: dict) -> str:
        """Возвращает путь к аудиофайлу для ключа audio_*."""
        audio_path = audio_path_by_key.get(audio_key)
        if audio_path and Path(str(audio_path)).exists():
            return str(audio_path)

        audio_data = tech_info.get(audio_key)
        if isinstance(audio_data, dict):
            file_path = audio_data.get('file_path')
            if file_path and Path(str(file_path)).exists():
                return str(file_path)

        return None

    def _match_tp_audio_by_source_pdf(self, pdf_data: dict, tech_info: dict, audio_path_by_key: dict) -> tuple:
        """Пытается найти точный audio_* по имени исходного PDF/аудио файла."""
        source_pdf = str(pdf_data.get('source_pdf') or '').strip()
        if not source_pdf:
            return None, None

        source_norm = normalize_stem(source_pdf)
        if not source_norm:
            return None, None

        for audio_key in ('audio_20_c', 'audio_20_uc', 'audio_51_c', 'audio_51_uc'):
            audio_data = tech_info.get(audio_key)
            if not isinstance(audio_data, dict):
                continue

            candidates = [
                audio_data.get('file_name'),
                audio_data.get('file_path'),
                audio_path_by_key.get(audio_key),
            ]
            for candidate in candidates:
                if not candidate:
                    continue
                if normalize_stem(candidate) == source_norm:
                    resolved_path = self._resolve_tp_audio_path(audio_key, tech_info, audio_path_by_key)
                    if resolved_path:
                        return audio_key, resolved_path

        return None, None

    def _find_tp_verify_items(self, tech_info: dict, audio_path_by_key: dict) -> list:
        """Находит все pdf_* записи с аудиофайлами для точного измерения True Peak."""
        logger.info("=== TP VERIFY: поиск аудиофайлов для измерения ===")
        logger.info(f"  tech_info keys: {list(tech_info.keys())}")
        logger.info(f"  audio_path_by_key: {list(audio_path_by_key.keys())}")

        self.status_update.emit("🔍 TP verify: поиск аудиофайлов...")

        items = []
        seen_audio = set()  # чтобы не измерять один файл дважды

        for key in ('pdf_20_c', 'pdf_20_uc', 'pdf_51_c', 'pdf_51_uc', 'pdf_20', 'pdf_51'):
            pdf_data = tech_info.get(key)
            if not isinstance(pdf_data, dict):
                continue

            tp = pdf_data.get('true_peak')
            # Конвертация в float (защита от строк из JSON)
            if tp is not None:
                try:
                    tp = float(tp)
                except (ValueError, TypeError):
                    tp = None

            # Находим соответствующий аудиофайл
            audio_file = None
            tried_keys = []
            matched_audio_key = None

            # 0) Приоритет: матчим по source_pdf -> audio file_name/file_path
            matched_audio_key, audio_file = self._match_tp_audio_by_source_pdf(
                pdf_data, tech_info, audio_path_by_key
            )
            if matched_audio_key and audio_file:
                tried_keys.append(f"source_pdf:{matched_audio_key}")
                logger.info(
                    f"  {key}: точное сопоставление по source_pdf -> "
                    f"{matched_audio_key} ({Path(audio_file).name})"
                )

            # 1) Прямой маппинг pdf_* → audio_*
            if not audio_file:
                for audio_key in self._PDF_TO_AUDIO_KEYS.get(key, []):
                    tried_keys.append(audio_key)
                    resolved_path = self._resolve_tp_audio_path(audio_key, tech_info, audio_path_by_key)
                    if resolved_path:
                        matched_audio_key = audio_key
                        audio_file = resolved_path
                        break

            # 2) Fallback: file_path в tech_info[audio_*]
            if not audio_file:
                channel_type = '51' if '51' in key else '20'
                for audio_key, audio_path in audio_path_by_key.items():
                    if channel_type in audio_key:
                        matched_audio_key = audio_key
                        audio_file = str(audio_path)
                        break

            # 3) Последний fallback: file_path по типу канала
            if not audio_file:
                channel_type = '51' if '51' in key else '20'
                for ti_key, ti_val in tech_info.items():
                    if channel_type in ti_key and isinstance(ti_val, dict):
                        fp = ti_val.get('file_path')
                        if fp and Path(str(fp)).exists():
                            matched_audio_key = ti_key
                            audio_file = str(fp)
                            break

            if not audio_file or not Path(audio_file).exists():
                logger.warning(f"  ⚠️ {key}: аудиофайл не найден (tried: {tried_keys})")
                continue

            if audio_file in seen_audio:
                logger.info(f"  {key}: аудиофайл уже в списке, пропускаем")
                continue

            seen_audio.add(audio_file)
            logger.info(
                f"  ✅ {key}: {Path(audio_file).name} "
                f"(audio_key={matched_audio_key}, Youlean TP={tp})"
            )

            items.append({
                'key': key,
                'label': self._PDF_KEY_LABELS.get(key, key),
                'current_value': tp,
                'audio_file': audio_file,
            })

        logger.info(f"=== TP VERIFY: найдено {len(items)} файлов для измерения ===")
        if items:
            labels = ', '.join(item['label'] for item in items)
            self.status_update.emit(f"🎯 TP verify: {len(items)} файлов: {labels}")
        else:
            self.status_update.emit("⚠️ TP verify: аудиофайлы не найдены")
        return items

    def receive_tp_verification(self, results: dict):
        """Вызывается из UI-потока после закрытия диалога."""
        self._tp_verify_results = results
        if self._tp_verify_loop and self._tp_verify_loop.isRunning():
            self._tp_verify_loop.quit()



class MacTitleBar(QWidget):
    """Compact in-app header styled to match the macOS-inspired UI."""

    def __init__(self, parent=None, title=""):
        super().__init__(parent)
        self._parent = parent
        self._drag_pos = None
        self.setFixedHeight(18)
        self.setStyleSheet("""
            background-color: #F7F7F8;
            border: none;
            border-bottom: 1px solid #E5E5EA;
        """)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(16, 0, 16, 0)
        layout.setSpacing(0)

        layout.addStretch(1)

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._drag_pos = event.globalPos() - self._parent.frameGeometry().topLeft()
            event.accept()

    def mouseMoveEvent(self, event):
        if self._drag_pos is not None and event.buttons() == Qt.LeftButton:
            self._parent.move(event.globalPos() - self._drag_pos)
            event.accept()

    def mouseReleaseEvent(self, event):
        self._drag_pos = None

    def mouseDoubleClickEvent(self, event):
        if self._parent.isMaximized():
            self._parent.showNormal()
        else:
            self._parent.showMaximized()


class MacToggleSwitch(QPushButton):
    """Small toggle indicator rendered as a compact rounded-square icon."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setCheckable(True)
        self.setCursor(Qt.PointingHandCursor)
        self.setFixedSize(18, 18)
        self.setStyleSheet("background: transparent; border: none;")

    def sizeHint(self):
        return QSize(18, 18)

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)

        rect = self.rect().adjusted(1, 1, -1, -1)
        is_checked = self.isChecked()
        is_enabled = self.isEnabled()

        fill_color = QColor("#007AFF" if is_checked else "#FFFFFF")
        border_color = QColor("#007AFF" if is_checked else "#D2D2D7")
        if not is_enabled:
            fill_color.setAlpha(110)
            border_color.setAlpha(110)
        elif self.underMouse():
            border_color = border_color.lighter(108)

        pen = QPen(border_color)
        pen.setWidthF(1.2)
        painter.setPen(pen)
        painter.setBrush(fill_color)
        painter.drawRoundedRect(rect, 4.5, 4.5)

        if is_checked:
            pen = QPen(QColor("#FFFFFF"))
            pen.setWidthF(1.7)
            pen.setCapStyle(Qt.RoundCap)
            pen.setJoinStyle(Qt.RoundJoin)
            painter.setPen(pen)
            path = QPainterPath()
            path.moveTo(rect.x() + rect.width() * 0.24, rect.y() + rect.height() * 0.54)
            path.lineTo(rect.x() + rect.width() * 0.43, rect.y() + rect.height() * 0.72)
            path.lineTo(rect.x() + rect.width() * 0.76, rect.y() + rect.height() * 0.34)
            painter.drawPath(path)


class PreviewDialog(QDialog):
    """Floating preview window that keeps the main form compact."""

    def __init__(self, parent=None):
        super().__init__(parent, Qt.Tool | Qt.WindowTitleHint | Qt.WindowCloseButtonHint)
        self._parent_window = parent
        self._positioned_once = False
        self.setModal(False)
        self.setWindowTitle("Предпросмотр")
        self.resize(760, 520)
        self.setMinimumSize(620, 380)
        self.setStyleSheet("""
            QDialog {
                background: #F7F7F8;
            }
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(8)

        title = QLabel("Предпросмотр перед генерацией")
        title.setFont(QFont(".AppleSystemUIFont", 12, QFont.DemiBold))
        title.setStyleSheet("color: #1D1D1F; background: transparent;")
        header_row.addWidget(title)

        header_row.addStretch()

        self.status_label = QLabel("Нет данных")
        self.status_label.setFont(QFont(".AppleSystemUIFont", 10))
        self.status_label.setStyleSheet("color: #86868B; background: transparent;")
        header_row.addWidget(self.status_label)

        layout.addLayout(header_row)

        self.preview_text = QTextBrowser()
        self.preview_text.setReadOnly(True)
        self.preview_text.setOpenExternalLinks(False)
        self.preview_text.setOpenLinks(False)
        self.preview_text.setStyleSheet("""
            QTextBrowser {
                background: #FFFFFF;
                border: 1px solid #E5E5EA;
                border-radius: 12px;
                padding: 8px;
                color: #1D1D1F;
                font-family: ".AppleSystemUIFont";
                font-size: 11px;
                selection-background-color: #DCEBFF;
            }
            QScrollBar:vertical {
                background: transparent; width: 7px; margin: 4px 0;
            }
            QScrollBar::handle:vertical {
                background: rgba(0,0,0,0.12); border-radius: 3px; min-height: 28px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0;
            }
        """)
        layout.addWidget(self.preview_text, 1)

    def showEvent(self, event):
        super().showEvent(event)
        if self._positioned_once or not self._parent_window:
            return
        parent_geom = self._parent_window.frameGeometry()
        x = parent_geom.center().x() - self.width() // 2
        y = parent_geom.center().y() - self.height() // 2
        self.move(max(x, 20), max(y, 20))
        self._positioned_once = True

    def closeEvent(self, event):
        super().closeEvent(event)
        if self._parent_window and hasattr(self._parent_window, "_on_preview_dialog_closed"):
            self._parent_window._on_preview_dialog_closed()


class BeastApp(QMainWindow):
    """macOS-style main window"""
    _ollama_status_signal = pyqtSignal(str, str)  # (status, detail) — thread-safe обновление индикатора
    _update_check_signal = pyqtSignal(object)  # UpdateInfo | None — результат проверки обновлений
    _update_download_signal = pyqtSignal(bool, str)  # (success, path_or_error)

    def __init__(self):
        super().__init__()
        self._ollama_status_signal.connect(self._set_ollama_dot)
        self._update_check_signal.connect(self._on_update_check_result)
        self._update_download_signal.connect(self._on_update_download_finished)
        self._pending_update_info = None
        self._update_progress_dialog = None
        self.config = load_app_config()

        self.report_gen = ExactReportGenerator()
        self.pdf_extractor = PDFExtractor()
        self.conclusion_gen = ConclusionGenerator(use_llm=False, config=self.config)
        self.csv_importer = CSVImporter()
        self.tech_extractor = TechnicalInfoExtractor()

        self.files_data = {
            'audio': [],
            'video': [],
            'csv': [],
            'pdf': [],
            'params': []
        }

        # Путь к автоматически созданному пустому CSV (если был сгенерирован)
        self._auto_created_csv_path = None
        self.last_output_folder = None
        self.last_report_docx_path = None
        self.preview_data = None
        self.preview_thread = None
        self._preview_busy = False
        self._preview_refresh_pending = False
        self._preview_epoch = 0
        self._preview_expanded = False
        self._preview_collapsed_height = 230
        self._preview_expanded_height = 430
        self.preview_dialog = None
        self.preview_card = None
        self.preview_text = None
        self.preview_status_label = None
        self.preview_expand_btn = None
        self._processing_active = False
        self.preview_timer = QTimer(self)
        self.preview_timer.setSingleShot(True)
        self.preview_timer.timeout.connect(self._run_preview_refresh)

        self.auto_reset_timer = QTimer(self)
        self.auto_reset_timer.setSingleShot(True)
        self.auto_reset_timer.timeout.connect(self._auto_reset_after_done)

        self._yandex_queue = YandexUploadQueueManager(
            get_token=SettingsDialog.get_yandex_token, get_roots=SettingsDialog.get_yandex_roots,
        )
        self._yandex_queue.queue_changed.connect(self._update_yandex_queue_badge)
        self._yandex_queue_offline = False
        self._yandex_queue.queue_paused_offline.connect(self._on_yandex_queue_offline_changed)

        # Куда именно уже отправлен отчёт (local_output_folder -> remote_folder_path) —
        # нужно, чтобы «Правка» знала, куда заливать изменения после сохранения.
        self._yandex_remote_by_local = {}
        self._yandex_queue.job_uploaded.connect(self._on_yandex_queue_job_uploaded)
        self._edit_sync = YandexEditSyncController(
            get_token=SettingsDialog.get_yandex_token,
            resolve_remote_path=self._resolve_edited_report_remote_path,
            parent=self,
        )
        self._edit_sync.status_changed.connect(self._on_edit_sync_status_changed)
        self._edit_sync.conflict.connect(self._on_edit_sync_conflict)

        # Пробел (как Quick Look в Finder) открывает превью последнего
        # сгенерированного отчёта — но только пока это окно активно и
        # фокус не в текстовом поле/на кнопке (чтобы не мешать обычному вводу).
        QApplication.instance().installEventFilter(self)

        self._integrity_check_thread = None
        self._integrity_checked = False
        QTimer.singleShot(5000, self._check_uploaded_reports_integrity)

        self.init_ui()

    def eventFilter(self, obj, event):
        if event.type() == QEvent.KeyPress and event.key() == Qt.Key_Space:
            if QApplication.activeWindow() is self:
                focus_widget = QApplication.focusWidget()
                editable_types = (QLineEdit, QPlainTextEdit, QTextEdit, QAbstractButton, QComboBox)
                if not isinstance(focus_widget, editable_types):
                    if self.last_report_docx_path and self.last_report_docx_path.exists():
                        _quick_look_preview(self.last_report_docx_path)
                        return True
        return super().eventFilter(obj, event)

    def closeEvent(self, event):
        # self.thread шадовит встроенный QObject.thread() только когда
        # уже был назначен как атрибут экземпляра — до первой генерации
        # отчёта getattr(self, "thread", ...) вернул бы этот метод, а не
        # None, поэтому здесь читаем строго из instance __dict__.
        processing_thread = self.__dict__.get("thread")
        if isinstance(processing_thread, QThread) and processing_thread.isRunning():
            reply = QMessageBox.question(
                self, "Идёт генерация отчёта",
                "Отчёт ещё формируется. Закрыть приложение и прервать генерацию?",
                QMessageBox.Yes | QMessageBox.No, QMessageBox.No,
            )
            if reply != QMessageBox.Yes:
                event.ignore()
                return
            _stop_thread(processing_thread)

        # Флаг ставится ДО остановки потоков: пока _stop_thread ждёт, поток
        # успевает эмитнуть финальный сигнал (queued), и его обработчик
        # исполнится уже после closeEvent — без флага он запустил бы новый
        # поток, который никто не остановит, и при teardown интерпретатора
        # процесс упал бы с "QThread: Destroyed while thread is still
        # running". То же с отложенными QTimer-колбэками (retry, проверка
        # целостности через 5с после старта).
        self._closing = True

        _stop_thread(getattr(self, "preview_thread", None))
        _stop_thread(getattr(self, "_yandex_find_thread", None))
        _stop_thread(getattr(self, "_yandex_versions_thread", None))
        _stop_thread(getattr(self, "_yandex_compare_thread", None))
        _stop_thread(getattr(self, "_yandex_upload_thread", None))
        _stop_thread(getattr(self, "_folder_picker_mkdir_thread", None))
        _stop_thread(getattr(self, "_yandex_fallback_find_thread", None))
        _stop_thread(getattr(self, "_integrity_check_thread", None))
        edit_sync = getattr(self, "_edit_sync", None)
        if edit_sync is not None:
            edit_sync.stop_all()
        queue_manager = getattr(self, "_yandex_queue", None)
        if queue_manager is not None:
            queue_manager.shutdown()

        super().closeEvent(event)

    def init_ui(self):
        self.setWindowTitle("Beast Auto Reporter")
        self.setGeometry(100, 100, 430, 560)
        self.setMinimumSize(400, 500)

        self.apply_macos_theme()

        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        root_layout = QVBoxLayout()
        root_layout.setSpacing(0)
        root_layout.setContentsMargins(0, 0, 0, 0)

        content_widget = QWidget()
        content_widget.setStyleSheet("background-color: #FFFFFF;")
        content_layout = QVBoxLayout(content_widget)
        content_layout.setSpacing(0)
        content_layout.setContentsMargins(0, 0, 0, 0)

        # Scroll body
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll_area.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QScrollBar:vertical {
                background: transparent; width: 4px; margin: 0;
            }
            QScrollBar::handle:vertical {
                background: #D2D2D7; border-radius: 2px; min-height: 30px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }
        """)
        scroll_content = QWidget()
        scroll_content.setStyleSheet("background: transparent;")
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setSpacing(6)
        scroll_layout.setContentsMargins(10, 4, 10, 10)

        # Sections
        scroll_layout.addWidget(self.create_mac_section("Тип отчёта", self.create_report_type_section()))
        scroll_layout.addWidget(self.create_mac_section("Параметры", self.create_options_section()))
        scroll_layout.addWidget(self.create_mac_section("Файлы", self.create_drop_section()))

        # Progress card
        scroll_layout.addWidget(self.create_progress_card())

        scroll_area.setWidget(scroll_content)
        content_layout.addWidget(scroll_area, 1)

        # Sticky footer
        footer = self._create_footer()
        content_layout.addWidget(footer)

        root_layout.addWidget(content_widget, 1)
        central_widget.setLayout(root_layout)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self._refresh_progress_status_text()

    def _create_footer(self):
        """Fixed footer with action buttons, matching the design."""
        footer = QWidget()
        footer.setFixedHeight(46)
        footer.setStyleSheet("""
            QWidget {
                background: #FAFAFA;
                border-top: 1px solid #E5E5EA;
            }
        """)
        layout = QHBoxLayout(footer)
        layout.setContentsMargins(10, 0, 10, 0)
        layout.setSpacing(7)

        icon_btn_style = """
            QPushButton {
                background: transparent;
                border: 1px solid #E5E5EA;
                border-radius: 8px;
                padding: 0px;
                color: #86868B;
            }
            QPushButton:hover { background: #F5F5F7; border-color: #D2D2D7; }
            QPushButton:pressed { background: #E8E8ED; }
        """

        clear_btn = QPushButton()
        clear_btn.setFixedSize(30, 30)
        clear_btn.setToolTip("Очистить файлы")
        clear_btn.setStyleSheet(icon_btn_style)
        clear_btn.setIcon(make_icon("trash", "#86868B", 15))
        clear_btn.setIconSize(QSize(15, 15))
        clear_btn.clicked.connect(self.clear_files)
        layout.addWidget(clear_btn)

        settings_btn = QPushButton()
        settings_btn.setFixedSize(30, 30)
        settings_btn.setToolTip("Настройки")
        settings_btn.setStyleSheet(icon_btn_style)
        settings_btn.setIcon(make_icon("gear", "#86868B", 15))
        settings_btn.setIconSize(QSize(15, 15))
        settings_btn.clicked.connect(self.open_settings)
        layout.addWidget(settings_btn)

        yandex_browse_btn = QPushButton()
        yandex_browse_btn.setFixedSize(30, 30)
        yandex_browse_btn.setToolTip("Файлы на Яндекс.Диске")
        yandex_browse_btn.setStyleSheet(icon_btn_style)
        yandex_browse_btn.setIcon(make_icon("folder_open", "#86868B", 15))
        yandex_browse_btn.setIconSize(QSize(15, 15))
        yandex_browse_btn.clicked.connect(self._open_yandex_disk_browser)
        layout.addWidget(yandex_browse_btn)

        layout.addStretch()

        self.preview_btn = QPushButton("Предпросмотр")
        self.preview_btn.setFont(QFont(".AppleSystemUIFont", 11))
        self.preview_btn.setFixedHeight(30)
        self.preview_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                border: 1px solid #E5E5EA;
                border-radius: 8px;
                padding: 0 11px;
                color: #86868B;
            }
            QPushButton:hover { background: #F5F5F7; border-color: #D2D2D7; color: #1D1D1F; }
            QPushButton:pressed { background: #E8E8ED; }
            QPushButton:disabled { background: transparent; border-color: #E5E5EA; color: #AEAEB2; }
        """)
        self.preview_btn.setIcon(make_icon("search", "#86868B", 13))
        self.preview_btn.setIconSize(QSize(13, 13))
        self.preview_btn.clicked.connect(self.toggle_preview_panel)
        self.preview_btn.setEnabled(False)
        layout.addWidget(self.preview_btn)

        self.generate_btn = QPushButton("Создать")
        self.generate_btn.setFont(QFont(".AppleSystemUIFont", 12, QFont.DemiBold))
        self.generate_btn.setFixedHeight(30)
        self.generate_btn.setStyleSheet("""
            QPushButton {
                background-color: #007AFF;
                border: none;
                border-radius: 8px;
                padding: 0 18px;
                color: white;
            }
            QPushButton:hover { background-color: #0063D1; }
            QPushButton:pressed { background-color: #004EA3; }
            QPushButton:disabled { background-color: #E8E8ED; color: #AEAEB2; }
        """)
        self.generate_btn.setIcon(make_icon("doc", "#FFFFFF", 14))
        self.generate_btn.setIconSize(QSize(14, 14))
        self.generate_btn.clicked.connect(self.start_processing)
        self.generate_btn.setEnabled(False)
        layout.addWidget(self.generate_btn)

        return footer

    def apply_macos_theme(self):
        palette = QPalette()
        palette.setColor(QPalette.Window, QColor(255, 255, 255))
        palette.setColor(QPalette.WindowText, QColor(29, 29, 31))
        palette.setColor(QPalette.Base, QColor(255, 255, 255))
        palette.setColor(QPalette.AlternateBase, QColor(245, 245, 247))
        palette.setColor(QPalette.ToolTipBase, QColor(50, 50, 50))
        palette.setColor(QPalette.ToolTipText, QColor(255, 255, 255))
        palette.setColor(QPalette.Text, QColor(29, 29, 31))
        palette.setColor(QPalette.Button, QColor(255, 255, 255))
        palette.setColor(QPalette.ButtonText, QColor(29, 29, 31))
        palette.setColor(QPalette.BrightText, QColor(255, 255, 255))
        palette.setColor(QPalette.Highlight, QColor(0, 122, 255))  # macOS blue
        palette.setColor(QPalette.HighlightedText, QColor(255, 255, 255))
        self.setPalette(palette)

    def _set_progress_bar_color(self, color_hex: str):
        self.progress_bar.setStyleSheet(f"""
            QProgressBar {{
                border: none;
                border-radius: 2px;
                background-color: rgba(0,0,0,0.06);
            }}
            QProgressBar::chunk {{
                background-color: {color_hex};
                border-radius: 2px;
            }}
        """)

    def _set_generate_button_processing(self, active: bool):
        self.generate_btn.setText("Создаётся..." if active else "Создать")
        self.generate_btn.setIcon(make_icon("sparkle" if active else "doc", "#FFFFFF", 14))

    def _set_progress_icon(self, icon_name: str, color: str):
        self.progress_icon_label.setPixmap(make_icon_pixmap(icon_name, color, 15))

    def _clean_status_text(self, text: str) -> str:
        if not text:
            return ""
        text = re.sub(r"^[^\wА-Яа-яA-Za-z0-9]+", "", text).strip()
        return text or "Ожидание"

    def _refresh_report_type_icons(self):
        for button, icon_name in self._report_type_icon_config:
            color = "white" if button.isChecked() else "#86868B"
            button.setIcon(make_icon(icon_name, color, 13))

    def _refresh_option_icons(self):
        for btn, icon_name in self._option_icon_config:
            color = "white" if btn.isChecked() else "#86868B"
            btn.setIcon(make_icon(icon_name, color, 13))

    def _set_report_type_by_key(self, report_type: str) -> bool:
        buttons = {
            "standard": self.report_type_main,
            "me": self.report_type_me,
            "me_ours": self.report_type_me_ours,
            "dcp": self.report_type_dcp,
        }
        button = buttons.get(report_type)
        if button is None:
            return False
        if not button.isChecked():
            button.setChecked(True)
        return True

    def _set_report_type_hint(self, message: str):
        if not getattr(self, "report_type_hint_label", None):
            return
        text = (message or "").strip()
        self.report_type_hint_label.setText(text)
        self.report_type_hint_label.setVisible(bool(text))

    def _set_progress_status_text(self, text: str):
        self._progress_status_full_text = text or ""
        self._refresh_progress_status_text()

    def _refresh_progress_status_text(self):
        if not getattr(self, "status_label", None):
            return
        full_text = getattr(self, "_progress_status_full_text", "") or ""
        if not full_text:
            self.status_label.setText("")
            return
        available_width = self.status_label.width()
        if available_width <= 0:
            self.status_label.setText(full_text)
            return
        elided = self.status_label.fontMetrics().elidedText(full_text, Qt.ElideRight, available_width)
        self.status_label.setText(elided)

    def _update_open_folder_button_state(self):
        if not getattr(self, "open_folder_btn", None):
            return
        enabled = self.last_output_folder is not None
        self.open_folder_btn.setVisible(True)
        self.open_folder_btn.setEnabled(enabled)
        self.open_folder_btn.setIcon(
            make_icon("folder_open", "#86868B" if enabled else "#B3B3BA", 13)
        )

    def _update_files_list_height(self):
        if not getattr(self, "files_list", None):
            return
        count = self.files_list.count()
        if count <= 0:
            self.files_list.setMaximumHeight(0)
            return
        row_height = self.files_list.sizeHintForRow(0)
        if row_height <= 0:
            row_height = 24
        visible_rows = min(count, 3)
        frame = self.files_list.frameWidth() * 2
        total_height = frame + (row_height * visible_rows) + 4
        self.files_list.setMaximumHeight(min(total_height, 86))

    def _on_report_type_selected(self, checked: bool):
        if checked:
            self._set_report_type_hint("")
            if self._has_loaded_files():
                self._invalidate_preview()

    def _get_all_loaded_files(self) -> list:
        ordered_groups = ("audio", "video", "pdf", "csv", "params")
        files = []
        for group in ordered_groups:
            files.extend(self.files_data.get(group, []))
        return files

    def _detect_report_type_from_filename(self, file_path: str):
        detected_type, analysis = detect_report_type_from_files([file_path])
        if analysis["reason"] == "ambiguous":
            return None
        return detected_type

    def _apply_auto_detected_report_type(self, files: list):
        if not getattr(self, "auto_report_type_checkbox", None):
            return
        if not self.auto_report_type_checkbox.isChecked():
            return
        if not files:
            self._set_report_type_hint("")
            return

        detected_type, analysis = detect_report_type_from_files(files)
        if not detected_type:
            if analysis["reason"] == "ambiguous":
                logger.warning(
                    "Автоопределение типа отчета отменено: конфликтующие сигналы %s",
                    analysis["scores"],
                )
                self._set_report_type_hint(
                    "Не удалось однозначно определить тип отчета по имени файла. Выберите его вручную."
                )
            else:
                logger.info("Автоопределение типа отчета: недостаточно сигналов в имени файла")
                self._set_report_type_hint(
                    "Автоопределение не сработало: в имени файла недостаточно сигналов. Выберите тип отчета вручную."
                )
            return

        previous_type = self.get_report_type()
        self._set_report_type_by_key(detected_type)
        self._set_report_type_hint("")
        if detected_type != previous_type:
            logger.info(
                "Тип отчета автоматически переключен: %s -> %s (%s)",
                previous_type,
                detected_type,
                analysis["scores"],
            )
        else:
            logger.info(
                "Автоопределение типа отчета подтвердило текущий выбор: %s (%s)",
                detected_type,
                analysis["scores"],
            )

    def _save_auto_detect_report_type_setting(self, enabled: bool):
        settings = SettingsDialog.load_settings()
        settings["auto_detect_report_type"] = enabled
        SettingsDialog.save_settings(settings)

    def _on_auto_report_type_detection_changed(self, state):
        enabled = state == Qt.Checked
        self._save_auto_detect_report_type_setting(enabled)
        logger.info(
            "Автоопределение типа отчета по имени файла: %s",
            "ВКЛЮЧЕНО" if enabled else "ВЫКЛЮЧЕНО",
        )
        if enabled:
            self._apply_auto_detected_report_type(self._get_all_loaded_files())
        else:
            self._set_report_type_hint("")

    def _has_loaded_files(self) -> bool:
        return any(self.files_data.get(key) for key in ('audio', 'video', 'csv', 'pdf', 'params'))

    def _update_action_buttons(self):
        has_files = self._has_loaded_files()
        allow_actions = has_files and not self._preview_busy and not self._processing_active
        if getattr(self, "preview_btn", None):
            self.preview_btn.setEnabled(has_files and not self._preview_busy and not self._processing_active)
        if getattr(self, "generate_btn", None):
            self.generate_btn.setEnabled(allow_actions)

    def _set_preview_button_processing(self, active: bool):
        if not getattr(self, "preview_btn", None):
            return
        if active:
            self.preview_btn.setText("Проверяем...")
        else:
            self.preview_btn.setText("Закрыть предпросмотр" if self._is_preview_open() else "Предпросмотр")
        self.preview_btn.setIcon(make_icon("sparkle" if active else "search", "#86868B", 13))

    def _is_preview_open(self) -> bool:
        return bool(getattr(self, "preview_dialog", None) and self.preview_dialog.isVisible())

    def _set_preview_visible(self, visible: bool):
        if not visible and self.preview_dialog is None:
            self._set_preview_button_processing(False)
            return
        dialog = self._ensure_preview_dialog()
        if visible:
            dialog.show()
            dialog.raise_()
            dialog.activateWindow()
        else:
            dialog.hide()
        self._set_preview_button_processing(False)

    def toggle_preview_panel(self):
        if self._is_preview_open():
            self._set_preview_visible(False)
            return
        if self._processing_active:
            return
        self._set_preview_visible(True)
        if self.preview_data is not None:
            self._populate_preview_widgets(self.preview_data)
            return
        if self._preview_busy:
            return
        self.request_preview_refresh(immediate=True)

    def _apply_preview_expanded_state(self):
        if getattr(self, "preview_text", None):
            self.preview_text.updateGeometry()
        if getattr(self, "preview_dialog", None) and self._is_preview_open():
            self.preview_dialog.adjustSize()

    def toggle_preview_expanded(self):
        self._preview_expanded = not self._preview_expanded
        self._apply_preview_expanded_state()

    def _set_preview_status(self, text: str, tone: str = "neutral"):
        if not getattr(self, "preview_status_label", None):
            return
        colors = {
            "neutral": "#86868B",
            "info": "#007AFF",
            "success": "#34C759",
            "warning": "#FF9F0A",
            "error": "#FF3B30",
        }
        self.preview_status_label.setText(text)
        self.preview_status_label.setStyleSheet(
            f"color: {colors.get(tone, '#86868B')}; background: transparent;"
        )

    def _populate_preview_widgets(self, preview: dict):
        if not getattr(self, "preview_text", None):
            return
        self.preview_text.setHtml(self._format_preview_html(preview))
        warning_count = len(filter_preview_warnings(preview.get("warnings", []), self.get_report_type()))
        if warning_count:
            self._set_preview_status(f"Проверено: {warning_count} предупрежд.", "warning")
        else:
            self._set_preview_status("Проверено: ошибок распознавания не найдено", "success")

    def _invalidate_preview(self):
        self.preview_timer.stop()
        self._preview_epoch += 1
        self.preview_data = None
        preview_was_open = self._is_preview_open()
        if getattr(self, "preview_text", None):
            self.preview_text.setHtml(self._preview_message_html("Предпросмотр обновляется после изменения файлов…"))
        self._set_preview_status("Ожидает проверки", "info" if self._has_loaded_files() else "neutral")
        self._preview_expanded = False
        self._apply_preview_expanded_state()
        if not preview_was_open:
            self._set_preview_visible(False)
        elif self._has_loaded_files() and not self._processing_active:
            self.request_preview_refresh()
        self._update_action_buttons()

    def request_preview_refresh(self, immediate: bool = False):
        if not self._has_loaded_files():
            self.reset_preview_card()
            return
        if self._processing_active:
            return
        if immediate:
            self.preview_timer.stop()
            self._run_preview_refresh()
        else:
            self.preview_timer.start(180)

    def _run_preview_refresh(self):
        if not self._has_loaded_files() or self._processing_active:
            return
        if self._preview_busy:
            self._preview_refresh_pending = True
            return

        self._preview_busy = True
        self._preview_refresh_pending = False
        if self._is_preview_open():
            self._set_preview_visible(True)
        self.preview_text.setHtml(self._preview_message_html("Анализирую файлы и вытаскиваю параметры из PDF…"))
        self._set_preview_status("Анализ файлов…", "info")
        self._set_preview_button_processing(True)
        self._update_action_buttons()

        self.preview_thread = PreviewAnalysisThread(self.files_data, self._preview_epoch, self.get_report_type())
        self.preview_thread.preview_ready.connect(self._on_preview_ready)
        self.preview_thread.preview_failed.connect(self._on_preview_failed)
        self.preview_thread.start()

    def _on_preview_ready(self, preview: dict, epoch: int):
        if epoch != self._preview_epoch:
            logger.info("Preview result ignored as stale (epoch %s != %s)", epoch, self._preview_epoch)
        else:
            self.preview_data = preview
            self._populate_preview_widgets(preview)

        self._preview_busy = False
        self._set_preview_button_processing(False)
        self._update_action_buttons()

        if self._preview_refresh_pending:
            self._preview_refresh_pending = False
            self.request_preview_refresh(immediate=True)

    def _on_preview_failed(self, error_text: str, epoch: int):
        if epoch == self._preview_epoch:
            self.preview_text.setHtml(
                self._preview_message_html(
                    f"Не удалось построить предпросмотр.<br><span style='font-size:11px;'>{html.escape(error_text)}</span>",
                    tone="error"
                )
            )
            self._set_preview_status("Ошибка предпросмотра", "error")
        self._preview_busy = False
        self._set_preview_button_processing(False)
        self._update_action_buttons()

        if self._preview_refresh_pending:
            self._preview_refresh_pending = False
            self.request_preview_refresh(immediate=True)

    def reset_preview_card(self):
        self.preview_timer.stop()
        self.preview_data = None
        self._preview_refresh_pending = False
        self._preview_expanded = False
        self._set_preview_visible(False)
        if getattr(self, "preview_text", None):
            self.preview_text.setHtml("")
        self._apply_preview_expanded_state()
        self._set_preview_status("Нет данных", "neutral")
        self._set_preview_button_processing(False)
        self._update_action_buttons()

    def _on_preview_dialog_closed(self):
        self._set_preview_button_processing(False)

    def _preview_message_html(self, message: str, tone: str = "neutral") -> str:
        colors = {
            "neutral": "#86868B",
            "info": "#007AFF",
            "success": "#34C759",
            "warning": "#FF9F0A",
            "error": "#FF3B30",
        }
        return (
            "<div style='font-family:\".AppleSystemUIFont\"; color:{color}; "
            "font-size:12px; padding:18px 6px; text-align:center;'>"
            "{message}</div>"
        ).format(color=colors.get(tone, "#86868B"), message=message)

    def _format_preview_html(self, preview: dict) -> str:
        current_report_type = self.get_report_type()
        selected_type = REPORT_TYPE_LABELS.get(current_report_type, current_report_type)
        detected_type = preview.get("report_type", {}).get("detected")
        detected_label = REPORT_TYPE_LABELS.get(detected_type, "не определён") if detected_type else "не определён"
        counts = preview.get("counts", {})
        slot_summary = preview.get("slot_summary", {})
        recognized = preview.get("recognized", [])
        warnings = filter_preview_warnings(preview.get("warnings", []), current_report_type)

        recognized_by_slot = {
            item.get("slot"): item for item in recognized if item.get("slot")
        }

        problems = []
        if detected_type and detected_type != self.get_report_type():
            problems.append(
                f"Тип отчета: выбран «{selected_type}», автоопределение видит «{detected_label}»."
            )
        for message in warnings:
            formatted_message = format_preview_warning(message, current_report_type)
            if formatted_message not in problems:
                problems.append(formatted_message)

        def slot_row(slot_name: str) -> str:
            file_name = slot_summary.get(slot_name)
            item = recognized_by_slot.get(slot_name, {})
            details = item.get("details") or []
            item_warnings = [
                format_preview_warning(message, current_report_type)
                for message in filter_preview_warnings(item.get("warnings") or [], current_report_type)
            ]
            metrics = item.get("metrics") or {}

            # Не показываем пустые строки без данных и без предупреждений.
            if not file_name and not item_warnings:
                return ""

            if item_warnings:
                status_label = "Проверить"
                status_bg = "#FFF7E8"
                status_color = "#A15C00"
                status_text = "<br>".join(html.escape(message) for message in item_warnings)
                row_bg = "#FFFDFC"
            else:
                status_label = "OK"
                status_bg = "#EEF8EF"
                status_color = "#248A3D"
                status_text = "Без замечаний"
                row_bg = "#FFFFFF"

            file_cell = html.escape(
                preview_display_filename(file_name, current_report_type, item.get("kind"))
            ) if file_name else "—"
            kind = item.get("kind")
            if kind == "pdf":
                details_lines = [
                    f"LUFS: {_preview_numeric(metrics.get('lufs'))}",
                    f"TRUE PEAK: {_preview_numeric(metrics.get('true_peak'))} dBTP",
                    f"LRA: {_preview_numeric(metrics.get('lra'))}",
                ]
            elif kind == "params":
                if current_report_type in {"me", "me_ours"}:
                    details_lines = [
                        f"TRUE PEAK: {_preview_numeric(metrics.get('true_peak'))} dBTP",
                    ]
                else:
                    details_lines = [
                        f"Target: {_preview_numeric(metrics.get('target_lufs'))} LUFS",
                        f"TRUE PEAK: {_preview_numeric(metrics.get('true_peak'))} dBTP",
                        f"LRA max: {_preview_numeric(metrics.get('lra_max'))}",
                    ]
            elif kind in {"audio", "video"}:
                details_lines = [f"Хронометраж: {_preview_duration_text(metrics.get('duration'))}"]
                details_lines.extend(part for part in details if part)
            else:
                details_lines = details

            details_cell = (
                "".join(
                    f"<div>{html.escape(str(line))}</div>"
                    for line in details_lines
                    if str(line).strip()
                )
                or "—"
            )
            return (
                f"<tr style='background:{row_bg};'>"
                f"<td style='padding:7px 8px; border-bottom:1px solid #F0F0F2; vertical-align:top; font-size:10px; color:#5C5C62;'><b>{html.escape(preview_slot_label(slot_name, current_report_type))}</b></td>"
                f"<td style='padding:7px 8px; border-bottom:1px solid #F0F0F2; vertical-align:top; font-size:11px; color:#1D1D1F;'><b>{file_cell}</b></td>"
                f"<td style='padding:7px 8px; border-bottom:1px solid #F0F0F2; vertical-align:top; font-size:10px; color:#5C5C62;'>{details_cell}</td>"
                "<td style='padding:7px 8px; border-bottom:1px solid #F0F0F2; vertical-align:top; font-size:10px;'>"
                f"<span style='display:inline-block; padding:2px 7px; border-radius:999px; background:{status_bg}; color:{status_color};'><b>{html.escape(status_label)}</b></span>"
                f"<div style='margin-top:4px; color:{status_color if status_label != 'OK' else '#5C5C62'};'>{status_text if status_label != 'OK' else html.escape(status_text)}</div>"
                "</td></tr>"
            )

        header_bg = "#FFF6F7" if problems else "#EEF8EF"
        header_border = "#FFD8DF" if problems else "#CFEBD4"
        header_title = "#C23A57" if problems else "#248A3D"
        header_text = (
            f"{len(problems)} проблем(ы) найдено"
            if problems else "Критичных проблем не найдено"
        )

        table_rows = "".join(slot_row(slot_name) for slot_name in (
            "video", "csv", "params",
            "audio_20_c", "audio_20_uc", "audio_51_c", "audio_51_uc",
            "pdf_20_c", "pdf_20_uc", "pdf_20", "pdf_51_c", "pdf_51_uc", "pdf_51",
        ))

        sections = [
            "<div style='font-family:\".AppleSystemUIFont\"; color:#1D1D1F; font-size:12px; line-height:1.35;'>",
            "<div style='padding:10px 12px; border-radius:12px; "
            f"background:{header_bg}; border:1px solid {header_border}; margin-bottom:10px;'>"
            f"<div style='font-size:12px; color:{header_title};'><b>{html.escape(header_text)}</b></div>"
            f"<div style='font-size:10px; color:#5C5C62; margin-top:3px;'>"
            f"Выбран: {html.escape(selected_type)} · Авто: {html.escape(detected_label)} · Файлов: {sum(counts.values())}</div>"
            "</div>",
        ]

        if problems:
            sections.extend([
                "<div style='font-size:11px; color:#C23A57; margin:0 0 6px 0;'><b>Проблемы</b></div>",
                "<div style='padding:10px 12px; border-radius:12px; background:#FFFFFF; border:1px solid #FFD8DF; margin-bottom:10px;'>",
                "".join(
                    f"<div style='font-size:11px; color:#7A2336; margin:0 0 6px 0;'>&bull; {html.escape(message)}</div>"
                    for message in problems[:8]
                ),
                "</div>",
            ])

        sections.extend([
            "<table style='width:100%; border-collapse:collapse; background:#FFFFFF; border:1px solid #ECECF0; border-radius:12px;'>",
            "<thead><tr>"
            "<th style='text-align:left; padding:7px 8px; font-size:10px; color:#86868B; border-bottom:1px solid #ECECF0;'>Слот</th>"
            "<th style='text-align:left; padding:7px 8px; font-size:10px; color:#86868B; border-bottom:1px solid #ECECF0;'>Файл</th>"
            "<th style='text-align:left; padding:7px 8px; font-size:10px; color:#86868B; border-bottom:1px solid #ECECF0;'>Данные</th>"
            "<th style='text-align:left; padding:7px 8px; font-size:10px; color:#86868B; border-bottom:1px solid #ECECF0;'>Статус</th>"
            "</tr></thead><tbody>",
            table_rows or "<tr><td colspan='4' style='padding:10px 8px; font-size:11px; color:#86868B;'>Нет распознанных файлов для показа</td></tr>",
            "</tbody></table>",
        ])

        sections.append("</div>")
        return "".join(sections)

    def _progress_visuals(self, value: int, message: str = ""):
        visuals = [
            (0, "copy", "#AEAEB2", "#F5F5F7", "Ожидание"),
            (10, "copy", "#007AFF", "#EBF5FF", "Копирование файлов..."),
            (20, "search", "#007AFF", "#EBF5FF", "Анализ файлов..."),
            (50, "pdf", "#007AFF", "#EBF5FF", "Извлечение данных из PDF..."),
            (70, "sparkle", "#007AFF", "#EBF5FF", "Генерация заключений..."),
            (85, "doc", "#007AFF", "#EBF5FF", "Создание отчета..."),
            (100, "check", "#34C759", "rgba(52, 199, 89, 0.12)", "Готово!"),
        ]
        selected = visuals[0]
        for item in visuals:
            if value >= item[0]:
                selected = item
        icon, icon_color, icon_bg, fallback_text = selected[1], selected[2], selected[3], selected[4]
        return icon, icon_color, icon_bg, message or fallback_text

    def create_progress_card(self):
        self.progress_card = QFrame()
        self.progress_card.setVisible(False)
        self.progress_card.setStyleSheet("""
            QFrame {
                background: #FFFFFF;
                border: 1px solid #E5E5EA;
                border-radius: 10px;
            }
        """)
        layout = QHBoxLayout(self.progress_card)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(8)

        # Icon chip
        self.progress_icon_label = QLabel()
        self.progress_icon_label.setAlignment(Qt.AlignCenter)
        self.progress_icon_label.setFixedSize(22, 22)
        self.progress_icon_label.setStyleSheet("background: #EBF5FF; border-radius: 6px;")
        self._set_progress_icon("copy", "#007AFF")
        layout.addWidget(self.progress_icon_label)

        center = QVBoxLayout()
        center.setContentsMargins(0, 0, 0, 0)
        center.setSpacing(4)

        self.status_label = QLabel("Ожидание")
        self.status_label.setFont(QFont(".AppleSystemUIFont", 10))
        self.status_label.setStyleSheet("color: #1D1D1F; background: transparent;")
        self.status_label.setMinimumWidth(0)
        self.status_label.setSizePolicy(QSizePolicy.Ignored, QSizePolicy.Preferred)
        center.addWidget(self.status_label)

        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setFixedHeight(3)
        self._set_progress_bar_color("#007AFF")
        center.addWidget(self.progress_bar)
        layout.addLayout(center, 1)

        self.progress_value_label = QLabel("0%")
        self.progress_value_label.setFont(QFont(".AppleSystemUIFont", 9))
        self.progress_value_label.setStyleSheet("color: #AEAEB2; background: transparent;")
        layout.addWidget(self.progress_value_label)

        self._progress_status_full_text = "Ожидание"
        self._refresh_progress_status_text()

        return self.progress_card

    def open_report_folder(self):
        if not self.last_output_folder:
            return
        try:
            subprocess.Popen(["open", str(self.last_output_folder)])
        except Exception as e:
            logger.warning(f"Не удалось открыть папку отчета: {e}")
    
    def create_mac_section(self, title, content_widget):
        section = QWidget()
        section.setStyleSheet("background: transparent;")
        section_layout = QVBoxLayout()
        section_layout.setContentsMargins(0, 0, 0, 0)
        section_layout.setSpacing(3)

        if title:
            section_title = QLabel(title)
            section_title.setFont(QFont(".AppleSystemUIFont", 11))
            section_title.setStyleSheet("color: #86868B; background: transparent; letter-spacing: 0.02em;")
            section_layout.addWidget(section_title)

        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: #FFFFFF;
                border-radius: 10px;
                border: 1px solid #E5E5EA;
            }
        """)
        card_layout = QVBoxLayout()
        card_layout.setContentsMargins(12, 10, 12, 10)
        card_layout.setSpacing(6)
        card_layout.addWidget(content_widget)
        card.setLayout(card_layout)

        section_layout.addWidget(card)
        section.setLayout(section_layout)
        return section
    
    def create_report_type_section(self):
        from PyQt5.QtWidgets import QButtonGroup
        widget = QWidget()
        widget.setStyleSheet("background: transparent;")
        layout = QVBoxLayout()
        layout.setSpacing(8)
        layout.setContentsMargins(0, 0, 0, 0)
        buttons_row = QHBoxLayout()
        buttons_row.setSpacing(2)
        buttons_row.setContentsMargins(0, 0, 0, 0)

        tab_style = """
            QPushButton {
                background: transparent;
                color: #86868B;
                spacing: 5px;
                padding: 6px 8px;
                border-radius: 8px;
                border: 1px solid transparent;
                font-weight: 500;
                font-family: ".AppleSystemUIFont";
                font-size: 11px;
                text-align: center;
            }
            QPushButton:hover {
                background-color: #F5F5F7;
                color: #1D1D1F;
            }
            QPushButton:checked {
                background-color: #007AFF;
                border: 1px solid #007AFF;
                color: white;
            }
        """

        self._report_type_group = QButtonGroup(widget)
        self._report_type_group.setExclusive(True)

        self.report_type_main = QPushButton("Основной")
        self.report_type_main.setFont(QFont(".AppleSystemUIFont", 11))
        self.report_type_main.setCheckable(True)
        self.report_type_main.setChecked(True)
        self.report_type_main.setStyleSheet(tab_style)
        self.report_type_main.setIconSize(QSize(14, 14))
        self.report_type_main.setMinimumWidth(98)
        self.report_type_main.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self._report_type_group.addButton(self.report_type_main)
        buttons_row.addWidget(self.report_type_main, 1)

        self.report_type_me = QPushButton("ME")
        self.report_type_me.setFont(QFont(".AppleSystemUIFont", 11))
        self.report_type_me.setCheckable(True)
        self.report_type_me.setStyleSheet(tab_style)
        self.report_type_me.setIconSize(QSize(14, 14))
        self.report_type_me.setMinimumWidth(58)
        self.report_type_me.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self._report_type_group.addButton(self.report_type_me)
        buttons_row.addWidget(self.report_type_me, 1)

        self.report_type_me_ours = QPushButton("ME (наши)")
        self.report_type_me_ours.setFont(QFont(".AppleSystemUIFont", 11))
        self.report_type_me_ours.setCheckable(True)
        self.report_type_me_ours.setStyleSheet(tab_style)
        self.report_type_me_ours.setIconSize(QSize(14, 14))
        self.report_type_me_ours.setMinimumWidth(112)
        self.report_type_me_ours.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self._report_type_group.addButton(self.report_type_me_ours)
        buttons_row.addWidget(self.report_type_me_ours, 1)

        self.report_type_dcp = QPushButton("DCP")
        self.report_type_dcp.setFont(QFont(".AppleSystemUIFont", 11))
        self.report_type_dcp.setCheckable(True)
        self.report_type_dcp.setStyleSheet(tab_style)
        self.report_type_dcp.setIconSize(QSize(14, 14))
        self.report_type_dcp.setMinimumWidth(58)
        self.report_type_dcp.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self._report_type_group.addButton(self.report_type_dcp)
        buttons_row.addWidget(self.report_type_dcp, 1)

        self._report_type_icon_config = [
            (self.report_type_main, "doc"),
            (self.report_type_me, "speaker"),
            (self.report_type_me_ours, "speaker"),
            (self.report_type_dcp, "film"),
        ]
        for button, _ in self._report_type_icon_config:
            button.toggled.connect(self._refresh_report_type_icons)
            button.toggled.connect(self._on_report_type_selected)
        self._refresh_report_type_icons()

        layout.addLayout(buttons_row)

        # Separator
        sep = QFrame()
        sep.setFixedHeight(1)
        sep.setStyleSheet("background: #E5E5EA; border: none;")
        layout.addWidget(sep)

        # Auto detect checkbox
        auto_row = QHBoxLayout()
        auto_row.setSpacing(7)
        auto_row.setContentsMargins(0, 0, 0, 0)

        self.auto_report_type_checkbox = MacToggleSwitch()
        self.auto_report_type_checkbox.setToolTip(
            "Автоматически определяет тип отчёта по имени файла.\n"
            "Если сигналов недостаточно или они конфликтуют, приложение предложит выбрать тип вручную."
        )
        auto_detect_enabled = SettingsDialog.load_settings().get("auto_detect_report_type", False)
        self.auto_report_type_checkbox.setChecked(auto_detect_enabled)
        self.auto_report_type_checkbox.toggled.connect(
            lambda checked: self._on_auto_report_type_detection_changed(Qt.Checked if checked else Qt.Unchecked)
        )
        auto_row.addWidget(self.auto_report_type_checkbox)
        auto_label = QLabel("Авто по имени файла")
        auto_label.setFont(QFont(".AppleSystemUIFont", 11))
        auto_label.setStyleSheet("color: #1D1D1F; background: transparent;")
        auto_label.setToolTip(self.auto_report_type_checkbox.toolTip())
        auto_row.addWidget(auto_label)
        auto_row.addStretch()
        layout.addLayout(auto_row)

        self.report_type_hint_label = QLabel("")
        self.report_type_hint_label.setVisible(False)
        self.report_type_hint_label.setWordWrap(True)
        self.report_type_hint_label.setFont(QFont(".AppleSystemUIFont", 10))
        self.report_type_hint_label.setStyleSheet(
            "color: #FF9F0A; background: transparent; padding-top: 2px;"
        )
        layout.addWidget(self.report_type_hint_label)

        widget.setLayout(layout)
        return widget
    
    def create_options_section(self):
        from PyQt5.QtWidgets import QLineEdit
        widget = QWidget()
        widget.setStyleSheet("background: transparent;")
        layout = QVBoxLayout()
        layout.setSpacing(9)
        layout.setContentsMargins(0, 0, 0, 0)

        option_tab_style = """
            QPushButton {
                background: transparent;
                color: #86868B;
                spacing: 4px;
                padding: 5px 7px;
                border-radius: 8px;
                border: 1px solid transparent;
                font-family: ".AppleSystemUIFont";
                font-size: 10px;
                font-weight: 500;
                text-align: center;
            }
            QPushButton:hover {
                background-color: #F5F5F7;
                color: #1D1D1F;
            }
            QPushButton:checked {
                background-color: #007AFF;
                border: 1px solid #007AFF;
                color: white;
            }
        """

        # Таймер для периодической проверки Ollama
        self.ollama_timer = QTimer()
        self.ollama_timer.timeout.connect(self._check_ollama_status)

        # Ollama dot indicator
        self.ollama_dot = QLabel()
        self.ollama_dot.setFixedSize(6, 6)
        self.ollama_dot.setStyleSheet("background-color: #8E8E93; border-radius: 3px;")
        self.ollama_dot.setToolTip("Ollama: не используется")

        # Видимое имя модели рядом с индикатором — пользователь должен видеть,
        # какая модель сконфигурирована, не наводя курсор на точку.
        self.ollama_model_label = QLabel(self.conclusion_gen.llm_model)
        self.ollama_model_label.setFont(QFont(".AppleSystemUIFont", 9))
        self.ollama_model_label.setStyleSheet("color: #86868B;")
        self.ollama_model_label.setToolTip("Ollama: не используется")

        self.ai_enabled_checkbox = QPushButton("AI генерация")
        self.ai_enabled_checkbox.setFont(QFont(".AppleSystemUIFont", 11))
        self.ai_enabled_checkbox.setCheckable(True)
        self.ai_enabled_checkbox.setChecked(False)
        self.ai_enabled_checkbox.setStyleSheet(option_tab_style)
        self.ai_enabled_checkbox.setIconSize(QSize(11, 11))
        self.ai_enabled_checkbox.setFixedHeight(28)
        self.ai_enabled_checkbox.setMinimumWidth(102)
        self.ai_enabled_checkbox.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self.ai_enabled_checkbox.toggled.connect(
            lambda checked: self.toggle_ai_generation(Qt.Checked if checked else Qt.Unchecked)
        )

        self.tp_verify_checkbox = QPushButton("TP verify")
        self.tp_verify_checkbox.setFont(QFont(".AppleSystemUIFont", 11))
        self.tp_verify_checkbox.setCheckable(True)
        self.tp_verify_checkbox.setChecked(False)
        self.tp_verify_checkbox.setStyleSheet(option_tab_style)
        self.tp_verify_checkbox.setIconSize(QSize(11, 11))
        self.tp_verify_checkbox.setFixedHeight(28)
        self.tp_verify_checkbox.setMinimumWidth(84)
        self.tp_verify_checkbox.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self.tp_verify_checkbox.setToolTip(
            "Точное измерение True Peak (4x oversampling, ITU-R BS.1770-4).\n"
            "Результат с точностью до сотых (как iZotope RX)."
        )

        self._option_icon_config = [
            (self.ai_enabled_checkbox, "sparkle"),
            (self.tp_verify_checkbox, "target"),
        ]
        for btn, _ in self._option_icon_config:
            btn.toggled.connect(self._refresh_option_icons)
        self._refresh_option_icons()

        # Tabs row
        options_row = QHBoxLayout()
        options_row.setSpacing(2)
        options_row.setContentsMargins(0, 0, 0, 0)

        ai_tab_widget = QWidget()
        ai_tab_widget.setStyleSheet("background: transparent;")
        ai_tab_widget.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        ai_tab_layout = QHBoxLayout(ai_tab_widget)
        ai_tab_layout.setSpacing(4)
        ai_tab_layout.setContentsMargins(0, 0, 0, 0)
        ai_tab_layout.addWidget(self.ai_enabled_checkbox)
        ai_tab_layout.addWidget(self.ollama_dot, 0, Qt.AlignVCenter)
        ai_tab_layout.addWidget(self.ollama_model_label, 0, Qt.AlignVCenter)

        options_row.addWidget(ai_tab_widget)
        options_row.addWidget(self.tp_verify_checkbox)
        options_row.addStretch(1)
        layout.addLayout(options_row)

        # Name + folder row
        self.name_input = QLineEdit()
        self.name_input.setPlaceholderText("Ваше имя")
        self.name_input.setFont(QFont(".AppleSystemUIFont", 11))
        self.name_input.setFixedHeight(30)
        self.name_input.setStyleSheet("""
            QLineEdit {
                background: #FFFFFF;
                border: 1px solid #D2D2D7;
                border-radius: 7px;
                padding: 0 8px;
                color: #1D1D1F;
                selection-background-color: #007AFF;
            }
            QLineEdit:focus { border: 1px solid #007AFF; }
        """)
        self.load_saved_name()
        self.name_input.textChanged.connect(self.save_name)
        self.output_folder_path = None

        self.output_folder_field = QLineEdit("Рабочий стол")
        self.output_folder_field.setReadOnly(True)
        self.output_folder_field.setFont(QFont(".AppleSystemUIFont", 11))
        self.output_folder_field.setFixedHeight(30)
        self.output_folder_field.setStyleSheet("""
            QLineEdit {
                background: #F5F5F7;
                border: 1px solid #E5E5EA;
                border-radius: 7px;
                padding: 0 8px;
                color: #86868B;
            }
        """)
        self.output_folder_field.setMaximumWidth(130)

        folder_btn_style = """
            QPushButton {
                background: #F5F5F7;
                border: 1px solid #E5E5EA;
                border-radius: 7px;
                padding: 0 10px;
                color: #86868B;
                height: 30px;
                font-size: 11px;
                font-family: ".AppleSystemUIFont";
            }
            QPushButton:hover { background: #ECECF0; border-color: #D2D2D7; }
            QPushButton:pressed { background: #E0E0E5; }
        """
        pick_btn = QPushButton("Папка")
        pick_btn.setFixedHeight(28)
        pick_btn.setMinimumWidth(78)
        pick_btn.setMaximumWidth(96)
        pick_btn.setStyleSheet(folder_btn_style)
        pick_btn.setIcon(make_icon("folder", "#86868B", 12))
        pick_btn.setIconSize(QSize(12, 12))
        pick_btn.setToolTip("Выбрать папку")
        pick_btn.clicked.connect(self._pick_folder_and_update_btn)
        self._folder_pick_btn = pick_btn

        fields_row = QHBoxLayout()
        fields_row.setSpacing(7)
        fields_row.setContentsMargins(0, 0, 0, 0)
        fields_row.addWidget(self.name_input, 1)
        fields_row.addWidget(pick_btn, 0)

        layout.addLayout(fields_row)
        widget.setLayout(layout)
        return widget

    def _pick_folder_and_update_btn(self):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для отчёта")
        if folder:
            self.output_folder_path = Path(folder)
            short = self.output_folder_path.name
            self._folder_pick_btn.setText(short[:9] + "…" if len(short) > 9 else short)
            self._folder_pick_btn.setToolTip(str(self.output_folder_path))
            self.output_folder_field.setText(short)
        else:
            self.output_folder_path = None
            self._folder_pick_btn.setText("Папка")
            self._folder_pick_btn.setToolTip("Выбрать папку")
            self.output_folder_field.setText("Рабочий стол")

    def create_name_section(self):
        from PyQt5.QtWidgets import QLineEdit

        widget = QWidget()
        widget.setStyleSheet("background: transparent;")
        layout = QVBoxLayout()
        layout.setSpacing(4)
        layout.setContentsMargins(0, 0, 0, 0)

        label = QLabel("Имя:")
        label.setFont(QFont(".AppleSystemUIFont", 11))
        label.setStyleSheet("background: transparent; border: none; color: #86868B;")
        layout.addWidget(label)

        self.name_input = QLineEdit()
        self.name_input.setPlaceholderText("Ваше имя")
        self.name_input.setFont(QFont(".AppleSystemUIFont", 12))
        self.name_input.setStyleSheet("""
            QLineEdit {
                background: #FFFFFF;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                padding: 7px 10px;
                color: #1D1D1F;
                selection-background-color: #007AFF;
            }
            QLineEdit:focus {
                border: 1px solid #007AFF;
                background: #FFFFFF;
            }
        """)

        self.load_saved_name()
        self.name_input.textChanged.connect(self.save_name)
        self.output_folder_path = None

        layout.addWidget(self.name_input)
        widget.setLayout(layout)
        return widget
    
    def create_output_folder_section(self):
        widget = QWidget()
        widget.setStyleSheet("background: transparent;")
        layout = QHBoxLayout()
        layout.setSpacing(6)
        layout.setContentsMargins(0, 0, 0, 0)

        small_btn_style = """
            QPushButton {
                background-color: #FFFFFF;
                border: 1px solid #E5E5EA;
                border-radius: 10px;
                padding: 0px;
                color: #86868B;
            }
            QPushButton:hover { background-color: #F5F5F7; border-color: #D2D2D7; }
            QPushButton:pressed { background-color: #E8E8ED; }
        """

        label = QLabel("Папка:")
        label.setFont(QFont(".AppleSystemUIFont", 11))
        label.setStyleSheet("background: transparent; border: none; color: #86868B;")
        layout.addWidget(label)

        from PyQt5.QtWidgets import QLineEdit
        self.output_folder_field = QLineEdit("Рабочий стол")
        self.output_folder_field.setReadOnly(True)
        self.output_folder_field.setFont(QFont(".AppleSystemUIFont", 11))
        self.output_folder_field.setStyleSheet("""
            QLineEdit {
                background: #F5F5F7;
                border: 1px solid #D2D2D7;
                border-radius: 8px;
                padding: 7px 10px;
                color: #86868B;
            }
        """)
        self.output_folder_field.setMaximumWidth(180)
        layout.addWidget(self.output_folder_field)

        pick_btn = QPushButton()
        pick_btn.setFixedSize(40, 40)
        pick_btn.setStyleSheet(small_btn_style)
        pick_btn.setIcon(make_icon("folder_open", "#86868B", 14))
        pick_btn.setIconSize(QSize(14, 14))
        pick_btn.clicked.connect(self.select_output_folder)
        layout.addWidget(pick_btn)

        clear_btn = QPushButton()
        clear_btn.setFixedSize(40, 40)
        clear_btn.setStyleSheet(small_btn_style)
        clear_btn.setIcon(make_icon("x", "#86868B", 14))
        clear_btn.setIconSize(QSize(14, 14))
        clear_btn.clicked.connect(self.clear_output_folder)
        layout.addWidget(clear_btn)

        layout.addStretch()
        widget.setLayout(layout)
        return widget
    
    def create_drop_section(self):
        widget = QWidget()
        widget.setStyleSheet("background: transparent;")
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(5)

        self.drop_zone = DropZone()
        self.drop_zone.files_dropped.connect(self.handle_dropped_files)
        self.drop_zone.folder_dropped.connect(self._on_report_folder_dropped)
        layout.addWidget(self.drop_zone)

        self.files_empty_label = QLabel("")
        self.files_empty_label.setVisible(False)

        self.files_list = QListWidget()
        self.files_list.setMaximumHeight(108)
        self.files_list.setVisible(False)
        self.files_list.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.files_list.setTextElideMode(Qt.ElideMiddle)
        self.files_list.setStyleSheet("""
            QListWidget {
                background-color: transparent;
                border: none;
                padding: 0px;
                font-family: ".AppleSystemUIFont";
                font-size: 11px;
                color: #1D1D1F;
                outline: none;
            }
            QListWidget::item {
                padding: 4px 0px;
                border-bottom: 1px solid #F5F5F7;
            }
            QListWidget::item:selected {
                background-color: #EBF5FF;
                color: #007AFF;
                border-radius: 8px;
            }
        """)
        layout.addWidget(self.files_list)

        action_btn_style = """
            QPushButton {
                background-color: #F5F5F7;
                border: 1px solid #E5E5EA;
                border-radius: 8px;
                color: #86868B;
                padding: 4px 6px;
            }
            QPushButton:hover {
                background-color: #EBF5FF;
                border-color: rgba(0, 122, 255, 0.18);
                color: #007AFF;
            }
            QPushButton:disabled {
                background-color: #F7F7F8;
                border-color: #E9E9ED;
                color: #B3B3BA;
            }
        """

        actions_row = QHBoxLayout()
        actions_row.setSpacing(6)

        self.open_folder_btn = QPushButton("Папка")
        self.open_folder_btn.setToolTip("Открыть папку с отчётом")
        self.open_folder_btn.setVisible(True)
        self.open_folder_btn.setEnabled(False)
        self.open_folder_btn.setFixedHeight(30)
        self.open_folder_btn.setStyleSheet(action_btn_style)
        self.open_folder_btn.setIcon(make_icon("folder_open", "#86868B", 13))
        self.open_folder_btn.setIconSize(QSize(13, 13))
        self.open_folder_btn.clicked.connect(self.open_report_folder)
        actions_row.addWidget(self.open_folder_btn)
        self._update_open_folder_button_state()

        self.send_to_disk_btn = QPushButton("Отправить")
        self.send_to_disk_btn.setToolTip("Отправить на Яндекс.Диск")
        self.send_to_disk_btn.setVisible(True)
        self.send_to_disk_btn.setEnabled(False)
        self.send_to_disk_btn.setFixedHeight(30)
        self.send_to_disk_btn.setStyleSheet(action_btn_style)
        self.send_to_disk_btn.setIcon(make_icon("folder_open", "#86868B", 13))
        self.send_to_disk_btn.setIconSize(QSize(13, 13))
        self.send_to_disk_btn.clicked.connect(self._send_report_to_disk)
        actions_row.addWidget(self.send_to_disk_btn)

        self.compare_with_disk_btn = QPushButton("Сравнить")
        self.compare_with_disk_btn.setToolTip("Сравнить с версией на Яндекс.Диске")
        self.compare_with_disk_btn.setVisible(True)
        self.compare_with_disk_btn.setEnabled(False)
        self.compare_with_disk_btn.setFixedHeight(30)
        self.compare_with_disk_btn.setStyleSheet(action_btn_style)
        self.compare_with_disk_btn.setIcon(make_icon("search", "#86868B", 13))
        self.compare_with_disk_btn.setIconSize(QSize(13, 13))
        self.compare_with_disk_btn.clicked.connect(self._compare_report_with_disk)
        actions_row.addWidget(self.compare_with_disk_btn)

        self.edit_report_btn = QPushButton("Правка")
        self.edit_report_btn.setToolTip(
            "Открыть отчёт для редактирования — после каждого сохранения\n"
            "правки автоматически уедут на Яндекс.Диск (если он уже был отправлен)."
        )
        self.edit_report_btn.setVisible(True)
        self.edit_report_btn.setEnabled(False)
        self.edit_report_btn.setFixedHeight(30)
        self.edit_report_btn.setStyleSheet(action_btn_style)
        self.edit_report_btn.setIcon(make_icon("doc", "#86868B", 13))
        self.edit_report_btn.setIconSize(QSize(13, 13))
        self.edit_report_btn.clicked.connect(self._edit_report)
        actions_row.addWidget(self.edit_report_btn)

        layout.addLayout(actions_row)

        # Отдельная строка под тремя кнопками — чтобы не ломать их ряд по
        # ширине узкого окна. Видна, только когда в очереди что-то есть.
        self.yandex_queue_btn = QPushButton("Очередь")
        self.yandex_queue_btn.setToolTip("Очередь автозагрузки на Яндекс.Диск")
        self.yandex_queue_btn.setVisible(False)
        self.yandex_queue_btn.setFixedHeight(26)
        self.yandex_queue_btn.setStyleSheet("""
            QPushButton {
                background-color: #FFF6E5;
                border: 1px solid #FFE2A8;
                border-radius: 8px;
                color: #C77700;
                padding: 3px 6px;
            }
            QPushButton:hover { background-color: #FFEFCC; }
        """)
        self.yandex_queue_btn.setIcon(make_icon("folder_open", "#C77700", 12))
        self.yandex_queue_btn.setIconSize(QSize(12, 12))
        self.yandex_queue_btn.clicked.connect(self._open_yandex_queue_dialog)
        layout.addWidget(self.yandex_queue_btn)

        self.edit_sync_status_label = QLabel("")
        self.edit_sync_status_label.setVisible(False)
        self.edit_sync_status_label.setFont(QFont(".AppleSystemUIFont", 10))
        self.edit_sync_status_label.setStyleSheet("color: #86868B; background: transparent;")
        layout.addWidget(self.edit_sync_status_label)

        widget.setLayout(layout)
        return widget

    def _ensure_preview_dialog(self):
        if self.preview_dialog is None:
            self.preview_dialog = PreviewDialog(self)
            self.preview_text = self.preview_dialog.preview_text
            self.preview_status_label = self.preview_dialog.status_label
        return self.preview_dialog
    
    def create_files_list_section(self):
        widget = QWidget()
        widget.setStyleSheet("background: transparent;")
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)

        self.files_empty_label = QLabel("Пусто")
        self.files_empty_label.setFont(QFont(".AppleSystemUIFont", 12))
        self.files_empty_label.setAlignment(Qt.AlignCenter)
        self.files_empty_label.setStyleSheet("color: #AEAEB2; background: transparent; padding: 14px;")
        layout.addWidget(self.files_empty_label)

        self.files_list = QListWidget()
        self.files_list.setMaximumHeight(132)
        self.files_list.setVisible(False)
        self.files_list.setStyleSheet("""
            QListWidget {
                background-color: transparent;
                border: none;
                padding: 0px;
                font-family: ".AppleSystemUIFont";
                font-size: 12px;
                color: #1D1D1F;
                outline: none;
            }
            QListWidget::item {
                padding: 6px 8px;
                border-bottom: 1px solid #F5F5F7;
            }
            QListWidget::item:selected {
                background-color: #EBF5FF;
                color: #007AFF;
                border-radius: 8px;
            }
        """)
        layout.addWidget(self.files_list)

        widget.setLayout(layout)
        return widget
    
    def create_buttons_section(self):
        # Buttons are now in the sticky footer created by _create_footer() in init_ui.
        widget = QWidget()
        widget.setStyleSheet("background: transparent;")
        widget.setVisible(False)
        widget.setFixedHeight(0)
        return widget

    def _create_buttons_section_legacy(self):
        widget = QWidget()
        widget.setStyleSheet("background: transparent;")
        outer = QVBoxLayout()
        outer.setContentsMargins(0, 6, 0, 0)
        outer.setSpacing(0)

        sep = QFrame()
        sep.setFixedHeight(1)
        sep.setStyleSheet("background-color: #E5E5EA; border: none;")
        outer.addWidget(sep)

        layout = QHBoxLayout()
        layout.setSpacing(10)
        layout.setContentsMargins(0, 10, 0, 0)

        clear_btn = QPushButton()
        clear_btn.setFixedSize(40, 40)
        clear_btn.setToolTip("Очистить файлы")
        clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #FFFFFF;
                border: 1px solid #E5E5EA;
                border-radius: 10px;
                padding: 0px;
                color: #86868B;
            }
            QPushButton:hover { background-color: #F5F5F7; border-color: #D2D2D7; }
            QPushButton:pressed { background-color: #E8E8ED; }
        """)
        clear_btn.setIcon(make_icon("trash", "#86868B", 16))
        clear_btn.setIconSize(QSize(16, 16))
        clear_btn.clicked.connect(self.clear_files)
        layout.addWidget(clear_btn)

        settings_btn = QPushButton()
        settings_btn.setFixedSize(40, 40)
        settings_btn.setToolTip("Настройки")
        settings_btn.setStyleSheet("""
            QPushButton {
                background-color: #FFFFFF;
                border: 1px solid #E5E5EA;
                border-radius: 10px;
                padding: 0px;
                color: #86868B;
            }
            QPushButton:hover { background-color: #F5F5F7; border-color: #D2D2D7; }
            QPushButton:pressed { background-color: #E8E8ED; }
        """)
        settings_btn.setIcon(make_icon("gear", "#86868B", 16))
        settings_btn.setIconSize(QSize(16, 16))
        settings_btn.clicked.connect(self.open_settings)
        layout.addWidget(settings_btn)

        layout.addStretch()

        self.preview_btn = QPushButton("Предпросмотр")
        self.preview_btn.setFont(QFont(".AppleSystemUIFont", 12, QFont.DemiBold))
        self.preview_btn.setFixedHeight(34)
        self.preview_btn.setStyleSheet("""
            QPushButton {
                background-color: #FFFFFF;
                border: 1px solid #BFD9FF;
                border-radius: 8px;
                padding: 6px 16px;
                color: #007AFF;
            }
            QPushButton:hover {
                background-color: #EBF5FF;
                border-color: #8DBEFF;
            }
            QPushButton:pressed {
                background-color: #DCEBFF;
            }
            QPushButton:disabled {
                background-color: #F5F5F7;
                border-color: #E5E5EA;
                color: #AEAEB2;
            }
        """)
        self.preview_btn.setIcon(make_icon("search", "#007AFF", 14))
        self.preview_btn.setIconSize(QSize(14, 14))
        self.preview_btn.clicked.connect(self.toggle_preview_panel)
        self.preview_btn.setEnabled(False)
        layout.addWidget(self.preview_btn)

        self.generate_btn = QPushButton("Создать")
        self.generate_btn.setFont(QFont(".AppleSystemUIFont", 13, QFont.DemiBold))
        self.generate_btn.setFixedHeight(34)
        self.generate_btn.setStyleSheet("""
            QPushButton {
                background-color: #007AFF;
                border: none;
                border-radius: 8px;
                padding: 6px 22px;
                color: white;
            }
            QPushButton:hover {
                background-color: #0063D1;
            }
            QPushButton:pressed {
                background-color: #004EA3;
            }
            QPushButton:disabled {
                background-color: #E8E8ED;
                color: #AEAEB2;
            }
        """)
        self.generate_btn.setIcon(make_icon("doc", "#FFFFFF", 14))
        self.generate_btn.setIconSize(QSize(14, 14))
        self.generate_btn.clicked.connect(self.start_processing)
        self.generate_btn.setEnabled(False)
        layout.addWidget(self.generate_btn)

        outer.addLayout(layout)
        widget.setLayout(outer)
        return widget
    
    def load_saved_name(self):
        """Загрузка сохраненного имени"""
        settings = SettingsDialog.load_settings()
        name = settings.get("name", "")
        if name:
            self.name_input.setText(name)

    def save_name(self):
        """Сохранение имени"""
        settings = SettingsDialog.load_settings()
        settings["name"] = self.name_input.text()
        SettingsDialog.save_settings(settings)
    
    def select_output_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для отчёта")
        if folder:
            self.output_folder_path = Path(folder)
            short = self.output_folder_path.name
            if getattr(self, "output_folder_field", None):
                self.output_folder_field.setText(short)
            if getattr(self, "_folder_pick_btn", None):
                self._folder_pick_btn.setText(short[:18] + "…" if len(short) > 18 else short)
            logger.info(f"Output folder selected: {self.output_folder_path}")

    def clear_output_folder(self):
        self.output_folder_path = None
        if getattr(self, "output_folder_field", None):
            self.output_folder_field.setText("Рабочий стол")
        if getattr(self, "_folder_pick_btn", None):
            self._folder_pick_btn.setText("Рабочий стол")
        logger.info("Output folder reset to Desktop")

    def open_settings(self):
        """Открыть диалог настроек"""
        dialog = SettingsDialog(parent=self)
        dialog.exec_()

    def handle_dropped_files(self, files):
        """Обработка перетащенных файлов"""
        logger.info(f"Получено файлов: {len(files)}")

        had_csv_in_batch = False

        for file_path in files:
            file_name = os.path.basename(file_path)
            file_ext = Path(file_path).suffix.lower()
            filename_lower = Path(file_path).stem.lower()
            item_label = file_name

            # Определяем тип файла
            if file_ext in ['.wav', '.mp3', '.flac', '.aac', '.mxf', '.caf', '.aiff', '.aif']:
                self.files_data['audio'].append(file_path)
                # Определяем тип аудио для отображения
                if '20' in filename_lower and 'cens' in filename_lower and 'uncens' not in filename_lower:
                    icon_name = "music"
                    item_label = f"{file_name} · 2.0 C"
                elif '20' in filename_lower and 'uncens' in filename_lower:
                    icon_name = "music"
                    item_label = f"{file_name} · 2.0 U"
                elif '51' in filename_lower and 'cens' in filename_lower and 'uncens' not in filename_lower:
                    icon_name = "music"
                    item_label = f"{file_name} · 5.1 C"
                elif '51' in filename_lower and 'uncens' in filename_lower:
                    icon_name = "music"
                    item_label = f"{file_name} · 5.1 U"
                else:
                    icon_name = "music"
            elif file_ext in ['.mp4', '.mov', '.avi', '.mkv', '.m4v', '.webm']:
                self.files_data['video'].append(file_path)
                icon_name = "video"
            elif file_ext == '.csv':
                self.files_data['csv'].append(file_path)
                had_csv_in_batch = True
                icon_name = "csv"
            elif file_ext == '.txt' and ('параметры' in filename_lower or 'parametry' in filename_lower):
                self.files_data['params'].append(file_path)
                icon_name = "params"
            elif file_ext == '.pdf':
                self.files_data['pdf'].append(file_path)
                # Определяем тип PDF для отображения
                if '20' in filename_lower and 'cens' in filename_lower and 'uncens' not in filename_lower:
                    icon_name = "pdf"
                    item_label = f"{file_name} · 2.0 C"
                elif '20' in filename_lower and 'uncens' in filename_lower:
                    icon_name = "pdf"
                    item_label = f"{file_name} · 2.0 U"
                elif '51' in filename_lower and 'cens' in filename_lower and 'uncens' not in filename_lower:
                    icon_name = "pdf"
                    item_label = f"{file_name} · 5.1 C"
                elif '51' in filename_lower and 'uncens' in filename_lower:
                    icon_name = "pdf"
                    item_label = f"{file_name} · 5.1 U"
                else:
                    icon_name = "pdf"
            else:
                continue

            # Добавляем в список
            item = QListWidgetItem(item_label)
            item.setIcon(make_icon(icon_name, "#86868B", 15))
            item.setToolTip(file_path)
            self.files_list.addItem(item)

        self._apply_auto_detected_report_type(self._get_all_loaded_files())

        # Если в этом батче пришёл реальный CSV — убираем ранее авто-созданный пустой
        if had_csv_in_batch and self._auto_created_csv_path:
            self._remove_auto_created_csv()

        # Если CSV нет, но есть аудио/видео — создаём пустой CSV автоматически
        if (not self.files_data['csv']
                and (self.files_data['audio'] or self.files_data['video'])
                and not self._auto_created_csv_path):
            self._create_empty_csv()

        # Toggle empty label
        has_files = self.files_list.count() > 0
        self.files_list.setVisible(has_files)
        self._update_files_list_height()
        self.drop_zone.set_files_loaded(has_files)
        self.drop_zone.set_completed_state(False)
        self._update_open_folder_button_state()

        # Обновляем доступность кнопки
        self._invalidate_preview()
        self._update_action_buttons()
        
        # Показываем детальную статистику
        stats = []
        if self.files_data['audio']:
            stats.append(f"Аудио {len(self.files_data['audio'])}")
        if self.files_data['video']:
            stats.append(f"Видео {len(self.files_data['video'])}")
        if self.files_data['csv']:
            stats.append(f"CSV {len(self.files_data['csv'])}")
        if self.files_data['pdf']:
            stats.append(f"PDF {len(self.files_data['pdf'])}")
        if self.files_data['params']:
            stats.append(f"Параметры {len(self.files_data['params'])}")
        
        if not self.progress_card.isVisible():
            self._set_progress_status_text(" | ".join(stats) if stats else "Нет файлов")

    def _adopt_report_folder(self, folder: Path, docx_path: Path) -> None:
        """Принимает готовую папку отчёта как текущую — общая логика для

        обычной генерации (processing_finished) и для перетащенной готовой
        папки отчёта (_on_report_folder_dropped): включает кнопки
        отправки/сравнения/правки и запускает слежение за файлами папки
        для автосинхронизации правок.
        """
        self.last_output_folder = folder
        self.last_report_docx_path = docx_path
        self._update_open_folder_button_state()
        self._edit_sync.watch_many(str(p) for p in folder.iterdir() if p.is_file())
        if getattr(self, "send_to_disk_btn", None):
            self.send_to_disk_btn.setEnabled(True)
        if getattr(self, "compare_with_disk_btn", None):
            self.compare_with_disk_btn.setEnabled(True)
        if getattr(self, "edit_report_btn", None):
            self.edit_report_btn.setEnabled(True)

    def _on_report_folder_dropped(self, folder_path: str):
        """Перетащили папку уже готового отчёта прямо в окно — принимаем её

        и сразу запускаем тот же поток, что кнопка «Отправить» (сравнение
        с предыдущей версией + отправка на Диск). Работает и для отчётов
        из прошлых сессий, не только для только что сгенерированного.
        """
        if self._processing_active:
            return
        folder = Path(folder_path)
        docx_candidates = sorted(folder.glob("отчет_*.docx"))
        if not docx_candidates:
            QMessageBox.warning(
                self, "Не похоже на отчёт",
                "В перетащенной папке не найден файл отчёта (отчет_*.docx)."
            )
            return
        self._adopt_report_folder(folder, docx_candidates[0])
        self._send_report_to_disk()

    # Заголовки пустого маркер-листа (совпадают с шапкой таблицы в exact_report_generator)
    _EMPTY_CSV_HEADERS = [
        "Timecode In", "Timecode Out", "Description",
        "2.0 C", "2.0 UC", "5.1 C", "5.1 UC",
        "БЛОКЕР", "ТРЕБУЕТ ИСПРАВЛЕНИЯ", "ТРЕБУЕТ КОММЕНТАРИЯ", "КОММЕНТАРИИ",
    ]

    def _create_empty_csv(self):
        """Создаёт пустой CSV (только шапка) и добавляет его в список файлов."""
        source = (self.files_data['audio'] or self.files_data['video'])[0]
        base_name = sanitize_base_name(source)
        tmp_dir = Path(tempfile.gettempdir()) / "beast_auto_reporter_empty_csv"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        csv_path = tmp_dir / f"{base_name}.csv"

        try:
            with open(csv_path, 'w', encoding='utf-8', newline='') as f:
                writer = csv.writer(f, delimiter='\t')
                writer.writerow(self._EMPTY_CSV_HEADERS)
        except Exception as e:
            logger.error(f"❌ Не удалось создать пустой CSV: {e}")
            return

        csv_path_str = str(csv_path)
        self.files_data['csv'].append(csv_path_str)
        self._auto_created_csv_path = csv_path_str

        item = QListWidgetItem(csv_path.name)
        item.setIcon(make_icon("csv", "#007AFF", 15))
        item.setToolTip(f"Автоматически создан пустой CSV:\n{csv_path_str}")
        self.files_list.addItem(item)
        self._update_files_list_height()

        logger.info(f"✅ Создан пустой CSV: {csv_path}")

    def _remove_auto_created_csv(self):
        """Удаляет ранее автоматически созданный пустой CSV (из списка, данных и диска)."""
        auto_path = self._auto_created_csv_path
        if not auto_path:
            return

        if auto_path in self.files_data['csv']:
            self.files_data['csv'].remove(auto_path)

        # Убираем соответствующий элемент из UI-списка по tooltip
        for idx in range(self.files_list.count() - 1, -1, -1):
            item = self.files_list.item(idx)
            if item is not None and item.toolTip().endswith(auto_path):
                self.files_list.takeItem(idx)
                break

        try:
            Path(auto_path).unlink(missing_ok=True)
        except Exception as e:
            logger.warning(f"Не удалось удалить авто-CSV {auto_path}: {e}")

        self._auto_created_csv_path = None
        logger.info("🗑  Авто-CSV удалён (загружен реальный CSV)")

    def _auto_reset_after_done(self):
        """Автоматически готовит приложение к следующему отчёту после успешной генерации."""
        if self._processing_active or not self.drop_zone._is_done:
            return
        self.clear_files()

    def clear_files(self):
        self.auto_reset_timer.stop()
        self._preview_epoch += 1
        if self._auto_created_csv_path:
            try:
                Path(self._auto_created_csv_path).unlink(missing_ok=True)
            except Exception:
                pass
            self._auto_created_csv_path = None

        self.files_data = {
            'audio': [],
            'video': [],
            'csv': [],
            'pdf': [],
            'params': []
        }
        self.files_list.clear()
        self._update_files_list_height()
        self.files_list.setVisible(False)
        self.generate_btn.setEnabled(False)
        self._set_generate_button_processing(False)
        self._set_progress_bar_color("#007AFF")
        self.progress_card.setVisible(False)
        self._set_progress_status_text("Ожидание")
        self._set_progress_icon("copy", "#AEAEB2")
        self.progress_icon_label.setStyleSheet("""
            background: #F5F5F7;
            border-radius: 7px;
        """)
        self.progress_value_label.setText("0%")
        self.progress_bar.setValue(0)
        self.drop_zone.set_processing_state(False)
        self.drop_zone.set_completed_state(False)
        self.drop_zone.set_files_loaded(False)
        self._set_report_type_hint("")
        self.last_output_folder = None
        self.last_report_docx_path = None
        self._update_open_folder_button_state()
        if getattr(self, "send_to_disk_btn", None):
            self.send_to_disk_btn.setEnabled(False)
        if getattr(self, "compare_with_disk_btn", None):
            self.compare_with_disk_btn.setEnabled(False)
        if getattr(self, "edit_report_btn", None):
            self.edit_report_btn.setEnabled(False)
        if getattr(self, "_edit_sync", None):
            self._edit_sync.unwatch_all()
        if getattr(self, "edit_sync_status_label", None):
            self.edit_sync_status_label.setVisible(False)
        self._processing_active = False
        self.reset_preview_card()
    
    def toggle_ai_generation(self, state):
        """Переключение AI генерации"""
        enabled = state == Qt.Checked
        self.conclusion_gen.use_llm = enabled
        logger.info(f"AI генерация: {'ВКЛЮЧЕНА' if enabled else 'ВЫКЛЮЧЕНА'}")
        if enabled:
            self._set_ollama_dot("checking", "")
            self._check_ollama_status()
            self.ollama_timer.start(5000)
        else:
            self.ollama_timer.stop()
            self._set_ollama_dot("disabled", "")

    def _set_ollama_dot(self, status: str, detail: str):
        """
        Обновить цвет, подсказку и подпись модели у индикатора Ollama.
        detail — доп. контекст для тултипа (например, список установленных моделей).
        """
        model_name = self.conclusion_gen.llm_model
        styles = {
            "disabled": ("background-color: #8E8E93; border-radius: 4px;",
                         "Ollama: не используется"),
            "checking": ("background-color: #FF9F0A; border-radius: 4px;",
                         "Ollama: проверка соединения..."),
            "ok": ("background-color: #30D158; border-radius: 4px;",
                   f"Ollama подключена ✓\nМодель «{model_name}» установлена и готова к работе"),
            "no_service": ("background-color: #FF3B30; border-radius: 4px;",
                           "Ollama не отвечает.\nЗапустите приложение Ollama "
                           f"(host: {self.conclusion_gen.ollama_host})"),
            "model_missing": ("background-color: #FF9F0A; border-radius: 4px;",
                               f"Ollama подключена, но модель «{model_name}» не установлена.\n"
                               f"Выполните в терминале: ollama pull {model_name}"
                               + (f"\n\nУстановлены: {detail}" if detail else "")),
            "error": ("background-color: #FF3B30; border-radius: 4px;",
                      "Ollama: не удалось проверить статус"),
        }
        style, tooltip = styles.get(status, styles["disabled"])
        self.ollama_dot.setStyleSheet(style)
        self.ollama_dot.setToolTip(tooltip)

        label_colors = {
            "disabled": "#86868B",
            "checking": "#86868B",
            "ok": "#30D158",
            "no_service": "#FF3B30",
            "model_missing": "#FF9F0A",
            "error": "#FF3B30",
        }
        self.ollama_model_label.setText(model_name)
        self.ollama_model_label.setStyleSheet(f"color: {label_colors.get(status, '#86868B')};")
        self.ollama_model_label.setToolTip(tooltip)

    def _check_ollama_status(self):
        """Проверить доступность Ollama и наличие модели в фоновом потоке (thread-safe через сигнал)."""
        import threading

        def _ping():
            try:
                info = self.conclusion_gen.get_ollama_status()
                if not info["reachable"]:
                    new_status, detail = "no_service", ""
                elif not info["model_installed"]:
                    new_status, detail = "model_missing", ", ".join(info["installed_models"][:5])
                else:
                    new_status, detail = "ok", ""
            except Exception:
                new_status, detail = "error", ""
            self._ollama_status_signal.emit(new_status, detail)

        threading.Thread(target=_ping, daemon=True).start()

    def _check_for_updates(self, silent: bool = True):
        """Проверить наличие новой версии на GitHub в фоновом потоке.

        silent=True (автопроверка при запуске): при отсутствии обновлений — тишина.
        silent=False (ручная проверка из настроек): всегда показывает результат,
        в том числе "у вас последняя версия", и игнорирует ранее пропущенную версию.
        """
        import threading

        self._update_check_silent = silent

        def _worker():
            try:
                info = check_for_update(APP_VERSION)
            except Exception as exc:
                logger.info("Проверка обновлений завершилась с ошибкой: %s", exc)
                info = None
            self._update_check_signal.emit(info)

        threading.Thread(target=_worker, daemon=True).start()

    def _on_update_check_result(self, info):
        silent = getattr(self, "_update_check_silent", True)

        if info is None:
            if not silent:
                QMessageBox.information(
                    self, "Обновление",
                    f"У вас установлена последняя версия ({APP_VERSION})."
                )
            return
        if silent and info.version == load_skipped_version():
            return

        box = QMessageBox(self)
        box.setWindowTitle("Доступно обновление")
        box.setText(
            f"Вышла новая версия Beast Auto Reporter {info.version}.\n"
            f"У вас установлена версия {APP_VERSION}."
        )
        if info.notes:
            box.setInformativeText(info.notes[:500])
        download_label = "Скачать и открыть" if info.is_direct_download else "Открыть страницу релиза"
        download_btn = box.addButton(download_label, QMessageBox.AcceptRole)
        box.addButton("Напомнить позже", QMessageBox.RejectRole)
        skip_btn = box.addButton("Пропустить версию", QMessageBox.DestructiveRole)
        box.setDefaultButton(download_btn)
        box.exec_()

        clicked = box.clickedButton()
        if clicked is download_btn:
            self._download_and_open_update(info)
        elif clicked is skip_btn:
            save_skipped_version(info.version)

    def _download_and_open_update(self, info):
        import threading

        if not info.download_url:
            QMessageBox.warning(
                self, "Обновление",
                f"Не удалось найти файл обновления.\nОткройте страницу релиза вручную: {info.html_url}"
            )
            return

        if not info.is_direct_download:
            # Нет файла под текущую архитектуру (arm64/Intel) — открываем
            # страницу релиза, чтобы пользователь выбрал файл сам, вместо
            # того чтобы молча скачать несовместимую сборку.
            try:
                subprocess.run(["open", info.download_url], check=False)
            except Exception as exc:
                logger.error("Не удалось открыть страницу релиза: %s", exc)
            return

        self._pending_update_info = info
        self._update_progress_dialog = QProgressDialog("Загрузка обновления…", None, 0, 0, self)
        self._update_progress_dialog.setWindowTitle("Обновление")
        self._update_progress_dialog.setCancelButton(None)
        self._update_progress_dialog.setWindowModality(Qt.WindowModal)
        self._update_progress_dialog.show()

        def _worker():
            try:
                path = download_update_asset(info.download_url)
                self._update_download_signal.emit(True, path)
            except Exception as exc:
                logger.error("Не удалось скачать обновление: %s", exc, exc_info=True)
                self._update_download_signal.emit(False, str(exc))

        threading.Thread(target=_worker, daemon=True).start()

    def _on_update_download_finished(self, success: bool, payload: str):
        if self._update_progress_dialog:
            self._update_progress_dialog.close()
            self._update_progress_dialog = None

        info = self._pending_update_info
        self._pending_update_info = None

        if success:
            try:
                subprocess.run(["open", payload], check=False)
            except Exception as exc:
                logger.error("Не удалось открыть скачанное обновление: %s", exc)
        else:
            html_url = info.html_url if info else ""
            QMessageBox.warning(
                self, "Обновление",
                f"Не удалось скачать обновление автоматически ({payload}).\n"
                f"Откройте страницу релиза вручную: {html_url}"
            )
    
    def get_report_type(self):
        """Получение выбранного типа отчета"""
        if self.report_type_me.isChecked():
            return "me"
        elif self.report_type_me_ours.isChecked():
            return "me_ours"
        elif self.report_type_dcp.isChecked():
            return "dcp"
        else:
            return "standard"
    
    def start_processing(self):
        """Запуск обработки файлов"""
        logger.info("=== ЗАПУСК ОБРАБОТКИ ===")
        self.preview_timer.stop()

        if self._preview_busy:
            QMessageBox.information(self, "Предпросмотр", "Дождитесь завершения предпросмотра файлов.")
            return

        # Проверка соответствия файлов (можно отключить в настройках)
        _settings = SettingsDialog.load_settings()
        if _settings.get("check_file_consistency", True):
            mismatch = validate_file_consistency(self.files_data)
            if mismatch:
                msg = QMessageBox(self)
                msg.setIcon(QMessageBox.Warning)
                msg.setWindowTitle("Несоответствие файлов")
                msg.setText("Похоже, что добавленные аудио, PDF и CSV относятся к разным материалам.")
                msg.setInformativeText("Продолжить генерацию отчёта?")
                msg.setDetailedText("\n".join(mismatch))
                msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                msg.setDefaultButton(QMessageBox.No)
                msg.button(QMessageBox.Yes).setText("Продолжить")
                msg.button(QMessageBox.No).setText("Отмена")
                details_btn = None
                for _btn in msg.findChildren(QPushButton):
                    if _btn.text() in ("Show Details...", "Hide Details..."):
                        details_btn = _btn
                if details_btn is not None:
                    def _relabel_details(btn=details_btn):
                        btn.setText("Детали")
                    details_btn.setText("Детали")
                    details_btn.clicked.connect(lambda: QTimer.singleShot(0, _relabel_details))
                if msg.exec_() != QMessageBox.Yes:
                    logger.info("Генерация отменена пользователем из-за несоответствия файлов")
                    return

        if getattr(self, "auto_report_type_checkbox", None) and self.auto_report_type_checkbox.isChecked():
            self._apply_auto_detected_report_type(self._get_all_loaded_files())
        
        # Получаем базовое имя для папки
        audio_files = self.files_data.get('audio', [])
        video_files = self.files_data.get('video', [])
        csv_files = self.files_data.get('csv', [])
        
        base_name = "отчет"
        if audio_files:
            base_name = sanitize_base_name(audio_files[0])
        elif video_files:
            base_name = sanitize_base_name(video_files[0])
        elif csv_files:
            base_name = sanitize_base_name(csv_files[0])
        
        # Папка вывода: выбранная пользователем или Рабочий стол
        if self.output_folder_path:
            output_folder = self.create_output_folder(base_name, base_dir=Path(self.output_folder_path))
            logger.info(f"Папка отчета: {output_folder}")
        else:
            output_folder = self.create_output_folder(base_name)
            logger.info(f"Папка отчета на Desktop: {output_folder}")
        self.last_output_folder = output_folder
        self.last_report_docx_path = Path(output_folder) / f"отчет_{base_name}.docx"

        # Отключаем кнопки
        self.auto_reset_timer.stop()
        self._processing_active = True
        self._set_generate_button_processing(True)
        self._update_action_buttons()
        self.progress_card.setVisible(True)
        self.progress_bar.setValue(0)
        self._set_progress_bar_color("#007AFF")
        self.drop_zone.set_completed_state(False)
        self.drop_zone.set_processing_state(True)
        self._update_open_folder_button_state()
        
        # Сохраняем имя
        self.prepared_by = self.name_input.text().strip() or "Не указано"
        
        # Запускаем поток с Desktop папкой
        report_type = self.get_report_type()
        _saved_settings = SettingsDialog.load_settings()
        pyloudnorm_enabled = _saved_settings.get("extended_analysis_enabled", False)
        tp_verify_enabled = self.tp_verify_checkbox.isChecked()
        delete_sources = _saved_settings.get("delete_sources_after_copy", False)
        logger.info(f"PyLoudNorm: {pyloudnorm_enabled}, Report type: {report_type}")
        logger.info(f"TP verify: {tp_verify_enabled}, Delete sources: {delete_sources}")

        self.thread = ProcessingThread(
            self, self.files_data, report_type, str(output_folder),
            pyloudnorm_enabled, tp_verify_enabled, delete_sources,
        )
        self.thread.status_update.connect(self.on_thread_status_update)
        self.thread.progress_update.connect(self.on_thread_progress_update)
        self.thread.finished.connect(self.processing_finished)
        self.thread.true_peak_results_ready.connect(self._show_tp_results_dialog)
        self.thread.start()
        logger.info("Thread started")

    def on_thread_status_update(self, message: str):
        self.progress_card.setVisible(True)
        icon, text_color, icon_bg, display_text = self._progress_visuals(
            self.progress_bar.value(),
            self._clean_status_text(message),
        )
        self._set_progress_icon(icon, text_color)
        self.progress_icon_label.setStyleSheet(
            f"background: {icon_bg}; border-radius: 7px;"
        )
        self._set_progress_status_text(display_text)

    def on_thread_progress_update(self, value: int):
        self.progress_bar.setValue(value)
        self.progress_card.setVisible(True)
        icon, icon_color, icon_bg, display_text = self._progress_visuals(
            value,
            self._clean_status_text(getattr(self, "_progress_status_full_text", self.status_label.text())),
        )
        self._set_progress_icon(icon, icon_color)
        self.progress_icon_label.setStyleSheet(
            f"background: {icon_bg}; border-radius: 7px;"
        )
        self._set_progress_status_text(display_text)
        self.progress_value_label.setText(f"{value}%")
        if value >= 100:
            self._set_progress_bar_color("#34C759")
            self.drop_zone.set_processing_state(False)
            self.drop_zone.set_completed_state(True)
            self._update_open_folder_button_state()
        elif value > 0:
            self._set_progress_bar_color("#007AFF")
            self.drop_zone.set_processing_state(True)
            self._update_open_folder_button_state()

    def _show_tp_results_dialog(self, measurement_results: list):
        """Показывает диалог с результатами точного измерения True Peak."""
        dialog = TruePeakResultsDialog(measurement_results, parent=self)
        if dialog.exec_() == QDialog.Accepted:
            results = dialog.get_results()
            logger.info(f"TP precise results applied: {results}")
        else:
            results = None
            logger.info("TP precise results rejected, keeping Youlean values")
        self.thread.receive_tp_verification(results)
    
    def processing_finished(self, success, message):
        """Завершение обработки"""
        self._processing_active = False
        self._set_generate_button_processing(False)
        self._update_action_buttons()
        
        if success:
            self._set_progress_bar_color("#34C759")
            self.drop_zone.set_processing_state(False)
            self.drop_zone.set_completed_state(True)
            self._set_progress_status_text("Готово!")
            self._set_progress_icon("check", "#34C759")
            self.progress_icon_label.setStyleSheet(
                "background: rgba(52, 199, 89, 0.12); border-radius: 7px;"
            )
            self.progress_card.setVisible(True)
            self.progress_value_label.setText("100%")
            self._update_open_folder_button_state()
            report_ready = bool(self.last_report_docx_path and self.last_report_docx_path.exists())
            if report_ready:
                self._adopt_report_folder(Path(self.last_output_folder), self.last_report_docx_path)
            else:
                if getattr(self, "send_to_disk_btn", None):
                    self.send_to_disk_btn.setEnabled(False)
                if getattr(self, "compare_with_disk_btn", None):
                    self.compare_with_disk_btn.setEnabled(False)
                if getattr(self, "edit_report_btn", None):
                    self.edit_report_btn.setEnabled(False)
            logger.info(f"=== SUCCESS ===")
            logger.info(f"{message}")

            _settings = SettingsDialog.load_settings()
            if report_ready and _settings.get("yandex_auto_upload", False) and SettingsDialog.get_yandex_token():
                meta = parse_report_filename(self.last_report_docx_path.name)
                self._yandex_queue.enqueue(self.last_output_folder, meta)

            if _settings.get("auto_reset_after_done", True):
                self.auto_reset_timer.start(5000)
        else:
            error_text = message if isinstance(message, str) else "Неизвестная ошибка"
            self.progress_card.setVisible(True)
            self._set_progress_status_text(f"Ошибка: {error_text[:50]}...")
            self._set_progress_icon("x", "#FF3B30")
            self.progress_icon_label.setStyleSheet(
                "background: rgba(255, 59, 48, 0.12); border-radius: 7px;"
            )
            self._set_progress_bar_color("#007AFF")
            self.drop_zone.set_processing_state(False)
            self.drop_zone.set_completed_state(False)
            self.drop_zone.set_files_loaded(self.files_list.count() > 0)
            self._update_open_folder_button_state()
            if getattr(self, "send_to_disk_btn", None):
                self.send_to_disk_btn.setEnabled(False)
            if getattr(self, "compare_with_disk_btn", None):
                self.compare_with_disk_btn.setEnabled(False)
            if getattr(self, "edit_report_btn", None):
                self.edit_report_btn.setEnabled(False)
            logger.error(f"=== ERROR ===")
            logger.error(f"{error_text}")
            QMessageBox.critical(self, "Ошибка", error_text)

    def _update_yandex_queue_badge(self):
        if not getattr(self, "yandex_queue_btn", None):
            return
        jobs = self._yandex_queue.queue.jobs
        pending = [j for j in jobs if j.status != "done"]
        if not pending:
            self.yandex_queue_btn.setVisible(False)
            return
        self.yandex_queue_btn.setVisible(True)
        if getattr(self, "_yandex_queue_offline", False):
            self.yandex_queue_btn.setText("Очередь: нет сети")
        else:
            self.yandex_queue_btn.setText(f"Очередь ({len(pending)})")

    def _on_yandex_queue_offline_changed(self, offline: bool):
        self._yandex_queue_offline = offline
        self._update_yandex_queue_badge()

    def _on_yandex_queue_job_uploaded(self, local_folder: str, remote_path: str):
        self._yandex_remote_by_local[local_folder] = remote_path
        _play_sound()

    def _open_yandex_queue_dialog(self):
        dialog = YandexUploadQueueDialog(self._yandex_queue, parent=self)
        dialog.exec_()

    def _check_uploaded_reports_integrity(self):
        """Тихая проверка (раз в сессию), что недавно отправленные отчёты

        всё ещё на месте на Диске — если что-то пропало мимо приложения,
        предупреждаем уведомлением, не блокирующим диалогом.
        """
        if getattr(self, "_closing", False) or self._integrity_checked:
            return
        self._integrity_checked = True

        token = SettingsDialog.get_yandex_token()
        if not token:
            return
        from src.report_uploader import load_uploaded_reports
        entries = load_uploaded_reports()
        if not entries:
            return

        self._integrity_check_thread = _IntegrityCheckThread(token, entries)
        self._integrity_check_thread.finished_check.connect(self._on_integrity_check_finished)
        self._integrity_check_thread.start()

    def _on_integrity_check_finished(self, missing: list):
        if not missing:
            return
        logger.warning(
            "Проверка целостности: %d отчёт(ов) пропали с Яндекс.Диска: %s",
            len(missing), [e.get("remote_path") for e in missing],
        )
        _send_system_notification(
            "Проверка на Диске",
            f"{len(missing)} отчёт(ов) пропали с Диска — подробности в логе приложения",
        )

    def _send_report_to_disk(self):
        """Сравнивает отчёт с предыдущей версией и отправляет его на Яндекс.Диск."""
        self._start_yandex_flow(action="send")

    def _compare_report_with_disk(self):
        """Только сравнивает отчёт с версией на Яндекс.Диске, без отправки."""
        self._start_yandex_flow(action="compare")

    def _open_yandex_disk_browser(self):
        """Открывает просмотрщик файлов на Яндекс.Диске (папка «отчеты»)."""
        token = SettingsDialog.get_yandex_token()
        if not token:
            QMessageBox.warning(
                self, "Нет токена",
                "Укажите OAuth-токен Яндекс.Диска в настройках (кнопка «Настройки»)."
            )
            return
        dialog = YandexDiskBrowserDialog(token, roots=SettingsDialog.get_yandex_roots(), parent=self)
        dialog.exec_()

    def _set_yandex_buttons_enabled(self, enabled: bool):
        if getattr(self, "send_to_disk_btn", None):
            self.send_to_disk_btn.setEnabled(enabled)
        if getattr(self, "compare_with_disk_btn", None):
            self.compare_with_disk_btn.setEnabled(enabled)

    def _start_yandex_flow(self, action: str):
        """action: "send" (сравнить и отправить) или "compare" (только сравнить)."""
        if not self.last_report_docx_path or not self.last_report_docx_path.exists():
            QMessageBox.warning(self, "Нет отчёта", "Сначала сгенерируйте отчёт.")
            return

        token = SettingsDialog.get_yandex_token()
        if not token:
            QMessageBox.warning(
                self, "Нет токена",
                "Укажите OAuth-токен Яндекс.Диска в настройках (кнопка «Настройки»)."
            )
            return

        self._yandex_token = token
        self._yandex_action = action
        self._yandex_meta = parse_report_filename(self.last_report_docx_path.name)
        self._set_yandex_buttons_enabled(False)

        if self._yandex_meta is None:
            # Имя файла не содержит распознаваемых season/episode-меток —
            # прежде чем требовать ручного выбора папки, проверяем алиас/
            # нечёткое совпадение по имени файла отчёта (fallback_series_key),
            # чтобы повторная отправка того же ролика не заставляла выбирать
            # папку заново каждый раз.
            self._yandex_fallback_key = fallback_series_key(self.last_report_docx_path.name)
            self._yandex_fallback_find_thread = _FallbackFolderFindThread(
                token, self._yandex_fallback_key, series_roots=SettingsDialog.get_yandex_roots(),
            )
            self._yandex_fallback_find_thread.resolved.connect(self._on_yandex_fallback_folder_found)
            self._yandex_fallback_find_thread.failed.connect(self._on_yandex_failed)
            self._yandex_fallback_find_thread.start()
            return

        self._yandex_find_thread = YandexDiskFindVersionsThread(
            token, self._yandex_meta, series_roots=SettingsDialog.get_yandex_roots(),
        )
        self._yandex_find_thread.resolved.connect(self._on_yandex_versions_found)
        self._yandex_find_thread.failed.connect(self._on_yandex_failed)
        self._yandex_find_thread.start()

    def _on_yandex_failed(self, error_text: str):
        self._set_yandex_buttons_enabled(True)
        QMessageBox.critical(self, "Ошибка Яндекс.Диска", error_text)

    def _on_yandex_versions_found(self, result: dict):
        if getattr(self, "_closing", False):
            return
        series_path = result.get("series_path")
        episode_path = result.get("episode_path")
        versions = result.get("versions") or []

        if series_path is None:
            self._set_yandex_buttons_enabled(True)
            if self._yandex_action == "compare":
                QMessageBox.information(
                    self, "Сравнение",
                    f"Папка «{self._yandex_meta.series}» не найдена на Диске — сравнивать не с чем."
                )
                return
            msg = QMessageBox(self)
            msg.setWindowTitle("Папка не найдена")
            msg.setText(
                f"Папка «{self._yandex_meta.series}» не найдена на Диске.\n"
                "Создать автоматически по имени файла или выбрать папку вручную?"
            )
            auto_btn = msg.addButton("Создать автоматически", QMessageBox.AcceptRole)
            manual_btn = msg.addButton("Выбрать вручную", QMessageBox.ActionRole)
            msg.addButton("Отмена", QMessageBox.RejectRole)
            msg.exec_()
            clicked = msg.clickedButton()
            if clicked is auto_btn:
                self._start_yandex_upload(create_if_missing=True)
            elif clicked is manual_btn:
                self._open_yandex_folder_picker()
            return

        self._yandex_target_folder = episode_path
        remember_series_alias(self._yandex_meta.series, series_path)

        if self._yandex_action == "send":
            # Отправка — без сравнения, просто загрузка в найденную папку.
            self._start_yandex_upload(create_if_missing=False, target_folder_path=episode_path)
            return

        self._handle_yandex_versions(versions)

    def _open_yandex_folder_picker(self):
        """Открывает диалог ручного выбора/создания папки на Диске."""
        from src.yandex_disk_client import YandexDiskClient, YandexDiskError

        try:
            client = YandexDiskClient(self._yandex_token)
            dialog = YandexFolderPickerDialog(client, roots=SettingsDialog.get_yandex_roots(), parent=self)
        except YandexDiskError as exc:
            self._set_yandex_buttons_enabled(True)
            QMessageBox.critical(self, "Ошибка Яндекс.Диска", str(exc))
            return

        if dialog.exec_() != QDialog.Accepted or not dialog.selected_path:
            self._set_yandex_buttons_enabled(True)
            return

        episode_path, series_path = resolve_manual_pick_target(dialog.selected_path, self._yandex_meta)
        # Создание папки эпизода — асинхронно, чтобы не блокировать GUI-
        # поток на медленной сети.
        self._folder_picker_mkdir_thread = _MkdirThread(client, episode_path)
        self._folder_picker_mkdir_thread.finished_mkdir.connect(
            lambda success, message: self._on_folder_picker_mkdir_done(success, message, episode_path, series_path)
        )
        self._folder_picker_mkdir_thread.start()

    def _on_folder_picker_mkdir_done(self, success: bool, message: str, episode_path: str, series_path: str) -> None:
        if getattr(self, "_closing", False):
            return
        if not success:
            self._set_yandex_buttons_enabled(True)
            QMessageBox.critical(self, "Ошибка Яндекс.Диска", message)
            return

        self._yandex_target_folder = episode_path
        if self._yandex_meta is not None:
            remember_series_alias(self._yandex_meta.series, series_path)
        elif getattr(self, "_yandex_fallback_key", None):
            # Имя файла не распознано — запоминаем выбор по ключу на основе
            # имени файла отчёта (см. _start_yandex_flow), чтобы повторная
            # отправка того же ролика не требовала ручного выбора снова.
            remember_series_alias(self._yandex_fallback_key, series_path)

        if self._yandex_action == "send":
            # Отправка — без сравнения, просто загрузка в выбранную папку.
            self._start_yandex_upload(create_if_missing=False, target_folder_path=episode_path)
            return

        self._yandex_versions_thread = YandexDiskFolderVersionsThread(self._yandex_token, episode_path)
        self._yandex_versions_thread.resolved.connect(self._handle_yandex_versions)
        self._yandex_versions_thread.failed.connect(self._on_yandex_failed)
        self._yandex_versions_thread.start()

    def _on_yandex_fallback_folder_found(self, episode_path: str):
        if getattr(self, "_closing", False):
            return
        if not episode_path:
            # Алиаса нет и нечёткий поиск не нашёл совпадений — как раньше,
            # ручной выбор папки.
            self._open_yandex_folder_picker()
            return

        self._yandex_target_folder = episode_path
        if self._yandex_action == "send":
            self._start_yandex_upload(create_if_missing=False, target_folder_path=episode_path)
            return

        self._yandex_versions_thread = YandexDiskFolderVersionsThread(self._yandex_token, episode_path)
        self._yandex_versions_thread.resolved.connect(self._handle_yandex_versions)
        self._yandex_versions_thread.failed.connect(self._on_yandex_failed)
        self._yandex_versions_thread.start()

    def _handle_yandex_versions(self, versions: list):
        """Только для действия «Сравнить». Быстрый путь: сразу сравниваем

        текущий черновик с последней версией на Диске, без диалога выбора.
        В итоговом окне есть кнопка «Выбрать другую версию» — тогда
        открывается полный выбор (любые две версии между собой, например
        первую с четвёртой).
        """
        self._set_yandex_buttons_enabled(True)
        self._yandex_versions_cache = versions

        if not versions:
            QMessageBox.information(self, "Сравнение", "Предыдущих версий этого эпизода не найдено.")
            return

        latest = versions[-1]
        date_text = latest["date"].strftime("%d.%m.%Y") if latest["date"] else ""
        latest_label = f"{date_text}  {latest['label']}" if date_text else latest["label"]
        self._run_yandex_compare(
            latest["path"], YandexVersionPickerDialog.CURRENT_DRAFT,
            latest_label, "Текущий черновик (ещё не отправлен)",
        )

    def _open_yandex_version_picker(self):
        """Полный выбор двух версий — открывается по кнопке «Выбрать другую версию»."""
        versions = getattr(self, "_yandex_versions_cache", [])
        dialog = YandexVersionPickerDialog(versions, parent=self)
        if dialog.exec_() != QDialog.Accepted:
            return
        self._run_yandex_compare(
            dialog.selection_old, dialog.selection_new,
            dialog.selection_old_label, dialog.selection_new_label,
        )

    def _run_yandex_compare(self, old_path: str, new_path: str, old_label: str, new_label: str):
        self._yandex_old_label = old_label
        self._yandex_new_label = new_label
        self._set_yandex_buttons_enabled(False)
        self._yandex_compare_thread = YandexDiskCompareThread(
            self._yandex_token, old_path, new_path, self.last_report_docx_path,
        )
        self._yandex_compare_thread.resolved.connect(self._on_yandex_comparison_ready)
        self._yandex_compare_thread.failed.connect(self._on_yandex_failed)
        self._yandex_compare_thread.start()

    def _on_yandex_comparison_ready(self, comparison):
        self._set_yandex_buttons_enabled(True)

        if comparison is None:
            QMessageBox.warning(self, "Сравнение", "Не удалось прочитать выбранную версию отчёта.")
            return

        dialog = YandexUploadDiffDialog(
            comparison, parent=self, upload_mode=False,
            old_label=getattr(self, "_yandex_old_label", None),
            new_label=getattr(self, "_yandex_new_label", None),
            allow_pick_another=True,
        )
        result_code = dialog.exec_()

        if result_code == YandexUploadDiffDialog.PICK_ANOTHER:
            self._open_yandex_version_picker()

    def _start_yandex_upload(self, create_if_missing: bool, target_folder_path: str = None):
        if getattr(self, "_closing", False):
            return
        self._set_yandex_buttons_enabled(False)
        self.edit_sync_status_label.setText("Отправка на Диск: 0%")
        self.edit_sync_status_label.setStyleSheet("color: #86868B; background: transparent;")
        self.edit_sync_status_label.setVisible(True)
        self._yandex_upload_thread = YandexDiskUploadThread(
            self._yandex_token, self.last_output_folder,
            meta=self._yandex_meta, create_if_missing=create_if_missing,
            series_roots=SettingsDialog.get_yandex_roots(),
            target_folder_path=target_folder_path,
        )
        self._yandex_upload_thread.progress.connect(self._on_yandex_upload_progress)
        self._yandex_upload_thread.finished_upload.connect(self._on_yandex_upload_finished)
        # На практике недостижимо в этом вызове (create_if_missing уже решён
        # выше по потоку, либо target_folder_path уже известен) — подключено
        # для отказоустойчивости, чтобы кнопки не остались заблокированными
        # молча, если сюда всё же попадёт SeriesFolderNotFoundError.
        self._yandex_upload_thread.needs_folder.connect(lambda message: self._on_yandex_upload_finished(False, message))
        self._yandex_upload_thread.start()

    def _on_yandex_upload_progress(self, sent: int, total: int):
        percent = int(sent * 100 / total) if total else 0
        self.edit_sync_status_label.setText(f"Отправка на Диск: {percent}%")
        self.edit_sync_status_label.setStyleSheet("color: #86868B; background: transparent;")
        self.edit_sync_status_label.setVisible(True)

    def _on_yandex_upload_finished(self, success: bool, message: str):
        self._set_yandex_buttons_enabled(True)
        self.edit_sync_status_label.setVisible(False)
        if success:
            if self.last_output_folder:
                self._yandex_remote_by_local[str(self.last_output_folder)] = message
                remember_uploaded_report(str(self.last_output_folder), message)
            _play_sound()
            QMessageBox.information(self, "Готово", f"Папка с отчётом отправлена на Диск:\n{message}")
        else:
            QMessageBox.critical(self, "Ошибка Яндекс.Диска", message)

    def _edit_report(self):
        """Открывает отчёт во внешнем редакторе; после каждого сохранения,

        если отчёт уже был отправлен на Диск, правки автоматически уедут
        туда же (перезапись того же файла). Слежение за файлами папки отчёта
        уже включено с момента генерации (см. processing_finished) — здесь
        только открываем файл, watch() заново не переподписываем, чтобы не
        сбросить слежение за остальными файлами папки.
        """
        if not self.last_report_docx_path or not self.last_report_docx_path.exists():
            QMessageBox.warning(self, "Нет отчёта", "Сначала сгенерируйте отчёт.")
            return
        subprocess.Popen(["open", str(self.last_report_docx_path)])
        self.edit_sync_status_label.setText("Отслеживаем изменения — сохраните файл в редакторе, чтобы обновить на Диске")
        self.edit_sync_status_label.setStyleSheet("color: #86868B; background: transparent;")
        self.edit_sync_status_label.setVisible(True)

    def _resolve_edited_report_remote_path(self, path: str) -> str | None:
        remote_folder = self._yandex_remote_by_local.get(str(self.last_output_folder))
        if not remote_folder:
            return None
        return f"{remote_folder}/{Path(path).name}"

    def _on_edit_sync_status_changed(self, status: str):
        is_error = status.startswith("Не удалось") or status.startswith("Правки НЕ")
        is_warning = "повтор через" in status or status.startswith("Конфликт")
        color = "#FF3B30" if is_error else ("#FF9500" if is_warning else (
            "#34C759" if status.startswith("Обновлено") else "#86868B"
        ))
        self.edit_sync_status_label.setText(status)
        self.edit_sync_status_label.setStyleSheet(f"color: {color}; background: transparent;")
        self.edit_sync_status_label.setVisible(True)

    def _on_edit_sync_conflict(self, path: str, actual_modified: str):
        choice = QMessageBox.question(
            self, "Конфликт версий",
            f"Файл на Яндекс.Диске был изменён после последней синхронизации "
            f"({_format_disk_modified_date(actual_modified)}).\n\n"
            "Перезаписать его вашей версией?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No,
        )
        self._edit_sync.resolve_conflict(path, overwrite=(choice == QMessageBox.Yes))

    def extract_base_name(self, filename: str) -> str:
        """Извлечение базового названия из имени файла (из v5.11)"""
        return sanitize_base_name(filename)
    
    def find_files_in_folder(self, folder: Path) -> dict:
        """Поиск файлов в папке (полная логика из v5.11)"""
        import re
        import subprocess
        
        def find_ffprobe():
            """Поиск ffprobe"""
            import shutil
            ffprobe_path = shutil.which('ffprobe')
            if ffprobe_path:
                return ffprobe_path
            return 'ffprobe'
        
        files = {
            'audio_20_c': None,
            'audio_20_uc': None,
            'audio_51_c': None,
            'audio_51_uc': None,
            'video': None,
            'csv': None,
            'pdf_20': None,
            'pdf_51': None,
            'pdf_20_c': None,
            'pdf_20_uc': None,
            'pdf_51_c': None,
            'pdf_51_uc': None,
            'params': None,
            'all_files': []
        }
        
        for file_path in folder.iterdir():
            if file_path.is_file():
                files['all_files'].append(file_path)
                
                name_lower = file_path.name.lower()
                
                # CSV файл
                if file_path.suffix.lower() == '.csv':
                    files['csv'] = file_path
                
                # Параметры.txt
                elif 'параметры' in name_lower or 'parametry' in name_lower:
                    files['params'] = file_path
                
                # Аудио файлы (ВАЖНО: проверяем 5.1 ПЕРЕД 2.0!)
                elif file_path.suffix.lower() in ['.wav', '.mp3', '.aiff', '.flac']:
                    channel_hint = detect_channel_from_name(name_lower)
                    channels = None
                    try:
                        ffprobe_cmd = find_ffprobe()
                        result = subprocess.run(
                            [ffprobe_cmd, '-v', 'error', '-select_streams', 'a:0', 
                             '-show_entries', 'stream=channels', '-of', 
                             'default=noprint_wrappers=1:nokey=1', str(file_path)],
                            capture_output=True, text=True, timeout=5
                        )
                        if result.returncode == 0 and result.stdout.strip():
                            channels = int(result.stdout.strip())
                    except Exception:
                        channels = None
                    
                    channel_meta = "51" if channels and channels >= 6 else ("20" if channels == 2 else None)
                    channel_final = channel_meta or channel_hint
                    cens_state = detect_cens_state(name_lower)
                    suffix = "uc" if cens_state == "uc" else "c"
                    
                    if channel_final == "51":
                        files[f'audio_51_{suffix}'] = file_path
                    elif channel_final == "20":
                        files[f'audio_20_{suffix}'] = file_path
                
                # Видео
                elif file_path.suffix.lower() in ['.mp4', '.mov', '.mkv', '.avi', '.mxf', '.m4v', '.webm', '.flv']:
                    files['video'] = file_path
                
                # PDF файлы
                elif file_path.suffix.lower() == '.pdf':
                    # Проверяем на 5.1
                    if ('5.1' in file_path.name or 
                        '_51_' in name_lower or '_51.' in name_lower or
                        name_lower.endswith('_51') or
                        '_50_' in name_lower or '5.0' in file_path.name):
                        
                        if 'uncens' in name_lower or '_uc' in name_lower:
                            files['pdf_51_uc'] = file_path
                        else:
                            files['pdf_51_c'] = file_path
                        if not files['pdf_51']:
                            files['pdf_51'] = file_path
                    
                    # Проверяем на 2.0
                    elif ('2.0' in file_path.name or 
                          '_20_' in name_lower or '_20.' in name_lower or
                          name_lower.endswith('_20') or
                          '_21_' in name_lower or '2.1' in file_path.name):
                        
                        if 'uncens' in name_lower or '_uc' in name_lower:
                            files['pdf_20_uc'] = file_path
                        else:
                            files['pdf_20_c'] = file_path
                        if not files['pdf_20']:
                            files['pdf_20'] = file_path
                    
                    # Автоопределение
                    else:
                        is_uncens = 'uncens' in name_lower or '_uc' in name_lower
                        if not files['pdf_20_c'] and not files['pdf_20_uc']:
                            files['pdf_20_uc' if is_uncens else 'pdf_20_c'] = file_path
                            if not files['pdf_20']:
                                files['pdf_20'] = file_path
                        elif not files['pdf_51_c'] and not files['pdf_51_uc']:
                            files['pdf_51_uc' if is_uncens else 'pdf_51_c'] = file_path
                            if not files['pdf_51']:
                                files['pdf_51'] = file_path
        
        return files
    
    def create_output_folder(self, base_name: str, base_dir: Path = None) -> Path:
        """Создание выходной папки (из v5.11)"""
        folder_name = f"отчет_{base_name}"
        
        desktop = Path.home() / "Desktop"
        root_dir = base_dir or desktop
        output_folder = root_dir / folder_name
        
        counter = 1
        while output_folder.exists():
            output_folder = root_dir / f"{folder_name}_{counter}"
            counter += 1
        
        output_folder.mkdir(parents=True, exist_ok=True)
        
        return output_folder


def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    app.setApplicationName("Beast Auto Reporter")

    # Global macOS-like stylesheet
    app.setStyleSheet("""
        QToolTip {
            background-color: #1D1D1F;
            color: #FFFFFF;
            border: none;
            border-radius: 4px;
            padding: 4px 8px;
            font-family: ".AppleSystemUIFont";
            font-size: 11px;
        }
        QMessageBox {
            background-color: #F5F5F7;
        }
        QMessageBox QLabel {
            color: #1D1D1F;
            font-family: ".AppleSystemUIFont";
            font-size: 13px;
        }
        QMessageBox QPushButton {
            background-color: #007AFF;
            color: white;
            border: none;
            border-radius: 6px;
            padding: 6px 20px;
            font-family: ".AppleSystemUIFont";
            font-size: 13px;
        }
        QMessageBox QPushButton:hover {
            background-color: #0063D1;
        }
    """)

    window = BeastApp()
    window.show()
    QTimer.singleShot(1500, window._check_for_updates)

    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
    
 