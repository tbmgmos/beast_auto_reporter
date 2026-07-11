"""
Table Parser Module

Парсинг CSV/HTML файлов и заполнение Word таблиц параметрами.

Использование:
    parser = TableParser(config)
    parser.fill_table_from_csv('/path/to/export.csv', word_table, column_map)
"""

import csv
import os
import logging
from typing import Dict, List, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class TableParser:
    """
    Парсер параметров из CSV/HTML в Word таблицы
    
    Поддерживаемые источники:
    - CSV: Табличный файл
    - HTML: Визуальный отчет
    """
    
    # Стандартный маппинг колонок
    DEFAULT_COLUMN_MAP = {
        'file_name': 0,
        'duration_seconds': 1,
        'integrated_lufs': 2,
        'true_peak_dbtp': 3,
        'lra': 4,
        'overall_compliant': 5,
        'sample_rate': 6,
        'channels': 7,
        'analysis_method': 8,
        'timestamp': 9
    }
    
    def __init__(self, config: Dict = None):
        """
        Инициализация парсера
        
        Args:
            config: Словарь с настройками
        """
        self.config = config or {}
        self.export_config = self.config.get('export', {})
        
        logger.info("TableParser инициализирован")
    
    def fill_table_from_csv(
        self,
        csv_path: str,
        table,
        column_mapping: Dict[str, int] = None,
        start_row: int = 0,
        data_row_offset: int = 0
    ) -> int:
        """
        Заполнение Word таблицы данными из CSV
        
        Args:
            csv_path: Путь к CSV файлу
            table: Word таблица (docx Table object)
            column_mapping: Маппинг полей на колонки
            start_row: Начальная строка для заполнения
            data_row_offset: Смещение данных (если первые строки - заголовки)
            
        Returns:
            Количество заполненных строк
        """
        if column_mapping is None:
            column_mapping = self.DEFAULT_COLUMN_MAP
        
        try:
            with open(csv_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
            
            filled_rows = 0
            for i, row in enumerate(rows):
                if i < data_row_offset:
                    continue
                
                table_row_idx = start_row + (i - data_row_offset)
                if table_row_idx >= len(table.rows):
                    logger.warning(f"Недостаточно строк в таблице: {table_row_idx}")
                    break
                
                table_row = table.rows[table_row_idx]
                
                # Заполняем каждую колонку
                for field, col_idx in column_mapping.items():
                    if col_idx >= len(table_row.cells):
                        continue
                    
                    value = row.get(field, '')
                    table_row.cells[col_idx].text = str(value) if value else '—'
                
                # Цветовая индикация статуса
                compliant_col = column_mapping.get('overall_compliant')
                if compliant_col is not None and compliant_col < len(table_row.cells):
                    is_compliant = row.get('overall_compliant', '').lower() == 'true'
                    self._set_row_color(table_row, is_compliant)
                
                filled_rows += 1
            
            logger.info(f"Заполнено {filled_rows} строк из {csv_path}")
            return filled_rows
            
        except Exception as e:
            logger.error(f"Ошибка парсинга CSV: {e}")
            return 0
    
    def fill_table_from_html(
        self,
        html_path: str,
        table,
        column_mapping: Dict[str, int] = None,
        start_row: int = 0
    ) -> int:
        """
        Заполнение Word таблицы данными из HTML
        
        Args:
            html_path: Путь к HTML файлу
            table: Word таблица
            column_mapping: Маппинг полей на колонки
            start_row: Начальная строка
            
        Returns:
            Количество заполненных строк
        """
        import re
        
        if column_mapping is None:
            column_mapping = self.DEFAULT_COLUMN_MAP
        
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Парсинг таблицы из HTML
            # Ищем строки <tr>...</tr>
            tr_pattern = r'<tr[^>]*>(.*?)</tr>'
            td_pattern = r'<td[^>]*>(.*?)</td>'
            
            rows_data = []
            for tr_match in re.finditer(tr_pattern, html_content, re.DOTALL | re.IGNORECASE):
                tr_content = tr_match.group(1)
                if '<th' in tr_content:
                    continue  # Пропускаем заголовок
                
                cells = []
                for td_match in re.finditer(td_pattern, tr_content, re.DOTALL | re.IGNORECASE):
                    cell_text = re.sub(r'<[^>]+>', '', td_match.group(1)).strip()
                    cells.append(cell_text)
                
                if cells and not cells[0].startswith('Summary') and not cells[0].startswith('Total'):
                    rows_data.append(cells)
            
            filled_rows = 0
            for i, cells in enumerate(rows_data):
                table_row_idx = start_row + i
                if table_row_idx >= len(table.rows):
                    break
                
                table_row = table.rows[table_row_idx]
                
                # Заполняем колонки
                for field, col_idx in column_mapping.items():
                    if col_idx < len(cells) and col_idx < len(table_row.cells):
                        table_row.cells[col_idx].text = cells[col_idx]
                
                filled_rows += 1
            
            logger.info(f"Заполнено {filled_rows} строк из {html_path}")
            return filled_rows
            
        except Exception as e:
            logger.error(f"Ошибка парсинга HTML: {e}")
            return 0
    
    def fill_table_from_params(
        self,
        params_list: List[Dict],
        table,
        column_mapping: Dict[str, int] = None,
        start_row: int = 0
    ) -> int:
        """
        Заполнение Word таблицы напрямую из списка параметров
        
        Args:
            params_list: Список словарей с параметрами
            table: Word таблица
            column_mapping: Маппинг полей на колонки
            start_row: Начальная строка
            
        Returns:
            Количество заполненных строк
        """
        if column_mapping is None:
            column_mapping = self.DEFAULT_COLUMN_MAP
        
        filled_rows = 0
        
        for i, params in enumerate(params_list):
            table_row_idx = start_row + i
            if table_row_idx >= len(table.rows):
                break
            
            table_row = table.rows[table_row_idx]
            
            # Форматируем значения
            for field, col_idx in column_mapping.items():
                value = self._format_field(params, field)
                if col_idx < len(table_row.cells):
                    table_row.cells[col_idx].text = value
            
            # Цветовая индикация
            is_compliant = params.get('compliance', {}).get('overall_compliant', False)
            self._set_row_color(table_row, is_compliant)
            
            filled_rows += 1
        
        logger.info(f"Заполнено {filled_rows} строк из параметров")
        return filled_rows
    
    def _format_field(self, params: Dict, field: str) -> str:
        """Форматирование поля для отображения"""
        if params.get('error'):
            return f"ERROR: {params.get('error', 'Unknown')}"
        
        measurements = params.get('measurements', {})
        compliance = params.get('compliance', {})
        
        field_formatters = {
            'file_name': lambda: params.get('file_name', '—'),
            'file_path': lambda: params.get('file_path', '—'),
            'duration_seconds': lambda: f"{params.get('duration', 0):.1f}s",
            'sample_rate': lambda: str(params.get('sample_rate', '—')),
            'channels': lambda: self._get_channels(params.get('channel_layout', '')),
            'format': lambda: Path(params.get('file_path', '')).suffix[1:] or '—',
            'integrated_lufs': lambda: f"{measurements.get('lufs', '—'):.1f}" if measurements.get('lufs') else '—',
            'true_peak_dbtp': lambda: f"{measurements.get('true_peak', '—'):.1f}" if measurements.get('true_peak') else '—',
            'lra': lambda: f"{measurements.get('lra', '—'):.1f}" if measurements.get('lra') else '—',
            'short_term_loudness': lambda: f"{measurements.get('short_term_loudness', '—'):.1f}" if measurements.get('short_term_loudness') else '—',
            'momentary_loudness': lambda: f"{measurements.get('momentary_loudness', '—'):.1f}" if measurements.get('momentary_loudness') else '—',
            'overall_compliant': lambda: '✅' if compliance.get('overall_compliant') else '❌',
            'analysis_method': lambda: params.get('analysis_method', '—'),
            'timestamp': lambda: params.get('timestamp', '—')[:19].replace('T', ' ')
        }
        
        formatter = field_formatters.get(field)
        if formatter:
            return formatter()
        return str(params.get(field, ''))
    
    def _get_channels(self, channel_layout: str) -> str:
        """Форматирование количества каналов"""
        if not channel_layout:
            return '—'
        try:
            channels = int(channel_layout.split('.')[0])
            return str(channels)
        except (ValueError, IndexError):
            return '—'
    
    def _set_row_color(self, table_row, is_compliant: bool) -> None:
        """
        Установка цвета строки по статусу
        
        Args:
            table_row: Строка таблицы
            is_compliant: True - зеленый, False - красный
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        compliant_color = self.export_config.get('csv', {}).get('compliant_color', '#90EE90')
        non_compliant_color = self.export_config.get('csv', {}).get('non_compliant_color', '#FFB6C1')
        
        color = compliant_color if is_compliant else non_compliant_color
        
        try:
            for cell in table_row.cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), color)
                cell._element.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            logger.debug(f"Не удалось установить цвет: {e}")
    
    def auto_detect_columns(self, csv_path: str) -> Dict[str, int]:
        """
        Автоматическое определение колонок в CSV
        
        Args:
            csv_path: Путь к CSV файлу
            
        Returns:
            {'field_name': column_index}
        """
        try:
            with open(csv_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                headers = reader.fieldnames
            
            column_map = {}
            for i, header in enumerate(headers):
                # Маппинг по названию
                header_lower = header.lower()
                if 'file' in header_lower and 'name' in header_lower:
                    column_map['file_name'] = i
                elif 'duration' in header_lower:
                    column_map['duration_seconds'] = i
                elif 'lufs' in header_lower or 'loudness' in header_lower:
                    column_map['integrated_lufs'] = i
                elif 'peak' in header_lower and 'true' in header_lower:
                    column_map['true_peak_dbtp'] = i
                elif 'lra' in header_lower or 'range' in header_lower:
                    column_map['lra'] = i
                elif 'compliant' in header_lower or 'status' in header_lower:
                    column_map['overall_compliant'] = i
                elif 'sample' in header_lower and 'rate' in header_lower:
                    column_map['sample_rate'] = i
                elif 'channel' in header_lower:
                    column_map['channels'] = i
                elif 'method' in header_lower or 'analy' in header_lower:
                    column_map['analysis_method'] = i
                elif 'time' in header_lower or 'date' in header_lower:
                    column_map['timestamp'] = i
            
            logger.info(f"Автоопределено {len(column_map)} колонок")
            return column_map
            
        except Exception as e:
            logger.error(f"Ошибка автоопределения колонок: {e}")
            return {}
    
    def create_table_from_csv(
        self,
        csv_path: str,
        output_path: str,
        add_header: bool = True,
        styles: Dict = None
    ) -> bool:
        """
        Создание Word таблицы из CSV файла
        
        Args:
            csv_path: Путь к CSV
            output_path: Путь для сохранения .docx
            add_header: Добавить заголовок
            styles: Стили
            
        Returns:
            True если успешно
        """
        from docx import Document
        from docx.shared import Pt
        
        try:
            with open(csv_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
            
            doc = Document()
            
            # Заголовок документа
            if add_header:
                doc.add_heading('Audio Parameters Report', 0)
            
            # Чтение данных для таблицы
            if not rows:
                logger.warning("CSV файл пуст")
                return False
            
            # Создание таблицы
            num_cols = len(rows[0])
            num_rows = len(rows)
            
            # Стили
            table_style = styles.get('style', 'Table Grid') if styles else 'Table Grid'
            font_name = styles.get('font_name', 'Arial') if styles else 'Arial'
            font_size = styles.get('font_size', 11) if styles else 11
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = table_style
            
            # Заполнение
            headers = reader.fieldnames
            for i, row in enumerate(rows):
                table_row = table.rows[i]
                for j, header in enumerate(headers):
                    if j < len(table_row.cells):
                        table_row.cells[j].text = str(row.get(header, ''))
            
            # Сохранение
            doc.save(output_path)
            logger.info(f"Таблица создана: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Ошибка создания таблицы: {e}")
            return False


if __name__ == "__main__":
    # Тестирование
    logging.basicConfig(level=logging.INFO)
    
    parser = TableParser()
    
    # Тест автоопределения колонок
    sample_csv = """"file_name,duration_seconds,integrated_lufs,true_peak_dbtp,lra,overall_compliant
test.wav,100.5,-23.1,-1.5,12.5,True
"""
    
    with open('/tmp/test.csv', 'w') as f:
        f.write(sample_csv)
    
    columns = parser.auto_detect_columns('/tmp/test.csv')
    print(f"Автоопределенные колонки: {columns}")
