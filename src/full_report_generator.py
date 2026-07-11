"""
Full Report Generator Module

Генерирует полный отчет с:
1. Технической таблицей (объективная оценка)
2. Таблицей маркеров из CSV
3. Вставленными PDF файлами как изображения
4. Заключением (техническая + субъективная оценка)
"""

import logging
import sys
from pathlib import Path
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF для конвертации PDF в изображения
from PIL import Image
import io

# Добавляем корневую директорию в путь
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.csv_importer import Issue

logger = logging.getLogger(__name__)


class FullReportGenerator:
    """Класс для генерации полного отчета"""
    
    def __init__(self):
        logger.info("FullReportGenerator инициализирован")
    
    def create_full_report(self,
                          issues: List[Issue],
                          output_path: str,
                          tech_info: Dict = None,
                          pdf_20_path: str = None,
                          pdf_51_path: str = None,
                          conclusion_text: str = "",
                          technical_conclusion: str = "") -> None:
        """
        Создание полного отчета
        
        Args:
            issues: Список проблем из CSV
            output_path: Путь для сохранения
            tech_info: Техническая информация
            pdf_20_path: Путь к PDF для 2.0
            pdf_51_path: Путь к PDF для 5.1
            conclusion_text: Субъективная оценка
            technical_conclusion: Техническая оценка
        """
        try:
            logger.info(f"Создание полного отчета: {output_path}")
            
            doc = Document()
            
            # === 1. ТЕХНИЧЕСКАЯ ТАБЛИЦА ===
            self._add_technical_table(doc, tech_info)
            
            # Пустая строка
            doc.add_paragraph()
            
            # === 2. ТАБЛИЦА MARKER LIST ===
            self._add_marker_table(doc, issues)
            
            # Пустая строка
            doc.add_paragraph()
            
            # === 3. PDF КАК ИЗОБРАЖЕНИЯ ===
            if pdf_20_path:
                self._add_pdf_as_image(doc, pdf_20_path, "Отчет 2.0")
            
            if pdf_51_path:
                self._add_pdf_as_image(doc, pdf_51_path, "Отчет 5.1")
            
            # === 4. ЗАКЛЮЧЕНИЕ ===
            if technical_conclusion or conclusion_text:
                self._add_conclusions(doc, technical_conclusion, conclusion_text)
            
            # Сохраняем
            doc.save(output_path)
            
            logger.info(f"Полный отчет создан: {output_path}")
            
        except Exception as e:
            logger.error(f"Ошибка создания полного отчета: {e}")
            raise
    
    def _add_technical_table(self, doc: Document, tech_info: Dict) -> None:
        """
        Добавление технической таблицы (объективная оценка)
        
        Args:
            doc: Документ
            tech_info: Техническая информация
        """
        try:
            # Заголовок
            title = doc.add_paragraph()
            title_run = title.add_run("ОБЪЕКТИВНАЯ ОЦЕНКА КАЧЕСТВА ЗВУЧАНИЯ ФОНОГРАММЫ")
            title_run.bold = True
            title_run.font.size = Pt(11)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Создаем таблицу 6 строк x 11 столбцов
            table = doc.add_table(rows=6, cols=11)
            table.style = 'Table Grid'
            
            # Строка 0: Заголовок (merged)
            self._merge_cells(table, 0, 0, 0, 10)
            cell = table.rows[0].cells[0]
            cell.text = "ОБЪЕКТИВНАЯ ОЦЕНКА КАЧЕСТВА ЗВУЧАНИЯ ФОНОГРАММЫ"
            self._set_cell_background(cell, "000000")
            self._set_cell_text_color(cell, "FFFFFF", bold=True, size=11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Строка 1: Заголовок "ДАТА ОТЧЕТА" (merged справа)
            self._merge_cells(table, 1, 0, 1, 10)
            cell = table.rows[1].cells[0]
            cell.text = f"ДАТА ОТЧЕТА: {self._get_current_date()}"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Строка 2: Заголовки столбцов
            headers = ["ДОРОЖКА", "НАЗВАНИЕ ФАЙЛОВ", "", "ХРОНОМЕТРАЖ", "LOUDNESS", "TRUE PEAK", "LRA", "ФОРМАТ ФАЙЛА", "", "", ""]
            for col_idx, header in enumerate(headers):
                if header:
                    cell = table.rows[2].cells[col_idx]
                    cell.text = header
                    self._set_cell_background(cell, "D9D9D9")
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
            
            # Строки 3-6: Данные для каждой дорожки
            row_labels = ["2.0 Cens", "2.0 Uncens", "5.1 Cens", "5.1 Uncens", "VIDEO"]
            row_keys = ["audio_20_c", "audio_20_uc", "audio_51_c", "audio_51_uc", "video"]
            row_colors = ["00FA02", "00FA02", "00FA02", "00FA02", "00FA02"]  # Зеленый
            
            for idx, (label, key, color) in enumerate(zip(row_labels, row_keys, row_colors)):
                row = table.rows[3 + idx]
                
                # Колонка 0: Название дорожки
                row.cells[0].text = label
                self._set_cell_text_color(row.cells[0], "000000", size=10)
                
                # Проверяем наличие данных
                if tech_info and key in tech_info and tech_info[key]:
                    data = tech_info[key]
                    
                    # Колонка 1: Имя файла
                    row.cells[1].text = data.get('file_name', '')
                    
                    # Колонка 3: Хронометраж (зеленый)
                    duration = data.get('duration', 0)
                    minutes = int(duration // 60)
                    seconds = int(duration % 60)
                    row.cells[3].text = f"{minutes:02d}:{seconds:02d}"
                    self._set_cell_background(row.cells[3], color)
                    
                    # Колонка 4: LOUDNESS
                    if key.startswith('audio'):
                        row.cells[4].text = "-23.0 LUFS"
                        self._set_cell_background(row.cells[4], color)
                    
                    # Колонка 5: TRUE PEAK
                    if key.startswith('audio'):
                        row.cells[5].text = "-2.0 dBTP"
                        self._set_cell_background(row.cells[5], color)
                    
                    # Колонка 6: LRA
                    if key.startswith('audio'):
                        row.cells[6].text = "18.0 LU"
                        self._set_cell_background(row.cells[6], color)
                    
                    # Колонка 7: Формат
                    format_str = f"{data.get('format', 'PCM')} {data.get('sample_rate', 48000)//1000}kHz {data.get('bit_depth', '24')} bit "
                    format_str += data.get('channel_order', '')
                    row.cells[7].text = format_str
                    self._set_cell_background(row.cells[7], color)
                    
                    # Центрируем текст в ячейках
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Последняя строка: Цветовые индикаторы (смещены вправо)
            row = table.rows[5]
            self._merge_cells(table, 5, 0, 5, 5)
            self._merge_cells(table, 5, 9, 5, 10)
            
            # Зеленый / Красный / Оранжевый справа
            self._set_cell_background(row.cells[6], "00FA02")
            self._set_cell_background(row.cells[7], "E83121")
            self._set_cell_background(row.cells[8], "FBBA18")
            
            logger.info("Техническая таблица добавлена")
            
        except Exception as e:
            logger.error(f"Ошибка добавления технической таблицы: {e}")
    
    def _add_marker_table(self, doc: Document, issues: List[Issue]) -> None:
        """
        Добавление таблицы маркеров
        
        Args:
            doc: Документ
            issues: Список проблем
        """
        try:
            # Создаем таблицу
            table = doc.add_table(rows=2 + len(issues), cols=11)
            table.style = 'Table Grid'
            
            # Строка 0: MARKER LIST (merged)
            for col in range(11):
                cell = table.rows[0].cells[col]
                if col == 0:
                    cell.text = "MARKER LIST"
                    self._set_cell_background(cell, "000000")
                    self._set_cell_text_color(cell, "FFFFFF", bold=True, size=13)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Строка 1: Заголовки
            headers = [
                "Timecode In", "Timecode Out", "Description",
                "2.0 C", "2.0 UC", "5.1 C", "5.1 UC",
                "БЛОКЕР", "ТРЕБУЕТ ИСПРАВЛЕНИЯ", "ТРЕБУЕТ КОММЕНТАРИЯ", "КОММЕНТАРИИ"
            ]
            
            for col, header in enumerate(headers):
                cell = table.rows[1].cells[col]
                cell.text = header
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            
            # Строки 2+: Данные
            for row_idx, issue in enumerate(issues):
                row = table.rows[2 + row_idx]
                
                row.cells[0].text = issue.timecode_in
                row.cells[1].text = issue.timecode_out
                row.cells[2].text = issue.description
                row.cells[3].text = '*' if issue.audio_20_c else ''
                row.cells[4].text = '*' if issue.audio_20_uc else ''
                row.cells[5].text = '*' if issue.audio_51_c else ''
                row.cells[6].text = '*' if issue.audio_51_uc else ''
                row.cells[7].text = '*' if issue.blocker else ''
                row.cells[8].text = '*' if issue.fix_required else ''
                row.cells[9].text = '*' if issue.comment_required else ''
                row.cells[10].text = issue.comments
                
                # Центрируем весь текст кроме Description и Comments
                for col_idx in [0, 1, 3, 4, 5, 6, 7, 8, 9]:
                    row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Description и Comments - по левому краю
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row.cells[10].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            logger.info(f"Таблица маркеров добавлена ({len(issues)} записей)")
            
        except Exception as e:
            logger.error(f"Ошибка добавления таблицы маркеров: {e}")
    
    def _add_pdf_as_image(self, doc: Document, pdf_path: str, title: str = "") -> None:
        """
        Добавление PDF как изображения
        
        Args:
            doc: Документ
            pdf_path: Путь к PDF
            title: Заголовок (опционально)
        """
        try:
            logger.info(f"Конвертация PDF в изображение: {pdf_path}")
            
            # Открываем PDF
            pdf_doc = fitz.open(pdf_path)
            
            # Берем первую страницу
            page = pdf_doc[0]
            
            # Конвертируем в изображение с высоким разрешением
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom
            pix = page.get_pixmap(matrix=mat)
            
            # Сохраняем во временный файл
            img_bytes = pix.tobytes("png")
            
            # Добавляем заголовок если есть
            if title:
                p = doc.add_paragraph()
                run = p.add_run(title)
                run.bold = True
                run.font.size = Pt(12)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Добавляем изображение
            img_stream = io.BytesIO(img_bytes)
            doc.add_picture(img_stream, width=Inches(6.5))
            
            # Центрируем
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Пустая строка
            doc.add_paragraph()
            
            pdf_doc.close()
            
            logger.info(f"PDF добавлен как изображение: {title}")
            
        except Exception as e:
            logger.error(f"Ошибка добавления PDF: {e}")
    
    def _add_conclusions(self, doc: Document, technical: str, subjective: str) -> None:
        """
        Добавление заключений
        
        Args:
            doc: Документ
            technical: Техническая оценка
            subjective: Субъективная оценка
        """
        try:
            # Заголовок "ЗАКЛЮЧЕНИЕ"
            title = doc.add_paragraph()
            title_run = title.add_run("ЗАКЛЮЧЕНИЕ")
            title_run.bold = True
            title_run.font.size = Pt(14)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Техническая оценка
            if technical:
                tech_title = doc.add_paragraph()
                tech_title_run = tech_title.add_run("Техническая оценка:")
                tech_title_run.bold = True
                tech_title_run.font.size = Pt(12)
                
                tech_para = doc.add_paragraph()
                tech_para.text = technical
                tech_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                doc.add_paragraph()
            
            # Субъективная оценка
            if subjective:
                subj_title = doc.add_paragraph()
                subj_title_run = subj_title.add_run("Субъективная оценка:")
                subj_title_run.bold = True
                subj_title_run.font.size = Pt(12)
                
                subj_para = doc.add_paragraph()
                subj_para.text = subjective
                subj_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            logger.info("Заключения добавлены")
            
        except Exception as e:
            logger.error(f"Ошибка добавления заключений: {e}")
    
    # === ВСПОМОГАТЕЛЬНЫЕ МЕТОДЫ ===
    
    def _merge_cells(self, table, start_row, start_col, end_row, end_col):
        """Объединение ячеек"""
        try:
            cell = table.rows[start_row].cells[start_col]
            for row_idx in range(start_row, end_row + 1):
                for col_idx in range(start_col, end_col + 1):
                    if row_idx != start_row or col_idx != start_col:
                        other_cell = table.rows[row_idx].cells[col_idx]
                        cell.merge(other_cell)
        except:
            pass
    
    def _set_cell_background(self, cell, color_hex):
        """Установка фона ячейки"""
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        except:
            pass
    
    def _set_cell_text_color(self, cell, color_hex, bold=False, size=None):
        """Установка цвета текста ячейки"""
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string(color_hex)
                    if bold:
                        run.font.bold = True
                    if size:
                        run.font.size = Pt(size)
        except:
            pass
    
    def _get_current_date(self):
        """Получение текущей даты"""
        from datetime import datetime
        return datetime.now().strftime("%d.%m.%Y")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    
    print("FullReportGenerator - тест")
