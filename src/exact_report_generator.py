"""
Exact Report Generator

Точная копия референсного документа + техническая таблица + заключение
"""

import logging
import sys
from pathlib import Path
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
import fitz
import io

sys.path.insert(0, str(Path(__file__).parent.parent))

from src.csv_importer import Issue
from src.technical_info_extractor import format_fps

logger = logging.getLogger(__name__)


class ExactReportGenerator:
    """Точная копия референса"""
    
    def __init__(self):
        logger.info("ExactReportGenerator инициализирован")

    @staticmethod
    def _fmt_metric(value, decimals: int = 2) -> str:
        """Формат метрик с заданной точностью."""
        try:
            return f"{float(value):.{int(decimals)}f}"
        except (TypeError, ValueError):
            return str(value)

    @staticmethod
    def _peak_decimals(pdf_data: dict, default_decimals: int = 1) -> int:
        """Количество знаков для True Peak: 2 если точное измерение, иначе default."""
        if isinstance(pdf_data, dict) and pdf_data.get('true_peak_source') == 'precise_4x_oversampling':
            return 2
        return default_decimals

    @staticmethod
    def _collect_pdf_paths(pdf_paths: List[str] = None,
                           pdf_20_path: str = None,
                           pdf_51_path: str = None) -> List[str]:
        """Собирает итоговый список PDF для вставки в отчет без дубликатов."""
        ordered_paths = []
        seen = set()

        for pdf_path in pdf_paths or []:
            if not pdf_path:
                continue
            normalized = str(pdf_path)
            if normalized in seen:
                continue
            seen.add(normalized)
            ordered_paths.append(normalized)

        for pdf_path in (pdf_20_path, pdf_51_path):
            if not pdf_path:
                continue
            normalized = str(pdf_path)
            if normalized in seen:
                continue
            seen.add(normalized)
            ordered_paths.append(normalized)

        return ordered_paths
    
    def create_exact_report(self,
                           issues: List[Issue],
                           output_path: str,
                           tech_info: Dict = None,
                           pdf_20_path: str = None,
                           pdf_51_path: str = None,
                           pdf_paths: List[str] = None,
                           conclusion_technical: str = "",
                           conclusion_subjective: str = "",
                           report_type: str = "main",
                           prepared_by: str = "",
                           ) -> None:
        """Создание точного отчета

        Args:
            report_type: Тип отчета - "main" (основной) или "me" (M&E)
        """
        try:
            logger.info(f"Создание отчета ({report_type}): {output_path}")
            
            doc = Document()
            
            # Настройки страницы как в РЫБА ОСНОВНОЙ НОВАЯ 2025.docx
            section = doc.sections[0]
            section.page_width = Cm(42.02)   # A3 ширина
            section.page_height = Cm(29.70)  # A3 высота
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.orientation = 1  # 1 = Landscape (альбомная)
            
            # === СТРАНИЦА 1: ТЕХНИЧЕСКАЯ ТАБЛИЦА ===
            # (остается на первой странице)
            
            # === 1. ТЕХНИЧЕСКАЯ ТАБЛИЦА с заключением ===
            if tech_info:
                if report_type == "me" or report_type == "me_ours":
                    # Используем M&E таблицу для обычных M&E и наших работ
                    self._add_me_technical_table(
                        doc, tech_info, conclusion_technical, conclusion_subjective, prepared_by
                    )
                else:
                    self._add_technical_table_with_conclusion(
                        doc,
                        tech_info,
                        conclusion_technical,
                        conclusion_subjective,
                        prepared_by,
                        report_type=report_type,
                    )
            
            # === PAGE BREAK (переход на страницу 2) ===
            para_break = doc.add_paragraph()
            run = para_break.add_run()
            run.add_break(WD_BREAK.PAGE)  # Page break
            
            # === СТРАНИЦА 2: MARKER LIST ===
            # Маркер-лист добавляется ВСЕГДА (даже если пустой - с заголовком и шапкой)
            if issues and len(issues) > 0:
                logger.info(f"Добавление маркер-листа с {len(issues)} проблемами")
            else:
                logger.info("Добавление пустого маркер-листа (только заголовок и шапка)")
            
            # === 2. MARKER LIST (точно как в референсе) ===
            self._add_marker_list_exact(doc, issues, report_type=report_type)

            # === 3. PDF КАК ИЗОБРАЖЕНИЯ (всегда, если файлы существуют) ===
            # PDF добавляются на отдельной странице (страница 3+)
            report_pdf_paths = self._collect_pdf_paths(
                pdf_paths=pdf_paths,
                pdf_20_path=pdf_20_path,
                pdf_51_path=pdf_51_path,
            )

            if report_pdf_paths:
                # Page break перед PDF (переход на страницу 3)
                para_break_pdf = doc.add_paragraph()
                run_pdf = para_break_pdf.add_run()
                run_pdf.add_break(WD_BREAK.PAGE)
                logger.info("Page break перед PDF добавлен")

            for idx, pdf_path in enumerate(report_pdf_paths):
                if idx > 0:
                    doc.add_paragraph()  # Отступ между PDF
                self._add_pdf_image(doc, pdf_path)
                logger.info(f"PDF добавлен в приложение отчета: {pdf_path}")
            
            doc.save(output_path)
            logger.info(f"Отчет создан: {output_path}")
            
        except Exception as e:
            logger.error(f"Ошибка создания отчета: {e}")
            raise
    
    def _add_technical_table_with_conclusion(self, doc: Document, tech_info: Dict,
                                             conclusion_tech: str, conclusion_subj: str,
                                             prepared_by: str = "", report_type: str = "standard") -> None:
        """Техническая таблица с заключением (точно по скриншоту)"""
        try:
            # Импорты для работы с форматированием
            from docx.shared import Pt, Cm, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
            from datetime import datetime

            OK_BG = "72F649"
            BAD_BG = "E73322"
            WARN_BG = "FCBA19"
            HEADER_BG = "D5D5D5"
            LIGHT_BG = "F5F5F5"
            OFFWHITE_BG = "FEFFFE"
            is_dcp_report = report_type == "dcp"
            peak_header = "SAMPLE PEAK" if is_dcp_report else "TRUE PEAK"
            peak_value_key = "sample_peak" if is_dcp_report else "true_peak"
            peak_unit = "dBFS" if is_dcp_report else "dBTP"
            peak_label_for_logs = "SAMPLE PEAK" if is_dcp_report else "TRUE PEAK"
            peak_decimals = 2 if is_dcp_report else 1

            if is_dcp_report:
                rows_data = [
                    ("51 DCP", "audio_51_c", "pdf_51_c"),
                    ("Video ref", "video", None),
                    ("Left", "audio_dcp_split_l", None),
                    ("Right", "audio_dcp_split_r", None),
                    ("Center", "audio_dcp_split_c", None),
                    ("LFE", "audio_dcp_split_lfe", None),
                    ("Left surround", "audio_dcp_split_ls", None),
                    ("Right surround", "audio_dcp_split_rs", None),
                ]
            else:
                rows_data = [
                    ("2.0 cens", "audio_20_c", "pdf_20_c"),
                    ("2.0 uncens", "audio_20_uc", "pdf_20_uc"),
                    ("5.1 cens", "audio_51_c", "pdf_51_c"),
                    ("5.1 uncens", "audio_51_uc", "pdf_51_uc"),
                    ("VIDEO", "video", None),
                ]
            
            # Таблица: 13 строк x 7 столбцов (как в референсном основном отчете)
            # Строки: 0=заголовок+дата, 1=отчет подготовил, 2=колонки,
            #         3-7=данные, 8=пустая, 9=индикаторы, 10=пустая, 11=ЗАКЛЮЧЕНИЕ, 12=текст
            table = doc.add_table(rows=len(rows_data) + 8, cols=7)
            table.style = 'Table Grid'
            
            # Устанавливаем ширину колонок (точно по референсу)
            # ДОРОЖКА | НАЗВАНИЕ ФАЙЛОВ | ХРОНОМЕТРАЖ | LOUDNESS | TRUE PEAK | LRA | ФОРМАТ ФАЙЛА
            widths_cm = [2.61, 14.27, 4.22, 2.63, 2.47, 1.89, 9.89]  # Итого: ~38.0 см
            
            # Метод 1: Устанавливаем через table.columns
            for idx, width_cm in enumerate(widths_cm):
                if idx < len(table.columns):
                    for cell in table.columns[idx].cells:
                        cell.width = Cm(width_cm)
            
            # Метод 2: Устанавливаем через XML напрямую
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Удаляем autofit
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
            
            # Устанавливаем ширину через tblGrid
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is not None:
                for idx, gridCol in enumerate(tblGrid.findall(qn('w:gridCol'))):
                    if idx < len(widths_cm):
                        gridCol.set(qn('w:w'), str(int(widths_cm[idx] * 567)))  # cm to twips
            
            # Строка 0: Шапка (как в M&E референсе)
            row0 = table.rows[0]
            
            # Левая часть: заголовок (объединяем первые 4 колонки)
            cell_left = row0.cells[0]
            for col in range(1, 4):
                cell_left.merge(row0.cells[col])
            cell_left.text = "ОБЪЕКТИВНАЯ ОЦЕНКА КАЧЕСТВА ЗВУЧАНИЯ ФОНОГРАММЫ"
            self._format_cell(cell_left, bg=HEADER_BG, bold=True, align="center", font_size=11, vertical_align="center")
            
            # Правая часть строки 0: дата
            cell_date_label = row0.cells[4]
            cell_date_label.merge(row0.cells[5])
            cell_date_label.text = "ДАТА ОТЧЕТА:"
            self._format_cell(cell_date_label, bold=True, align="left", font_size=10, vertical_align="center")

            cell_date_value = row0.cells[6]
            cell_date_value.text = datetime.now().strftime('%d.%m.%Y')
            self._format_cell(cell_date_value, bold=True, align="right", font_size=10, vertical_align="center")

            # Строка 1: "ОТЧЁТ ПОДГОТОВИЛ"
            row1 = table.rows[1]
            cell_blank = row1.cells[0]
            for col in range(1, 4):
                cell_blank.merge(row1.cells[col])
            self._format_cell(cell_blank, bg=OFFWHITE_BG, align="left", vertical_align="center")

            cell_prepared_label = row1.cells[4]
            cell_prepared_label.merge(row1.cells[5])
            cell_prepared_label.text = "ОТЧЁТ ПОДГОТОВИЛ:"
            self._format_cell(cell_prepared_label, bg=LIGHT_BG, bold=True, align="left", font_size=10, vertical_align="center")

            cell_prepared_value = row1.cells[6]
            cell_prepared_value.text = prepared_by or "Не указано"
            self._format_cell(cell_prepared_value, bg=LIGHT_BG, bold=True, align="right", font_size=10, vertical_align="center")
            
            # Строка 2: Заголовки столбцов
            headers = ["ДОРОЖКА", "НАЗВАНИЕ ФАЙЛОВ", "ХРОНОМЕТРАЖ", 
                      "LOUDNESS", peak_header, "LRA", "ФОРМАТ ФАЙЛА"]
            for col, header in enumerate(headers):
                cell = table.rows[2].cells[col]
                cell.text = header
                self._format_cell(cell, bg=HEADER_BG, bold=True, align="left", font_size=9, vertical_align="center")
            
            # Fallback маппинг: если точный ключ не найден, пробуем альтернативы
            pdf_key_fallbacks = {
                'pdf_20_c': ['pdf_20_c', 'pdf_20'],
                'pdf_20_uc': ['pdf_20_uc'],
                'pdf_51_c': ['pdf_51_c', 'pdf_51'],
                'pdf_51_uc': ['pdf_51_uc'],
                'pdf_20': ['pdf_20'],
                'pdf_51': ['pdf_51']
            }
            
            # Трекинг использованных источников для предотвращения дубликатов
            _used_pdf_sources = set()
            _used_audio_files = set()

            logger.info(f"rows_data: {rows_data}")
            logger.info(f"pdf_key_fallbacks: {pdf_key_fallbacks}")
            logger.info(f"Total table rows: {len(table.rows)}")
            logger.info(f"Row indices for data: {[3 + i for i in range(len(rows_data))]}")
            
            # Номинальные значения (берем из tech_info['params'] если есть)
            params = tech_info.get('params', {}) if tech_info else {}
            target_lufs = params.get('target_lufs', -23.0)
            lufs_tolerance = params.get('lufs_tolerance', 0.5)
            target_peak = params.get('true_peak', -2.0)
            target_lra = params.get('lra_max', 18.0)
            
            # Debug: Log what PDF data is available
            logger.info("=== PDF DATA AVAILABILITY CHECK ===")
            logger.info(f"tech_info type: {type(tech_info)}")
            logger.info(f"tech_info keys: {list(tech_info.keys()) if tech_info else 'None'}")
            
            # Логируем ВСЕ ключи начинающиеся с pdf_
            logger.info("=== ALL pdf_ keys in tech_info ===")
            for k, v in tech_info.items():
                if isinstance(k, str) and k.startswith('pdf_'):
                    logger.info(f"  {k}: {v}")
            logger.info("=== END pdf_ keys ===")
            
            # Check each pdf key explicitly
            for label, key, pdf_key in rows_data:
                has_audio = tech_info and key in tech_info and tech_info.get(key) is not None
                has_pdf = pdf_key and tech_info and pdf_key in tech_info
                pdf_data = tech_info.get(pdf_key) if has_pdf else None
                
                logger.info(f"  {label}:")
                logger.info(f"    - audio_key={key}, has_audio={has_audio}")
                logger.info(f"    - pdf_key={pdf_key}, has_pdf={has_pdf}")
                
                if has_pdf and pdf_data:
                    logger.info(f"    - PDF DATA FOUND: {pdf_data}")
                    logger.info(
                        f"    - lufs={pdf_data.get('lufs')}, "
                        f"peak={pdf_data.get(peak_value_key)}, lra={pdf_data.get('lra')}"
                    )
                elif has_pdf:
                    logger.warning(f"    - WARNING: pdf_key '{pdf_key}' exists but pdf_data is None or empty!")
                else:
                    logger.info(f"    - NO PDF DATA")
            logger.info("=== END PDF DATA CHECK ===")
            
            # Определяем FPS из видео или параметров
            fps = 25  # По умолчанию 25 fps
            if 'video' in tech_info and tech_info['video']:
                video_data = tech_info['video']
                fps = video_data.get('fps', 25)
            elif params and 'fps' in params:
                fps = params.get('fps', 25)
            
            logger.info(f"Используется FPS: {fps} для проверки кратности кадру")

            # Функция сравнения хронометража по отображаемым миллисекундам
            def durations_match(dur1, dur2):
                if dur1 is None or dur2 is None or dur1 == 0 or dur2 == 0:
                    return False
                # round() вместо int() — int() truncates, что даёт -1мс при float вроде 1376.27999...
                return int(dur1 * 1000) == int(dur2 * 1000)

            # Функция проверки кратности кадру (с допуском 0.5 мс)
            def is_frame_aligned(duration_seconds, fps=25):
                """Проверяет, кратна ли длительность одному кадру"""
                if duration_seconds <= 0:
                    return True
                ms = duration_seconds * 1000
                frame_ms = 1000.0 / fps  # 40.0 для 25fps, ~41.67 для 24fps
                nearest_frame = round(ms / frame_ms)
                return abs(ms - nearest_frame * frame_ms) < 0.5

            # Собираем длительности всех файлов для сравнения
            durations = {}
            for label, key, _ in rows_data:
                if tech_info and key in tech_info and tech_info[key]:
                    durations[key] = tech_info[key].get('duration', 0)

            # Проверяем, совпадают ли все аудиофайлы между собой
            audio_durations_ms = []
            for _label, key, _pdf_key in rows_data:
                if not key.startswith('audio'):
                    continue
                if key in durations and durations[key] > 0:
                    audio_durations_ms.append(int(durations[key] * 1000))
            all_audio_match = len(audio_durations_ms) <= 1 or len(set(audio_durations_ms)) == 1
            audio_durations_ms_set = set(audio_durations_ms)
            video_duration_ms = None
            if durations.get('video', 0) > 0:
                video_duration_ms = int(durations['video'] * 1000)

            for idx, (label, key, pdf_key) in enumerate(rows_data):
                logger.info(f"\n=== ОБРАБОТКА СТРОКИ {idx}: {label} ===")
                logger.info(f"  key={key}, pdf_key={pdf_key}")
                
                row = table.rows[3 + idx]
                logger.info(f"  Row index in table: {3 + idx}")
                
                # Установить label в первую ячейку (серый фон как у заголовков)
                row.cells[0].text = label
                logger.info(f"  ✓ Set cell[0] = '{label}'")
                self._format_cell(row.cells[0], bg=HEADER_BG, bold=True, align="left")
                
                # Проверяем наличие аудио файла
                has_audio = tech_info and key in tech_info and tech_info.get(key) is not None
                # Проверяем наличие PDF файла
                has_pdf = pdf_key and pdf_key in tech_info and tech_info.get(pdf_key) is not None
                
                logger.info(f"  has_audio={has_audio}, has_pdf={has_pdf}")

                # Важно: поддержка отчета без аудиофайлов (только PDF + CSV).
                # Раньше эта ветка была вложена в has_audio и не исполнялась.
                if (not has_audio) and pdf_key:
                    pdf_data = None
                    used_pdf_key = None
                    fallback_keys = pdf_key_fallbacks.get(pdf_key, [pdf_key])
                    for fallback_key in fallback_keys:
                        if fallback_key in tech_info and tech_info.get(fallback_key) is not None:
                            pdf_data = tech_info[fallback_key]
                            used_pdf_key = fallback_key
                            logger.info(
                                f"  ✓ PDF-ONLY: Found data using key '{fallback_key}' (wanted: '{pdf_key}')"
                            )
                            break

                    # Дедупликация: пропускаем строку если этот PDF-источник уже использован
                    if pdf_data:
                        source_pdf = str(pdf_data.get('source_pdf') or '')
                        if source_pdf and source_pdf in _used_pdf_sources:
                            logger.info(f"  ⏭ Пропуск дубликата PDF: {source_pdf} (уже использован)")
                            continue
                        if source_pdf:
                            _used_pdf_sources.add(source_pdf)

                    if pdf_data:
                        source_name = source_pdf.rsplit('.', 1)[0] if source_pdf and '.' in source_pdf else source_pdf
                        if source_name:
                            row.cells[1].text = source_name
                            self._format_cell(row.cells[1], align="left")

                        if pdf_data.get('lufs') is not None:
                            value = pdf_data['lufs']
                            row.cells[3].text = f"{self._fmt_metric(value, decimals=1)} LUFS"
                            if is_dcp_report:
                                lufs_bg = OK_BG if abs(value - target_lufs) <= lufs_tolerance else WARN_BG
                            else:
                                lufs_bg = OK_BG if abs(value - target_lufs) <= lufs_tolerance else BAD_BG
                            self._format_cell(row.cells[3], bg=lufs_bg, align="center")

                        peak_value = pdf_data.get(peak_value_key)
                        if peak_value is None and is_dcp_report:
                            peak_value = pdf_data.get('true_peak')
                        if peak_value is not None:
                            value = peak_value
                            sign = "+" if value > 0 else ""
                            dp = self._peak_decimals(pdf_data, peak_decimals)
                            row.cells[4].text = f"{sign}{self._fmt_metric(value, decimals=dp)} {peak_unit}"
                            if is_dcp_report:
                                peak_bg = OK_BG if value <= 0 else BAD_BG
                            else:
                                peak_bg = OK_BG if value <= target_peak else BAD_BG
                            self._format_cell(row.cells[4], bg=peak_bg, align="center")

                        if pdf_data.get('lra') is not None:
                            value = pdf_data['lra']
                            row.cells[5].text = f"{self._fmt_metric(value, decimals=1)} LU"
                            if is_dcp_report:
                                lra_bg = OK_BG if value <= target_lra else WARN_BG
                            else:
                                lra_bg = OK_BG if value <= target_lra else BAD_BG
                            self._format_cell(row.cells[5], bg=lra_bg, align="center")

                        logger.info(
                            f"  PDF-ONLY values set for {label} "
                            f"(key={used_pdf_key or pdf_key}): "
                            f"LUFS={pdf_data.get('lufs')}, "
                            f"{peak_label_for_logs}={pdf_data.get(peak_value_key)}, "
                            f"LRA={pdf_data.get('lra')}"
                        )
                    else:
                        logger.warning(f"  ⚠️ PDF-ONLY: No data found for {pdf_key} (tried: {fallback_keys})")
                    continue
                
                # ДИАГНОСТИКА: показываем что есть в tech_info для этого pdf_key
                if pdf_key:
                    if pdf_key in tech_info:
                        pdf_data = tech_info[pdf_key]
                        logger.info(f"  ✓ pdf_key '{pdf_key}' FOUND in tech_info")
                        logger.info(f"     pdf_data type: {type(pdf_data)}")
                        if isinstance(pdf_data, dict):
                            logger.info(
                                f"     lufs={pdf_data.get('lufs')}, "
                                f"peak={pdf_data.get(peak_value_key)}, lra={pdf_data.get('lra')}"
                            )
                        else:
                            logger.info(f"     pdf_data value: {pdf_data}")
                    else:
                        logger.warning(f"  ⚠️ pdf_key '{pdf_key}' NOT FOUND in tech_info!")
                        # Показываем какие ключи начинаются с pdf_
                        pdf_keys = [k for k in tech_info.keys() if isinstance(k, str) and k.startswith('pdf_')]
                        logger.warning(f"     Available pdf_ keys: {pdf_keys}")
                
                if has_audio:
                    data = tech_info[key]
                    file_name = data.get('file_name', '')

                    # Дедупликация: пропускаем если этот аудиофайл уже использован
                    if file_name and file_name in _used_audio_files:
                        logger.info(f"  ⏭ Пропуск дубликата аудио: {file_name} (уже использован)")
                        continue
                    if file_name:
                        _used_audio_files.add(file_name)

                    logger.info(f"  Audio data found: {file_name}")

                    # Имя файла с проверкой на некорректный тип (убираем расширение)
                    file_name_without_ext = file_name.rsplit('.', 1)[0] if '.' in file_name else file_name
                    row.cells[1].text = file_name_without_ext
                    logger.info(f"  ✓ Set cell[1] = '{file_name_without_ext}'")
                    
                    # Проверяем на некорректный тип фонограммы в названии
                    has_incorrect_type = self._check_incorrect_audio_type(file_name)
                    name_mismatch = bool(data.get('name_channel_mismatch')) or bool(data.get('name_incorrect_type'))
                    name_error = has_incorrect_type or name_mismatch
                    
                    if name_error:
                        self._format_cell(row.cells[1], align="left", bg=BAD_BG, text_color="FFFFFF")
                        logger.warning(f"⚠️  Название файла с ошибкой: {file_name}")
                    else:
                        self._format_cell(row.cells[1], align="left")
                    
                    # Хронометраж (формат: H:MM:SS.mmm) с цветовой индикацией
                    duration = data.get('duration', 0)  # в секундах
                    total_ms = int(duration * 1000)
                    hours = total_ms // 3600000
                    mins = (total_ms % 3600000) // 60000
                    secs = (total_ms % 60000) // 1000
                    millis = total_ms % 1000
                    row.cells[2].text = f"{hours}:{mins:02d}:{secs:02d}.{millis:03d}"
                    
                    # Цветовая индикация с проверкой кратности кадру
                    if duration > 0:
                        if key.startswith('audio'):
                            # Если есть видео, оцениваем дорожку относительно него.
                            # Так синхронная с видео дорожка остается зеленой, даже если
                            # другая аудиодорожка имеет иной хронометраж.
                            if video_duration_ms is not None:
                                chrono_bg = OK_BG if total_ms == video_duration_ms else BAD_BG
                            elif not all_audio_match:
                                chrono_bg = BAD_BG  # Красный - аудиофайлы не совпадают
                            else:
                                chrono_bg = OK_BG  # Зеленый - аудиофайлы совпадают
                        elif key == 'video':
                            # Для видео: сравниваем с аудио
                            if audio_durations_ms:
                                video_ms = int(duration * 1000)
                                if video_ms in audio_durations_ms_set:
                                    chrono_bg = OK_BG  # Зеленый - совпадает с аудио
                                else:
                                    chrono_bg = BAD_BG  # Красный - не совпадает с аудио
                            else:
                                frame_ok = is_frame_aligned(duration, fps)
                                chrono_bg = OK_BG if frame_ok else WARN_BG
                        else:
                            frame_ok = is_frame_aligned(duration, fps)
                            chrono_bg = OK_BG if frame_ok else WARN_BG
                    else:
                        chrono_bg = "FFFFFF"  # Белый - нет данных
                    
                    self._format_cell(row.cells[2], bg=chrono_bg, align="center")
                    
                    # Технические параметры (для аудио с PDF или только PDF)
                    # Используем fallback ключи если точный ключ не найден
                    pdf_data = None
                    used_pdf_key = None
                    if pdf_key:
                        fallback_keys = pdf_key_fallbacks.get(pdf_key, [pdf_key])
                        for fallback_key in fallback_keys:
                            if fallback_key in tech_info and tech_info.get(fallback_key) is not None:
                                pdf_data = tech_info[fallback_key]
                                used_pdf_key = fallback_key
                                logger.info(f"  ✓ Found PDF data using key '{fallback_key}' (wanted: '{pdf_key}')")
                                break
                        if pdf_data is None:
                            logger.warning(f"  ⚠️ PDF data not found for {pdf_key} (tried: {fallback_keys})")

                    # DCP Sample Peak измеряется приложением непосредственно по
                    # аудио, поэтому он доступен и без внешнего PDF-отчёта.
                    direct_sample_peak = data.get('sample_peak') if is_dcp_report else None
                    if direct_sample_peak is not None:
                        sign = "+" if direct_sample_peak > 0 else ""
                        row.cells[4].text = (
                            f"{sign}{self._fmt_metric(direct_sample_peak, decimals=peak_decimals)} {peak_unit}"
                        )
                        direct_peak_bg = OK_BG if direct_sample_peak <= 0 else BAD_BG
                        self._format_cell(row.cells[4], bg=direct_peak_bg, align="center")
                    
                    if pdf_data is not None:
                        # Трекинг PDF-источника для предотвращения дублей
                        _src = str(pdf_data.get('source_pdf') or '')
                        if _src:
                            _used_pdf_sources.add(_src)

                        # LUFS с цветовой индикацией
                        if pdf_data.get('lufs') is not None:
                            row.cells[3].text = f"{self._fmt_metric(pdf_data['lufs'], decimals=1)} LUFS"
                            if is_dcp_report:
                                lufs_bg = OK_BG if abs(pdf_data['lufs'] - target_lufs) <= lufs_tolerance else WARN_BG
                            elif abs(pdf_data['lufs'] - target_lufs) <= lufs_tolerance:
                                lufs_bg = OK_BG
                            else:
                                lufs_bg = BAD_BG
                            self._format_cell(row.cells[3], bg=lufs_bg, align="center")
                            logger.info(f"  LUFS added to cell: {self._fmt_metric(pdf_data['lufs'], decimals=1)}")
                        else:
                            logger.warning(f"  ⚠️ LUFS is None in pdf_data for {used_pdf_key or pdf_key}")

                        # Peak metric (TRUE PEAK / SAMPLE PEAK)
                        peak_value = pdf_data.get(peak_value_key)
                        if peak_value is None and is_dcp_report:
                            peak_value = pdf_data.get('true_peak')
                        if peak_value is not None:
                            sign = "+" if peak_value > 0 else ""
                            dp = self._peak_decimals(pdf_data, peak_decimals)
                            row.cells[4].text = f"{sign}{self._fmt_metric(peak_value, decimals=dp)} {peak_unit}"
                            if is_dcp_report:
                                peak_bg = OK_BG if peak_value <= 0 else BAD_BG
                            else:
                                peak_bg = OK_BG if peak_value <= target_peak else BAD_BG
                            self._format_cell(row.cells[4], bg=peak_bg, align="center")
                            logger.info(
                                f"  {peak_label_for_logs} added to cell: "
                                f"{sign}{self._fmt_metric(peak_value, decimals=dp)}"
                            )
                        else:
                            logger.warning(f"  ⚠️ {peak_label_for_logs} is None in pdf_data for {used_pdf_key or pdf_key}")

                        # LRA
                        if pdf_data.get('lra') is not None:
                            row.cells[5].text = f"{self._fmt_metric(pdf_data['lra'], decimals=1)} LU"
                            if is_dcp_report:
                                lra_bg = OK_BG if pdf_data['lra'] <= target_lra else WARN_BG
                            else:
                                lra_bg = OK_BG if pdf_data['lra'] <= target_lra else BAD_BG
                            self._format_cell(row.cells[5], bg=lra_bg, align="center")
                            logger.info(f"  LRA added to cell: {self._fmt_metric(pdf_data['lra'], decimals=1)}")
                        else:
                            logger.warning(f"  ⚠️ LRA is None in pdf_data for {used_pdf_key or pdf_key}")
                        
                        logger.info(
                            f"  PDF data processed for {label}: LUFS={pdf_data.get('lufs')}, "
                            f"{peak_label_for_logs}={pdf_data.get(peak_value_key)}, LRA={pdf_data.get('lra')}"
                        )
                        
                        # ДОБАВЛЯЕМ формат файла из audio data (mediainfo)
                        if key in tech_info and isinstance(tech_info[key], dict):
                            audio_data = tech_info[key]
                            sr = audio_data.get('sample_rate', 48000) // 1000
                            bd = str(audio_data.get('bit_depth', 'PCM_24')).replace('PCM_', '')
                            co = audio_data.get('channel_order', '')
                            format_text = f"PCM {sr}kHz {bd} bit {co}"
                            row.cells[6].text = format_text

                            # Определяем цвет ячейки формата
                            format_bg = self._get_format_bg(sr, bd, co, key, OK_BG, BAD_BG, WARN_BG)

                            self._format_cell(row.cells[6], bg=format_bg, align="left")
                            logger.info(f"  Format added from audio: {format_text}")
                    elif is_dcp_report and key.startswith('audio_dcp_split_'):
                        peak_value = data.get('sample_peak')
                        if peak_value is not None:
                            sign = "+" if peak_value > 0 else ""
                            row.cells[4].text = (
                                f"{sign}{self._fmt_metric(peak_value, decimals=peak_decimals)} {peak_unit}"
                            )
                            peak_bg = OK_BG if peak_value <= 0 else BAD_BG
                            self._format_cell(row.cells[4], bg=peak_bg, align="center")

                        sr = int(data.get('sample_rate', 48000) or 48000) // 1000
                        bd = str(data.get('bit_depth', 'PCM_24')).replace('PCM_', '')
                        row.cells[6].text = f"PCM {sr}kHz {bd} bit split"
                        format_bg = OK_BG if sr == 48 and bd == "24" else WARN_BG
                        self._format_cell(row.cells[6], bg=format_bg, align="left")
                    
                    # Обработка случая когда есть только PDF, но нет аудио (с fallback логикой)
                    elif not has_audio and pdf_key:
                        # Пробуем fallback ключи
                        pdf_data = None
                        used_pdf_key = None
                        fallback_keys = pdf_key_fallbacks.get(pdf_key, [pdf_key])
                        for fallback_key in fallback_keys:
                            if fallback_key in tech_info and tech_info.get(fallback_key) is not None:
                                pdf_data = tech_info[fallback_key]
                                used_pdf_key = fallback_key
                                logger.info(f"  ✓ PDF-ONLY: Found data using key '{fallback_key}' (wanted: '{pdf_key}')")
                                break

                        # Дедупликация
                        if pdf_data:
                            _src2 = str(pdf_data.get('source_pdf') or '')
                            if _src2 and _src2 in _used_pdf_sources:
                                logger.info(f"  ⏭ Пропуск дубликата PDF (block2): {_src2}")
                                continue
                            if _src2:
                                _used_pdf_sources.add(_src2)

                        if pdf_data:
                            logger.info(f"  === PDF-ONLY BLOCK ENTERED for {label} ===")
                            logger.info(f"  pdf_key={pdf_key}, used_key={used_pdf_key}, has_audio={has_audio}")
                            logger.info(f"  pdf_data={pdf_data}")
                            
                            # LUFS
                            if pdf_data.get('lufs') is not None:
                                value = pdf_data['lufs']
                                row.cells[3].text = f"{self._fmt_metric(value, decimals=1)} LUFS"
                                logger.info(f"  ✓ SET cell[3] to '{self._fmt_metric(value, decimals=1)} LUFS'")
                                if is_dcp_report:
                                    lufs_bg = OK_BG if abs(value - target_lufs) <= lufs_tolerance else WARN_BG
                                elif abs(value - target_lufs) <= lufs_tolerance:
                                    lufs_bg = OK_BG
                                else:
                                    lufs_bg = BAD_BG
                                self._format_cell(row.cells[3], bg=lufs_bg, align="center")
                            else:
                                logger.warning(f"  ✗ lufs is None, cell[3] not set")

                            # Peak metric (TRUE PEAK / SAMPLE PEAK)
                            peak_value = pdf_data.get(peak_value_key)
                            if peak_value is None and is_dcp_report:
                                peak_value = pdf_data.get('true_peak')
                            if peak_value is not None:
                                value = peak_value
                                sign = "+" if value > 0 else ""
                                dp = self._peak_decimals(pdf_data, peak_decimals)
                                row.cells[4].text = f"{sign}{self._fmt_metric(value, decimals=dp)} {peak_unit}"
                                logger.info(
                                    f"  ✓ SET cell[4] to "
                                    f"'{sign}{self._fmt_metric(value, decimals=dp)} {peak_unit}'"
                                )
                                if is_dcp_report:
                                    peak_bg = OK_BG if value <= 0 else BAD_BG
                                else:
                                    peak_bg = OK_BG if value <= target_peak else BAD_BG
                                self._format_cell(row.cells[4], bg=peak_bg, align="center")
                            else:
                                logger.warning(f"  ✗ {peak_value_key} is None, cell[4] not set")

                            # LRA
                            if pdf_data.get('lra') is not None:
                                value = pdf_data['lra']
                                row.cells[5].text = f"{self._fmt_metric(value, decimals=1)} LU"
                                logger.info(f"  ✓ SET cell[5] to '{self._fmt_metric(value, decimals=1)} LU'")
                                if is_dcp_report:
                                    lra_bg = OK_BG if value <= target_lra else WARN_BG
                                else:
                                    lra_bg = OK_BG if value <= target_lra else BAD_BG
                                self._format_cell(row.cells[5], bg=lra_bg, align="center")
                            else:
                                logger.warning(f"  ✗ lra is None, cell[5] not set")
                        else:
                            logger.warning(f"  ⚠️ PDF-ONLY: No data found for {pdf_key} (tried: {fallback_keys})")
                    
                    # Для VIDEO показываем формат
                    elif key == 'video':
                        if key in tech_info and isinstance(tech_info[key], dict):
                            video_data = tech_info[key]
                            video_format = video_data.get('format', 'MOV')
                            fps_video = video_data.get('fps', 25)
                            format_text = f"{video_format} {format_fps(fps_video)}fps"
                            row.cells[6].text = format_text
                            self._format_cell(row.cells[6], bg=OK_BG, align="left")
                            logger.info(f"  VIDEO format added: {format_text}")
                    
                    elif key.startswith('audio'):
                        # Для uncens - только формат
                        sr = data.get('sample_rate', 48000) // 1000
                        bd = str(data.get('bit_depth', 'PCM_24')).replace('PCM_', '')
                        co = data.get('channel_order', '')
                        format_text = f"PCM {sr}kHz {bd} bit {co}"
                        row.cells[6].text = format_text

                        format_bg = self._get_format_bg(sr, bd, co, key, OK_BG, BAD_BG, WARN_BG)
                        self._format_cell(row.cells[6], bg=format_bg, align="left")
                    
                    elif key == 'video':
                        # Для видео - показываем только формат и FPS
                        video_format = data.get('format', 'MOV')
                        fps_video = data.get('fps', 25)
                        format_text = f"{video_format} {format_fps(fps_video)}fps"
                        row.cells[6].text = format_text
                        self._format_cell(row.cells[6], bg=OK_BG, align="left")
                    
                    # Summary of what was set
                    logger.info(f"  === SUMMARY for {label} ===")
                    logger.info(f"    cell[0] (label): '{row.cells[0].text}'")
                    logger.info(f"    cell[1] (file): '{row.cells[1].text}'")
                    logger.info(f"    cell[2] (duration): '{row.cells[2].text}'")
                    logger.info(f"    cell[3] (lufs): '{row.cells[3].text}'")
                    logger.info(f"    cell[4] (peak): '{row.cells[4].text}'")
                    logger.info(f"    cell[5] (lra): '{row.cells[5].text}'")
                    logger.info(f"    cell[6] (format): '{row.cells[6].text}'")
                    logger.info(f"  === END SUMMARY ===\n")
            
            service_row = 3 + len(rows_data)

            # Пустая строка (объединенная)
            row = table.rows[service_row]
            cell = row.cells[0]
            for col in range(1, 7):
                cell.merge(row.cells[col])
            self._format_cell(cell, align="left")

            # Цветовые индикаторы (точно как в референсе)
            row = table.rows[service_row + 1]
            left_block = row.cells[0]
            left_block.merge(row.cells[1])
            left_block.merge(row.cells[2])
            self._format_cell(left_block, bg=LIGHT_BG, align="center")
            self._format_cell(row.cells[3], bg="00FA00", align="center")
            self._format_cell(row.cells[4], bg="E73322", align="center")
            self._format_cell(row.cells[5], bg="FCBA19", align="center")
            self._format_cell(row.cells[6], bg=LIGHT_BG, align="center")

            # Пустая строка (объединенная)
            row = table.rows[service_row + 2]
            cell = row.cells[0]
            for col in range(1, 7):
                cell.merge(row.cells[col])
            self._format_cell(cell, align="left")

            # Заголовок "ЗАКЛЮЧЕНИЕ:" (объединенная)
            row = table.rows[service_row + 3]
            cell = row.cells[0]
            for col in range(1, 7):
                cell.merge(row.cells[col])
            cell.text = "ЗАКЛЮЧЕНИЕ:"
            self._format_cell(cell, bg=HEADER_BG, bold=True, font_size=10, align="left")
            
            # Область для заключений (объединенная, 2 подзаголовка)
            row = table.rows[service_row + 4]
            cell = row.cells[0]
            for col in range(1, 7):
                cell.merge(row.cells[col])
            
            # ЛОГИРОВАНИЕ ЗАКЛЮЧЕНИЙ В ГЕНЕРАТОРЕ
            logger.info("=== ЗАКЛЮЧЕНИЯ В ГЕНЕРАТОРЕ ОТЧЕТА (STANDARD) ===")
            logger.info(f"conclusion_tech: '{conclusion_tech}' (длина: {len(conclusion_tech)})")
            logger.info(f"conclusion_subj: '{conclusion_subj}' (длина: {len(conclusion_subj)})")
            logger.info("=" * 50)
            
            # Добавляем оба заключения в одну ячейку:
            # сначала оценки, затем пустая строка и список проблем
            para = cell.paragraphs[0]
            para.clear()

            conclusion_text = self._compose_conclusion_text(conclusion_tech, conclusion_subj)
            if conclusion_text:
                run = para.add_run(conclusion_text)
                run.font.name = "Helvetica Neue"
                run.font.size = Pt(10)
            
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Добавляем межстрочный интервал для лучшей читаемости
            pPr = para._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '360')  # 1.5 интервал
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)

            logger.info("Техническая таблица с областью заключения добавлена")
            
        except Exception as e:
            logger.error(f"Ошибка добавления технической таблицы: {e}")
    
    def _add_me_technical_table(self, doc: Document, tech_info: Dict,
                                conclusion_tech: str, conclusion_subj: str, prepared_by: str = "") -> None:
        """Техническая таблица M&E (упрощенная версия)"""
        # Импорты для работы с форматированием
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from datetime import datetime
        
        logger.info("=== ГЕНЕРАЦИЯ M&E ТАБЛИЦЫ ===")
        logger.info(f"tech_info keys: {list(tech_info.keys()) if tech_info else 'None'}")
        
        # Логируем доступные pdf_ ключи
        logger.info("=== M&E: Доступные pdf_ ключи ===")
        if tech_info:
            for k in tech_info:
                if isinstance(k, str) and k.startswith('pdf_'):
                    v = tech_info[k]
                    if isinstance(v, dict):
                        logger.info(f"  {k}: lufs={v.get('lufs')}, true_peak={v.get('true_peak')}, lra={v.get('lra')}")
                    else:
                        logger.info(f"  {k}: {v}")
        logger.info("=== END ===")
        
        if tech_info and 'video' in tech_info:
            video_data = tech_info['video']
            logger.info(f"VIDEO DATA: {video_data}")
            logger.info(f"  - file_name: {video_data.get('file_name')}")
            logger.info(f"  - duration: {video_data.get('duration')}")
            logger.info(f"  - fps: {video_data.get('fps')}")
            logger.info(f"  - format: {video_data.get('format')}")
        else:
            logger.warning("VIDEO НЕ НАЙДЕНО в tech_info!")
        
        try:
            HEADER_BG = "D5D5D5"
            LIGHT_BG = "F5F5F5"
            OFFWHITE_BG = "FEFFFE"

            from src.me_tracks import iter_dynamic_me_pairs

            def pick_audio_key(prefix: str):
                for k in (f"audio_{prefix}_c", f"audio_{prefix}_uc", f"audio_{prefix}"):
                    if tech_info and k in tech_info and tech_info[k]:
                        return k
                return None

            def pick_pdf_key(prefix: str, audio_key: str):
                if not tech_info:
                    return None
                if audio_key:
                    if audio_key.endswith("_uc") and f"pdf_{prefix}_uc" in tech_info:
                        return f"pdf_{prefix}_uc"
                    if audio_key.endswith("_c") and f"pdf_{prefix}_c" in tech_info:
                        return f"pdf_{prefix}_c"
                for k in (f"pdf_{prefix}_c", f"pdf_{prefix}_uc", f"pdf_{prefix}"):
                    if k in tech_info:
                        return k
                return None

            audio_20_key = pick_audio_key("20")
            audio_51_key = pick_audio_key("51")
            rows_data = [
                ("20 ME", audio_20_key or "audio_20_c", pick_pdf_key("20", audio_20_key)),
                ("51 ME", audio_51_key or "audio_51_c", pick_pdf_key("51", audio_51_key)),
            ]
            rows_data.extend(iter_dynamic_me_pairs(tech_info or {}))
            rows_data.append(("VIDEO", "video", None))

            # 3 строки шапки + динамические дорожки + 5 служебных строк.
            table = doc.add_table(rows=len(rows_data) + 8, cols=7)
            table.style = 'Table Grid'
            
            # Базовая сетка колонок (точно как в референсе):
            # ДОРОЖКА | НАЗВАНИЕ | ХРОНО(1) | ХРОНО(2) | ХРОНО(3) | TRUE PEAK | ФОРМАТ
            widths_cm = [2.61, 15.95, 2.53, 2.63, 1.81, 2.54, 9.89]  # ~38 см
            
            for idx, width_cm in enumerate(widths_cm):
                if idx < len(table.columns):
                    for cell in table.columns[idx].cells:
                        cell.width = Cm(width_cm)
            
            # XML настройки таблицы
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
            
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is not None:
                for idx, gridCol in enumerate(tblGrid.findall(qn('w:gridCol'))):
                    if idx < len(widths_cm):
                        gridCol.set(qn('w:w'), str(int(widths_cm[idx] * 567)))
            
            # Строка 0: Шапка (предыдущий вариант)
            row0 = table.rows[0]
            cell_left = row0.cells[0]
            for col in range(1, 4):
                cell_left.merge(row0.cells[col])
            cell_left.text = "ОБЪЕКТИВНАЯ ОЦЕНКА КАЧЕСТВА ЗВУЧАНИЯ ФОНОГРАММЫ"
            self._format_cell(cell_left, bg="D5D5D5", bold=True, align="center", font_size=11, vertical_align="center")

            cell_date_label = row0.cells[4]
            cell_date_label.merge(row0.cells[5])
            cell_date_label.text = "ДАТА ОТЧЕТА:"
            self._format_cell(cell_date_label, bg="FFFFFF", bold=True, align="left", font_size=10, vertical_align="center")

            cell_date_value = row0.cells[6]
            cell_date_value.text = datetime.now().strftime('%d.%m.%Y')
            self._format_cell(cell_date_value, bg="FFFFFF", bold=True, align="right", font_size=10, vertical_align="center")
            
            # Строка 1: "ОТЧЁТ ПОДГОТОВИЛ"
            row1 = table.rows[1]
            cell_blank = row1.cells[0]
            for col in range(1, 4):
                cell_blank.merge(row1.cells[col])
            self._format_cell(cell_blank, bg="FEFFFE", align="left", vertical_align="center")

            cell_prepared_label = row1.cells[4]
            cell_prepared_label.merge(row1.cells[5])
            cell_prepared_label.text = "ОТЧЁТ ПОДГОТОВИЛ:"
            self._format_cell(cell_prepared_label, bg="F5F5F5", bold=True, align="left", font_size=10, vertical_align="center")

            cell_prepared_value = row1.cells[6]
            cell_prepared_value.text = prepared_by or "Не указано"
            self._format_cell(cell_prepared_value, bg="F5F5F5", bold=True, align="right", font_size=10, vertical_align="center")
            
            # Строка 2: Заголовки столбцов
            row2 = table.rows[2]
            row2.cells[0].text = "ДОРОЖКА"
            self._format_cell(row2.cells[0], bg="D5D5D5", bold=True, align="center", font_size=9, vertical_align="center")

            row2.cells[1].text = "НАЗВАНИЕ ФАЙЛОВ"
            self._format_cell(row2.cells[1], bg="D5D5D5", bold=True, align="center", font_size=9, vertical_align="center")

            chrono_header = row2.cells[2]
            chrono_header.merge(row2.cells[3])
            chrono_header.merge(row2.cells[4])
            chrono_header.text = "ХРОНОМЕТРАЖ"
            self._format_cell(chrono_header, bg="D5D5D5", bold=True, align="center", font_size=9, vertical_align="center")

            row2.cells[5].text = "TRUE PEAK"
            self._format_cell(row2.cells[5], bg="D5D5D5", bold=True, align="center", font_size=9, vertical_align="center")

            row2.cells[6].text = "ФОРМАТ ФАЙЛА"
            self._format_cell(row2.cells[6], bg="D5D5D5", bold=True, align="center", font_size=9, vertical_align="center")
            
            # Fallback маппинг для M&E (как в основной таблице)
            pdf_key_fallbacks = {
                'pdf_20_c': ['pdf_20_c', 'pdf_20'],
                'pdf_20_uc': ['pdf_20_uc'],
                'pdf_51_c': ['pdf_51_c', 'pdf_51'],
                'pdf_51_uc': ['pdf_51_uc'],
                'pdf_20': ['pdf_20'],
                'pdf_51': ['pdf_51']
            }
            
            # Параметры
            params = tech_info.get('params', {}) if tech_info else {}
            target_peak = params.get('true_peak', -2.0)
            
            # FPS
            fps = 25
            if 'video' in tech_info and tech_info['video']:
                fps = tech_info['video'].get('fps', 25)

            def durations_match(dur1, dur2):
                return dur1 and dur2 and int(dur1 * 1000) == int(dur2 * 1000)

            def is_frame_aligned(duration_seconds, fps=25):
                if duration_seconds <= 0:
                    return True
                ms = duration_seconds * 1000
                frame_ms = 1000.0 / fps
                nearest_frame = round(ms / frame_ms)
                return abs(ms - nearest_frame * frame_ms) < 0.5

            # Хронометраж
            durations = {}
            for label, key, _ in rows_data:
                if tech_info and key in tech_info and tech_info[key]:
                    durations[key] = tech_info[key].get('duration', 0)

            # Проверяем, совпадают ли все аудиофайлы между собой
            audio_keys = [key for _label, key, _pdf_key in rows_data if key.startswith('audio')]
            audio_durations_ms = []
            for key in audio_keys:
                if key in durations and durations[key] > 0:
                    audio_durations_ms.append(int(durations[key] * 1000))
            all_audio_match = len(audio_durations_ms) <= 1 or len(set(audio_durations_ms)) == 1
            audio_durations_ms_set = set(audio_durations_ms)
            video_duration_ms = None
            if durations.get('video', 0) > 0:
                video_duration_ms = int(durations['video'] * 1000)

            for idx, (label, key, pdf_key) in enumerate(rows_data):
                row = table.rows[3 + idx]
                # В M&E хронометраж занимает 3 базовые колонки
                row.cells[2].merge(row.cells[3])
                row.cells[2].merge(row.cells[4])
                row.cells[0].text = label
                self._format_cell(row.cells[0], bg=HEADER_BG, bold=True, align="left")

                has_data = tech_info and key in tech_info and tech_info[key]
                has_pdf_only = not has_data and pdf_key and pdf_key in tech_info
                
                if not has_data and not has_pdf_only:
                    logger.warning(f"M&E: Нет данных для {label} (key={key})")
                    # Оставляем ячейки пустыми
                    for col in range(1, 7):
                        row.cells[col].text = ""
                        self._format_cell(row.cells[col], bg="FFFFFF", align="center")
                    continue
                
                # Если есть только PDF без аудио (для M&E) - с fallback
                if has_pdf_only and key.startswith('audio'):
                    logger.info(f"M&E: Обработка {label} только с PDF (нет аудиофайла)")
                    
                    # Пробуем fallback ключи
                    pdf_data = None
                    used_pdf_key = None
                    if pdf_key:
                        fallback_keys = pdf_key_fallbacks.get(pdf_key, [pdf_key])
                        for fb_key in fallback_keys:
                            if fb_key in tech_info and tech_info.get(fb_key) is not None:
                                pdf_data = tech_info[fb_key]
                                used_pdf_key = fb_key
                                logger.info(f"  Found PDF data using key '{fb_key}'")
                                break
                    
                    if pdf_data:
                        # Название файла (берем из source_pdf)
                        source_pdf = pdf_data.get('source_pdf', '')
                        if source_pdf:
                            file_name_without_ext = source_pdf.rsplit('.', 1)[0] if '.' in source_pdf else source_pdf
                            row.cells[1].text = file_name_without_ext
                        else:
                            row.cells[1].text = ""
                        self._format_cell(row.cells[1], align="left")
                        
                        # Хронометраж (пустой)
                        row.cells[2].text = ""
                        self._format_cell(row.cells[2], align="center")
                        
                        # TRUE PEAK из PDF
                        true_peak = pdf_data.get('true_peak')
                        if true_peak is not None:
                            sign = "+" if true_peak > 0 else ""
                            dp = self._peak_decimals(pdf_data, 1)
                            row.cells[5].text = f"{sign}{self._fmt_metric(true_peak, decimals=dp)} dBTP"
                            peak_bg = "00FA02" if true_peak <= target_peak else "E83121"
                            self._format_cell(row.cells[5], bg=peak_bg, align="center")
                            logger.info(f"  TRUE PEAK: {sign}{self._fmt_metric(true_peak, decimals=dp)} dBTP")
                        else:
                            row.cells[5].text = ""
                            self._format_cell(row.cells[5], bg="FFFFFF", align="center")

                        # Формат файла (если есть данные)
                        sr = pdf_data.get('sample_rate')
                        bd = pdf_data.get('bit_depth')
                        if sr:
                            sr_khz = sr // 1000
                            format_text = f"PCM {sr_khz}kHz"
                            if bd:
                                format_text += f" {bd} bit"
                            row.cells[6].text = format_text
                            self._format_cell(row.cells[6], bg="00FA02", align="left")
                        else:
                            row.cells[6].text = ""
                            self._format_cell(row.cells[6], bg="FFFFFF", align="left")
                    else:
                        logger.warning(f"  No PDF data found for {label}")
                    
                    continue
                
                if has_data:
                    data = tech_info[key]
                    
                    # Имя файла с проверкой на некорректный тип (убираем расширение)
                    file_name = data.get('file_name', '')
                    file_name_without_ext = file_name.rsplit('.', 1)[0] if '.' in file_name else file_name
                    row.cells[1].text = file_name_without_ext
                    
                    # Проверяем на некорректный тип фонограммы в названии
                    has_incorrect_type = self._check_incorrect_audio_type(file_name)
                    name_mismatch = bool(data.get('name_channel_mismatch')) or bool(data.get('name_incorrect_type'))
                    name_error = has_incorrect_type or name_mismatch
                    
                    if name_error:
                        self._format_cell(row.cells[1], align="left", bg="E83121", text_color="FFFFFF")
                        logger.warning(f"⚠️  M&E: Название файла с ошибкой: {file_name}")
                    else:
                        self._format_cell(row.cells[1], align="left")
                    
                    if key == 'video':
                        logger.info(f"M&E: Обработка VIDEO")
                        logger.info(f"  - файл: {data.get('file_name', 'N/A')}")
                        logger.info(f"  - duration: {data.get('duration', 0)}")
                        logger.info(f"  - format: {data.get('format', 'N/A')}")
                        logger.info(f"  - fps: {data.get('fps', 'N/A')}")
                    
                    # Хронометраж
                    duration = data.get('duration', 0)
                    total_ms = int(duration * 1000)
                    hours = total_ms // 3600000
                    mins = (total_ms % 3600000) // 60000
                    secs = (total_ms % 60000) // 1000
                    millis = total_ms % 1000
                    row.cells[2].text = f"{hours}:{mins:02d}:{secs:02d}.{millis:03d}"

                    # Цветовая индикация
                    if duration > 0:
                        if key.startswith('audio'):
                            # Если есть видео, оцениваем дорожку относительно него.
                            # Так синхронная с видео дорожка остается зеленой, даже если
                            # другая аудиодорожка имеет иной хронометраж.
                            if video_duration_ms is not None:
                                chrono_bg = "00FA02" if total_ms == video_duration_ms else "E83121"
                            elif not all_audio_match:
                                chrono_bg = "E83121"  # Красный - аудиофайлы не совпадают
                            else:
                                chrono_bg = "00FA02"  # Зеленый - аудиофайлы совпадают
                        elif key == 'video':
                            # Для видео: сравниваем с аудио
                            if audio_durations_ms:
                                video_ms = int(duration * 1000)
                                if video_ms in audio_durations_ms_set:
                                    chrono_bg = "00FA02"
                                else:
                                    chrono_bg = "E83121"
                            else:
                                frame_ok = is_frame_aligned(duration, fps)
                                chrono_bg = "00FA02" if frame_ok else "FBBA18"
                        else:
                            frame_ok = is_frame_aligned(duration, fps)
                            chrono_bg = "00FA02" if frame_ok else "FBBA18"
                    else:
                        chrono_bg = "FFFFFF"
                    
                    self._format_cell(row.cells[2], bg=chrono_bg, align="center")
                    
                    # TRUE PEAK (только для аудио с PDF) - с fallback
                    if key.startswith('audio') and pdf_key:
                        # Пробуем fallback ключи
                        pdf_data = None
                        used_pdf_key = None
                        fallback_keys = pdf_key_fallbacks.get(pdf_key, [pdf_key])
                        for fb_key in fallback_keys:
                            if fb_key in tech_info and tech_info.get(fb_key) is not None:
                                pdf_data = tech_info[fb_key]
                                used_pdf_key = fb_key
                                logger.info(f"M&E: {label} Found PDF data using key '{fb_key}' (wanted: '{pdf_key}')")
                                break
                        
                        if pdf_data:
                            true_peak = pdf_data.get('true_peak')
                            if true_peak is not None:
                                # Добавляем знак + для положительных значений
                                sign = "+" if true_peak > 0 else ""
                                dp = self._peak_decimals(pdf_data, 1)
                                row.cells[5].text = f"{sign}{self._fmt_metric(true_peak, decimals=dp)} dBTP"
                                peak_bg = "00FA02" if true_peak <= target_peak else "E83121"
                                self._format_cell(row.cells[5], bg=peak_bg, align="center")
                                logger.info(
                                    f"M&E: {label} TRUE PEAK форматирован "
                                    f"({sign}{self._fmt_metric(true_peak, decimals=dp)} dBTP)"
                                )
                            else:
                                row.cells[5].text = ""
                                self._format_cell(row.cells[5], bg="FFFFFF", align="center")
                                logger.info(f"M&E: {label} TRUE PEAK пустой (нет данных)")
                        else:
                            row.cells[5].text = ""
                            self._format_cell(row.cells[5], bg="FFFFFF", align="center")
                            logger.warning(f"M&E: {label} PDF data not found (tried: {fallback_keys})")
                    elif key == 'video':
                        # Для VIDEO ячейка TRUE PEAK пустая с белым фоном
                        row.cells[5].text = ""
                        self._format_cell(row.cells[5], bg="FFFFFF", align="center")
                    
                    # Формат файла
                    if key.startswith('audio'):
                        sr = data.get('sample_rate', 48000) // 1000
                        bd = str(data.get('bit_depth', 'PCM_24')).replace('PCM_', '')
                        co = data.get('channel_order', '')
                        format_text = f"PCM {sr}kHz {bd} bit {co}"
                        row.cells[6].text = format_text

                        format_bg = self._get_format_bg(sr, bd, co, key, "00FA02", "E83121", "FBBA18")
                        self._format_cell(row.cells[6], bg=format_bg, align="left")
                    elif key == 'video':
                        # Формат видео - только формат и FPS
                        video_format = data.get('format', 'MOV')
                        fps_video = data.get('fps', 25)
                        format_text = f"{video_format} {format_fps(fps_video)}fps"
                        row.cells[6].text = format_text
                        self._format_cell(row.cells[6], bg="00FA02", align="left")
            
            service_row = 3 + len(rows_data)

            # Пустой отступ перед индикаторами
            row = table.rows[service_row]
            spacer_top = row.cells[0]
            for col in range(1, 7):
                spacer_top.merge(row.cells[col])
            self._format_cell(spacer_top, align="left")

            # Цветовые индикаторы (точно как в референсном M&E)
            row = table.rows[service_row + 1]
            left_block = row.cells[0]
            left_block.merge(row.cells[1])
            left_block.merge(row.cells[2])
            left_block.text = ""
            self._format_cell(left_block, bg="F5F5F5", align="center", border=True)

            row.cells[3].text = ""
            self._format_cell(row.cells[3], bg="00FA00", align="center", border=True)

            row.cells[4].text = ""
            self._format_cell(row.cells[4], bg="E73322", align="center", border=True)

            row.cells[5].text = ""
            self._format_cell(row.cells[5], bg="FCBA19", align="center", border=True)

            row.cells[6].text = ""
            self._format_cell(row.cells[6], bg="F5F5F5", align="center", border=True)
            
            # Пустой отступ после индикаторов
            row = table.rows[service_row + 2]
            spacer_bottom = row.cells[0]
            for col in range(1, 7):
                spacer_bottom.merge(row.cells[col])
            self._format_cell(spacer_bottom, align="left")

            # Заголовок "ЗАКЛЮЧЕНИЕ:"
            row = table.rows[service_row + 3]
            cell = row.cells[0]
            for col in range(1, 7):
                cell.merge(row.cells[col])
            cell.text = "ЗАКЛЮЧЕНИЕ:"
            self._format_cell(cell, bold=True, font_size=10, align="left")
            
            # Текст заключения (заголовки уже в тексте заключений)
            row = table.rows[service_row + 4]
            cell = row.cells[0]
            for col in range(1, 7):
                cell.merge(row.cells[col])
            
            # ЛОГИРОВАНИЕ ЗАКЛЮЧЕНИЙ В ГЕНЕРАТОРЕ
            logger.info("=== ЗАКЛЮЧЕНИЯ В ГЕНЕРАТОРЕ ОТЧЕТА (M&E) ===")
            logger.info(f"conclusion_tech: '{conclusion_tech}' (длина: {len(conclusion_tech)})")
            logger.info(f"conclusion_subj: '{conclusion_subj}' (длина: {len(conclusion_subj)})")
            logger.info("=" * 50)
            
            # Сначала оценки, затем пустая строка и список проблем
            conclusion_text = self._compose_conclusion_text(conclusion_tech, conclusion_subj)
            cell.text = conclusion_text
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Helvetica Neue'
                    run.font.size = Pt(10)
                paragraph.paragraph_format.line_spacing = 1.5
            
            logger.info("M&E таблица добавлена")
            
        except Exception as e:
            logger.error(f"Ошибка добавления M&E таблицы: {e}")

    def _compose_conclusion_text(self, conclusion_tech: str, conclusion_subj: str) -> str:
        """Формирует блок заключения (как на референсных скриншотах).

        Вариант A — техн. без проблем:
            По технической оценке нареканий не обнаружено.
                                                          ← пустая строка
            По субъективной оценке выявлены следующие недочеты:
                                                          ← пустая строка
            -  проблема 1
            -  проблема 2

        Вариант B — техн. с проблемами:
            По техническим характеристикам выявлены следующие недочёты:
            - Параметр ...                                ← сразу, без пустой строки
            - Хронометраж ...
                                                          ← пустая строка
            По субъективной оценке выявлены следующие недочёты:
                                                          ← пустая строка
            [ЗАПОЛНИТЬ ВРУЧНУЮ]
        """
        def ensure_period(text: str) -> str:
            """Гарантирует точку в конце строки (кроме двоеточия)."""
            text = text.rstrip()
            if text and text[-1] not in '.!:':
                text += '.'
            return text

        def render_block(text: str) -> str:
            if not text:
                return ""
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            if not lines:
                return ""

            heading = ensure_period(lines[0])
            problems = lines[1:]

            if not problems:
                # Нет проблем — только заголовок
                return heading
            # Есть проблемы — пустая строка перед списком
            return heading + "\n\n" + "\n".join(problems)

        parts = []
        tech_block = render_block(conclusion_tech)
        subj_block = render_block(conclusion_subj)

        if tech_block:
            parts.append(tech_block)
        if subj_block:
            parts.append(subj_block)

        # Между блоками — одна пустая строка (\n\n)
        return "\n\n".join(parts)
    
    def _check_incorrect_audio_type(self, description: str) -> bool:
        """Проверка на некорректный тип фонограммы в названии файла
        
        Args:
            description: Текст описания из маркера (БЕЗ расширения .wav)
            
        Returns:
            True если найдены некорректные типы (50, 21 вместо 51, 20)
        """
        import re
        
        # УМНАЯ ПРОВЕРКА: ищем неправильные типы только в контексте аудио файлов
        # После _50_ или _21_ должно идти: _cens, _uncens, _mix, _c, _uc или конец
        incorrect_patterns = [
            r'_50_(cens|uncens|mix|c|uc)',  # _50_cens, _50_uncens, etc.
            r'_50$',                         # _50 в конце
            r'_21_(cens|uncens|mix|c|uc)',  # _21_cens, _21_uncens, etc.
            r'_21$',                         # _21 в конце
            r'_5\.0_(cens|uncens|mix)',     # _5.0_cens, etc.
            r'_2\.1_(cens|uncens|mix)',     # _2.1_cens, etc.
        ]
        
        for pattern in incorrect_patterns:
            if re.search(pattern, description, re.IGNORECASE):
                return True
        return False
    
    def _add_marker_list_exact(self, doc: Document, issues: List[Issue], report_type: str = "main") -> None:
        """MARKER LIST - точно как в референсе"""
        try:
            is_me_report = report_type in ("me", "me_ours")

            # Добавляем заголовок "MARKER LIST"
            header_para = doc.add_paragraph()
            header_run = header_para.add_run("MARKER LIST")
            header_run.font.name = "Helvetica Neue"
            header_run.font.size = Pt(14)
            header_run.font.bold = True
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Отступ после заголовка
            doc.add_paragraph()
            
            if is_me_report:
                headers = [
                    "ID", "Timecode In", "Timecode Out", "Description",
                    "2.0 ME", "5.1 ME", "2.0 DX", "5.1 DX", "2.0 OPT", "5.1 OPT",
                    "БЛОКЕР", "ТРЕБУЕТ ИСПРАВЛЕНИЯ", "ТРЕБУЕТ КОММЕНТАРИЯ", "КОММЕНТАРИИ"
                ]
                widths_cm = [1.50, 2.45, 2.40, 8.31, 1.27, 1.32, 1.27, 1.32, 1.27, 1.32, 1.98, 3.55, 3.48, 6.57]
            else:
                headers = [
                    "ID", "Timecode In", "Timecode Out", "Description",
                    "2.0 C", "2.0 UC", "5.1 C", "5.1 UC",
                    "БЛОКЕР", "ТРЕБУЕТ ИСПРАВЛЕНИЯ", "ТРЕБУЕТ КОММЕНТАРИЯ", "КОММЕНТАРИИ"
                ]
                widths_cm = [1.50, 2.45, 2.40, 10.43, 1.27, 1.56, 1.32, 1.50, 1.98, 3.55, 3.48, 6.57]

            # Создаем таблицу БЕЗ заголовочной строки с повторяющимся текстом
            table = doc.add_table(rows=1 + len(issues), cols=len(headers))
            table.style = 'Table Grid'
            
            # Метод 1: Устанавливаем через table.columns
            for idx, width_cm in enumerate(widths_cm):
                if idx < len(table.columns):
                    for cell in table.columns[idx].cells:
                        cell.width = Cm(width_cm)
            
            # Метод 2: Устанавливаем через XML напрямую
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Удаляем autofit
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
            
            # Строка 0: Заголовки столбцов (единственная заголовочная строка)
            for col, header in enumerate(headers):
                cell = table.rows[0].cells[col]
                cell.text = header
                # Фон #bdc0bf, шрифт Helvetica Neue, жирный, центрирование
                self._format_cell(cell, bg="bdc0bf", font_name="Helvetica Neue", bold=True, align="center", font_size=10)
            
            # Строки 1+: Данные с чередующимся фоном
            for row_idx, issue in enumerate(issues):
                row = table.rows[1 + row_idx]
                
                # Проверяем на "НОВЫЙ МАРКЕР" в комментариях
                is_new_marker = issue.comments and "НОВЫЙ МАРКЕР" in issue.comments.upper()
                
                # Убираем расширение .wav из описания
                description_clean = issue.description
                if description_clean.lower().endswith('.wav'):
                    description_clean = description_clean[:-4]
                
                # Проверяем на некорректный тип фонограммы в описании
                has_incorrect_type = self._check_incorrect_audio_type(description_clean)
                
                # Заполняем данные
                row.cells[0].text = issue.marker_id or ''
                row.cells[1].text = issue.timecode_in
                row.cells[2].text = issue.timecode_out
                row.cells[3].text = description_clean
                if is_me_report:
                    track_values = dict(getattr(issue, 'me_tracks', {}) or {})
                    track_values.setdefault('me_20', bool(issue.audio_20_c))
                    track_values.setdefault('me_51', bool(issue.audio_51_c))
                    for offset, track_key in enumerate(('me_20', 'me_51', 'dx_20', 'dx_51', 'opt_20', 'opt_51')):
                        row.cells[4 + offset].text = '*' if track_values.get(track_key) else ''
                    row.cells[10].text = '*' if issue.blocker else ''
                    row.cells[11].text = '*' if issue.fix_required else ''
                    row.cells[12].text = '*' if issue.comment_required else ''
                    row.cells[13].text = issue.comments
                else:
                    row.cells[4].text = '*' if issue.audio_20_c else ''
                    row.cells[5].text = '*' if issue.audio_20_uc else ''
                    row.cells[6].text = '*' if issue.audio_51_c else ''
                    row.cells[7].text = '*' if issue.audio_51_uc else ''
                    row.cells[8].text = '*' if issue.blocker else ''
                    row.cells[9].text = '*' if issue.fix_required else ''
                    row.cells[10].text = '*' if issue.comment_required else ''
                    row.cells[11].text = issue.comments
                
                # Определяем цвет фона
                if is_new_marker:
                    # Цвет для "НОВЫЙ МАРКЕР"
                    bg_color = "FF9999"
                    logger.info(f"Строка {row_idx + 1}: НОВЫЙ МАРКЕР - коралловый фон")
                else:
                    # Чередующийся фон для нечетных строк (начиная с row_idx=1)
                    bg_color = "f5f5f5" if row_idx % 2 == 1 else "auto"
                
                # Форматируем каждую ячейку
                for col_idx, cell in enumerate(row.cells):
                    # Определяем цвет текста для Description (col 3)
                    text_color = "FF0000" if col_idx == 3 and has_incorrect_type else None
                    
                    # Фон только для некоторых столбцов (как в референсе)
                    if is_me_report:
                        shaded_columns = [0, 1, 2] + list(range(4, 14))
                    else:
                        shaded_columns = [0, 1, 2, 4, 5, 6, 7, 8, 9, 10, 11]

                    if col_idx in shaded_columns:
                        self._format_cell(
                            cell, bg=bg_color, font_name="Helvetica Neue",
                            font_size=10, text_color=text_color,
                        )
                    else:
                        # Description (col 3) без фона (если не НОВЫЙ МАРКЕР)
                        desc_bg = bg_color if is_new_marker else "auto"
                        self._format_cell(
                            cell, bg=desc_bg, font_name="Helvetica Neue",
                            font_size=10, text_color=text_color,
                        )
                
                if has_incorrect_type:
                    logger.info(f"Строка {row_idx + 1}: Некорректный тип фонограммы - красный текст")
            
            logger.info(f"MARKER LIST добавлен ({len(issues)} записей)")
            
        except Exception as e:
            logger.error(f"Ошибка добавления MARKER LIST: {e}")
    
    def _add_pdf_image(self, doc: Document, pdf_path: str) -> None:
        """Добавление PDF как изображение (ширина = ширине таблицы)"""
        try:
            logger.info(f"Конвертация PDF: {pdf_path}")
            
            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[0]
            
            # Конвертация в изображение с высоким разрешением
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            
            # Добавляем изображение с шириной равной таблице (38.01 см)
            img_stream = io.BytesIO(img_bytes)
            doc.add_picture(img_stream, width=Cm(38.0))
            
            # Центрируем
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            pdf_doc.close()
            
            logger.info("PDF добавлен как изображение (38 см)")
            
        except Exception as e:
            logger.error(f"Ошибка добавления PDF: {e}")
    
    def _add_conclusion_OLD(self, doc: Document, technical: str, subjective: str) -> None:
        """Добавление заключения"""
        try:
            # Заголовок
            title = doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.name = 'Helvetica Neue'
                run.font.size = Pt(14)
                run.font.bold = True
            
            doc.add_paragraph()
            
            # Техническая оценка
            if technical:
                tech_title = doc.add_paragraph("Техническая оценка:")
                for run in tech_title.runs:
                    run.font.name = 'Helvetica Neue'
                    run.font.size = Pt(12)
                    run.font.bold = True
                
                tech_para = doc.add_paragraph(technical)
                tech_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in tech_para.runs:
                    run.font.name = 'Helvetica Neue'
                    run.font.size = Pt(11)
                
                doc.add_paragraph()
            
            # Субъективная оценка
            if subjective:
                subj_title = doc.add_paragraph("Субъективная оценка:")
                for run in subj_title.runs:
                    run.font.name = 'Helvetica Neue'
                    run.font.size = Pt(12)
                    run.font.bold = True
                
                subj_para = doc.add_paragraph(subjective)
                subj_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in subj_para.runs:
                    run.font.name = 'Helvetica Neue'
                    run.font.size = Pt(11)
            
            logger.info("Заключение добавлено")
            
        except Exception as e:
            logger.error(f"Ошибка добавления заключения: {e}")
    
    @staticmethod
    def _get_format_bg(sr, bd, co, key, ok_bg, bad_bg, warn_bg) -> str:
        """Определение цвета фона ячейки формата файла.

        Правила для 5.1:
        - Нет метаданных порядка каналов ('channels' в co) → RED
        - Порядок каналов не соответствует стандартному → WARN
        Для 2.0: порядок каналов НЕ проверяется.
        Для всех: Sample rate != 48 или bit depth != 24 → RED
        """
        STANDARD_51_ORDER = "L R C LFE Ls Rs"

        # 1. Проверка sample rate и bit depth (стандарт 48/24)
        try:
            bd_int = int(str(bd).replace('bit', '').strip())
        except (ValueError, TypeError):
            bd_int = 0
        if sr != 48 or bd_int != 24:
            return bad_bg

        # 2. Порядок каналов — только для 5.1
        if "51" in key:
            if not co or "channels" in co.lower() or co.lower() == "unknown":
                return bad_bg
            if co.strip() != STANDARD_51_ORDER:
                return warn_bg

        return ok_bg

    def _format_cell(self, cell, bg="auto", font_name="Helvetica Neue",
                    font_size=None, font_color=None, bold=False, align=None, 
                    border=True, vertical_align="center", text_color=None):
        """Форматирование ячейки
        
        Args:
            text_color: Цвет текста (синоним font_color для удобства)
        """
        try:
            # text_color является синонимом для font_color
            if text_color and not font_color:
                font_color = text_color
            
            tc_pr = cell._element.get_or_add_tcPr()
            
            # Фон
            if bg != "auto":
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), bg)
                tc_pr.append(shading_elm)
            
            # Вертикальное выравнивание
            v_align = OxmlElement('w:vAlign')
            v_align.set(qn('w:val'), vertical_align)
            tc_pr.append(v_align)
            
            # Отступы внутри ячейки (padding)
            tc_mar = OxmlElement('w:tcMar')
            for side in ['top', 'left', 'bottom', 'right']:
                mar = OxmlElement(f'w:{side}')
                mar.set(qn('w:w'), '100')  # 100 twips = примерно 1.76 мм
                mar.set(qn('w:type'), 'dxa')
                tc_mar.append(mar)
            tc_pr.append(tc_mar)
            
            # Границы (если нужны)
            if border:
                tc_borders = OxmlElement('w:tcBorders')
                for side in ['top', 'left', 'bottom', 'right']:
                    border_elm = OxmlElement(f'w:{side}')
                    border_elm.set(qn('w:val'), 'single')
                    border_elm.set(qn('w:sz'), '4')  # Размер границы
                    border_elm.set(qn('w:space'), '0')
                    border_elm.set(qn('w:color'), '000000')  # Черный
                    tc_borders.append(border_elm)
                tc_pr.append(tc_borders)
            
            # Текст
            for para in cell.paragraphs:
                # Убираем отступы до/после параграфа
                pPr = para._element.get_or_add_pPr()
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
                spacing.set(qn('w:line'), '240')
                spacing.set(qn('w:lineRule'), 'auto')
                pPr.append(spacing)
                
                if align:
                    if align == "center":
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align == "left":
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif align == "right":
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                for run in para.runs:
                    run.font.name = font_name
                    if font_size:
                        run.font.size = Pt(font_size)
                    else:
                        run.font.size = Pt(9)  # Размер по умолчанию
                    if font_color:
                        run.font.color.rgb = RGBColor.from_string(font_color)
                    if bold:
                        run.font.bold = True
        except Exception as e:
            logger.error(f"Ошибка форматирования ячейки: {e}")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    print("ExactReportGenerator готов")
