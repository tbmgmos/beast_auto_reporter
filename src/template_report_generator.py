"""
Template Report Generator

Генерация отчетов ТОЧНО по шаблону отчет_petr_2_s1_e2_2024_09_12_rus.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pathlib import Path
from typing import List, Dict
import logging

logger = logging.getLogger(__name__)


class TemplateReportGenerator:
    """Генератор отчетов по шаблону"""
    
    def __init__(self):
        """Инициализация генератора"""
        self.template_path = None
        logger.info("TemplateReportGenerator инициализирован")
    
    def create_report_from_template(
        self,
        defects: List,
        output_path: str,
        params_info: Dict = None,
        timecode_info: Dict = None
    ) -> str:
        """
        Создание отчета по шаблону
        
        Args:
            defects: Список дефектов
            output_path: Путь для сохранения
            params_info: Информация из Параметры.txt
            timecode_info: Информация о хронометраже
            
        Returns:
            Путь к созданному файлу
        """
        try:
            logger.info("Создание отчета по шаблону...")
            
            # Создаем новый документ
            doc = Document()
            
            # Настройка шрифта по умолчанию
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial Unicode MS'
            font.size = Pt(13)
            
            # Создаем таблицу
            # Количество строк = 1 (заголовок MARKER LIST) + 1 (заголовки столбцов) + количество дефектов
            num_rows = 2 + len(defects)
            num_cols = 11  # Количество столбцов
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            # === Первая строка: объединенные ячейки "MARKER LIST" ===
            first_row = table.rows[0]
            
            # Объединяем все ячейки первой строки
            for i in range(num_cols - 1):
                first_row.cells[0].merge(first_row.cells[i + 1])
            
            # Заполняем "MARKER LIST"
            cell = first_row.cells[0]
            cell.text = "MARKER LIST"
            
            # Форматирование
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Arial Unicode MS'
                    run.font.size = Pt(13)
                    run.font.bold = True
            
            # === Вторая строка: заголовки столбцов ===
            headers = [
                'Timecode In',
                'Timecode Out',
                'Description',
                '2.0 C',
                '2.0 UC',
                '5.1 C',
                '5.1 UC',
                'БЛОКЕР',
                'ТРЕБУЕТ ИСПРАВЛЕНИЯ',
                'ТРЕБУЕТ КОММЕНТАРИЯ',
                'КОММЕНТАРИИ'
            ]
            
            header_row = table.rows[1]
            for i, header_text in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header_text
                
                # Форматирование заголовков
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Arial Unicode MS'
                        run.font.size = Pt(13)
                        run.font.bold = True
            
            # === Строки данных ===
            for idx, defect in enumerate(defects):
                row_idx = idx + 2
                row = table.rows[row_idx]
                
                # Timecode In
                row.cells[0].text = getattr(defect, 'timecode_in', '')
                
                # Timecode Out
                row.cells[1].text = getattr(defect, 'timecode_out', '') or ''
                
                # Description
                row.cells[2].text = getattr(defect, 'description', '')
                
                # Определяем в каких каналах присутствует дефект
                channels = getattr(defect, 'channels', [])
                has_20 = any(c in ['*', '2.0'] for c in channels)
                has_51 = any(c in ['*', '5.1'] for c in channels)
                
                # 2.0 C (censored)
                row.cells[3].text = '*' if has_20 else ''
                
                # 2.0 UC (uncensored) - обычно пусто
                row.cells[4].text = ''
                
                # 5.1 C
                row.cells[5].text = '*' if has_51 else ''
                
                # 5.1 UC - обычно пусто
                row.cells[6].text = ''
                
                # Severity columns
                severity = getattr(defect, 'severity', 'comment_required')
                
                row.cells[7].text = '*' if severity == 'blocker' else ''
                row.cells[8].text = '*' if severity == 'fix_required' else ''
                row.cells[9].text = '*' if severity == 'comment_required' else ''
                
                # КОММЕНТАРИИ - пусто
                row.cells[10].text = ''
                
                # Форматирование ячеек
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial Unicode MS'
                            run.font.size = Pt(13)
            
            # Применяем стиль таблицы (границы)
            self._apply_table_style(table)
            
            # Сохраняем документ
            doc.save(output_path)
            
            logger.info(f"Отчет создан: {output_path}")
            logger.info(f"Всего маркеров: {len(defects)}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Ошибка создания отчета: {e}")
            raise
    
    def _apply_table_style(self, table):
        """Применение стиля к таблице (границы и т.д.)"""
        try:
            # Добавляем границы ко всем ячейкам
            for row in table.rows:
                for cell in row.cells:
                    # Границы
                    tcPr = cell._element.get_or_add_tcPr()
                    tcBorders = parse_xml(
                        f'<w:tcBorders {nsdecls("w")}>'
                        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '</w:tcBorders>'
                    )
                    tcPr.append(tcBorders)
        except Exception as e:
            logger.warning(f"Не удалось применить стиль таблицы: {e}")
    
    def load_template(self, template_path: str):
        """
        Загрузка существующего шаблона
        
        Args:
            template_path: Путь к файлу шаблона
        """
        self.template_path = template_path
        logger.info(f"Шаблон загружен: {template_path}")
    
    def create_from_existing_template(
        self,
        template_path: str,
        defects: List,
        output_path: str
    ) -> str:
        """
        Создание отчета на основе существующего шаблона
        
        Args:
            template_path: Путь к шаблону
            defects: Список дефектов
            output_path: Путь для сохранения
            
        Returns:
            Путь к созданному файлу
        """
        try:
            # Загружаем шаблон
            doc = Document(template_path)
            
            # Находим таблицу
            if not doc.tables:
                raise ValueError("Таблица не найдена в шаблоне")
            
            table = doc.tables[0]
            
            # Удаляем все строки кроме первых двух (MARKER LIST и заголовки)
            rows_to_delete = len(table.rows) - 2
            for _ in range(rows_to_delete):
                table._element.remove(table.rows[-1]._element)
            
            # Добавляем новые строки с дефектами
            for defect in defects:
                row = table.add_row()
                
                # Заполняем ячейки
                row.cells[0].text = getattr(defect, 'timecode_in', '')
                row.cells[1].text = getattr(defect, 'timecode_out', '') or ''
                row.cells[2].text = getattr(defect, 'description', '')
                
                channels = getattr(defect, 'channels', [])
                has_20 = any(c in ['*', '2.0'] for c in channels)
                has_51 = any(c in ['*', '5.1'] for c in channels)
                
                row.cells[3].text = '*' if has_20 else ''
                row.cells[4].text = ''
                row.cells[5].text = '*' if has_51 else ''
                row.cells[6].text = ''
                
                severity = getattr(defect, 'severity', 'comment_required')
                row.cells[7].text = '*' if severity == 'blocker' else ''
                row.cells[8].text = '*' if severity == 'fix_required' else ''
                row.cells[9].text = '*' if severity == 'comment_required' else ''
                row.cells[10].text = ''
                
                # Копируем форматирование из шаблона
                for i, cell in enumerate(row.cells):
                    if i < len(table.rows[2].cells):
                        # Копируем стиль из первой строки данных (если есть)
                        template_cell = table.rows[2].cells[i] if len(table.rows) > 2 else None
                        if template_cell:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Arial Unicode MS'
                                    run.font.size = Pt(13)
            
            # Сохраняем
            doc.save(output_path)
            
            logger.info(f"Отчет создан на основе шаблона: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Ошибка создания отчета из шаблона: {e}")
            raise


if __name__ == "__main__":
    # Тестирование
    logging.basicConfig(level=logging.INFO)
    
    generator = TemplateReportGenerator()
    print("TemplateReportGenerator готов к использованию!")

