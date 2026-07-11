"""
Template Processor Module

Обработчик шаблонов DOCX для генерации отчетов
Позволяет использовать готовые шаблоны Word вместо программного создания
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class TemplateProcessor:
    """Обработчик шаблонов DOCX"""
    
    def __init__(self, template_path: str):
        """
        Инициализация процессора
        
        Args:
            template_path: Путь к файлу шаблона DOCX
        """
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")
        
        self.doc = Document(str(self.template_path))
        logger.info(f"Загружен шаблон: {self.template_path}")
    
    def replace_placeholders(self, replacements: Dict[str, str]) -> int:
        """
        Замена плейсхолдеров в документе
        
        Args:
            replacements: Словарь замен {плейсхолдер: значение}
            
        Returns:
            Количество произведенных замен
        """
        count = 0
        
        # Замена в параграфах
        for paragraph in self.doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    # Сохраняем форматирование
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
                            count += 1
        
        # Замена в таблицах
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                # Сохраняем форматирование
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(value))
                                        count += 1
        
        logger.info(f"Заменено плейсхолдеров: {count}")
        return count
    
    def replace_in_cell(self, table_index: int, row_index: int, cell_index: int, 
                       placeholder: str, value: str, 
                       color: Optional[str] = None) -> bool:
        """
        Замена плейсхолдера в конкретной ячейке
        
        Args:
            table_index: Индекс таблицы (0-based)
            row_index: Индекс строки (0-based)
            cell_index: Индекс ячейки (0-based)
            placeholder: Плейсхолдер для замены
            value: Новое значение
            color: RGB цвет фона ячейки (hex, например "00FA02")
            
        Returns:
            True если замена произведена, False иначе
        """
        try:
            table = self.doc.tables[table_index]
            cell = table.rows[row_index].cells[cell_index]
            
            # Замена текста
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
            
            # Установка цвета фона
            if color:
                self._set_cell_background(cell, color)
            
            return True
        
        except (IndexError, AttributeError) as e:
            logger.error(f"Ошибка замены в ячейке [{table_index}][{row_index}][{cell_index}]: {e}")
            return False
    
    def duplicate_row(self, table_index: int, row_index: int, count: int) -> List:
        """
        Дублирование строки таблицы
        
        Args:
            table_index: Индекс таблицы (0-based)
            row_index: Индекс строки-шаблона (0-based)
            count: Количество копий
            
        Returns:
            Список созданных строк
        """
        try:
            table = self.doc.tables[table_index]
            template_row = table.rows[row_index]
            
            new_rows = []
            for _ in range(count):
                new_row = table.add_row()
                
                # Копируем содержимое и форматирование
                for i, template_cell in enumerate(template_row.cells):
                    new_cell = new_row.cells[i]
                    
                    # Копируем текст
                    for para in template_cell.paragraphs:
                        new_para = new_cell.paragraphs[0] if new_cell.paragraphs else new_cell.add_paragraph()
                        new_para.text = para.text
                        new_para.alignment = para.alignment
                        
                        # Копируем форматирование runs
                        for run in para.runs:
                            new_run = new_para.runs[0] if new_para.runs else new_para.add_run()
                            new_run.text = run.text
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.name:
                                new_run.font.name = run.font.name
                
                new_rows.append(new_row)
            
            logger.info(f"Продублирована строка {row_index} в таблице {table_index}: {count} копий")
            return new_rows
        
        except (IndexError, AttributeError) as e:
            logger.error(f"Ошибка дублирования строки [{table_index}][{row_index}]: {e}")
            return []
    
    def get_table_count(self) -> int:
        """Получить количество таблиц в документе"""
        return len(self.doc.tables)
    
    def get_row_count(self, table_index: int) -> int:
        """Получить количество строк в таблице"""
        try:
            return len(self.doc.tables[table_index].rows)
        except IndexError:
            return 0
    
    def _set_cell_background(self, cell, color: str):
        """
        Установка цвета фона ячейки
        
        Args:
            cell: Ячейка таблицы
            color: RGB цвет (hex строка, например "00FA02")
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def save(self, output_path: str):
        """
        Сохранение документа
        
        Args:
            output_path: Путь для сохранения результата
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        logger.info(f"Документ сохранен: {output_path}")


def create_template_from_report(report_path: str, output_path: str, 
                                placeholders: Dict[str, str]) -> bool:
    """
    Создание шаблона из существующего отчета
    
    Заменяет реальные значения на плейсхолдеры
    
    Args:
        report_path: Путь к существующему отчету
        output_path: Путь для сохранения шаблона
        placeholders: Словарь замен {реальное_значение: плейсхолдер}
        
    Returns:
        True если успешно, False иначе
    """
    try:
        doc = Document(report_path)
        
        # Обратная замена: значения -> плейсхолдеры
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for value, placeholder in placeholders.items():
                                if value in run.text:
                                    run.text = run.text.replace(value, placeholder)
        
        doc.save(output_path)
        logger.info(f"Шаблон создан: {output_path}")
        return True
    
    except Exception as e:
        logger.error(f"Ошибка создания шаблона: {e}")
        return False


# Примеры использования
if __name__ == "__main__":
    # Настройка логирования
    logging.basicConfig(level=logging.INFO)
    
    # Пример 1: Замена плейсхолдеров
    processor = TemplateProcessor("templates/docx/standard_template.docx")
    
    replacements = {
        '{{ДАТА}}': '27.01.2026',
        '{{ИМЯ}}': 'Влад',
        '{{ФАЙЛ_20_CENS}}': 'audio_20_cens.wav',
        '{{ХРОНО_20_CENS}}': '00:53:48.000',
        '{{LUFS_20_CENS}}': '-23.0 LUFS',
        '{{PEAK_20_CENS}}': '-1.0 dBTP',
        '{{LRA_20_CENS}}': '10.0 LU',
    }
    
    processor.replace_placeholders(replacements)
    processor.save("output/test_report.docx")
    
    # Пример 2: Дублирование строк для маркер-листа
    processor2 = TemplateProcessor("templates/docx/standard_template.docx")
    new_rows = processor2.duplicate_row(table_index=1, row_index=1, count=10)
    processor2.save("output/test_marker_list.docx")
