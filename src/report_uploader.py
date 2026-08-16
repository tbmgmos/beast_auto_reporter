"""
Report Uploader Module

Оркестрация отправки готового отчёта на Яндекс.Диск: определение целевой
папки (сериал/серия) по имени файла, поиск предыдущей версии того же
эпизода и краткое сравнение по маркерам/параметрам.
"""

from __future__ import annotations

import io
import json
import logging
import re
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from pathlib import Path


from src.app_paths import CONFIG_DIR, atomic_write_text, migrate_legacy_config_file
from src.file_matching import _bases_match
from src.marker_identity import read_marker_csv, read_marker_csv_bytes
from src.report_filename import ReportMeta, extract_date_loosely, parse_report_filename
from src.yandex_disk_client import YandexDiskClient, YandexDiskError

logger = logging.getLogger(__name__)

REPORTS_ROOT = "/отчеты"


class SeriesFolderNotFoundError(ValueError):
    """Папка сериала/фильма не найдена на Диске и create_if_missing=False.

    Наследуется от ValueError для обратной совместимости с уже
    существующими `pytest.raises(ValueError)`; вызывающий код (UI) должен
    ловить именно этот тип, чтобы не парсить текст сообщения об ошибке.
    """


def _folder_name_slug(name: str) -> str:
    return name.strip().lower().replace(" ", "_")


def _list_folder_if_exists(client: YandexDiskClient, folder_path: str) -> list[dict]:
    """list_folder, но 404 (папка ещё не создана) считается пустой папкой,

    а не ошибкой — так бывает для ещё не залитых ранее эпизодов/серий.
    """
    try:
        return client.list_folder(folder_path)
    except YandexDiskError as exc:
        if exc.status_code == 404:
            return []
        raise


def _folder_exists(client: YandexDiskClient, folder_path: str) -> bool:
    """В отличие от _list_folder_if_exists, различает «папка пустая» от

    «папки нет» — нужно только для проверки живости алиаса сериала.
    """
    try:
        client.list_folder(folder_path)
        return True
    except YandexDiskError as exc:
        if exc.status_code == 404:
            return False
        raise


SERIES_ALIASES_FILE = CONFIG_DIR / "series_aliases.json"
_LEGACY_SERIES_ALIASES_FILE = Path.home() / ".beast_auto_reporter_series_aliases.json"

# Алиасы .npr-проектов Nuendo (сезон -> папка) — отдельный файл от алиасов
# серий отчётов: разные корни на Диске и разная семантика ключа (см.
# fallback_series_key), пересечение по имени было бы совпадением, не связью.
NPR_ALIASES_FILE = CONFIG_DIR / "npr_aliases.json"


def _series_alias_key(series: str) -> str:
    """Нормализованный ключ алиаса — тот же слаг, что и для сравнения папок."""
    return _folder_name_slug(series)


_FALLBACK_KEY_VERSION_SUFFIX = re.compile(r"_\d{4}_\d{2}_\d{2}(_v\d+)?(_\d+)?$", re.IGNORECASE)


def fallback_series_key(docx_name: str) -> str:
    """Ключ алиаса для отчётов, чьё имя файла не распознал parse_report_filename

    (нет sNN/eNN-меток — например, преролл/рекламный ролик), а также для
    .npr-файлов проектов Nuendo (см. NPR_ALIASES_FILE), у которых своя,
    не привязанная к meta.series схема имени. В обоих случаях
    в качестве стабильного идентификатора (повторяющегося при повторных
    отправках того же ролика/проекта) используется само имя файла — без
    "отчет_"-префикса и расширения. Хвост вида "_2026_03_27_v1"/
    "_2026_03_27_v1_2" (дата + версия/дубль, та же форма, что
    REPORT_PATTERN отрезает у обычных отчётов) тоже отбрасывается — иначе
    каждая новая версия получала бы свой ключ, и ручной выбор папки
    никогда бы не запоминался на будущее.
    """
    stem = docx_name
    for ext in (".docx", ".pdf", ".csv", ".npr"):
        if stem.lower().endswith(ext):
            stem = stem[: -len(ext)]
            break
    if stem.startswith("отчет_"):
        stem = stem[len("отчет_"):]
    stem = _FALLBACK_KEY_VERSION_SUFFIX.sub("", stem)
    return stem


def load_series_aliases(path: Path = SERIES_ALIASES_FILE) -> dict:
    """key (нормализованный) -> подтверждённый путь папки сериала на Диске.

    Отсутствующий/битый файл — не ошибка, просто пустой словарь.
    """
    if path == SERIES_ALIASES_FILE:
        # Миграция только для реального пути по умолчанию — не трогаем
        # произвольные пути, которые передают тесты (tmp_path и т.п.).
        migrate_legacy_config_file(_LEGACY_SERIES_ALIASES_FILE, path)
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        logger.warning(f"Не удалось прочитать алиасы серий: {exc}")
        return {}
    if not isinstance(data, dict):
        return {}
    return {str(k): str(v) for k, v in data.items()}


def save_series_aliases(aliases: dict, path: Path = SERIES_ALIASES_FILE) -> None:
    try:
        atomic_write_text(
            path,
            json.dumps(aliases, ensure_ascii=False, indent=2, sort_keys=True),
        )
    except OSError as exc:
        logger.warning(f"Не удалось сохранить алиасы серий: {exc}")


def remember_series_alias(series: str, series_folder_path: str, path: Path = SERIES_ALIASES_FILE) -> None:
    """Записывает/обновляет подтверждённое сопоставление сериал -> папка."""
    aliases = load_series_aliases(path)
    aliases[_series_alias_key(series)] = series_folder_path
    save_series_aliases(aliases, path)


def forget_series_alias(series: str, path: Path = SERIES_ALIASES_FILE) -> None:
    """Убирает более не действительный алиас (папка удалена/переименована

    на Диске мимо приложения) — используется при обнаружении 404.
    """
    aliases = load_series_aliases(path)
    key = _series_alias_key(series)
    if key in aliases:
        del aliases[key]
        save_series_aliases(aliases, path)


def list_series_aliases(path: Path = SERIES_ALIASES_FILE) -> list[tuple[str, str]]:
    """Список (ключ, путь_папки) алиасов серий, отсортированный по ключу —

    для отображения в UI-диалоге управления алиасами.
    """
    return sorted(load_series_aliases(path).items())


def alias_key_for_path(path: str, aliases_path: Path = SERIES_ALIASES_FILE) -> str | None:
    """Обратный поиск: есть ли алиас, указывающий именно на этот путь.

    Используется для подсказки в дереве просмотрщика Диска — «эта папка
    уже привязана к серии/проекту такому-то», без похода в отдельный
    диалог управления алиасами.
    """
    for key, aliased_path in load_series_aliases(aliases_path).items():
        if aliased_path == path:
            return key
    return None


VARIANT_OVERRIDES_FILE = CONFIG_DIR / "variant_overrides.json"


def load_variant_overrides(path: Path = VARIANT_OVERRIDES_FILE) -> dict:
    """remote_path -> вручную назначенный вариант ("ME"/"AD"/"VO"/"MAIN").

    Назначается через ПКМ в просмотрщике Диска — на случай, когда
    автоопределение по имени папки (см. YandexDiskBrowserDialog._report_variant)
    ошибается или не справляется вовсе (нестандартное имя без узнаваемых
    маркеров). "MAIN" — явное «без варианта», отличается от отсутствия
    записи вообще («Авто» — работает автоопределение). Отсутствующий/битый
    файл — не ошибка, просто пустой словарь.
    """
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        logger.warning(f"Не удалось прочитать ручные назначения типа отчёта: {exc}")
        return {}
    if not isinstance(data, dict):
        return {}
    return {str(k): str(v) for k, v in data.items()}


def save_variant_overrides(overrides: dict, path: Path = VARIANT_OVERRIDES_FILE) -> None:
    try:
        atomic_write_text(
            path,
            json.dumps(overrides, ensure_ascii=False, indent=2, sort_keys=True),
        )
    except OSError as exc:
        logger.warning(f"Не удалось сохранить ручные назначения типа отчёта: {exc}")


def set_variant_override(remote_path: str, variant: str | None, path: Path = VARIANT_OVERRIDES_FILE) -> None:
    """variant=None ("Авто") убирает override — снова работает автоопределение."""
    overrides = load_variant_overrides(path)
    if variant is None:
        overrides.pop(remote_path, None)
    else:
        overrides[remote_path] = variant
    save_variant_overrides(overrides, path)


UPLOADED_REPORTS_FILE = CONFIG_DIR / "uploaded_reports.json"
_MAX_REMEMBERED_UPLOADS = 20


def load_uploaded_reports(path: Path = UPLOADED_REPORTS_FILE) -> list[dict]:
    """Список последних успешно отправленных на Диск отчётов (новые сверху) —

    {"local_folder": str, "remote_path": str, "uploaded_at": ISO-строка}.
    Отсутствующий/битый файл — не ошибка, просто пустой список. Используется
    для тихой проверки целостности при старте приложения (не пропала ли
    папка с Диска мимо приложения).
    """
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        logger.warning(f"Не удалось прочитать список отправленных отчётов: {exc}")
        return []
    if not isinstance(data, list):
        return []
    return data


def save_uploaded_reports(entries: list[dict], path: Path = UPLOADED_REPORTS_FILE) -> None:
    try:
        atomic_write_text(path, json.dumps(entries, ensure_ascii=False, indent=2))
    except OSError as exc:
        logger.warning(f"Не удалось сохранить список отправленных отчётов: {exc}")


def remember_uploaded_report(local_folder: str, remote_path: str, path: Path = UPLOADED_REPORTS_FILE) -> None:
    """Запоминает успешно отправленный отчёт (новый — в начало списка),

    храня не более _MAX_REMEMBERED_UPLOADS последних записей.
    """
    from datetime import datetime, timezone

    entries = load_uploaded_reports(path)
    entries = [e for e in entries if e.get("remote_path") != remote_path]
    entries.insert(0, {
        "local_folder": local_folder,
        "remote_path": remote_path,
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    })
    save_uploaded_reports(entries[:_MAX_REMEMBERED_UPLOADS], path)


def forget_uploaded_reports(entries_to_forget: list[dict], path: Path = UPLOADED_REPORTS_FILE) -> None:
    """Убирает записи (по remote_path) из списка отправленных отчётов.

    Вызывается после уведомления «отчёты пропали с Диска»: без этого те же
    протухшие записи находились бы проверкой целостности заново при каждом
    запуске, и одно и то же уведомление всплывало бы каждую сессию.
    """
    forget_paths = {e.get("remote_path") for e in entries_to_forget}
    entries = load_uploaded_reports(path)
    remaining = [e for e in entries if e.get("remote_path") not in forget_paths]
    if len(remaining) != len(entries):
        save_uploaded_reports(remaining, path)


_EPISODE_FOLDER_NAME = re.compile(r"^e\d{2}$")


def resolve_manual_pick_target(picked_path: str, meta: ReportMeta | None) -> tuple[str, str]:
    """(episode_path, series_path) для вручную выбранной пользователем папки.

    Если meta не распознана (имя файла не разобралось) — единственный
    доступный вариант: пользователь указал непосредственно целевую папку,
    структуру определить нечем; episode_path == series_path == picked_path.

    Иначе: если выбранная папка называется "eNN" (в форме папки эпизода
    вообще, не обязательно с номером ИМЕННО этого отчёта — пользователь
    мог по ошибке выбрать папку другого, ранее созданного эпизода) —
    папкой сериала считается её родитель, а не сама она: это защищает от
    "матрёшки" (papка_нового_эпизода внутри papки_старого_эпизода), если
    пользователь при повторной отправке выбрал уже существующую эпизодную
    папку вместо папки сериала. В остальных случаях выбранная папка
    считается папкой сериала, а папка эпизода — её подпапкой "eNN"
    (создаётся вызывающим кодом при необходимости — mkdir идемпотентен).

    Известное ограничение: эвристика определяет "похоже на папку эпизода"
    чисто по имени (regex "eNN"). Если у пользователя реально существует
    папка, буквально названная так не по смыслу "эпизод" (например,
    сокращение, случайное совпадение), функция ошибочно сочтёт её папкой
    эпизода и возьмёт её родителя за папку сериала. Это осознанный
    компромисс (такие имена крайне маловероятны в реальной структуре
    /отчеты/...), а не баг.
    """
    if meta is None:
        return picked_path, picked_path
    expected_name = f"e{meta.episode:02d}"
    picked_name = picked_path.rstrip("/").rsplit("/", 1)[-1]
    if _EPISODE_FOLDER_NAME.match(picked_name):
        series_path = picked_path.rsplit("/", 1)[0] if "/" in picked_path else picked_path
        if picked_name == expected_name:
            return picked_path, series_path
        return f"{series_path}/{expected_name}", series_path
    return f"{picked_path}/{expected_name}", picked_path


def upload_folder(client: YandexDiskClient, local_folder: Path, remote_folder: str,
                   progress_callback=None) -> list[str]:
    """Загружает все файлы из локальной папки отчёта в папку на Диске.

    Не рекурсивно — папка отчёта на выходе плоская (docx, pdf, csv и т.п.).
    Возвращает список путей загруженных файлов на Диске.

    progress_callback(bytes_sent, total_bytes), если передан, отражает
    совокупный прогресс по всем файлам папки, а не по одному файлу.
    """
    # Скрытые файлы (.DS_Store и прочие дотфайлы Finder/редакторов) и
    # lock-файлы MS Office (~$отчет_....docx, создаются, пока документ
    # открыт в Word) — служебный мусор, ему нечего делать в папке отчёта
    # на Диске.
    entries = [
        e for e in sorted(Path(local_folder).iterdir())
        if e.is_file() and not e.name.startswith(".") and not e.name.startswith("~$")
    ]
    total_bytes = sum(e.stat().st_size for e in entries) or 1
    uploaded = []
    sent_before = 0
    for entry in entries:
        remote_path = f"{remote_folder}/{entry.name}"
        file_sent_before = sent_before

        def _on_file_progress(sent, _total, _base=file_sent_before):
            progress_callback(_base + sent, total_bytes)

        client.upload(entry, remote_path, progress_callback=_on_file_progress if progress_callback else None)
        uploaded.append(remote_path)
        sent_before += entry.stat().st_size
    return uploaded


def upload_paths_recursive(client: YandexDiskClient, local_paths: list[Path], remote_folder: str,
                            progress_callback=None) -> None:
    """Загружает произвольный список локальных файлов и папок в remote_folder,

    рекурсивно (в отличие от upload_folder, который специально не
    рекурсивен — папка готового отчёта плоская). Используется для
    перетаскивания из Finder, который отдаёт смешанный список файлов и
    папок одним drop'ом. Недостающие подпапки на Диске создаются по мере
    необходимости (mkdir идемпотентен, но каждая папка создаётся не более
    одного раза за вызов).

    progress_callback(done, total), если передан, считает файлы (не байты).
    """
    all_files = []  # [(local_file_path, relative_path_from_remote_folder)]
    for raw_path in local_paths:
        path = Path(raw_path)
        if path.is_dir():
            for f in sorted(path.rglob("*")):
                if f.is_file() and not f.name.startswith(".") and not f.name.startswith("~$"):
                    all_files.append((f, Path(path.name) / f.relative_to(path)))
        elif path.is_file() and not path.name.startswith(".") and not path.name.startswith("~$"):
            all_files.append((path, Path(path.name)))

    total = len(all_files) or 1
    created_dirs = set()

    def _ensure_dir(remote_dir: str) -> None:
        # mkdir Диска требует существования родителя (см. docstring mkdir) —
        # создаём уровни по порядку от корня, а не сразу самый глубокий.
        if remote_dir in created_dirs:
            return
        parent = remote_dir.rsplit("/", 1)[0]
        if parent and parent != remote_folder:
            _ensure_dir(parent)
        client.mkdir(remote_dir)
        created_dirs.add(remote_dir)

    for done, (local_file, rel_path) in enumerate(all_files, start=1):
        rel_parts = rel_path.parts[:-1]
        if rel_parts:
            remote_dir = "/".join([remote_folder, *rel_parts])
            _ensure_dir(remote_dir)
            remote_path = f"{remote_dir}/{rel_path.name}"
        else:
            remote_path = f"{remote_folder}/{rel_path.name}"
        client.upload(local_file, remote_path)
        if progress_callback:
            progress_callback(done, total)


def find_series_folder(
    client: YandexDiskClient, series_slug: str, root: str = REPORTS_ROOT,
    *, roots: list[str] = None, aliases_path: Path = SERIES_ALIASES_FILE,
) -> str | None:
    """Ищет папку сериала, похожую на series_slug, во всех корневых папках.

    Сначала проверяется ранее подтверждённый алиас (сериал -> папка) — если
    он есть и папка всё ещё существует, возвращается сразу, без нечёткого
    скана корней. Протухший алиас (папка удалена/переименована мимо
    приложения) забывается, и поиск продолжается по прежней нечёткой логике,
    перебором всех корней по порядку (первое совпадение побеждает).
    `roots`, если передан, используется вместо одиночного `root` (который
    остаётся для обратной совместимости — старый однокорневой вызов).
    Возвращает полный путь найденной папки или None, если совпадений нет.
    """
    search_roots = list(roots) if roots else [root]

    aliases = load_series_aliases(aliases_path)
    aliased_path = aliases.get(_series_alias_key(series_slug))
    if aliased_path is not None:
        if _folder_exists(client, aliased_path):
            logger.info("find_series_folder: «%s» -> алиас «%s»", series_slug, aliased_path)
            return aliased_path
        logger.info("find_series_folder: алиас «%s» -> «%s» устарел, ищем заново", series_slug, aliased_path)
        forget_series_alias(series_slug, aliases_path)

    target = _folder_name_slug(series_slug)
    for search_root in search_roots:
        items = _list_folder_if_exists(client, search_root)
        for item in items:
            if item.get("type") != "dir":
                continue
            name = item.get("name", "")
            if _bases_match(_folder_name_slug(name), target):
                found = item.get("path") or f"{search_root}/{name}"
                logger.info("find_series_folder: «%s» -> найдено «%s»", series_slug, found)
                return found
    logger.info("find_series_folder: «%s» -> не найдено среди корней %s", series_slug, search_roots)
    return None


def resolve_target_path(
    client: YandexDiskClient,
    meta: ReportMeta,
    *,
    create_if_missing: bool = False,
    root: str = REPORTS_ROOT,
    roots: list[str] = None,
    aliases_path: Path = SERIES_ALIASES_FILE,
) -> tuple[str, bool]:
    """Определяет путь папки серии на Диске для данного отчёта.

    Возвращает (путь_к_папке_серии, папка_сериала_создана_заново).
    Если папка сериала не найдена и create_if_missing=False, поднимает
    ValueError — вызывающий код (UI) должен подтвердить создание новой
    папки перед повторным вызовом с create_if_missing=True.
    Найденный или созданный путь запоминается как алиас сериала — следующий
    отчёт того же сериала находит папку мгновенно, без повторного поиска.
    При создании новой папки сериала (не найдена ни в одном из корней)
    используется первый корень из `roots` — остальные считаются
    дополнительными местами для поиска, не для создания.
    """
    search_roots = list(roots) if roots else [root]
    series_path = find_series_folder(client, meta.series, roots=search_roots, aliases_path=aliases_path)
    created = False

    if series_path is None:
        if not create_if_missing:
            raise SeriesFolderNotFoundError(f"Папка для «{meta.series}» не найдена на Диске")
        create_root = search_roots[0]
        series_path = f"{create_root}/{meta.series}"
        client.mkdir(create_root)
        client.mkdir(series_path)
        created = True

    remember_series_alias(meta.series, series_path, aliases_path)

    episode_folder = f"e{meta.episode:02d}"
    episode_path = f"{series_path}/{episode_folder}"
    client.mkdir(episode_path)

    return episode_path, created


def _is_report_entry(item: dict) -> bool:
    """Подпапка отчёта (текущая структура) или .docx-файл прямо в папке

    эпизода (старая структура, до введения подпапок) — поддерживаем обе,
    чтобы не терять уже загруженные ранее отчёты.
    """
    name = item.get("name", "")
    if not name.startswith("отчет_"):
        return False
    if item.get("type") == "dir":
        return True
    return item.get("type") == "file" and name.lower().endswith(".docx")


def list_report_versions(
    client: YandexDiskClient,
    folder_path: str,
    meta: ReportMeta = None,
) -> list[dict]:
    """Возвращает все версии отчёта в папке, отсортированные от старой к новой.

    Каждый элемент: {"path": str, "label": str, "date": date | None}.
    Если meta передан — фильтрует по season/episode (авто-маршрут по имени
    файла); без meta — берёт все версии отчёта в папке как есть (папка уже
    выбрана вручную, дополнительная фильтрация не нужна).
    Отчёт может быть подпапкой «отчет_...» (текущая структура) либо, для
    уже загруженных ранее файлов, отдельным .docx (старая структура).
    """
    items = _list_folder_if_exists(client, folder_path)
    logger.info(
        "list_report_versions: %s -> %d элементов: %s",
        folder_path, len(items), [item.get("name") for item in items],
    )

    versions = []
    for item in items:
        if not _is_report_entry(item):
            continue
        name = item.get("name", "")
        candidate_meta = parse_report_filename(name)
        if meta is not None:
            if candidate_meta is None:
                logger.info("list_report_versions: не распознано имя «%s»", name)
                continue
            if candidate_meta.season != meta.season or candidate_meta.episode != meta.episode:
                continue
            if candidate_meta.variant != meta.variant:
                # Один эпизод может лежать в одной папке вместе с ME-версией
                # (variant="MnE" в имени) — это не версии друг друга, а
                # параллельные отдельные отчёты, их нельзя сравнивать/путать
                # местами при автоподборе «предыдущей версии».
                continue
        # Если имя не разобрал REPORT_PATTERN целиком (нестандартная схема —
        # другой формат сезона/эпизода, составной тег вроде "cens_AD" и
        # т.п.), не сдаёмся сразу на дату ЗАГРУЗКИ на Диск — сперва пробуем
        # найти дату НАПИСАНИЯ отчёта прямо в имени лёгким сканированием.
        # Иначе только что загруженный (но написанный давно) отчёт с
        # нестандартным именем сортировался бы как самый новый по времени
        # загрузки, хотя по содержанию — самый старый (реальный регресс:
        # "программа не определяет где новый и где старый отчёт").
        date_value = candidate_meta.date if candidate_meta else extract_date_loosely(name)
        sort_key = date_value.isoformat() if date_value else item.get("modified", "")
        versions.append({
            "path": item.get("path") or f"{folder_path}/{name}",
            "label": name,
            "date": date_value,
            "_sort_key": sort_key,
        })

    versions.sort(key=lambda v: v["_sort_key"])
    for version in versions:
        del version["_sort_key"]
    return versions


def group_versions_by_category(versions: list, overrides: dict = None) -> dict:
    """Группирует версии отчёта (см. list_report_versions) по категории

    варианта (см. report_filename.categorize_variant/extract_variant_loosely)
    — "main"/"me"/"vo"/"dub"/"other". overrides — {remote_path: "ME"/"AD"/
    .../"MAIN"}, ручные назначения через ПКМ в просмотрщике Диска (см.
    load_variant_overrides) — имеют приоритет над автоопределением по имени.

    Один эпизод нередко лежит в одной папке вместе с ME/AD/VO-версией того
    же отчёта — это не версии друг друга, а параллельные отдельные отчёты
    (см. YandexDiskBrowserDialog._label_report_version_siblings). Группировка
    — через лёгкое сканирование имени (extract_variant_loosely), а не
    строгий REPORT_PATTERN целиком: многие реальные папки не совпадают со
    строгой схемой вообще (другой формат даты, составной тег вроде
    "cens_AD" и т.п.), и раньше list_report_versions с заданным meta молча
    отбрасывал такие версии вместо того, чтобы дать им шанс на сравнение
    (см. YandexDiskFindVersionsThread — та же категория, что и meta.variant
    у текущего черновика, теперь ищется среди ВСЕХ версий папки, а не
    только среди тех, что совпали со строгим парсером).
    """
    from src.report_filename import categorize_variant, extract_variant_loosely

    overrides = overrides or {}
    groups: dict = {}
    for version in versions:
        override = overrides.get(version.get("path", ""))
        variant = (None if override == "MAIN" else override) if override is not None \
            else extract_variant_loosely(version.get("label", ""))
        category = categorize_variant(variant)
        groups.setdefault(category, []).append(version)
    return groups


def find_previous_report(
    client: YandexDiskClient,
    folder_path: str,
    meta: ReportMeta,
) -> str | None:
    """Путь самой свежей версии отчёта того же сезона/эпизода, либо None."""
    versions = list_report_versions(client, folder_path, meta)
    return versions[-1]["path"] if versions else None


def find_latest_report_in_folder(client: YandexDiskClient, folder_path: str) -> str | None:
    """Путь самой свежей версии отчёта в указанной папке (без фильтра по

    season/episode — используется, когда папка выбрана вручную), либо None.
    """
    versions = list_report_versions(client, folder_path)
    return versions[-1]["path"] if versions else None


@dataclass
class DocumentSummary:
    marker_count: int = 0
    blocker_count: int = 0
    new_marker_count: int = 0
    parameters: dict = field(default_factory=dict)  # {"2.0 cens": {"LOUDNESS": "-23.1 LUFS", ...}, ...}
    parameter_status: dict = field(default_factory=dict)  # {"2.0 cens": {"LOUDNESS": "bad"|"warn"}, ...}
    # Содержимое строк MARKER LIST для попозиционного diff'а между версиями:
    # [{"tc_in", "tc_out", "description", "blocker": bool, "comments"}, ...]
    markers: list = field(default_factory=list)


# Цвета заливки, которыми ExactReportGenerator подсвечивает несоответствие
# норме в технической таблице (см. src/exact_report_generator.py: BAD_BG/WARN_BG).
_BAD_CELL_FILL = "E73322"
_WARN_CELL_FILL = "FCBA19"


def _cell_status(cell) -> str | None:
    """Статус ячейки по её заливке в отчёте: "bad" (красная), "warn"

    (жёлтая) или None (в норме/не размечена) — переиспользуем цветовую
    маркировку, которую уже расставляет генератор отчёта, вместо того
    чтобы заново пересчитывать пороги (target_lufs/lufs_tolerance и т.п.).
    """
    from docx.oxml.ns import qn

    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = (shd.get(qn("w:fill")) or "").upper()
    if fill == _BAD_CELL_FILL:
        return "bad"
    if fill == _WARN_CELL_FILL:
        return "warn"
    return None


def _iter_all_tables(doc) -> list:
    """Все таблицы документа на любой глубине вложенности, включая

    таблицы внутри текстовых блоков/фигур (w:txbxContent) — именно так
    экспортирует таблицы Pages при сохранении в .docx. Штатный doc.tables
    возвращает только таблицы верхнего уровня тела документа и полностью
    пропускает вложенные, из-за чего у старых отчётов (и таблица маркеров,
    и техническая таблица параметров лежат в текстовых блоках) сравнение
    не видело вообще никаких данных и показывало 0 маркеров / «без
    изменений» по параметрам, хотя данные в документе есть.
    """
    from docx.oxml.ns import qn
    from docx.table import Table

    tables = []
    seen = set()
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        if id(tbl_el) in seen:
            continue
        seen.add(id(tbl_el))
        tables.append(Table(tbl_el, doc))
    return tables


def _summarize_document(doc) -> DocumentSummary:
    """Считает маркеры/блокеров/новых маркеров и снимает значения технических

    параметров (LOUDNESS/TRUE PEAK/LRA и т.п. по дорожкам) из таблиц отчёта.
    """
    summary = DocumentSummary()
    all_tables = _iter_all_tables(doc)

    for table in all_tables:
        if not table.rows:
            continue
        # Строка заголовка ищется сканированием (как и в таблице параметров
        # ниже), а не берётся жёстко row[0]: в старом формате отчёта первая
        # строка — объединённый титул «MARKER LIST» на всю ширину, а сами
        # заголовки колонок («Timecode In»/«Timecode Out»/…) лежат во второй
        # строке. Новые отчёты приложения кладут заголовки прямо в row[0].
        # Раньше проверялась только row[0] и притом регистрозависимо, поэтому
        # для старых отчётов таблица маркеров молча не находилась вообще —
        # сравнение показывало 0 → 0 маркеров/блокеров для обеих версий сразу.
        header_row_idx = None
        header_cells = None
        for idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if "TIMECODE IN" in {text.upper() for text in cells}:
                header_row_idx = idx
                header_cells = cells
                break
        if header_row_idx is None:
            continue

        # Полностью пустые строки после заголовка не считаем маркерами:
        # в старом формате (экспорт Pages) под заголовком таблицы маркеров
        # часто лежит пустая строка-заглушка — без фильтра «чистый» отчёт
        # без замечаний показывал бы 1 маркер вместо 0.
        data_rows = [
            row for row in table.rows[header_row_idx + 1:]
            if any(cell.text.strip() for cell in row.cells)
        ]
        summary.marker_count = len(data_rows)

        blocker_idx = next((i for i, text in enumerate(header_cells) if text.upper() == "БЛОКЕР"), None)
        if blocker_idx is not None:
            summary.blocker_count = sum(
                1 for row in data_rows if blocker_idx < len(row.cells) and row.cells[blocker_idx].text.strip() == "*"
            )

        comments_idx = next((i for i, text in enumerate(header_cells) if text.upper() == "КОММЕНТАРИИ"), None)
        if comments_idx is not None:
            summary.new_marker_count = sum(
                1 for row in data_rows
                if comments_idx < len(row.cells)
                and "НОВЫЙ МАРКЕР" in row.cells[comments_idx].text.strip().upper()
            )

        def _column_index(*names):
            for i, text in enumerate(header_cells):
                if text.upper() in names:
                    return i
            return None

        tc_in_idx = _column_index("TIMECODE IN")
        tc_out_idx = _column_index("TIMECODE OUT")
        description_idx = _column_index("DESCRIPTION", "ОПИСАНИЕ ПРОБЛЕМЫ", "ОПИСАНИЕ")
        marker_id_idx = _column_index("ID")

        def _cell_text(row, idx):
            if idx is None or idx >= len(row.cells):
                return ""
            return row.cells[idx].text.strip()

        summary.markers = [
            {
                "id": _cell_text(row, marker_id_idx),
                "tc_in": _cell_text(row, tc_in_idx),
                "tc_out": _cell_text(row, tc_out_idx),
                "description": _cell_text(row, description_idx),
                "blocker": _cell_text(row, blocker_idx) == "*",
                "comments": _cell_text(row, comments_idx),
            }
            for row in data_rows
        ]
        break

    for table in all_tables:
        header_row_idx = None
        header_cells = None
        for idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if "ДОРОЖКА" in cells and any(peak in cells for peak in ("TRUE PEAK", "SAMPLE PEAK")):
                header_row_idx = idx
                header_cells = cells
                break
        if header_row_idx is None:
            continue
        for row in table.rows[header_row_idx + 1:]:
            cells = row.cells
            texts = [cell.text.strip() for cell in cells]
            if not texts or not texts[0]:
                break
            label = texts[0]
            summary.parameters[label] = {}
            summary.parameter_status[label] = {}
            for header, cell, text in zip(header_cells[1:], cells[1:], texts[1:]):
                if not header or not text:
                    continue
                summary.parameters[label][header] = text
                status = _cell_status(cell)
                if status:
                    summary.parameter_status[label][header] = status
        break

    return summary


# Технические заголовки колонок -> человекочитаемые названия для UI.
_PARAMETER_LABELS = {
    "LOUDNESS": "Громкость",
    "TRUE PEAK": "True Peak",
    "SAMPLE PEAK": "Sample Peak",
    "LRA": "Диапазон (LRA)",
    "ХРОНОМЕТРАЖ": "Хронометраж",
    "ФОРМАТ ФАЙЛА": "Формат файла",
    "НАЗВАНИЕ ФАЙЛОВ": "Файл",
}


def _humanize_parameter_name(name: str) -> str:
    return _PARAMETER_LABELS.get(name.strip().upper(), name.strip().capitalize())


def _format_parameter_changes(old_params: dict, new_params: dict, new_status: dict = None) -> list[dict]:
    """Список изменившихся дорожек:

    [{"label": "2.0 cens", "changes": [
        {"field": "Громкость", "old": "-23.0 LUFS", "new": "-24.5 LUFS", "status": "bad"|"warn"|None},
        ...
    ]}, ...]

    "status" — читается из заливки ячейки НОВОГО значения в самом отчёте
    (та же цветовая маркировка несоответствия норме, что и в MARKER LIST/
    технической таблице), чтобы не пересчитывать пороги заново.
    """
    new_status = new_status or {}
    result = []
    for label in sorted(set(old_params) | set(new_params)):
        old_row = old_params.get(label)
        new_row = new_params.get(label)
        if old_row is None:
            result.append({"label": label, "changes": [{"field": "Дорожка", "old": "—", "new": "добавлена", "status": None}]})
            continue
        if new_row is None:
            result.append({"label": label, "changes": [{"field": "Дорожка", "old": "была", "new": "удалена", "status": None}]})
            continue
        label_status = new_status.get(label, {})
        changes = [
            {
                "field": _humanize_parameter_name(field_name),
                "old": old_row.get(field_name, "—"),
                "new": new_row.get(field_name, "—"),
                "status": label_status.get(field_name),
            }
            for field_name in sorted(set(old_row) | set(new_row))
            if old_row.get(field_name) != new_row.get(field_name)
        ]
        if changes:
            result.append({"label": label, "changes": changes})
    return result


def diff_markers(old_markers: list, new_markers: list) -> dict:
    """Попозиционный diff маркеров двух версий отчёта.

    Постоянный Beast ID — главный ключ. Для старых отчётов без ID
    fallback — Timecode In + номер дубля.
    """
    old_ids = {marker.get("id", "").strip() for marker in old_markers if marker.get("id", "").strip()}
    new_ids = {marker.get("id", "").strip() for marker in new_markers if marker.get("id", "").strip()}
    shared_ids = old_ids & new_ids

    def _keyed(markers):
        occurrences: dict = {}
        result = {}
        for marker in markers:
            marker_id = marker.get("id", "").strip()
            if marker_id in shared_ids:
                key = ("id", marker_id)
            else:
                tc = marker.get("tc_in", "")
                occurrences[tc] = occurrences.get(tc, 0) + 1
                key = ("tc", tc, occurrences[tc])
            # Повтор ID в повреждённом/ручном CSV не затирает строку.
            if key in result:
                tc = marker.get("tc_in", "")
                occurrences[tc] = occurrences.get(tc, 0) + 1
                key = ("tc", tc, occurrences[tc])
            result[key] = marker
        return result

    old_by_key = _keyed(old_markers)
    new_by_key = _keyed(new_markers)
    added = [marker for key, marker in new_by_key.items() if key not in old_by_key]
    removed = [marker for key, marker in old_by_key.items() if key not in new_by_key]

    field_labels = [
        ("tc_in", "Timecode In"),
        ("description", "Описание"),
        ("tc_out", "Timecode Out"),
        ("blocker", "Блокер"),
        ("comments", "Комментарии"),
    ]

    def _display(field_name, value):
        if field_name == "blocker":
            return "да" if value else "нет"
        return value or "—"

    changed = []
    for key, new_marker in new_by_key.items():
        old_marker = old_by_key.get(key)
        if old_marker is None:
            continue
        field_changes = [
            {
                "field": label,
                "old": _display(field_name, old_marker.get(field_name)),
                "new": _display(field_name, new_marker.get(field_name)),
            }
            for field_name, label in field_labels
            if old_marker.get(field_name) != new_marker.get(field_name)
        ]
        if field_changes:
            changed.append({
                "id": new_marker.get("id", ""),
                "tc_in": new_marker.get("tc_in", ""),
                "changes": field_changes,
            })

    return {"added": added, "removed": removed, "changed": changed}


@dataclass
class ReportComparison:
    marker_count_old: int
    marker_count_new: int
    blocker_count_old: int
    blocker_count_new: int
    new_marker_count_old: int
    new_marker_count_new: int
    parameter_changes: list
    # Результат diff_markers (added/removed/changed) либо None, если
    # таблицы маркеров не найдены ни в одной из версий.
    marker_diff: dict = None
    # Снимок технических параметров НОВОЙ версии (label дорожки -> {параметр:
    # значение}) и их статусов ("bad"/"warn") — не только diff, а текущее
    # состояние как есть. Используется LLM-сводкой, чтобы говорить и о том,
    # что не менялось (хронометраж, формат файла, текущая громкость и т.п.).
    new_parameters: dict = None
    new_parameter_status: dict = None


def _find_docx_in_report_folder(client: YandexDiskClient, report_folder_path: str) -> str | None:
    """Находит файл отчёта (.docx) внутри подпапки отчёта.

    Предпочитает файл с префиксом «отчет_» (текущее соглашение об имени),
    но если такого нет — берёт любой .docx как есть (кроме lock-файлов
    MS Office «~$...» и скрытых файлов). Нужно для старых папок,
    загруженных ещё ДО того, как это соглашение появилось — там и папка,
    и сам .docx внутри могут быть названы без префикса (реальный случай:
    «one_last_sin_s01_e01_Master_uncens_2025_05_14/one_last_sin_..._2025_05_14.docx»),
    а строгое требование префикса у файла давало «Не удалось прочитать
    выбранную версию отчёта» на совершенно реальных, читаемых версиях.
    """
    items = _list_folder_if_exists(client, report_folder_path)
    logger.info(
        "_find_docx_in_report_folder: %s -> %d элементов: %s",
        report_folder_path, len(items), [(i.get("name"), i.get("type")) for i in items],
    )
    fallback_path = None
    for item in items:
        name = item.get("name", "")
        if item.get("type") != "file" or not name.lower().endswith(".docx"):
            continue
        if name.startswith("~$") or name.startswith("."):
            continue
        path = item.get("path") or f"{report_folder_path}/{name}"
        if name.startswith("отчет_"):
            return path
        if fallback_path is None:
            fallback_path = path
    return fallback_path


def _find_csv_in_report_folder(client: YandexDiskClient, report_folder_path: str) -> str | None:
    """Find the marker-list CSV stored beside a report DOCX."""
    if report_folder_path.lower().endswith(".csv"):
        return report_folder_path
    if report_folder_path.lower().endswith(".docx"):
        return None
    for item in _list_folder_if_exists(client, report_folder_path):
        name = item.get("name", "")
        if (
            item.get("type") == "file"
            and name.lower().endswith(".csv")
            and not name.startswith((".", "~$"))
        ):
            return item.get("path") or f"{report_folder_path}/{name}"
    return None


def _find_local_marker_csv(docx_path: Path) -> Path | None:
    candidates = sorted(
        path for path in docx_path.parent.glob("*.csv")
        if not path.name.startswith((".", "~$"))
    )
    return candidates[0] if candidates else None


def _replace_summary_markers_from_csv(summary: DocumentSummary, markers: list[dict]) -> None:
    """Prefer original marker-list data while retaining DOCX parameters."""
    summary.markers = markers
    summary.marker_count = len(markers)
    summary.blocker_count = sum(1 for marker in markers if marker.get("blocker"))
    summary.new_marker_count = sum(
        1 for marker in markers
        if "НОВЫЙ МАРКЕР" in (marker.get("comments") or "").upper()
    )


def _resolve_docx_path(client: YandexDiskClient, report_path: str) -> str | None:
    """report_path — путь к подпапке отчёта (текущая структура) либо

    напрямую к .docx-файлу (старая плоская структура). Возвращает путь
    к .docx или None, если файл не найден.
    """
    if report_path.lower().endswith(".docx"):
        return report_path
    return _find_docx_in_report_folder(client, report_path)


def _build_comparison(old_summary: DocumentSummary, new_summary: DocumentSummary) -> ReportComparison:
    marker_diff = None
    if old_summary.markers or new_summary.markers:
        marker_diff = diff_markers(old_summary.markers, new_summary.markers)
    return ReportComparison(
        marker_count_old=old_summary.marker_count,
        marker_count_new=new_summary.marker_count,
        blocker_count_old=old_summary.blocker_count,
        blocker_count_new=new_summary.blocker_count,
        new_marker_count_old=old_summary.new_marker_count,
        new_marker_count_new=new_summary.new_marker_count,
        parameter_changes=_format_parameter_changes(
            old_summary.parameters, new_summary.parameters, new_summary.parameter_status
        ),
        marker_diff=marker_diff,
        new_parameters=new_summary.parameters,
        new_parameter_status=new_summary.parameter_status,
    )


# Обрезка длинных списков маркеров в брифе для LLM: на отчёте со 100+
# маркерами полный перечень съест весь контекст модели, а для сводки
# «что поменялось» хватает представительной выборки + счётчиков.
_VERSION_BRIEF_MAX_MARKERS_PER_GROUP = 30

# status из _format_parameter_changes (заливка ячейки в отчёте) ->
# явная пометка для LLM, чтобы нарушения нормы не потерялись в сводке.
_BRIEF_PARAMETER_STATUS_HINTS = {"bad": "НЕ В НОРМЕ", "warn": "на грани нормы"}


def _brief_marker_line(marker: dict) -> str:
    description = marker.get("description") or "—"
    blocker_suffix = " (блокер)" if marker.get("blocker") else ""
    return f"- {marker.get('tc_in', '')} — {description}{blocker_suffix}"


def format_comparison_brief(
    comparison: ReportComparison,
    old_label: str = None,
    new_label: str = None,
    max_markers_per_group: int = _VERSION_BRIEF_MAX_MARKERS_PER_GROUP,
) -> str:
    """Компактный текстовый бриф различий двух версий отчёта — вход для
    LLM-промпта сводки (см. ConclusionGenerator.summarize_version_changes).

    Чистая функция: без Qt, без сети — легко тестируется. Детальные группы
    (добавленные/удалённые/изменённые маркеры, параметры) обрезаются до
    max_markers_per_group строк с хвостом «…и ещё N» — счётчики в заголовке
    при этом всегда полные.
    """
    header = []
    if old_label:
        header.append(f"Старая версия: {old_label}")
    if new_label:
        header.append(f"Новая версия: {new_label}")
    header.append(
        f"Маркеров: {comparison.marker_count_old} → {comparison.marker_count_new} "
        f"(отмечено «новый маркер»: {comparison.new_marker_count_old} → {comparison.new_marker_count_new}), "
        f"блокеров: {comparison.blocker_count_old} → {comparison.blocker_count_new}."
    )

    # Текущее состояние новой версии (не diff): даёт LLM материал говорить
    # и о том, что НЕ менялось — хронометраж, формат файла, текущая громкость.
    body = []
    new_parameters = comparison.new_parameters or {}
    if new_parameters:
        body.append("Текущие технические параметры новой версии:")
        new_status = comparison.new_parameter_status or {}
        for label in sorted(new_parameters):
            label_status = new_status.get(label, {})
            parts = []
            for param, value in new_parameters[label].items():
                hint = _BRIEF_PARAMETER_STATUS_HINTS.get(label_status.get(param))
                parts.append(f"{param}: {value} ({hint})" if hint else f"{param}: {value}")
            body.append(f"- {label}: " + "; ".join(parts))

    changes = []
    diff = comparison.marker_diff or {}

    def _marker_group(title: str, markers: list) -> None:
        if not markers:
            return
        changes.append(f"{title} ({len(markers)}):")
        shown = markers[:max_markers_per_group]
        changes.extend(_brief_marker_line(marker) for marker in shown)
        rest = len(markers) - len(shown)
        if rest:
            changes.append(f"…и ещё {rest}.")

    _marker_group("Добавленные маркеры", diff.get("added", []))
    _marker_group("Удалённые маркеры", diff.get("removed", []))

    changed = diff.get("changed", [])
    if changed:
        changes.append(f"Изменённые маркеры ({len(changed)}):")
        shown = changed[:max_markers_per_group]
        for entry in shown:
            for item in entry.get("changes", []):
                changes.append(
                    f"- {entry.get('tc_in', '')}: {item['field']}: «{item['old']}» → «{item['new']}»"
                )
        rest = len(changed) - len(shown)
        if rest:
            changes.append(f"…и ещё {rest}.")

    if comparison.parameter_changes:
        changes.append("Изменения технических параметров:")
        for change in comparison.parameter_changes:
            for item in change.get("changes", []):
                hint = _BRIEF_PARAMETER_STATUS_HINTS.get(item.get("status"))
                suffix = f" ({hint})" if hint else ""
                changes.append(
                    f"- {change['label']}, {item['field']}: {item['old']} → {item['new']}{suffix}"
                )

    if changes:
        body.append("Различия между версиями:")
        body.extend(changes)
    else:
        body.append(
            "Изменений в маркерах и технических параметрах между версиями не обнаружено — "
            "отличаются только общие счётчики выше (или версии идентичны)."
        )
    return "\n".join(header + body)


def compare_with_previous(
    client: YandexDiskClient,
    previous_report_path: str,
    new_local_docx_path: Path,
) -> ReportComparison | None:
    """Сравнивает новый (ещё не отправленный) черновик отчёта с версией на Диске.

    Количество маркеров/блокеров/новых маркеров + изменившиеся технические
    параметры по дорожкам. Возвращает None, если версия на Диске не найдена.
    """
    previous_docx_path = _resolve_docx_path(client, previous_report_path)
    if previous_docx_path is None:
        return None

    from docx import Document

    old_doc = Document(io.BytesIO(client.download_bytes(previous_docx_path)))
    new_doc = Document(str(new_local_docx_path))
    old_summary = _summarize_document(old_doc)
    new_summary = _summarize_document(new_doc)

    previous_csv_path = _find_csv_in_report_folder(client, previous_report_path)
    local_csv_path = _find_local_marker_csv(new_local_docx_path)
    try:
        if previous_csv_path is not None:
            _replace_summary_markers_from_csv(
                old_summary, read_marker_csv_bytes(client.download_bytes(previous_csv_path))
            )
    except Exception as exc:
        logger.warning("Не удалось прочитать CSV предыдущей версии: %s", exc)
    try:
        if local_csv_path is not None:
            _replace_summary_markers_from_csv(new_summary, read_marker_csv(local_csv_path))
    except Exception as exc:
        logger.warning("Не удалось прочитать CSV текущей версии: %s", exc)

    return _build_comparison(old_summary, new_summary)


def compare_two_versions(
    client: YandexDiskClient,
    old_report_path: str,
    new_report_path: str,
) -> ReportComparison | None:
    """Сравнивает две уже загруженные на Диск версии отчёта между собой

    (например, первую версию с четвёртой) — без участия локального черновика.
    Возвращает None, если хотя бы одна версия не найдена.

    Поиск .docx и скачивание обеих версий не зависят друг от друга, поэтому
    каждый из этих двух шагов выполняется параллельно (в двух потоках) —
    иначе вторая версия начинала бы искаться/скачиваться только после
    полного завершения первой, и окно сравнения открывалось бы заметно
    дольше. Скачивание запускается только если обе версии нашлись —
    если хотя бы одна не найдена, ни один файл не скачивается.
    """
    with ThreadPoolExecutor(max_workers=2) as executor:
        old_docx_path, new_docx_path = executor.map(
            lambda path: _resolve_docx_path(client, path), (old_report_path, new_report_path)
        )
    if old_docx_path is None or new_docx_path is None:
        return None

    with ThreadPoolExecutor(max_workers=2) as executor:
        old_bytes, new_bytes = executor.map(client.download_bytes, (old_docx_path, new_docx_path))

    from docx import Document

    old_doc = Document(io.BytesIO(old_bytes))
    new_doc = Document(io.BytesIO(new_bytes))
    old_summary = _summarize_document(old_doc)
    new_summary = _summarize_document(new_doc)

    csv_paths = (
        _find_csv_in_report_folder(client, old_report_path),
        _find_csv_in_report_folder(client, new_report_path),
    )
    for summary, csv_path, label in zip(
        (old_summary, new_summary), csv_paths, ("старой", "новой")
    ):
        try:
            if csv_path is not None:
                _replace_summary_markers_from_csv(
                    summary, read_marker_csv_bytes(client.download_bytes(csv_path))
                )
        except Exception as exc:
            logger.warning("Не удалось прочитать CSV %s версии: %s", label, exc)

    return _build_comparison(old_summary, new_summary)


@dataclass
class UploadJob:
    """Один отчёт в очереди автозагрузки на Яндекс.Диск."""

    id: int
    local_folder: Path
    meta: ReportMeta | None
    status: str = "queued"  # queued | uploading | needs_folder | failed | done
    attempts: int = 0
    error: str | None = None
    target_folder_path: str | None = None


class UploadQueue:
    """Статус-машина очереди автозагрузки — без Qt-зависимостей, чтобы

    логику можно было протестировать напрямую. Реальные сетевые вызовы
    (поиск папки, загрузка) делает вызывающий код (в главном файле,
    QObject-обёртка), эта очередь только хранит состояние и решает,
    что делать дальше — повторить, сдаться или продолжить.
    """

    MAX_ATTEMPTS = 3
    RETRY_DELAYS_MS = [5000, 20000]

    def __init__(self):
        self._jobs: list[UploadJob] = []
        self._next_id = 1

    @property
    def jobs(self) -> list[UploadJob]:
        return list(self._jobs)

    def enqueue(self, local_folder: Path, meta: ReportMeta | None) -> UploadJob:
        job = UploadJob(id=self._next_id, local_folder=local_folder, meta=meta)
        self._next_id += 1
        self._jobs.append(job)
        return job

    def next_queued(self) -> UploadJob | None:
        for job in self._jobs:
            if job.status == "queued":
                return job
        return None

    def is_uploading(self) -> bool:
        return any(job.status == "uploading" for job in self._jobs)

    def mark_uploading(self, job: UploadJob) -> None:
        job.status = "uploading"

    def mark_done(self, job: UploadJob) -> None:
        job.status = "done"

    def mark_needs_folder(self, job: UploadJob) -> None:
        job.status = "needs_folder"
        job.error = None

    def mark_failed_or_retry(self, job: UploadJob, error: str) -> int | None:
        """Увеличивает счётчик попыток и решает, что дальше.

        Возвращает задержку в мс перед следующей попыткой (статус
        остаётся/возвращается в "queued"), либо None — попытки исчерпаны,
        статус переведён в "failed".
        """
        job.attempts += 1
        job.error = error
        if job.attempts >= self.MAX_ATTEMPTS:
            job.status = "failed"
            return None
        job.status = "queued"
        return self.RETRY_DELAYS_MS[min(job.attempts - 1, len(self.RETRY_DELAYS_MS) - 1)]

    def retry(self, job: UploadJob) -> None:
        """Ручной повтор (из UI) — сбрасывает попытки и возвращает в очередь."""
        job.status = "queued"
        job.attempts = 0
        job.error = None

    def to_dicts(self) -> list[dict]:
        """Сериализует очередь для сохранения между запусками приложения.

        "done"-job'ы не сохраняются — они уже сделали своё дело. "uploading"
        записывается как "queued": сам поток загрузки не переживает
        перезапуск процесса, а начатая-но-незавершённая загрузка должна
        просто повториться заново при следующем запуске.
        """
        result = []
        for job in self._jobs:
            if job.status == "done":
                continue
            status = "queued" if job.status == "uploading" else job.status
            result.append({
                "id": job.id,
                "local_folder": str(job.local_folder),
                "status": status,
                "attempts": job.attempts,
                "error": job.error,
            })
        return result

    @classmethod
    def from_dicts(cls, data: list[dict]) -> "UploadQueue":
        """Восстанавливает очередь из to_dicts(). meta не сериализуется как

        есть (даты внутри ReportMeta усложняют JSON) — вместо этого заново
        разбирается из имени файла отчёта в local_folder, как при первом
        enqueue(). Папки, которых больше нет на диске, пропускаются.
        """
        queue = cls()
        max_id = 0
        for entry in data:
            local_folder = Path(entry.get("local_folder", ""))
            if not local_folder.is_dir():
                continue

            meta = None
            for candidate in sorted(local_folder.glob("отчет_*.docx")):
                meta = parse_report_filename(candidate.name)
                if meta is not None:
                    break

            job_id = entry.get("id", max_id + 1)
            job = UploadJob(
                id=job_id,
                local_folder=local_folder,
                meta=meta,
                status=entry.get("status", "queued"),
                attempts=entry.get("attempts", 0),
                error=entry.get("error"),
            )
            queue._jobs.append(job)
            max_id = max(max_id, job_id)
        queue._next_id = max_id + 1
        return queue


def save_queue_state(queue: UploadQueue, path: Path) -> None:
    """Сохраняет состояние очереди автозагрузки в JSON-файл."""
    try:
        atomic_write_text(path, json.dumps(queue.to_dicts(), ensure_ascii=False, indent=2))
    except OSError as exc:
        logger.warning(f"Не удалось сохранить состояние очереди Яндекс.Диска: {exc}")


def load_queue_state(path: Path) -> UploadQueue:
    """Загружает состояние очереди автозагрузки из JSON-файла.

    Отсутствующий или битый файл — не ошибка, просто пустая очередь.
    """
    if not path.exists():
        return UploadQueue()
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        logger.warning(f"Не удалось прочитать состояние очереди Яндекс.Диска: {exc}")
        return UploadQueue()
    return UploadQueue.from_dicts(data)
