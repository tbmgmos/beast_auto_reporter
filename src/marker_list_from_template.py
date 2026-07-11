"""
Marker List from Template Module

Добавление маркер-листа из готового шаблона DOCX
"""

import logging
from pathlib import Path
from typing import List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class MarkerListFromTemplate:
    """Добавление маркер-листа из шаблона"""
    
    def __init__(self, template_path: str):
        """
        Инициализация
        
        Args:
            template_path: Путь к шаблону маркер-листа
        """
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Шаблон маркер-листа не найден: {template_path}")
        
        # Загружаем шаблон для копирования структуры
        self.template_doc = Document(str(self.template_path))
        self.template_table = self.template_doc.tables[0]
        
        logger.info(f"Загружен шаблон маркер-листа: {self.template_path}")
    
    def add_to_document(self, doc: Document, issues: List) -> bool:
        """
        Добавление маркер-листа в документ из шаблона
        
        Args:
            doc: Документ Word
            issues: Список проблем (Issue объекты)
            
        Returns:
            True если успешно
        """
        try:
            logger.info(f"Добавление маркер-листа из шаблона ({len(issues)} проблем)")
            
            # Добавляем page break перед таблицей (чтобы таблица была на странице 3)
            doc.add_page_break()
            
            # Копируем только таблицу из шаблона
            table = self._copy_table_structure(doc)
            
            # Добавляем строки с данными
            for idx, issue in enumerate(issues):
                self._add_issue_row(table, issue, row_idx=idx)
            
            logger.info(f"Маркер-лист добавлен: {len(issues)} записей")
            return True
        
        except Exception as e:
            logger.error(f"Ошибка добавления маркер-листа: {e}")
            return False
    
    def _copy_table_structure(self, doc: Document):
        """
        Копирование структуры таблицы из шаблона
        
        Args:
            doc: Целевой документ
            
        Returns:
            Созданная таблица
        """
        # Создаем таблицу с 3 строками:
        # row 0: MARKER LIST (merged, green background)
        # row 1: column headers (green background)
        # row 2: template for formatting
        template_rows = min(2, len(self.template_table.rows))
        cols = len(self.template_table.columns)
        
        # Создаем 3 строки
        table = doc.add_table(rows=3, cols=cols)
        table.style = self.template_table.style
        
        # === Row 0: MARKER LIST (merged, green background) ===
        for i in range(cols - 1):
            table.rows[0].cells[0].merge(table.rows[0].cells[i + 1])
        
        cell = table.rows[0].cells[0]
        cell.text = "MARKER LIST"
        self._set_cell_background(cell, "70AD47")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Helvetica Neue"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # === Row 1: Column headers (green background) ===
        if template_rows >= 1:
            template_row = self.template_table.rows[0]
            target_row = table.rows[1]
            
            for j, template_cell in enumerate(template_row.cells):
                target_cell = target_row.cells[j]
                target_cell.text = template_cell.text
                self._set_cell_background(target_cell, "70AD47")
                
                for template_para, target_para in zip(template_cell.paragraphs, target_cell.paragraphs):
                    target_para.alignment = template_para.alignment
                    for template_run, target_run in zip(template_para.runs, target_para.runs):
                        target_run.bold = template_run.bold
                        target_run.italic = template_run.italic
                        if template_run.font.size:
                            target_run.font.size = template_run.font.size
                        if template_run.font.name:
                            target_run.font.name = template_run.font.name
                        target_run.font.color.rgb = RGBColor(255, 255, 255)
                
                self._copy_cell_background(template_cell, target_cell)
                self._copy_cell_borders(template_cell, target_cell)
        
        # === Row 2: Template for formatting (cleared text) ===
        if template_rows >= 2:
            template_row = self.template_table.rows[1]
            target_row = table.rows[2]
            
            for j, template_cell in enumerate(template_row.cells):
                target_cell = target_row.cells[j]
                target_cell.text = ""
                
                for template_para, target_para in zip(template_cell.paragraphs, target_cell.paragraphs):
                    target_para.alignment = template_para.alignment
                    for template_run, target_run in zip(template_para.runs, target_para.runs):
                        target_run.bold = template_run.bold
                        target_run.italic = template_run.italic
                        if template_run.font.size:
                            target_run.font.size = template_run.font.size
                        if template_run.font.name:
                            target_run.font.name = template_run.font.name
                
                self._copy_cell_background(template_cell, target_cell)
                self._copy_cell_borders(template_cell, target_cell)
        
        # Копируем ширину столбцов
        for i, template_col in enumerate(self.template_table.columns):
            if i < len(table.columns):
                table.columns[i].width = template_col.width
        
        return table
    
    def _add_issue_row(self, table, issue, row_idx=0):
        """
        Добавление строки с проблемой
        
        Args:
            table: Таблица
            issue: Объект Issue
            row_idx: Индекс строки (для чередования цветов)
        """
        # Копируем строку-шаблон (строка 2)
        template_row = self.template_table.rows[2] if len(self.template_table.rows) > 2 else self.template_table.rows[1]
        new_row = table.add_row()
        
        # Проверяем на "НОВЫЙ МАРКЕР" в комментариях
        is_new_marker = issue.comments and "НОВЫЙ МАРКЕР" in issue.comments.upper()
        
        # Определяем цвет фона
        if is_new_marker:
            bg_color = "FFFACD"  # LemonChiffon (желтый)
        else:
            # Чередование цветов: белый / светло-серый
            bg_color = "FFFFFF" if row_idx % 2 == 0 else "F5F5F5"
        
        # Сначала заполняем данными
        # Убираем расширение .wav из описания
        description_clean = issue.description or ""
        if description_clean.lower().endswith('.wav'):
            description_clean = description_clean[:-4]
        
        # Заполняем ячейки данными
        new_row.cells[0].text = issue.timecode_in or ""
        new_row.cells[1].text = issue.timecode_out or ""
        new_row.cells[2].text = description_clean
        new_row.cells[3].text = "●" if issue.audio_20_c else ""
        new_row.cells[4].text = "●" if issue.audio_20_uc else ""
        new_row.cells[5].text = "●" if issue.audio_51_c else ""
        new_row.cells[6].text = "●" if issue.audio_51_uc else ""
        new_row.cells[7].text = "●" if issue.blocker else ""
        new_row.cells[8].text = "●" if issue.fix_required else ""
        new_row.cells[9].text = "●" if issue.comment_required else ""
        
        # NOTES - комментарии БЕЗ "НОВЫЙ МАРКЕР"
        comments_text = issue.comments or ""
        if is_new_marker:
            comments_text = comments_text.replace("НОВЫЙ МАРКЕР", "").strip()
        new_row.cells[10].text = comments_text
        
        # Теперь применяем форматирование
        for i, cell in enumerate(new_row.cells):
            # Копируем фон и границы из шаблона
            template_cell = template_row.cells[i]
            self._copy_cell_background(template_cell, cell)
            self._copy_cell_borders(template_cell, cell)
            
            # Применяем цвет фона
            self._set_cell_background(cell, bg_color)
            
            # Применяем выравнивание и шрифт
            if cell.paragraphs:
                para = cell.paragraphs[0]
                # Description (колонка 2) - по левому краю, остальные - по центру
                if i == 2:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Убеждаемся что есть run с правильным шрифтом
                if para.runs:
                    run = para.runs[0]
                    run.font.name = "Helvetica Neue"
                    run.font.size = Pt(8)
    
    def _copy_cell_background(self, source_cell, target_cell):
        """Копирование фона ячейки"""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # Получаем shading элемент из источника
            source_tcPr = source_cell._element.tcPr
            if source_tcPr is not None:
                source_shd = source_tcPr.find(qn('w:shd'))
                if source_shd is not None:
                    # Копируем в целевую ячейку
                    target_tcPr = target_cell._element.get_or_add_tcPr()
                    target_shd = OxmlElement('w:shd')
                    target_shd.set(qn('w:fill'), source_shd.get(qn('w:fill')))
                    target_tcPr.append(target_shd)
        except Exception as e:
            logger.debug(f"Не удалось скопировать фон ячейки: {e}")
    
    def _copy_cell_borders(self, source_cell, target_cell):
        """Копирование границ ячейки"""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # Получаем borders элемент из источника
            source_tcPr = source_cell._element.tcPr
            if source_tcPr is not None:
                source_tcBorders = source_tcPr.find(qn('w:tcBorders'))
                if source_tcBorders is not None:
                    # Копируем в целевую ячейку
                    target_tcPr = target_cell._element.get_or_add_tcPr()
                    target_tcBorders = OxmlElement('w:tcBorders')
                    
                    # Копируем все границы
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        source_border = source_tcBorders.find(qn(f'w:{border_name}'))
                        if source_border is not None:
                            target_border = OxmlElement(f'w:{border_name}')
                            for attr in source_border.attrib:
                                target_border.set(attr, source_border.get(attr))
                            target_tcBorders.append(target_border)
                    
                    target_tcPr.append(target_tcBorders)
        except Exception as e:
            logger.debug(f"Не удалось скопировать границы ячейки: {e}")
    
    def _set_cell_background(self, cell, color_hex):
        """Установка фона ячейки"""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            tc_pr = cell._element.get_or_add_tcPr()
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_hex)
            tc_pr.append(shading_elm)
        except Exception as e:
            logger.debug(f"Не удалось установить фон ячейки: {e}")


# Пример использования
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    
    # Создаем тестовый документ
    from src.csv_importer import Issue
    
    # Тестовые проблемы
    issues = [
        Issue(
            timecode_in="01:00:00:00",
            timecode_out="01:00:05:12",
            description="Тестовая проблема 1",
            audio_20_c=True,
            audio_51_c=True,
            blocker=False,
            requires_fix=True,
            requires_comment=False
        ),
        Issue(
            timecode_in="01:00:10:00",
            timecode_out="01:00:15:00",
            description="Тестовая проблема 2",
            audio_20_c=False,
            audio_51_c=True,
            blocker=True,
            requires_fix=False,
            requires_comment=True
        ),
    ]
    
    # Создаем документ
    doc = Document()
    doc.add_heading("Тестовый отчет", level=1)
    
    # Добавляем маркер-лист из шаблона
    marker_list = MarkerListFromTemplate("templates/docx/standard_original.docx")
    marker_list.add_to_document(doc, issues)
    
    # Сохраняем
    doc.save("output/test_marker_list_from_template.docx")
    print("✅ Тестовый документ создан!")
