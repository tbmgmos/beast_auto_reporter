"""
New Report Generator Module

Генерирует отчет на основе шаблона РЫБА ОСНОВНОЙ НОВАЯ 2025.docx
Использует импортированные данные из CSV
"""

import logging
import sys
from pathlib import Path
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

# Добавляем корневую директорию в путь
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.csv_importer import Issue

logger = logging.getLogger(__name__)


class NewReportGenerator:
    """Класс для генерации отчета по новому шаблону"""
    
    def __init__(self, template_path: str = None):
        """
        Инициализация генератора
        
        Args:
            template_path: Путь к шаблону DOCX
        """
        self.template_path = template_path or '/Users/vladog/Desktop/ШАБЛОН/РЫБА ОСНОВНОЙ НОВАЯ 2025.docx'
        logger.info(f"NewReportGenerator инициализирован с шаблоном: {self.template_path}")
    
    def create_report(self,
                     issues: List[Issue],
                     output_path: str,
                     tech_info: Dict = None,
                     conclusion: str = "") -> None:
        """
        Создание отчета на основе шаблона
        
        Args:
            issues: Список проблем из CSV
            output_path: Путь для сохранения отчета
            tech_info: Техническая информация (опционально)
            conclusion: Заключение (опционально)
        """
        try:
            logger.info(f"Создание отчета: {output_path}")
            
            # Открываем шаблон
            doc = Document(self.template_path)
            
            # Добавляем техническую информацию в начало (если есть)
            if tech_info:
                self._add_technical_info(doc, tech_info)
            
            # Работаем с таблицей
            if len(doc.tables) > 0:
                table = doc.tables[0]
                
                # Таблица должна уже иметь заголовок (строки 0 и 1)
                # Строка 2 - это шаблон для данных
                
                # Удаляем пустую строку 2 если она есть
                if len(table.rows) > 2:
                    # Сохраняем стиль второй строки для новых строк
                    template_row = table.rows[2]
                    template_cells = [cell for cell in template_row.cells]
                    
                    # Удаляем шаблонную строку
                    table._element.remove(template_row._element)
                
                # Добавляем данные из issues
                for issue in issues:
                    row = table.add_row()
                    
                    # Заполняем ячейки
                    row.cells[0].text = issue.timecode_in
                    row.cells[1].text = issue.timecode_out
                    row.cells[2].text = issue.description
                    row.cells[3].text = '*' if issue.audio_20_c else ''
                    row.cells[4].text = '*' if issue.audio_20_uc else ''
                    row.cells[5].text = '*' if issue.audio_51_c else ''
                    row.cells[6].text = '*' if issue.audio_51_uc else ''
                    row.cells[7].text = '*' if issue.blocker else ''
                    row.cells[8].text = '*' if issue.fix_required else ''
                    row.cells[9].text = '*' if issue.comment_required else ''
                    row.cells[10].text = issue.comments
                    
                    # Копируем форматирование
                    self._format_table_row(row)
                
                logger.info(f"Добавлено {len(issues)} записей в таблицу")
            
            # Добавляем заключение в конец
            if conclusion:
                self._add_conclusion(doc, conclusion)
            
            # Сохраняем
            doc.save(output_path)
            
            logger.info(f"Отчет успешно создан: {output_path}")
            
        except Exception as e:
            logger.error(f"Ошибка создания отчета: {e}")
            raise
    
    def _add_technical_info(self, doc: Document, tech_info: Dict) -> None:
        """
        Добавление технической информации в начало документа
        
        Args:
            doc: Документ DOCX
            tech_info: Словарь с технической информацией
        """
        try:
            # Добавляем параграф перед таблицей
            if len(doc.tables) > 0:
                table_paragraph = doc.tables[0]._element
                parent = table_paragraph.getparent()
                
                # Создаем новые параграфы
                title = doc.add_paragraph()
                title.text = "ТЕХНИЧЕСКИЕ ПАРАМЕТРЫ"
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].bold = True
                title.runs[0].font.size = Pt(14)
                
                # Перемещаем в начало
                parent.insert(0, title._element)
                
                # Добавляем информацию
                info_lines = []
                
                for key, data in tech_info.items():
                    if isinstance(data, dict) and data:
                        file_name = data.get('file_name', key)
                        info_lines.append(f"\n{file_name}:")
                        info_lines.append(f"  Формат: {data.get('format', 'N/A')}")
                        info_lines.append(f"  Частота дискретизации: {data.get('sample_rate', 'N/A')} Hz")
                        info_lines.append(f"  Битность: {data.get('bit_depth', 'N/A')}")
                        info_lines.append(f"  Каналы: {data.get('channels', 'N/A')} ({data.get('channel_order', 'N/A')})")
                        info_lines.append(f"  Длительность: {data.get('duration', 0):.2f} сек")
                
                if info_lines:
                    info_para = doc.add_paragraph()
                    info_para.text = "\n".join(info_lines)
                    info_para.runs[0].font.size = Pt(11)
                    parent.insert(1, info_para._element)
                    
                    # Пустая строка
                    spacer = doc.add_paragraph()
                    parent.insert(2, spacer._element)
            
            logger.info("Техническая информация добавлена")
            
        except Exception as e:
            logger.error(f"Ошибка добавления технической информации: {e}")
    
    def _add_conclusion(self, doc: Document, conclusion: str) -> None:
        """
        Добавление заключения в конец документа
        
        Args:
            doc: Документ DOCX
            conclusion: Текст заключения
        """
        try:
            # Пустая строка
            doc.add_paragraph()
            
            # Заголовок
            title = doc.add_paragraph()
            title.text = "ЗАКЛЮЧЕНИЕ"
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].bold = True
            title.runs[0].font.size = Pt(14)
            
            # Текст заключения
            conclusion_para = doc.add_paragraph()
            conclusion_para.text = conclusion
            conclusion_para.runs[0].font.size = Pt(12)
            conclusion_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            logger.info("Заключение добавлено")
            
        except Exception as e:
            logger.error(f"Ошибка добавления заключения: {e}")
    
    def _format_table_row(self, row) -> None:
        """
        Форматирование строки таблицы
        
        Args:
            row: Строка таблицы
        """
        try:
            for cell in row.cells:
                # Устанавливаем шрифт
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Description (колонка 2) - выравнивание по левому краю
            if len(row.cells) > 2:
                for paragraph in row.cells[2].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Comments (колонка 10) - выравнивание по левому краю
            if len(row.cells) > 10:
                for paragraph in row.cells[10].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        except Exception as e:
            logger.error(f"Ошибка форматирования строки: {e}")


if __name__ == "__main__":
    # Тестирование
    logging.basicConfig(level=logging.INFO)
    
    from src.csv_importer import CSVImporter
    from src.conclusion_generator import ConclusionGenerator
    
    # Импортируем проблемы
    importer = CSVImporter()
    issues = importer.import_issues('/Users/vladog/Desktop/ШАБЛОН/petr_2_s1_e2_2024_09_12_rus.csv')
    categories = importer.categorize_issues(issues)
    
    # Генерируем заключение
    conclusion_gen = ConclusionGenerator(use_llm=False)
    conclusion = conclusion_gen.generate_conclusion(issues, categories)
    
    # Создаем отчет
    report_gen = NewReportGenerator()
    
    tech_info = {
        'audio_20': {
            'file_name': 'test_2.0.wav',
            'format': 'WAV',
            'sample_rate': 48000,
            'bit_depth': 'PCM_24',
            'channels': 2,
            'channel_order': 'L, R',
            'duration': 1450.04
        },
        'audio_51': {
            'file_name': 'test_5.1.wav',
            'format': 'WAV',
            'sample_rate': 48000,
            'bit_depth': 'PCM_24',
            'channels': 6,
            'channel_order': 'L, R, C, LFE, Ls, Rs',
            'duration': 1450.04
        }
    }
    
    output_path = '/tmp/test_report.docx'
    report_gen.create_report(issues, output_path, tech_info, conclusion)
    
    print(f"\n✅ Тестовый отчет создан: {output_path}")

