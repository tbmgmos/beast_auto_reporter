"""
Exact Report Generator

Точная копия референсного документа + техническая таблица + заключение
"""

import logging
import sys
from pathlib import Path
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
import fitz
import io

sys.path.insert(0, str(Path(__file__).parent.parent))

from src.csv_importer import Issue

logger = logging.getLogger(__name__)


class ExactReportGenerator:
    """Точная копия референса"""
    
    def __init__(self):
        logger.info("ExactReportGenerator инициализирован")
    
    def create_exact_report(self,
                           issues: List[Issue],
                           output_path: str,
                           tech_info: Dict = None,
                           pdf_20_path: str = None,
                           pdf_51_path: str = None,
                           conclusion_technical: str = "",
                           conclusion_subjective: str = "") -> None:
        """Создание точного отчета"""
        try:
            logger.info(f"Создание отчета: {output_path}")
            
            doc = Document()
            
            # Настройки страницы как в РЫБА ОСНОВНОЙ НОВАЯ 2025.docx
            section = doc.sections[0]
            section.page_width = Cm(42.02)   # A3 ширина
            section.page_height = Cm(29.70)  # A3 высота
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.orientation = 1  # 1 = Landscape (альбомная)
            
            # === СТРАНИЦА 1: ТЕХНИЧЕСКАЯ ТАБЛИЦА ===
            # (остается на первой странице)
            
            # === 1. ТЕХНИЧЕСКАЯ ТАБЛИЦА с заключением ===
            if tech_info:
                self._add_technical_table_with_conclusion(
                    doc, tech_info, conclusion_technical, conclusion_subjective
                )
            
            # === PAGE BREAK (переход на страницу 2, как в РЫБА) ===
            para_break = doc.add_paragraph()
            run = para_break.add_run()
            run.add_break(WD_BREAK.PAGE)  # Page break
            
            # === СТРАНИЦА 2: MARKER LIST ===
            
            # === 2. MARKER LIST (точно как в референсе) ===
            self._add_marker_list_exact(doc, issues)
            
            # === 3. PDF КАК ИЗОБРАЖЕНИЯ ===
            if pdf_20_path:
                doc.add_paragraph()
                self._add_pdf_image(doc, pdf_20_path)
            
            if pdf_51_path:
                doc.add_paragraph()
                self._add_pdf_image(doc, pdf_51_path)
            
            doc.save(output_path)
            logger.info(f"Отчет создан: {output_path}")
            
        except Exception as e:
            logger.error(f"Ошибка создания отчета: {e}")
            raise
    
    def _add_technical_table_with_conclusion(self, doc: Document, tech_info: Dict, 
                                             conclusion_tech: str, conclusion_subj: str) -> None:
        """Техническая таблица с заключением (точно по скриншоту)"""
        try:
            # Таблица: 11 строк x 7 столбцов
            # Строки: 0=заголовок+дата (объединенная), 1=пустая, 2=колонки, 
            #         3-7=данные (2.0 cens, 2.0 uncens, 5.1 cens, 5.1 uncens, VIDEO), 
            #         8=индикаторы, 9=ЗАКЛЮЧЕНИЕ, 10=текст заключения
            table = doc.add_table(rows=11, cols=7)
            table.style = 'Table Grid'
            
            # Устанавливаем ширину колонок (по скриншоту)
            # ДОРОЖКА | НАЗВАНИЕ ФАЙЛОВ | ХРОНОМЕТРАЖ | LOUDNESS | TRUE PEAK | LRA | ФОРМАТ ФАЙЛА
            widths_cm = [3.5, 12.0, 4.0, 4.0, 4.0, 3.5, 7.0]  # Итого: 38.0 см
            
            # Метод 1: Устанавливаем через table.columns
            for idx, width_cm in enumerate(widths_cm):
                if idx < len(table.columns):
                    for cell in table.columns[idx].cells:
                        cell.width = Cm(width_cm)
            
            # Метод 2: Устанавливаем через XML напрямую
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Удаляем autofit
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
            
            # Устанавливаем ширину через tblGrid
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is not None:
                for idx, gridCol in enumerate(tblGrid.findall(qn('w:gridCol'))):
                    if idx < len(widths_cm):
                        gridCol.set(qn('w:w'), str(int(widths_cm[idx] * 567)))  # cm to twips
            
            # Строка 0: Заголовок + Дата (разделенная)
            row0 = table.rows[0]
            
            # Левая часть: "ОБЪЕКТИВНАЯ ОЦЕНКА КАЧЕСТВА ЗВУЧАНИЯ ФОНОГРАММЫ" (объединяем первые 5 колонок)
            cell_left = row0.cells[0]
            for col in range(1, 5):
                cell_left.merge(row0.cells[col])
            cell_left.text = "ОБЪЕКТИВНАЯ ОЦЕНКА КАЧЕСТВА ЗВУЧАНИЯ ФОНОГРАММЫ"
            self._format_cell(cell_left, bg="E7E6E6", bold=True, align="center", font_size=11, vertical_align="center")
            
            # Правая часть: "ДАТА ОТЧЕТА:" (объединяем последние 2 колонки)
            from datetime import datetime
            cell_right = row0.cells[5]
            cell_right.merge(row0.cells[6])
            cell_right.text = f"ДАТА ОТЧЕТА:\n{datetime.now().strftime('%d.%m.%Y')}"
            self._format_cell(cell_right, bg="E7E6E6", bold=True, align="center", font_size=10, vertical_align="center")
            
            # Строка 1: Пустая
            # (оставляем пустой)
            
            # Строка 2: Заголовки столбцов
            headers = ["ДОРОЖКА", "НАЗВАНИЕ ФАЙЛОВ", "ХРОНОМЕТРАЖ", 
                      "LOUDNESS", "TRUE PEAK", "LRA", "ФОРМАТ ФАЙЛА"]
            for col, header in enumerate(headers):
                cell = table.rows[2].cells[col]
                cell.text = header
                self._format_cell(cell, bg="D9D9D9", bold=True, align="center", font_size=9, vertical_align="center")
            
            # Строки 3-7: Данные с цветовой индикацией
            rows_data = [
                ("2.0 cens", "audio_20_c", "pdf_20"),
                ("2.0 uncens", "audio_20_uc", None),
                ("5.1 cens", "audio_51_c", "pdf_51"),
                ("5.1 uncens", "audio_51_uc", None),
                ("VIDEO", "video", None)
            ]
            
            # Номинальные значения (берем из tech_info['params'] если есть)
            params = tech_info.get('params', {}) if tech_info else {}
            target_lufs = params.get('target_lufs', -23.0)
            lufs_tolerance = params.get('lufs_tolerance', 0.5)
            target_peak = params.get('true_peak', -2.0)
            target_lra = params.get('lra_max', 18.0)
            
            # Собираем длительности всех файлов для сравнения
            durations = {}
            for label, key, _ in rows_data:
                if tech_info and key in tech_info and tech_info[key]:
                    durations[key] = tech_info[key].get('duration', 0)
            
            # Определяем эталонную длительность (берем первый аудио файл)
            reference_duration = None
            for key in ['audio_20_c', 'audio_51_c', 'audio_20_uc', 'audio_51_uc']:
                if key in durations and durations[key] > 0:
                    reference_duration = durations[key]
                    break
            
            # Функция сравнения хронометража (допуск 0.1 сек = 100 мс)
            def durations_match(dur1, dur2, tolerance=0.1):
                if dur1 is None or dur2 is None or dur1 == 0 or dur2 == 0:
                    return False
                return abs(dur1 - dur2) <= tolerance
            
            for idx, (label, key, pdf_key) in enumerate(rows_data):
                row = table.rows[3 + idx]
                row.cells[0].text = label
                self._format_cell(row.cells[0], align="left")
                
                # Проверяем наличие аудио файла
                has_audio = tech_info and key in tech_info and tech_info[key]
                # Проверяем наличие PDF файла
                has_pdf = pdf_key and pdf_key in tech_info
                
                if has_audio:
                    data = tech_info[key]
                    
                    # Имя файла
                    row.cells[1].text = data.get('file_name', '')
                    self._format_cell(row.cells[1], align="left")
                    
                    # Хронометраж (формат: H:MM:SS.mmm) с цветовой индикацией
                    duration = data.get('duration', 0)  # в секундах
                    hours = int(duration // 3600)
                    mins = int((duration % 3600) // 60)
                    secs = int(duration % 60)
                    millis = round((duration % 1) * 1000)  # round для точности как в Nuendo!
                    row.cells[2].text = f"{hours}:{mins:02d}:{secs:02d}.{millis:03d}"
                    
                    # Цветовая индикация: зеленый если совпадает, красный если нет
                    if reference_duration and duration > 0:
                        if durations_match(duration, reference_duration):
                            chrono_bg = "00B050"  # Зеленый - совпадает
                        else:
                            chrono_bg = "FF0000"  # Красный - не совпадает
                    else:
                        chrono_bg = "FFC000"  # Оранжевый - нет эталона
                    
                    self._format_cell(row.cells[2], bg=chrono_bg, align="center")
                    
                    # Технические параметры (только для аудио с PDF)
                    if key.startswith('audio') and pdf_key and pdf_key in tech_info:
                        pdf_data = tech_info[pdf_key]
                        
                        # LUFS с цветовой индикацией
                        lufs = pdf_data.get('lufs')
                        if lufs is not None:
                            row.cells[3].text = f"{lufs:.1f} LUFS"
                            # Проверка нормы
                            if abs(lufs - target_lufs) <= lufs_tolerance:
                                lufs_bg = "00B050"  # Зеленый - норма
                            else:
                                lufs_bg = "FF0000"  # Красный - превышение
                            self._format_cell(row.cells[3], bg=lufs_bg, align="center")
                        else:
                            row.cells[3].text = ""
                        
                        # TRUE PEAK с цветовой индикацией
                        true_peak = pdf_data.get('true_peak')
                        if true_peak is not None:
                            row.cells[4].text = f"{true_peak:.1f} dBTP"
                            # Проверка нормы (должен быть МЕНЬШЕ target_peak)
                            if true_peak <= target_peak:
                                peak_bg = "00B050"  # Зеленый - норма
                            else:
                                peak_bg = "FF0000"  # Красный - превышение
                            self._format_cell(row.cells[4], bg=peak_bg, align="center")
                        else:
                            row.cells[4].text = ""
                        
                        # LRA с цветовой индикацией
                        lra = pdf_data.get('lra')
                        if lra is not None:
                            row.cells[5].text = f"{lra:.1f} LU"
                            # Проверка нормы (должен быть МЕНЬШЕ target_lra)
                            if lra <= target_lra:
                                lra_bg = "00B050"  # Зеленый - норма
                            else:
                                lra_bg = "FF0000"  # Красный - превышение
                            self._format_cell(row.cells[5], bg=lra_bg, align="center")
                        else:
                            row.cells[5].text = ""
                        
                        # Формат файла
                        sr = data.get('sample_rate', 48000) // 1000
                        bd = str(data.get('bit_depth', 'PCM_24')).replace('PCM_', '')
                        co = data.get('channel_order', '')
                        format_text = f"PCM {sr}kHz {bd} bit {co}"
                        row.cells[6].text = format_text
                        self._format_cell(row.cells[6], bg="00B050", align="left")
                    elif key.startswith('audio'):
                        # Для uncens - только формат
                        sr = data.get('sample_rate', 48000) // 1000
                        bd = str(data.get('bit_depth', 'PCM_24')).replace('PCM_', '')
                        co = data.get('channel_order', '')
                        format_text = f"PCM {sr}kHz {bd} bit {co}"
                        row.cells[6].text = format_text
                        self._format_cell(row.cells[6], align="left")
                
                # Если НЕТ аудио, но ЕСТЬ PDF - показываем данные из PDF
                elif has_pdf and key.startswith('audio'):
                    pdf_data = tech_info[pdf_key]
                    
                    # LUFS
                    lufs = pdf_data.get('lufs')
                    if lufs is not None:
                        row.cells[3].text = f"{lufs:.1f} LUFS"
                        if abs(lufs - target_lufs) <= lufs_tolerance:
                            lufs_bg = "00B050"
                        else:
                            lufs_bg = "FF0000"
                        self._format_cell(row.cells[3], bg=lufs_bg, align="center")
                    
                    # TRUE PEAK
                    true_peak = pdf_data.get('true_peak')
                    if true_peak is not None:
                        row.cells[4].text = f"{true_peak:.1f} dBTP"
                        if true_peak <= target_peak:
                            peak_bg = "00B050"
                        else:
                            peak_bg = "FF0000"
                        self._format_cell(row.cells[4], bg=peak_bg, align="center")
                    
                    # LRA
                    lra = pdf_data.get('lra')
                    if lra is not None:
                        row.cells[5].text = f"{lra:.1f} LU"
                        if lra <= target_lra:
                            lra_bg = "00B050"
                        else:
                            lra_bg = "FF0000"
                        self._format_cell(row.cells[5], bg=lra_bg, align="center")
                # else: нет ни аудио, ни PDF - оставляем ячейки пустыми
            
            # Строка 8: Цветовые индикаторы (объединяем первые 4 колонки, последние 3 - индикаторы)
            row = table.rows[8]
            cell_merged = row.cells[0]
            for col in range(1, 4):
                cell_merged.merge(row.cells[col])
            
            # Индикаторы в последних 3 ячейках
            self._format_cell(row.cells[4], bg="00B050")  # Зеленый - норма
            self._format_cell(row.cells[5], bg="FF0000")  # Красный - превышение
            self._format_cell(row.cells[6], bg="FFC000")  # Оранжевый - нет данных
            
            # Строка 9: Заголовок "ЗАКЛЮЧЕНИЕ:" (объединенная)
            row = table.rows[9]
            cell = row.cells[0]
            for col in range(1, 7):
                cell.merge(row.cells[col])
            cell.text = "ЗАКЛЮЧЕНИЕ:"
            self._format_cell(cell, bold=True, font_size=10, align="left")
            
            # Строка 10: Область для заключений (объединенная, 2 подзаголовка)
            row = table.rows[10]
            cell = row.cells[0]
            for col in range(1, 7):
                cell.merge(row.cells[col])
            
            # Добавляем оба заключения в одну ячейку
            para = cell.paragraphs[0]
            para.clear()
            
            # Заголовок 1: По техническим характеристикам
            run = para.add_run("По техническим характеристикам")
            run.font.name = "Helvetica Neue"
            run.font.size = Pt(11)  # Увеличен с 9 до 11
            run.font.bold = True
            
            # Текст заключения (техническое)
            if conclusion_tech:
                para.add_run("\n")
                run = para.add_run(conclusion_tech)
                run.font.name = "Helvetica Neue"
                run.font.size = Pt(10)  # Увеличен с 9 до 10
            
            # Заголовок 2: По субъективной оценке
            para.add_run("\n\n")
            run = para.add_run("По субъективной оценке")
            run.font.name = "Helvetica Neue"
            run.font.size = Pt(11)  # Увеличен с 9 до 11
            run.font.bold = True
            
            # Текст заключения (субъективное)
            if conclusion_subj:
                para.add_run("\n")
                run = para.add_run(conclusion_subj)
                run.font.name = "Helvetica Neue"
                run.font.size = Pt(10)  # Увеличен с 9 до 10
            
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Добавляем межстрочный интервал для лучшей читаемости
            pPr = para._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '360')  # 1.5 интервал
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)
            
            logger.info("Техническая таблица с областью заключения добавлена")
            
        except Exception as e:
            logger.error(f"Ошибка добавления технической таблицы: {e}")
    
    def _add_marker_list_exact(self, doc: Document, issues: List[Issue]) -> None:
        """MARKER LIST - точно как в референсе"""
        try:
            # Создаем таблицу БЕЗ заголовочной строки с повторяющимся текстом
            table = doc.add_table(rows=1 + len(issues), cols=11)
            table.style = 'Table Grid'
            
            # Устанавливаем ширину колонок из РЫБА ОСНОВНОЙ НОВАЯ 2025.docx
            widths_cm = [2.45, 2.40, 11.93, 1.27, 1.56, 1.32, 1.50, 1.98, 3.55, 3.48, 6.57]
            
            # Метод 1: Устанавливаем через table.columns
            for idx, width_cm in enumerate(widths_cm):
                if idx < len(table.columns):
                    for cell in table.columns[idx].cells:
                        cell.width = Cm(width_cm)
            
            # Метод 2: Устанавливаем через XML напрямую
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Удаляем autofit
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
            
            # Строка 0: Заголовки столбцов (единственная заголовочная строка)
            headers = [
                "Timecode In", "Timecode Out", "Description",
                "2.0 C", "2.0 UC", "5.1 C", "5.1 UC",
                "БЛОКЕР", "ТРЕБУЕТ ИСПРАВЛЕНИЯ", "ТРЕБУЕТ КОММЕНТАРИЯ", "КОММЕНТАРИИ"
            ]
            
            for col, header in enumerate(headers):
                cell = table.rows[0].cells[col]
                cell.text = header
                # Фон #bdc0bf, шрифт Helvetica Neue, жирный, центрирование
                self._format_cell(cell, bg="bdc0bf", font_name="Helvetica Neue", bold=True, align="center", font_size=9)
            
            # Строки 1+: Данные с чередующимся фоном
            for row_idx, issue in enumerate(issues):
                row = table.rows[1 + row_idx]
                
                # Заполняем данные
                row.cells[0].text = issue.timecode_in
                row.cells[1].text = issue.timecode_out
                row.cells[2].text = issue.description
                row.cells[3].text = '*' if issue.audio_20_c else ''
                row.cells[4].text = '*' if issue.audio_20_uc else ''
                row.cells[5].text = '*' if issue.audio_51_c else ''
                row.cells[6].text = '*' if issue.audio_51_uc else ''
                row.cells[7].text = '*' if issue.blocker else ''
                row.cells[8].text = '*' if issue.fix_required else ''
                row.cells[9].text = '*' if issue.comment_required else ''
                row.cells[10].text = issue.comments
                
                # Чередующийся фон для нечетных строк (начиная с row_idx=1)
                bg_color = "f5f5f5" if row_idx % 2 == 1 else "auto"
                
                # Форматируем каждую ячейку
                for col_idx, cell in enumerate(row.cells):
                    # Фон только для некоторых столбцов (как в референсе)
                    if col_idx in [0, 1, 3, 4, 5, 6, 7, 8, 9, 10]:
                        self._format_cell(cell, bg=bg_color, font_name="Helvetica Neue")
                    else:
                        # Description (col 2) без фона
                        self._format_cell(cell, bg="auto", font_name="Helvetica Neue")
            
            logger.info(f"MARKER LIST добавлен ({len(issues)} записей)")
            
        except Exception as e:
            logger.error(f"Ошибка добавления MARKER LIST: {e}")
    
    def _add_pdf_image(self, doc: Document, pdf_path: str) -> None:
        """Добавление PDF как изображение (ширина = ширине таблицы)"""
        try:
            logger.info(f"Конвертация PDF: {pdf_path}")
            
            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[0]
            
            # Конвертация в изображение с высоким разрешением
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            
            # Добавляем изображение с шириной равной таблице (38.01 см)
            img_stream = io.BytesIO(img_bytes)
            doc.add_picture(img_stream, width=Cm(38.0))
            
            # Центрируем
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            pdf_doc.close()
            
            logger.info("PDF добавлен как изображение (38 см)")
            
        except Exception as e:
            logger.error(f"Ошибка добавления PDF: {e}")
    
    def _add_conclusion_OLD(self, doc: Document, technical: str, subjective: str) -> None:
        """Добавление заключения"""
        try:
            # Заголовок
            title = doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.name = 'Helvetica Neue'
                run.font.size = Pt(14)
                run.font.bold = True
            
            doc.add_paragraph()
            
            # Техническая оценка
            if technical:
                tech_title = doc.add_paragraph("Техническая оценка:")
                for run in tech_title.runs:
                    run.font.name = 'Helvetica Neue'
                    run.font.size = Pt(12)
                    run.font.bold = True
                
                tech_para = doc.add_paragraph(technical)
                tech_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in tech_para.runs:
                    run.font.name = 'Helvetica Neue'
                    run.font.size = Pt(11)
                
                doc.add_paragraph()
            
            # Субъективная оценка
            if subjective:
                subj_title = doc.add_paragraph("Субъективная оценка:")
                for run in subj_title.runs:
                    run.font.name = 'Helvetica Neue'
                    run.font.size = Pt(12)
                    run.font.bold = True
                
                subj_para = doc.add_paragraph(subjective)
                subj_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in subj_para.runs:
                    run.font.name = 'Helvetica Neue'
                    run.font.size = Pt(11)
            
            logger.info("Заключение добавлено")
            
        except Exception as e:
            logger.error(f"Ошибка добавления заключения: {e}")
    
    def _format_cell(self, cell, bg="auto", font_name="Helvetica Neue", 
                    font_size=None, font_color=None, bold=False, align=None, 
                    border=True, vertical_align="center"):
        """Форматирование ячейки"""
        try:
            tc_pr = cell._element.get_or_add_tcPr()
            
            # Фон
            if bg != "auto":
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), bg)
                tc_pr.append(shading_elm)
            
            # Вертикальное выравнивание
            v_align = OxmlElement('w:vAlign')
            v_align.set(qn('w:val'), vertical_align)
            tc_pr.append(v_align)
            
            # Отступы внутри ячейки (padding)
            tc_mar = OxmlElement('w:tcMar')
            for side in ['top', 'left', 'bottom', 'right']:
                mar = OxmlElement(f'w:{side}')
                mar.set(qn('w:w'), '100')  # 100 twips = примерно 1.76 мм
                mar.set(qn('w:type'), 'dxa')
                tc_mar.append(mar)
            tc_pr.append(tc_mar)
            
            # Границы (если нужны)
            if border:
                tc_borders = OxmlElement('w:tcBorders')
                for side in ['top', 'left', 'bottom', 'right']:
                    border_elm = OxmlElement(f'w:{side}')
                    border_elm.set(qn('w:val'), 'single')
                    border_elm.set(qn('w:sz'), '4')  # Размер границы
                    border_elm.set(qn('w:space'), '0')
                    border_elm.set(qn('w:color'), '000000')  # Черный
                    tc_borders.append(border_elm)
                tc_pr.append(tc_borders)
            
            # Текст
            for para in cell.paragraphs:
                # Убираем отступы до/после параграфа
                pPr = para._element.get_or_add_pPr()
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
                spacing.set(qn('w:line'), '240')
                spacing.set(qn('w:lineRule'), 'auto')
                pPr.append(spacing)
                
                if align:
                    if align == "center":
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align == "left":
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif align == "right":
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                for run in para.runs:
                    run.font.name = font_name
                    if font_size:
                        run.font.size = Pt(font_size)
                    else:
                        run.font.size = Pt(9)  # Размер по умолчанию
                    if font_color:
                        run.font.color.rgb = RGBColor.from_string(font_color)
                    if bold:
                        run.font.bold = True
        except Exception as e:
            logger.error(f"Ошибка форматирования ячейки: {e}")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    print("ExactReportGenerator готов")

