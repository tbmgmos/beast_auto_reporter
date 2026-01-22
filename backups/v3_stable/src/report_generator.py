"""
Report Generator Module

Генерация отчетов в форматах:
- CSV (таблица маркеров с дефектами)
- PDF (с графиками и измерениями)
- DOCX (итоговый документ)
"""

import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Для работы без GUI
import seaborn as sns
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
from typing import Dict, List
import logging
import io

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Класс для генерации отчетов"""
    
    def __init__(self, config: Dict = None):
        """
        Инициализация генератора отчетов
        
        Args:
            config: Словарь с настройками
        """
        self.config = config or {}
        report_cfg = self.config.get('report', {})
        
        self.company_name = report_cfg.get('company_name', 'Beast Audio QC')
        self.include_graphs = report_cfg.get('include_graphs', True)
        self.output_prefix = report_cfg.get('output_prefix', 'report')
        self.date_format = report_cfg.get('date_format', '%Y_%m_%d')
        
        # Severity levels
        severity_cfg = report_cfg.get('severity_levels', {})
        self.severity_labels = {
            'blocker': severity_cfg.get('blocker', 'БЛОКЕР'),
            'fix_required': severity_cfg.get('fix_required', 'ТРЕБУЕТ ИСПРАВЛЕНИЯ'),
            'comment_required': severity_cfg.get('comment_required', 'ТРЕБУЕТ КОММЕНТАРИЯ')
        }
        
        logger.info("ReportGenerator инициализирован")
    
    def generate_csv(
        self, 
        defects: List, 
        output_path: str,
        track_name: str = "MARKERS DATA 1"
    ) -> str:
        """
        Генерация CSV файла с маркерами
        
        Args:
            defects: Список дефектов
            output_path: Путь для сохранения
            track_name: Название трека
            
        Returns:
            Путь к созданному файлу
        """
        try:
            logger.info(f"Генерация CSV отчета: {output_path}")
            
            # Подготовка данных
            rows = []
            for defect in defects:
                # Определение маркеров для каналов и severity
                channels_20 = "*" if "2.0" in getattr(defect, 'channels', ["*"]) or "*" in getattr(defect, 'channels', ["*"]) else ""
                channels_51 = "*" if "5.1" in getattr(defect, 'channels', ["*"]) or "*" in getattr(defect, 'channels', ["*"]) else ""
                
                severity = getattr(defect, 'severity', 'comment_required')
                blocker = "*" if severity == 'blocker' else ""
                fix_required = "*" if severity == 'fix_required' else ""
                comment_required = "*" if severity == 'comment_required' else ""
                
                row = {
                    'Track name': track_name,
                    'Timecode In': getattr(defect, 'timecode_in', ''),
                    'Timecode Out': getattr(defect, 'timecode_out', ''),
                    'Description': getattr(defect, 'description', ''),
                    'Length': int(getattr(defect, 'duration', 0)),
                    '2.0 C': channels_20,
                    '5.1 C': channels_51,
                    'БЛОКЕР': blocker,
                    'ТРЕБУЕТ ИСПРАВЛЕНИЯ': fix_required,
                    'ТРЕБУЕТ КОММЕНТАРИЯ': comment_required
                }
                rows.append(row)
            
            # Создание DataFrame
            df = pd.DataFrame(rows)
            
            # Сохранение в CSV
            df.to_csv(output_path, index=False, sep='\t', encoding='utf-8-sig')
            
            logger.info(f"CSV отчет создан: {len(rows)} записей")
            return output_path
            
        except Exception as e:
            logger.error(f"Ошибка генерации CSV: {e}")
            raise
    
    def generate_pdf(
        self,
        analysis_results: Dict,
        defects: List,
        conclusion: str,
        output_path: str
    ) -> str:
        """
        Генерация PDF отчета с графиками
        
        Args:
            analysis_results: Результаты анализа
            defects: Список дефектов
            conclusion: Заключение от LLM
            output_path: Путь для сохранения
            
        Returns:
            Путь к созданному файлу
        """
        try:
            logger.info(f"Генерация PDF отчета: {output_path}")
            
            # Создание документа
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Заголовок
            title_style = styles['Title']
            title = Paragraph(f"<b>{self.company_name}</b><br/>Отчет о качестве аудио", title_style)
            story.append(title)
            story.append(Spacer(1, 12))
            
            # Дата
            date_text = Paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles['Normal'])
            story.append(date_text)
            story.append(Spacer(1, 12))
            
            # Информация о файле
            file_info_text = f"""
            <b>Информация о файле:</b><br/>
            Название: {analysis_results.get('file_name', 'N/A')}<br/>
            Длительность: {analysis_results.get('duration', 0)} сек<br/>
            Формат: {analysis_results.get('channel_layout', 'N/A')}<br/>
            Sample Rate: {analysis_results.get('sample_rate', 'N/A')} Hz
            """
            story.append(Paragraph(file_info_text, styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Измерения
            measurements = analysis_results.get('measurements', {})
            compliance = analysis_results.get('compliance', {})
            
            measurements_data = [
                ['Параметр', 'Значение', 'Норма', 'Статус'],
                [
                    'LUFS',
                    f"{measurements.get('lufs', 'N/A')} dB",
                    '-23.0 ±0.5 LUFS',
                    '✓' if compliance.get('lufs_compliant') else '✗'
                ],
                [
                    'TRUE PEAK',
                    f"{measurements.get('true_peak', 'N/A')} dBTP",
                    '≤ -2.0 dBTP',
                    '✓' if compliance.get('true_peak_compliant') else '✗'
                ],
                [
                    'LRA',
                    f"{measurements.get('lra', 'N/A')} LU",
                    '≤ 18 LU',
                    '✓' if compliance.get('lra_compliant') else '✗'
                ]
            ]
            
            measurements_table = Table(measurements_data)
            measurements_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(measurements_table)
            story.append(Spacer(1, 12))
            
            # Статистика дефектов
            defect_summary = self._summarize_defects_for_pdf(defects)
            defect_text = f"""
            <b>Обнаруженные дефекты:</b><br/>
            Всего: {defect_summary['total']}<br/>
            Критичные (блокеры): {defect_summary['blocker']}<br/>
            Требуют исправления: {defect_summary['fix_required']}<br/>
            Требуют комментария: {defect_summary['comment_required']}
            """
            story.append(Paragraph(defect_text, styles['Normal']))
            story.append(Spacer(1, 12))
            
            # График (если включен)
            if self.include_graphs and measurements.get('lufs'):
                graph_path = self._create_measurements_graph(measurements, analysis_results.get('targets', {}))
                if graph_path:
                    img = Image(graph_path, width=400, height=200)
                    story.append(img)
                    story.append(Spacer(1, 12))
            
            # Заключение
            story.append(Paragraph("<b>Заключение:</b>", styles['Heading2']))
            conclusion_paragraphs = conclusion.split('\n')
            for para in conclusion_paragraphs:
                if para.strip():
                    story.append(Paragraph(para, styles['Normal']))
                    story.append(Spacer(1, 6))
            
            # Сборка документа
            doc.build(story)
            
            logger.info("PDF отчет создан успешно")
            return output_path
            
        except Exception as e:
            logger.error(f"Ошибка генерации PDF: {e}")
            raise
    
    def generate_docx(
        self,
        analysis_results: Dict,
        defects: List,
        conclusion: str,
        output_path: str
    ) -> str:
        """
        Генерация DOCX отчета
        
        Args:
            analysis_results: Результаты анализа
            defects: Список дефектов
            conclusion: Заключение от LLM
            output_path: Путь для сохранения
            
        Returns:
            Путь к созданному файлу
        """
        try:
            logger.info(f"Генерация DOCX отчета: {output_path}")
            
            doc = Document()
            
            # Заголовок
            title = doc.add_heading(f'{self.company_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('Отчет о качестве аудио', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Дата
            date_para = doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_paragraph()  # Пустая строка
            
            # Информация о файле
            doc.add_heading('1. Информация о файле', level=2)
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Light List Accent 1'
            
            info_data = [
                ('Название:', analysis_results.get('file_name', 'N/A')),
                ('Длительность:', f"{analysis_results.get('duration', 0)} сек"),
                ('Формат:', analysis_results.get('channel_layout', 'N/A')),
                ('Sample Rate:', f"{analysis_results.get('sample_rate', 'N/A')} Hz")
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = label
                info_table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
            
            # Измерения
            doc.add_heading('2. Измерения громкости (EBU R128)', level=2)
            
            measurements = analysis_results.get('measurements', {})
            compliance = analysis_results.get('compliance', {})
            targets = analysis_results.get('targets', {})
            
            meas_table = doc.add_table(rows=4, cols=4)
            meas_table.style = 'Light Grid Accent 1'
            
            # Заголовки
            headers = meas_table.rows[0].cells
            headers[0].text = 'Параметр'
            headers[1].text = 'Значение'
            headers[2].text = 'Норма'
            headers[3].text = 'Статус'
            
            # Данные
            meas_data = [
                ('LUFS', f"{measurements.get('lufs', 'N/A')} dB", 
                 f"{targets.get('target_lufs', -23)} ±{targets.get('lufs_tolerance', 0.5)} LUFS",
                 '✓ Соответствует' if compliance.get('lufs_compliant') else '✗ Не соответствует'),
                ('TRUE PEAK', f"{measurements.get('true_peak', 'N/A')} dBTP",
                 f"≤ {targets.get('true_peak_threshold', -2.0)} dBTP",
                 '✓ Соответствует' if compliance.get('true_peak_compliant') else '✗ Не соответствует'),
                ('LRA', f"{measurements.get('lra', 'N/A')} LU",
                 f"≤ {targets.get('lra_max', 18)} LU",
                 '✓ Соответствует' if compliance.get('lra_compliant') else '✗ Не соответствует')
            ]
            
            for i, row_data in enumerate(meas_data, start=1):
                cells = meas_table.rows[i].cells
                for j, value in enumerate(row_data):
                    cells[j].text = str(value)
            
            doc.add_paragraph()
            
            # Статистика дефектов
            doc.add_heading('3. Обнаруженные дефекты', level=2)
            
            defect_summary = self._summarize_defects_for_pdf(defects)
            
            doc.add_paragraph(f"Всего обнаружено: {defect_summary['total']} дефектов")
            doc.add_paragraph(f"  • Критичные (блокеры): {defect_summary['blocker']}")
            doc.add_paragraph(f"  • Требуют исправления: {defect_summary['fix_required']}")
            doc.add_paragraph(f"  • Требуют комментария: {defect_summary['comment_required']}")
            
            doc.add_paragraph()
            
            # Заключение
            doc.add_heading('4. Заключение', level=2)
            
            for line in conclusion.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            # Сохранение
            doc.save(output_path)
            
            logger.info("DOCX отчет создан успешно")
            return output_path
            
        except Exception as e:
            logger.error(f"Ошибка генерации DOCX: {e}")
            raise
    
    def generate_all_reports(
        self,
        analysis_results: Dict,
        defects: List,
        conclusion: str,
        output_dir: str,
        base_name: str = None
    ) -> Dict[str, str]:
        """
        Генерация всех отчетов (CSV, PDF, DOCX)
        
        Args:
            analysis_results: Результаты анализа
            defects: Список дефектов
            conclusion: Заключение
            output_dir: Директория для сохранения
            base_name: Базовое имя файлов
            
        Returns:
            Словарь с путями к созданным файлам
        """
        try:
            # Создание директории если не существует
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            
            # Генерация базового имени
            if not base_name:
                timestamp = datetime.now().strftime(self.date_format)
                file_base = analysis_results.get('file_name', 'audio').split('.')[0]
                base_name = f"{self.output_prefix}_{file_base}_{timestamp}"
            
            paths = {}
            
            # CSV
            csv_path = Path(output_dir) / f"{base_name}.csv"
            paths['csv'] = self.generate_csv(defects, str(csv_path))
            
            # PDF
            pdf_path = Path(output_dir) / f"{base_name}.pdf"
            paths['pdf'] = self.generate_pdf(analysis_results, defects, conclusion, str(pdf_path))
            
            # DOCX
            docx_path = Path(output_dir) / f"{base_name}.docx"
            paths['docx'] = self.generate_docx(analysis_results, defects, conclusion, str(docx_path))
            
            logger.info(f"Все отчеты созданы в: {output_dir}")
            return paths
            
        except Exception as e:
            logger.error(f"Ошибка генерации отчетов: {e}")
            raise
    
    def _summarize_defects_for_pdf(self, defects: List) -> Dict:
        """Подсчет дефектов по категориям"""
        summary = {
            'total': len(defects),
            'blocker': 0,
            'fix_required': 0,
            'comment_required': 0
        }
        
        for defect in defects:
            severity = getattr(defect, 'severity', 'comment_required')
            if severity in summary:
                summary[severity] += 1
        
        return summary
    
    def _create_measurements_graph(self, measurements: Dict, targets: Dict) -> str:
        """Создание графика измерений"""
        try:
            fig, ax = plt.subplots(1, 1, figsize=(8, 4))
            
            # Данные
            params = ['LUFS', 'TRUE\nPEAK', 'LRA']
            values = [
                measurements.get('lufs', 0),
                measurements.get('true_peak', 0),
                measurements.get('lra', 0)
            ]
            
            target_values = [
                targets.get('target_lufs', -23),
                targets.get('true_peak_threshold', -2),
                targets.get('lra_max', 18)
            ]
            
            x = range(len(params))
            
            # Столбцы
            bars = ax.bar(x, values, color=['#3498db', '#e74c3c', '#2ecc71'], alpha=0.7, label='Измерено')
            ax.axhline(y=0, color='black', linestyle='-', linewidth=0.5)
            
            # Целевые линии
            for i, (val, target) in enumerate(zip(values, target_values)):
                ax.plot([i-0.3, i+0.3], [target, target], 'r--', linewidth=2)
            
            ax.set_xticks(x)
            ax.set_xticklabels(params)
            ax.set_ylabel('Значение')
            ax.set_title('Измерения аудио параметров')
            ax.legend(['Целевое значение', 'Измеренное значение'])
            ax.grid(True, alpha=0.3)
            
            # Сохранение во временный файл
            temp_path = '/tmp/measurements_graph.png'
            plt.tight_layout()
            plt.savefig(temp_path, dpi=150, bbox_inches='tight')
            plt.close()
            
            return temp_path
            
        except Exception as e:
            logger.error(f"Ошибка создания графика: {e}")
            return None


if __name__ == "__main__":
    # Тестирование
    logging.basicConfig(level=logging.INFO)
    
    config = {
        'report': {
            'company_name': 'Beast Audio QC',
            'include_graphs': True,
            'output_prefix': 'report'
        }
    }
    
    generator = ReportGenerator(config)
    print("ReportGenerator готов к использованию!")

