#!/usr/bin/env python3
"""
Анализ шаблона отчета для понимания структуры
"""

from docx import Document
import sys

def analyze_docx(file_path):
    """Анализирует структуру DOCX файла"""
    
    print(f"Анализ файла: {file_path}\n")
    print("=" * 80)
    
    doc = Document(file_path)
    
    # Информация о документе
    print(f"\n📄 ОБЩАЯ ИНФОРМАЦИЯ:")
    print(f"   Всего параграфов: {len(doc.paragraphs)}")
    print(f"   Всего таблиц: {len(doc.tables)}")
    print(f"   Всего секций: {len(doc.sections)}")
    
    # Стили параграфов
    print(f"\n🎨 СТИЛИ ПАРАГРАФОВ:")
    styles_used = set()
    for para in doc.paragraphs:
        if para.style.name:
            styles_used.add(para.style.name)
    for style in sorted(styles_used):
        print(f"   - {style}")
    
    # Структура документа
    print(f"\n📋 СТРУКТУРА ДОКУМЕНТА:")
    print(f"\n" + "-" * 80)
    
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        if text:  # Только непустые параграфы
            style = para.style.name
            # Форматирование
            bold = any(run.bold for run in para.runs)
            italic = any(run.italic for run in para.runs)
            font_size = None
            if para.runs:
                font_size = para.runs[0].font.size
            
            format_str = []
            if bold:
                format_str.append("BOLD")
            if italic:
                format_str.append("ITALIC")
            if font_size:
                format_str.append(f"{font_size.pt}pt")
            
            format_info = f" [{', '.join(format_str)}]" if format_str else ""
            
            # Показываем первые 100 символов
            text_preview = text[:100] + "..." if len(text) > 100 else text
            
            print(f"{i:3d}. [{style}]{format_info}")
            print(f"      {text_preview}")
            print()
    
    # Анализ таблиц
    if doc.tables:
        print(f"\n📊 ТАБЛИЦЫ:")
        print("-" * 80)
        for i, table in enumerate(doc.tables, 1):
            print(f"\nТаблица {i}:")
            print(f"   Строк: {len(table.rows)}")
            print(f"   Столбцов: {len(table.columns)}")
            
            # Показываем заголовки (первая строка)
            if table.rows:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                print(f"   Заголовки: {headers}")
                
                # Показываем первые 3 строки данных
                print(f"\n   Первые строки:")
                for row_idx, row in enumerate(table.rows[1:4], 1):
                    cells = [cell.text.strip()[:30] for cell in row.cells]
                    print(f"      {row_idx}. {cells}")
    
    # Анализ форматирования
    print(f"\n🎨 ДЕТАЛИ ФОРМАТИРОВАНИЯ:")
    print("-" * 80)
    
    # Цвета
    colors_used = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.color and run.font.color.rgb:
                colors_used.add(str(run.font.color.rgb))
    
    if colors_used:
        print(f"   Цвета текста: {colors_used}")
    
    # Размеры шрифтов
    font_sizes = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                font_sizes.add(run.font.size.pt)
    
    if font_sizes:
        print(f"   Размеры шрифтов: {sorted(font_sizes)}")
    
    # Шрифты
    fonts_used = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts_used.add(run.font.name)
    
    if fonts_used:
        print(f"   Шрифты: {sorted(fonts_used)}")

if __name__ == "__main__":
    template_path = "/Users/vladog/Desktop/отчет_petr_2_s1_e2_2024_09_12_rus/отчет_petr_2_s1_e2_2024_09_12_rus.docx"
    
    try:
        analyze_docx(template_path)
    except Exception as e:
        print(f"Ошибка: {e}")
        import traceback
        traceback.print_exc()

