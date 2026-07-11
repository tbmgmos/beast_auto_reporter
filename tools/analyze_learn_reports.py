from __future__ import annotations

import csv
import json
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document


LEARN_ROOT = Path("/Users/vlad/Desktop/отчеты/отчеты_learn")
OUTPUT_JSON = Path("/Users/vlad/Desktop/Code_projects/Beast_auto_reporter/tools/learn_reports_analysis.json")


@dataclass
class ReportPair:
    directory: Path
    csv_path: Path | None
    docx_path: Path | None
    report_type: str


def infer_report_type(name: str) -> str:
    lower = name.lower()
    if any(token in lower for token in ["m&e", "mn&e", "mne", "_me_", " me_", "_me.", " me.", "ме", "m&e_"]):
        return "me"
    return "main"


def find_pairs(root: Path) -> list[ReportPair]:
    pairs: list[ReportPair] = []
    for directory in sorted(p for p in root.rglob("*") if p.is_dir()):
        files = [p for p in directory.iterdir() if p.is_file()]
        if not files:
            continue
        csv_files = [p for p in files if p.suffix.lower() == ".csv" and "explicit_words" not in p.name.lower()]
        docx_files = [p for p in files if p.suffix.lower() == ".docx" and not p.name.startswith("~$")]
        if not csv_files and not docx_files:
            continue
        csv_path = sorted(csv_files)[0] if csv_files else None
        docx_path = sorted(docx_files)[0] if docx_files else None
        type_hint = infer_report_type(directory.name + " " + (csv_path.name if csv_path else "") + " " + (docx_path.name if docx_path else ""))
        pairs.append(ReportPair(directory=directory, csv_path=csv_path, docx_path=docx_path, report_type=type_hint))
    return pairs


def load_csv_issues(csv_path: Path | None) -> list[dict]:
    if csv_path is None:
        return []
    with csv_path.open("r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.DictReader(fh)
        rows = []
        for row in reader:
            description = (row.get("Description") or row.get("ОПИСАНИЕ ПРОБЛЕМЫ") or row.get("ОПИСАНИЕ") or "").strip()
            tc_in = (row.get("Timecode In") or row.get("TC IN") or row.get("TC_IN") or "").strip()
            if not tc_in and not description:
                continue
            rows.append(
                {
                    "timecode_in": tc_in,
                    "description": description,
                    "blocker": ((row.get("БЛОКЕР") or row.get("BLOCKER") or "").strip() == "*"),
                    "fix_required": ((row.get("ТРЕБУЕТ ИСПРАВЛЕНИЯ") or row.get("FIX REQUIRED") or "").strip() == "*"),
                }
            )
        return rows


def iter_docx_text(docx_path: Path) -> Iterable[str]:
    doc = Document(str(docx_path))
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        yield text


def extract_subjective_conclusion(docx_path: Path | None) -> dict:
    if docx_path is None:
        return {"header": None, "items": []}

    text = "\n".join(iter_docx_text(docx_path))
    marker = "По субъективной оценке выявлены следующие недочёты:"
    alt_marker = "По субъективной оценке выявлены следующие недочеты:"
    start = text.find(marker)
    header = marker
    if start < 0:
        start = text.find(alt_marker)
        header = alt_marker
    if start < 0:
        return {"header": None, "items": []}

    tail = text[start + len(header):]
    end_markers = [
        "\nMARKER LIST",
        "\nTimecode In",
        "\nTC IN",
        "\nПо техническим характеристикам",
    ]
    end_pos = len(tail)
    for end_marker in end_markers:
        pos = tail.find(end_marker)
        if pos >= 0:
            end_pos = min(end_pos, pos)
    tail = tail[:end_pos]

    item_pattern = re.compile(r"(?:^|\n)-\s{0,4}(.*?)(?=(?:\n-\s{0,4})|\Z)", re.S)
    items = []
    for match in item_pattern.finditer(tail):
        item = re.sub(r"\s+", " ", match.group(1)).strip()
        if item:
            items.append(item)
    return {"header": header, "items": items}


def classify_issue_description(description: str) -> str:
    desc = description.lower()
    if any(token in desc for token in ["реплик", "голос", "слышны реплики", "слышна реплика"]):
        return "реплики"
    if any(token in desc for token in ["вздох", "выдох", "дыхан", "голос одного из актеров", "тональн"]):
        return "вздохи_тональные"
    if any(token in desc for token in ["не хватает", "отсутствует звук", "отсутствуют", "недоста"]):
        return "отсутствие_синхронов"
    if any(token in desc for token in ["фонового шума", "фоновго шума", "звучание фона", "искажение фона", "скачок фона", "скачек фона"]):
        return "фон_атмосфера"
    if any(token in desc for token in ["несинхрон", "не синхр", "рассинхр"]):
        return "несинхрон"
    if any(token in desc for token in ["щелч", "слюн", "шип", "свист", "треск"]):
        return "артефакты"
    return "прочее"


def analyze_pairs(pairs: list[ReportPair]) -> dict:
    summary = {
        "total_pairs": len(pairs),
        "by_report_type": Counter(),
        "missing_csv": [],
        "missing_docx": [],
        "conclusion_line_counter": Counter(),
        "issue_category_counter": Counter(),
        "report_samples": [],
        "me_conclusion_patterns": Counter(),
        "main_conclusion_patterns": Counter(),
    }

    line_buckets = defaultdict(list)

    for pair in pairs:
        summary["by_report_type"][pair.report_type] += 1
        if pair.csv_path is None:
            summary["missing_csv"].append(str(pair.directory))
        if pair.docx_path is None:
            summary["missing_docx"].append(str(pair.directory))

        issues = load_csv_issues(pair.csv_path)
        conclusion = extract_subjective_conclusion(pair.docx_path)

        for issue in issues:
            summary["issue_category_counter"][classify_issue_description(issue["description"])] += 1

        for item in conclusion["items"]:
            summary["conclusion_line_counter"][item] += 1
            if pair.report_type == "me":
                summary["me_conclusion_patterns"][item] += 1
            else:
                summary["main_conclusion_patterns"][item] += 1
            line_buckets[item].append(
                {
                    "directory": str(pair.directory),
                    "csv": str(pair.csv_path) if pair.csv_path else None,
                    "docx": str(pair.docx_path) if pair.docx_path else None,
                }
            )

        summary["report_samples"].append(
            {
                "directory": str(pair.directory),
                "report_type": pair.report_type,
                "csv_path": str(pair.csv_path) if pair.csv_path else None,
                "docx_path": str(pair.docx_path) if pair.docx_path else None,
                "issue_count": len(issues),
                "blocker_count": sum(1 for issue in issues if issue["blocker"]),
                "conclusion_items": conclusion["items"],
                "issue_descriptions": [issue["description"] for issue in issues],
            }
        )

    summary["by_report_type"] = dict(summary["by_report_type"])
    summary["conclusion_line_counter"] = dict(summary["conclusion_line_counter"].most_common(200))
    summary["issue_category_counter"] = dict(summary["issue_category_counter"].most_common())
    summary["me_conclusion_patterns"] = dict(summary["me_conclusion_patterns"].most_common(120))
    summary["main_conclusion_patterns"] = dict(summary["main_conclusion_patterns"].most_common(120))
    summary["line_examples"] = {line: refs[:10] for line, refs in line_buckets.items() if len(refs) >= 2}
    return summary


def main() -> None:
    pairs = find_pairs(LEARN_ROOT)
    analysis = analyze_pairs(pairs)
    OUTPUT_JSON.write_text(json.dumps(analysis, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"pairs={analysis['total_pairs']}")
    print(f"by_report_type={analysis['by_report_type']}")
    print(f"missing_csv={len(analysis['missing_csv'])}")
    print(f"missing_docx={len(analysis['missing_docx'])}")
    print("top_issue_categories:")
    for key, value in list(analysis["issue_category_counter"].items())[:12]:
        print(f"  {key}: {value}")
    print("top_me_patterns:")
    for key, value in list(analysis["me_conclusion_patterns"].items())[:12]:
        print(f"  {value}x {key}")
    print("top_main_patterns:")
    for key, value in list(analysis["main_conclusion_patterns"].items())[:12]:
        print(f"  {value}x {key}")
    print(f"json={OUTPUT_JSON}")


if __name__ == "__main__":
    main()
