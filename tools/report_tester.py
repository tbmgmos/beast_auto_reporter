#!/usr/bin/env python3
"""
Standalone report tester:
- PDF parsing
- DOCX generation
- Table filling verification (logs table rows)
"""

import argparse
import logging
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

from src.exact_report_generator import ExactReportGenerator
from src.technical_info_extractor import TechnicalInfoExtractor
from src.csv_importer import CSVImporter
from src.pdf_extractor import PDFExtractor


def setup_logging(log_path: Path) -> logging.Logger:
    logger = logging.getLogger("report_tester")
    logger.setLevel(logging.INFO)
    logger.handlers.clear()

    fmt = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")

    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setLevel(logging.INFO)
    fh.setFormatter(fmt)
    logger.addHandler(fh)

    sh = logging.StreamHandler()
    sh.setLevel(logging.INFO)
    sh.setFormatter(fmt)
    logger.addHandler(sh)

    return logger


def detect_pdf_type(filename: str):
    name = filename.lower()
    is_51 = any(marker in name for marker in ['_51_', '_5.1_', ' 51 ', ' 5.1 ', '-51-', '-5.1-', '_51.', '51_', '5.1_', '51.', '.51', '.5.1'])
    is_20 = any(marker in name for marker in ['_20_', '_2.0_', ' 20 ', ' 2.0 ', '-20-', '-2.0-', '_20.', '20_', '2.0_', '20.', '.20', '.2.0'])
    is_cens = 'cens' in name and 'uncens' not in name
    is_uncens = 'uncens' in name

    if is_51 and is_cens:
        return "51_c"
    if is_51 and is_uncens:
        return "51_uc"
    if is_51:
        return "51"
    if is_20 and is_cens:
        return "20_c"
    if is_20 and is_uncens:
        return "20_uc"
    if is_20:
        return "20"
    return None


def collect_files(input_dir: Path):
    audio_ext = {".wav", ".aif", ".aiff", ".flac"}
    video_ext = {".mov", ".mp4", ".mxf"}
    csv_ext = {".csv"}
    pdf_ext = {".pdf"}

    audio = []
    video = []
    csv = []
    pdf = []

    for p in sorted(input_dir.iterdir()):
        if not p.is_file():
            continue
        ext = p.suffix.lower()
        if ext in audio_ext:
            audio.append(p)
        elif ext in video_ext:
            video.append(p)
        elif ext in csv_ext:
            csv.append(p)
        elif ext in pdf_ext:
            pdf.append(p)

    return audio, video, csv, pdf


def log_table(docx_path: Path, logger: logging.Logger):
    try:
        doc = Document(str(docx_path))
        logger.info(f"Tables in docx: {len(doc.tables)}")
        if not doc.tables:
            return
        table = doc.tables[0]
        logger.info("=== TABLE 0 ROWS ===")
        for i, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            logger.info(f"Row {i}: {cells}")
        logger.info("=== END TABLE 0 ===")
    except Exception as e:
        logger.error(f"Failed to parse docx for table logging: {e}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input_dir", help="Folder with source files (audio/video/csv/pdf)")
    ap.add_argument("--output-dir", default=None, help="Folder to write report output")
    ap.add_argument("--report-type", default="main", choices=["main", "me", "me_ours"])
    ap.add_argument("--prepared-by", default="")
    args = ap.parse_args()

    input_dir = Path(args.input_dir)
    if not input_dir.exists():
        raise SystemExit(f"Input folder not found: {input_dir}")

    output_dir = Path(args.output_dir) if args.output_dir else (input_dir / f"report_test_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    output_dir.mkdir(parents=True, exist_ok=True)

    log_path = output_dir / "report_test.log"
    logger = setup_logging(log_path)

    logger.info(f"Input: {input_dir}")
    logger.info(f"Output: {output_dir}")

    audio_files, video_files, csv_files, pdf_files = collect_files(input_dir)
    logger.info(f"Audio: {len(audio_files)} | Video: {len(video_files)} | CSV: {len(csv_files)} | PDF: {len(pdf_files)}")

    # Copy files to output folder (for isolated run)
    for p in audio_files + video_files + csv_files + pdf_files:
        shutil.copy2(p, output_dir / p.name)

    # Extract tech info
    tech_extractor = TechnicalInfoExtractor()
    pdf_extractor = PDFExtractor()

    tech_info = {}

    for p in audio_files:
        info = tech_extractor.extract_audio_info(str(p))
        if info:
            tech_info_key = f"audio_{p.stem.lower()}"
            tech_info[tech_info_key] = info

    if video_files:
        vinfo = tech_extractor.extract_video_info(str(video_files[0]))
        if vinfo:
            tech_info["video"] = vinfo

    # Parse PDFs
    pdf_paths = {"20_c": None, "20_uc": None, "20": None, "51_c": None, "51_uc": None, "51": None}
    for p in pdf_files:
        kind = detect_pdf_type(p.stem)
        if kind:
            pdf_paths[kind] = str(output_dir / p.name)
        pdf_data = pdf_extractor.extract_technical_info(str(p))
        if pdf_data:
            if kind:
                tech_info[f"pdf_{kind}"] = pdf_data

    # Issues from CSV
    issues = []
    if csv_files:
        try:
            importer = CSVImporter()
            issues = importer.import_issues(str(csv_files[0]))
        except Exception as e:
            logger.error(f"CSV import error: {e}")

    # Generate report
    base_name = audio_files[0].stem if audio_files else (video_files[0].stem if video_files else "report")
    report_path = output_dir / f"report_test_{base_name}.docx"

    gen = ExactReportGenerator()
    gen.create_exact_report(
        issues=issues,
        output_path=str(report_path),
        tech_info=tech_info,
        pdf_20_c_path=pdf_paths["20_c"],
        pdf_20_uc_path=pdf_paths["20_uc"],
        pdf_51_c_path=pdf_paths["51_c"],
        pdf_51_uc_path=pdf_paths["51_uc"],
        pdf_20_path=pdf_paths["20"],
        pdf_51_path=pdf_paths["51"],
        report_type=args.report_type,
        prepared_by=args.prepared_by
    )

    logger.info(f"Report created: {report_path}")
    log_table(report_path, logger)
    logger.info(f"Log saved: {log_path}")


if __name__ == "__main__":
    main()
