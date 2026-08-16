import io
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.report_filename import parse_report_filename
from src.report_uploader import (
    alias_key_for_path,
    compare_two_versions,
    compare_with_previous,
    diff_markers,
    fallback_series_key,
    forget_uploaded_reports,
    find_latest_report_in_folder,
    find_previous_report,
    find_series_folder,
    forget_series_alias,
    group_versions_by_category,
    list_report_versions,
    list_series_aliases,
    load_series_aliases,
    load_uploaded_reports,
    load_variant_overrides,
    NPR_ALIASES_FILE,
    remember_series_alias,
    remember_uploaded_report,
    resolve_manual_pick_target,
    resolve_target_path,
    set_variant_override,
    SeriesFolderNotFoundError,
    upload_folder,
    upload_paths_recursive,
)
from src.yandex_disk_client import YandexDiskError


META = parse_report_filename("отчет_Nepreklonniy_vozrast_s01_e02_MnE_2025_07_14_rus.docx")


def test_find_series_folder_exact_match(tmp_path):
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "Nepreklonniy_vozrast", "type": "dir", "path": "disk:/отчеты/Nepreklonniy_vozrast"},
        {"name": "Other_show", "type": "dir", "path": "disk:/отчеты/Other_show"},
    ]

    result = find_series_folder(client, "Nepreklonniy_vozrast", aliases_path=tmp_path / "aliases.json")

    assert result == "disk:/отчеты/Nepreklonniy_vozrast"
    client.list_folder.assert_called_once_with("/отчеты")


def test_find_series_folder_fuzzy_match_typo(tmp_path):
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "Nepreklonniy_vozrst", "type": "dir", "path": "disk:/отчеты/Nepreklonniy_vozrst"},
    ]

    result = find_series_folder(client, "Nepreklonniy_vozrast", aliases_path=tmp_path / "aliases.json")

    assert result == "disk:/отчеты/Nepreklonniy_vozrst"


def test_find_series_folder_no_match_returns_none(tmp_path):
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "Completely_different", "type": "dir", "path": "disk:/отчеты/Completely_different"},
    ]

    assert find_series_folder(client, "Nepreklonniy_vozrast", aliases_path=tmp_path / "aliases.json") is None


def test_resolve_target_path_raises_when_series_missing_and_not_creating(tmp_path):
    client = MagicMock()
    client.list_folder.return_value = []

    with pytest.raises(ValueError):
        resolve_target_path(client, META, create_if_missing=False, aliases_path=tmp_path / "aliases.json")


def test_resolve_target_path_raises_specifically_series_folder_not_found(tmp_path):
    # Вызывающий код (UI) должен ловить именно этот тип, а не парсить текст
    # сообщения об ошибке, чтобы отличить "нужна папка" от прочих сбоев.
    client = MagicMock()
    client.list_folder.return_value = []

    with pytest.raises(SeriesFolderNotFoundError):
        resolve_target_path(client, META, create_if_missing=False, aliases_path=tmp_path / "aliases.json")


def test_resolve_target_path_creates_series_and_episode_folders(tmp_path):
    client = MagicMock()
    client.list_folder.return_value = []

    path, created = resolve_target_path(client, META, create_if_missing=True, aliases_path=tmp_path / "aliases.json")

    assert path == "/отчеты/Nepreklonniy_vozrast/e02"
    assert created is True
    client.mkdir.assert_any_call("/отчеты/Nepreklonniy_vozrast")
    client.mkdir.assert_any_call("/отчеты/Nepreklonniy_vozrast/e02")


def test_resolve_target_path_reuses_existing_series_folder(tmp_path):
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "Nepreklonniy_vozrast", "type": "dir", "path": "disk:/отчеты/Nepreklonniy_vozrast"},
    ]

    path, created = resolve_target_path(client, META, create_if_missing=False, aliases_path=tmp_path / "aliases.json")

    assert path == "disk:/отчеты/Nepreklonniy_vozrast/e02"
    assert created is False


def test_list_report_versions_returns_all_matching_oldest_to_newest():
    # META (см. выше) — вариант "MnE": в выдачу должны попасть только
    # другие MnE-версии того же эпизода, не обычный (без variant) отчёт
    # рядом в той же папке — это отдельная, не связанная цепочка версий.
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Nepreklonniy_vozrast_s01_e02_MnE_2025_06_23_rus", "type": "dir",
         "path": "disk:/e02/v2"},
        {"name": "отчет_Nepreklonniy_vozrast_s01_e02_MnE_2025_05_19_rus", "type": "dir",
         "path": "disk:/e02/v1"},
        {"name": "отчет_Nepreklonniy_vozrast_s01_e02_2025_05_20_rus", "type": "dir",
         "path": "disk:/e02/no_variant"},
        {"name": "отчет_Nepreklonniy_vozrast_s01_e01_2025_01_01_rus", "type": "dir",
         "path": "disk:/e02/other_episode"},
        {"name": "readme.txt", "type": "file", "path": "disk:/e02/readme.txt"},
    ]

    versions = list_report_versions(client, "отчеты/Nepreklonniy_vozrast/e02", META)

    assert [v["path"] for v in versions] == ["disk:/e02/v1", "disk:/e02/v2"]
    assert versions[0]["date"].isoformat() == "2025-05-19"
    assert versions[1]["date"].isoformat() == "2025-06-23"


def test_list_report_versions_excludes_mismatched_variant():
    meta_no_variant = parse_report_filename("отчет_Nepreklonniy_vozrast_s01_e02_2025_07_14_rus.docx")
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Nepreklonniy_vozrast_s01_e02_2025_05_19_rus", "type": "dir",
         "path": "disk:/e02/plain"},
        {"name": "отчет_Nepreklonniy_vozrast_s01_e02_MnE_2025_05_19_rus", "type": "dir",
         "path": "disk:/e02/mne"},
    ]

    versions = list_report_versions(client, "отчеты/Nepreklonniy_vozrast/e02", meta_no_variant)

    assert [v["path"] for v in versions] == ["disk:/e02/plain"]


def test_list_report_versions_without_meta_returns_all_report_entries():
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_KP_Orlov_2026_03_27_v1", "type": "dir",
         "path": "disk:/e48/new", "modified": "2026-03-27T10:00:00+00:00"},
        {"name": "отчет_KP_Orlov_2026_03_20_v1", "type": "dir",
         "path": "disk:/e48/old", "modified": "2026-03-20T10:00:00+00:00"},
        {"name": "readme.txt", "type": "file", "path": "disk:/e48/readme.txt"},
    ]

    versions = list_report_versions(client, "отчеты/Show/e48")

    assert [v["path"] for v in versions] == ["disk:/e48/old", "disk:/e48/new"]


def test_list_report_versions_sorts_by_embedded_date_not_upload_time_for_nonstandard_names():
    # Реальный регресс: отчёт с нестандартным именем (не совпадает с
    # REPORT_PATTERN — другой формат сезона/эпизода), написанный ДАВНО, но
    # ЗАГРУЖЕННЫЙ на Диск только что — раньше сортировался по времени
    # загрузки (modified) и выглядел "самым новым", хотя по дате в имени
    # он самый старый.
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_MAZHOR_DUBAI_2025_03_01_rus", "type": "dir",
         "path": "disk:/e08/old_content_recent_upload", "modified": "2026-07-22T10:00:00+00:00"},
        {"name": "отчет_MAZHOR_DUBAI_2025_11_28_rus", "type": "dir",
         "path": "disk:/e08/new_content_older_upload", "modified": "2026-01-01T10:00:00+00:00"},
    ]

    versions = list_report_versions(client, "отчеты/Show/e08")

    # По дате В ИМЕНИ (2025-03-01 vs 2025-11-28) "old_content_recent_upload"
    # старше, несмотря на то что modified (дата загрузки) у него куда позже.
    assert [v["path"] for v in versions] == ["disk:/e08/old_content_recent_upload", "disk:/e08/new_content_older_upload"]


def test_group_versions_by_category_separates_variants():
    versions = [
        {"path": "/a/main", "label": "отчет_Show_s01_e05_2026_04_05_rus"},
        {"path": "/a/me", "label": "отчет_Show_s01_e05_ME_2026_04_05_rus"},
        {"path": "/a/ad", "label": "отчет_Show_s01_e05_AD_2026_04_05_rus"},
    ]

    groups = group_versions_by_category(versions)

    assert [v["path"] for v in groups["main"]] == ["/a/main"]
    assert [v["path"] for v in groups["me"]] == ["/a/me"]
    assert [v["path"] for v in groups["ad"]] == ["/a/ad"]


def test_group_versions_by_category_finds_marker_inside_compound_tag():
    # Реальный регресс: "cens_AD" не разбирается REPORT_PATTERN целиком,
    # но лёгкое сканирование всё равно находит "AD" по границам слова —
    # эта версия не должна смешиваться с группой "main".
    versions = [
        {"path": "/a/ad1", "label": "отчет_besprintsipnye_v_pitere_s01_e08_cens_AD_2025_06_11_rus"},
        {"path": "/a/mne1", "label": "отчет_Nepreklonniy_vozrast_s01_e08_MnE_2025_08_11_rus"},
    ]

    groups = group_versions_by_category(versions)

    assert [v["path"] for v in groups["ad"]] == ["/a/ad1"]
    assert [v["path"] for v in groups["me"]] == ["/a/mne1"]


def test_group_versions_by_category_respects_manual_override():
    versions = [{"path": "/a/weird", "label": "weird_folder_name"}]

    groups = group_versions_by_category(versions, overrides={"/a/weird": "AD"})

    assert [v["path"] for v in groups["ad"]] == ["/a/weird"]


def test_group_versions_by_category_dcp_override():
    versions = [{"path": "/a/dcp", "label": "DCP +18"}]

    groups = group_versions_by_category(versions, overrides={"/a/dcp": "DCP"})

    assert [v["path"] for v in groups["dcp"]] == ["/a/dcp"]


def test_group_versions_by_category_merges_cens_uncens_into_main():
    # CENS/UNCENS — признак цензурирования основного отчёта, а не отдельный
    # параллельный тип поставки (см. categorize_variant) — версии должны
    # сравниваться в ОДНОЙ цепочке с обычными "main"-версиями, а не
    # теряться в отдельной несвязанной группе.
    versions = [
        {"path": "/a/v1", "label": "отчет_Show_s01_e08_2025_06_10_rus"},
        {"path": "/a/v2", "label": "отчет_Show_s01_e08_2025_06_11_rus"},
        {"path": "/a/v3_uncens", "label": "отчет_Show_s01_e08_uncens_2025_06_23_rus"},
    ]

    groups = group_versions_by_category(versions)

    assert [v["path"] for v in groups["main"]] == ["/a/v1", "/a/v2", "/a/v3_uncens"]
    assert "other" not in groups


def test_group_versions_by_category_override_main_beats_auto_detected_marker():
    versions = [{"path": "/a/x", "label": "отчет_Show_s01_e05_ME_2026_04_05_rus"}]

    groups = group_versions_by_category(versions, overrides={"/a/x": "MAIN"})

    assert [v["path"] for v in groups["main"]] == ["/a/x"]


def test_find_previous_report_picks_latest_same_episode():
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Nepreklonniy_vozrast_s01_e02_2025_05_19_rus", "type": "dir",
         "path": "disk:/e02/v1"},
        {"name": "отчет_Nepreklonniy_vozrast_s01_e02_MnE_2025_06_23_rus", "type": "dir",
         "path": "disk:/e02/v2"},
        {"name": "отчет_Nepreklonniy_vozrast_s01_e01_2025_01_01_rus", "type": "dir",
         "path": "disk:/e02/other_episode"},
        {"name": "readme.txt", "type": "file", "path": "disk:/e02/readme.txt"},
    ]

    result = find_previous_report(client, "отчеты/Nepreklonniy_vozrast/e02", META)

    assert result == "disk:/e02/v2"


def test_find_previous_report_picks_across_legacy_flat_and_new_folder_structure():
    # До введения подпапок отчёт лежал плоским .docx прямо в папке эпизода;
    # после — в подпапке «отчет_...». Обе версии одного эпизода должны
    # находиться и сравниваться между собой по дате.
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Nepreklonniy_vozrast_s01_e02_2025_05_19_rus.docx", "type": "file",
         "path": "disk:/e02/legacy.docx"},
        {"name": "отчет_Nepreklonniy_vozrast_s01_e02_MnE_2025_06_23_rus", "type": "dir",
         "path": "disk:/e02/new_folder"},
    ]

    result = find_previous_report(client, "/отчеты/Nepreklonniy_vozrast/e02", META)

    assert result == "disk:/e02/new_folder"


def test_find_previous_report_returns_none_when_empty():
    client = MagicMock()
    client.list_folder.return_value = []

    assert find_previous_report(client, "отчеты/Nepreklonniy_vozrast/e02", META) is None


def test_find_previous_report_returns_none_when_episode_folder_not_created_yet():
    # Новый эпизод, который никогда раньше не заливали на Диск — папки ещё
    # нет, и list_folder отвечает 404 DiskNotFoundError. Это не ошибка,
    # а просто "предыдущих версий нет".
    client = MagicMock()
    client.list_folder.side_effect = YandexDiskError("Resource not found.", status_code=404)

    assert find_previous_report(client, "/отчеты/Show/e06", META) is None


def test_find_series_folder_returns_none_when_root_not_created_yet(tmp_path):
    client = MagicMock()
    client.list_folder.side_effect = YandexDiskError("Resource not found.", status_code=404)

    assert find_series_folder(client, "Nepreklonniy_vozrast", aliases_path=tmp_path / "aliases.json") is None


def test_find_previous_report_propagates_non_404_errors():
    client = MagicMock()
    client.list_folder.side_effect = YandexDiskError("Server error", status_code=500)

    with pytest.raises(YandexDiskError):
        find_previous_report(client, "/отчеты/Show/e06", META)


def test_find_latest_report_in_folder_picks_most_recently_modified():
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_KP_Orlov_2026_03_20_v1", "type": "dir",
         "path": "disk:/e48/old", "modified": "2026-03-20T10:00:00+00:00"},
        {"name": "отчет_KP_Orlov_2026_03_27_v1", "type": "dir",
         "path": "disk:/e48/new", "modified": "2026-03-27T10:00:00+00:00"},
        {"name": "readme.txt", "type": "file", "path": "disk:/e48/readme.txt", "modified": "2026-03-28T10:00:00+00:00"},
    ]

    result = find_latest_report_in_folder(client, "отчеты/Show/e48")

    assert result == "disk:/e48/new"


def test_find_latest_report_in_folder_returns_none_when_no_reports():
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "readme.txt", "type": "file", "path": "disk:/e48/readme.txt", "modified": "2026-03-28T10:00:00+00:00"},
    ]

    assert find_latest_report_in_folder(client, "отчеты/Show/e48") is None


def test_compare_with_previous_finds_docx_inside_report_folder(tmp_path):
    new_docx_path = tmp_path / "отчет_Show.docx"
    Document().save(new_docx_path)

    old_bytes_io = io.BytesIO()
    Document().save(old_bytes_io)

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Show_s01_e02_2025_06_23_rus.docx", "type": "file",
         "path": "disk:/e02/v1/отчет_Show_s01_e02_2025_06_23_rus.docx"},
    ]
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e02/v1", new_docx_path)

    client.download_bytes.assert_called_once_with("disk:/e02/v1/отчет_Show_s01_e02_2025_06_23_rus.docx")
    assert comparison is not None


def test_compare_with_previous_works_with_legacy_flat_docx_path(tmp_path):
    # Старая структура: previous_report_path указывает прямо на .docx-файл,
    # а не на подпапку — сравнение должно работать и в этом случае, без
    # похода в list_folder за поиском файла внутри.
    new_docx_path = tmp_path / "отчет_Show.docx"
    Document().save(new_docx_path)

    old_bytes_io = io.BytesIO()
    Document().save(old_bytes_io)

    client = MagicMock()
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e02/legacy.docx", new_docx_path)

    client.download_bytes.assert_called_once_with("disk:/e02/legacy.docx")
    client.list_folder.assert_not_called()
    assert comparison is not None


def test_compare_with_previous_returns_none_when_no_docx_in_folder(tmp_path):
    new_docx_path = tmp_path / "отчет_Show.docx"
    Document().save(new_docx_path)

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "source.pdf", "type": "file", "path": "disk:/e02/v1/source.pdf"},
    ]

    assert compare_with_previous(client, "disk:/e02/v1", new_docx_path) is None
    client.download_bytes.assert_not_called()


def test_compare_with_previous_finds_docx_without_report_prefix_in_legacy_folder(tmp_path):
    # Реальный регресс: папки, загруженные ДО введения соглашения об имени
    # с префиксом «отчет_», могут содержать .docx без этого префикса и в
    # имени самого файла (не только папки) — например
    # «one_last_sin_s01_e01_Master_uncens_2025_05_14.docx». Раньше это
    # давало «Не удалось прочитать выбранную версию отчёта» на совершенно
    # читаемом файле.
    new_docx_path = tmp_path / "отчет_Show.docx"
    Document().save(new_docx_path)

    old_bytes_io = io.BytesIO()
    Document().save(old_bytes_io)

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "one_last_sin_s01_e01_Master_uncens_2025_05_14.docx", "type": "file",
         "path": "disk:/e01/legacy/one_last_sin_s01_e01_Master_uncens_2025_05_14.docx"},
        {"name": "~$one_last_sin_s01_e01_Master_uncens_2025_05_14.docx", "type": "file",
         "path": "disk:/e01/legacy/~$one_last_sin_s01_e01_Master_uncens_2025_05_14.docx"},
    ]
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e01/legacy", new_docx_path)

    client.download_bytes.assert_called_once_with(
        "disk:/e01/legacy/one_last_sin_s01_e01_Master_uncens_2025_05_14.docx"
    )
    assert comparison is not None


def test_compare_with_previous_prefers_report_prefixed_docx_over_legacy_one(tmp_path):
    # Если в папке ЕСТЬ файл с современным префиксом — он приоритетнее
    # любого другого .docx рядом, даже если тот встретился в списке раньше.
    new_docx_path = tmp_path / "отчет_Show.docx"
    Document().save(new_docx_path)

    old_bytes_io = io.BytesIO()
    Document().save(old_bytes_io)

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "unrelated_notes.docx", "type": "file", "path": "disk:/e01/v1/unrelated_notes.docx"},
        {"name": "отчет_Show_s01_e01_2025_06_23_rus.docx", "type": "file",
         "path": "disk:/e01/v1/отчет_Show_s01_e01_2025_06_23_rus.docx"},
    ]
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e01/v1", new_docx_path)

    client.download_bytes.assert_called_once_with("disk:/e01/v1/отчет_Show_s01_e01_2025_06_23_rus.docx")
    assert comparison is not None


def _set_cell_bg(cell, fill: str):
    """Имитирует заливку ячейки, как это делает ExactReportGenerator._format_cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _make_report_docx(path, *, markers, track_params, track_cell_bg=None):
    """Собирает .docx с таблицей маркеров (как MARKER LIST) и таблицей

    параметров (как в _add_technical_table_with_conclusion), для тестов
    подсчёта маркеров/параметров. markers — либо число (строки только с
    таймкодом), либо список словарей {tc_in, tc_out, description, blocker,
    comments}. track_cell_bg — необязательная заливка конкретных ячеек
    параметров: {"2.0 cens": {"LOUDNESS": "E73322"}}.
    """
    doc = Document()
    track_cell_bg = track_cell_bg or {}

    marker_rows = ([{"tc_in": f"00:00:{i:02d}"} for i in range(markers)]
                   if isinstance(markers, int) else markers)
    marker_table = doc.add_table(rows=1 + len(marker_rows), cols=8)
    headers = ["Timecode In", "Timecode Out", "Description", "2.0 C",
               "БЛОКЕР", "ТРЕБУЕТ ИСПРАВЛЕНИЯ", "ТРЕБУЕТ КОММЕНТАРИЯ", "КОММЕНТАРИИ"]
    for col, header in enumerate(headers):
        marker_table.rows[0].cells[col].text = header
    for row_idx, marker in enumerate(marker_rows):
        row = marker_table.rows[1 + row_idx]
        row.cells[0].text = marker.get("tc_in", "")
        row.cells[1].text = marker.get("tc_out", "")
        row.cells[2].text = marker.get("description", "")
        row.cells[4].text = "*" if marker.get("blocker") else ""
        row.cells[7].text = marker.get("comments", "")

    param_headers = ["ДОРОЖКА", "НАЗВАНИЕ ФАЙЛОВ", "ХРОНОМЕТРАЖ", "LOUDNESS", "TRUE PEAK", "LRA", "ФОРМАТ ФАЙЛА"]
    param_table = doc.add_table(rows=1 + len(track_params), cols=len(param_headers))
    for col, header in enumerate(param_headers):
        param_table.rows[0].cells[col].text = header
    for row_idx, (label, values) in enumerate(track_params.items()):
        row = param_table.rows[1 + row_idx]
        row.cells[0].text = label
        for col, header in enumerate(param_headers[1:], start=1):
            row.cells[col].text = values.get(header, "")
            bg = track_cell_bg.get(label, {}).get(header)
            if bg:
                _set_cell_bg(row.cells[col], bg)

    doc.save(path)


def test_compare_with_previous_counts_markers(tmp_path):
    new_docx_path = tmp_path / "отчет_Show.docx"
    _make_report_docx(
        new_docx_path, markers=3,
        track_params={"2.0 cens": {"LOUDNESS": "-23.0 LUFS", "TRUE PEAK": "-2.0 dBTP", "LRA": "8.0 LU"}},
    )

    old_bytes_io = io.BytesIO()
    _make_report_docx(
        old_bytes_io, markers=2,
        track_params={"2.0 cens": {"LOUDNESS": "-23.0 LUFS", "TRUE PEAK": "-2.0 dBTP", "LRA": "8.0 LU"}},
    )

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Show_s01_e02_2025_06_23_rus.docx", "type": "file",
         "path": "disk:/e02/v1/отчет_Show_s01_e02_2025_06_23_rus.docx"},
    ]
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e02/v1", new_docx_path)

    assert (comparison.marker_count_old, comparison.marker_count_new) == (2, 3)
    assert comparison.parameter_changes == []


def test_compare_with_previous_finds_marker_table_regardless_of_header_case(tmp_path):
    # Реальный регресс: проверка заголовка таблицы маркеров была
    # регистрозависимой ("Timecode In" ровно так), в отличие от всех
    # остальных сравнений заголовков в этой же функции (БЛОКЕР/КОММЕНТАРИИ/
    # TIMECODE OUT/DESCRIPTION — все через .upper()). Если в конкретном
    # .docx заголовок сохранился в другом регистре (например, из другого
    # источника, не сгенерированного этим приложением), таблица маркеров
    # молча не находилась вообще — сравнение показывало 0 маркеров для
    # обеих версий сразу, даже если реально они были.
    def _make_docx_with_uppercase_headers(path, marker_count):
        doc = Document()
        table = doc.add_table(rows=1 + marker_count, cols=3)
        headers = ["TIMECODE IN", "timecode out", "Description"]  # разный регистр
        for col, header in enumerate(headers):
            table.rows[0].cells[col].text = header
        for row_idx in range(marker_count):
            table.rows[1 + row_idx].cells[0].text = f"00:00:{row_idx:02d}"
        doc.save(path)

    new_docx_path = tmp_path / "отчет_Show.docx"
    _make_docx_with_uppercase_headers(new_docx_path, marker_count=3)

    old_bytes_io = io.BytesIO()
    _make_docx_with_uppercase_headers(old_bytes_io, marker_count=2)

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Show_s01_e02_2025_06_23_rus.docx", "type": "file",
         "path": "disk:/e02/v1/отчет_Show_s01_e02_2025_06_23_rus.docx"},
    ]
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e02/v1", new_docx_path)

    assert (comparison.marker_count_old, comparison.marker_count_new) == (2, 3)


def test_compare_with_previous_finds_marker_table_with_merged_title_row(tmp_path):
    # Реальный регресс (старый формат отчёта): первая строка таблицы —
    # объединённый титул «MARKER LIST» на всю ширину, а сами заголовки
    # колонок («Timecode In»/…) — во ВТОРОЙ строке. Раньше парсер смотрел
    # только row[0], видел «MARKER LIST», не находил «Timecode In» и молча
    # пропускал всю таблицу — сравнение показывало 0 маркеров для обеих
    # версий, хотя реально их были десятки.
    def _make_docx_with_merged_title_row(path, marker_count):
        doc = Document()
        table = doc.add_table(rows=2 + marker_count, cols=4)
        # row 0 — объединённый титул (эмулируем: одинаковый текст во всех ячейках)
        for cell in table.rows[0].cells:
            cell.text = "MARKER LIST"
        # row 1 — настоящие заголовки колонок
        for col, header in enumerate(["Timecode In", "Timecode Out", "Description", "БЛОКЕР"]):
            table.rows[1].cells[col].text = header
        # row 2+ — данные
        for row_idx in range(marker_count):
            table.rows[2 + row_idx].cells[0].text = f"00:00:{row_idx:02d}"
            table.rows[2 + row_idx].cells[3].text = "*" if row_idx == 0 else ""
        doc.save(path)

    new_docx_path = tmp_path / "отчет_Show.docx"
    _make_docx_with_merged_title_row(new_docx_path, marker_count=3)

    old_bytes_io = io.BytesIO()
    _make_docx_with_merged_title_row(old_bytes_io, marker_count=2)

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Show_s01_e02_2025_06_23_rus.docx", "type": "file",
         "path": "disk:/e02/v1/отчет_Show_s01_e02_2025_06_23_rus.docx"},
    ]
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e02/v1", new_docx_path)

    assert (comparison.marker_count_old, comparison.marker_count_new) == (2, 3)
    assert (comparison.blocker_count_old, comparison.blocker_count_new) == (1, 1)


def _nest_table_in_textbox(doc, table):
    # Переносит уже созданную таблицу внутрь текстового блока (w:txbxContent),
    # эмулируя структуру .docx, экспортированного из Pages: таблица перестаёт
    # быть прямым потомком тела документа, поэтому штатный doc.tables её не
    # видит (возвращает пустой список), а _iter_all_tables — находит.
    from src.report_uploader import _summarize_document  # noqa: F401 (гарантия импорта модуля)
    tbl_el = table._tbl
    body = doc.element.body
    body.remove(tbl_el)
    wrapper = OxmlElement("w:p")
    txbx = OxmlElement("w:txbxContent")
    txbx.append(tbl_el)
    wrapper.append(txbx)
    body.append(wrapper)


def test_summarize_document_reads_tables_nested_in_textboxes():
    # Реальный регресс: у старых отчётов (экспорт из Pages) и таблица
    # маркеров, и таблица параметров лежат ВНУТРИ текстовых блоков, а не на
    # верхнем уровне тела документа. python-docx doc.tables возвращает
    # только таблицы верхнего уровня и пропускает вложенные — сравнение
    # видело 0 маркеров и «без изменений» по параметрам, хотя данные есть.
    from src.report_uploader import _summarize_document

    doc = Document()
    # Таблица маркеров: титул «MARKER LIST», строка заголовков, 2 маркера +
    # пустая строка-заглушка (как в реальных отчётах Pages).
    marker_table = doc.add_table(rows=5, cols=4)
    for cell in marker_table.rows[0].cells:
        cell.text = "MARKER LIST"
    for col, header in enumerate(["Timecode In", "Timecode Out", "Description", "БЛОКЕР"]):
        marker_table.rows[1].cells[col].text = header
    marker_table.rows[2].cells[0].text = "01:00:00:00"
    marker_table.rows[2].cells[3].text = "*"
    marker_table.rows[3].cells[0].text = "01:00:05:00"
    # rows[4] оставляем полностью пустой — не должна считаться маркером
    _nest_table_in_textbox(doc, marker_table)

    # Таблица параметров — тоже вложенная.
    param_table = doc.add_table(rows=2, cols=4)
    for col, header in enumerate(["ДОРОЖКА", "ХРОНОМЕТРАЖ", "LOUDNESS", "TRUE PEAK"]):
        param_table.rows[0].cells[col].text = header
    for col, val in enumerate(["2.0 cens", "0:35:00", "-23.0 LUFS", "-2.0 dBTP"]):
        param_table.rows[1].cells[col].text = val
    _nest_table_in_textbox(doc, param_table)

    assert doc.tables == []  # штатный API вложенные таблицы не видит (репро бага)

    summary = _summarize_document(doc)

    assert summary.marker_count == 2  # 2 реальных маркера, пустая строка не в счёт
    assert summary.blocker_count == 1
    assert list(summary.parameters.keys()) == ["2.0 cens"]
    assert summary.parameters["2.0 cens"]["LOUDNESS"] == "-23.0 LUFS"


def test_compare_with_previous_reports_parameter_changes(tmp_path):
    new_docx_path = tmp_path / "отчет_Show.docx"
    _make_report_docx(
        new_docx_path, markers=0,
        track_params={"2.0 cens": {"LOUDNESS": "-24.5 LUFS", "TRUE PEAK": "-2.0 dBTP", "LRA": "8.0 LU"}},
    )

    old_bytes_io = io.BytesIO()
    _make_report_docx(
        old_bytes_io, markers=0,
        track_params={"2.0 cens": {"LOUDNESS": "-23.0 LUFS", "TRUE PEAK": "-2.0 dBTP", "LRA": "8.0 LU"}},
    )

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Show_s01_e02_2025_06_23_rus.docx", "type": "file",
         "path": "disk:/e02/v1/отчет_Show_s01_e02_2025_06_23_rus.docx"},
    ]
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e02/v1", new_docx_path)

    assert len(comparison.parameter_changes) == 1
    change = comparison.parameter_changes[0]
    assert change["label"] == "2.0 cens"
    loudness_change = next(c for c in change["changes"] if c["field"] == "Громкость")
    assert loudness_change["old"] == "-23.0 LUFS"
    assert loudness_change["new"] == "-24.5 LUFS"


def test_summarize_document_reads_true_peak_from_merged_chronometrage_me_table():
    # Реальная структура M&E-таблицы (см. exact_report_generator._generate_me_table):
    # заголовок «ХРОНОМЕТРАЖ» объединён на 3 сетевые колонки (python-docx
    # повторяет текст объединённой ячейки для каждой из них), а колонки
    # LOUDNESS/LRA в этом шаблоне вообще отсутствуют — только TRUE PEAK.
    # Раньше было подозрение, что из-за повтора заголовка парсер теряет
    # True Peak — на реальном файле (отчет_ulichnaya_eda..._ME_...) и в
    # этом тесте показано, что значение считывается корректно и не теряется.
    from src.report_uploader import _summarize_document

    doc = Document()
    table = doc.add_table(rows=2, cols=7)
    header_row = table.rows[0]
    header_row.cells[0].text = "ДОРОЖКА"
    header_row.cells[1].text = "НАЗВАНИЕ ФАЙЛОВ"
    chrono = header_row.cells[2]
    chrono.merge(header_row.cells[3])
    chrono.merge(header_row.cells[4])
    chrono.text = "ХРОНОМЕТРАЖ"
    header_row.cells[5].text = "TRUE PEAK"
    header_row.cells[6].text = "ФОРМАТ ФАЙЛА"

    data_row = table.rows[1]
    data_row.cells[0].text = "2.0 ME"
    data_row.cells[1].text = "ulichnaya_eda_s01_e05_20_ME_2026_04_05_v1"
    data_row.cells[2].text = "0:37:44.916"
    data_row.cells[5].text = "-0.5 dBTP"
    data_row.cells[6].text = "PCM 48kHz 24 bit L R"

    summary = _summarize_document(doc)

    assert summary.parameters["2.0 ME"]["TRUE PEAK"] == "-0.5 dBTP"
    assert "LOUDNESS" not in summary.parameters["2.0 ME"]
    assert "LRA" not in summary.parameters["2.0 ME"]


def test_summarize_document_reads_sample_peak_for_dcp_reports():
    # DCP-отчёты используют колонку «SAMPLE PEAK» вместо «TRUE PEAK» (см.
    # exact_report_generator.py:189, is_dcp_report) — должна читаться
    # так же надёжно, под своим собственным ключом.
    from src.report_uploader import _summarize_document

    doc = Document()
    table = doc.add_table(rows=2, cols=7)
    headers = ["ДОРОЖКА", "НАЗВАНИЕ ФАЙЛОВ", "ХРОНОМЕТРАЖ", "LOUDNESS", "SAMPLE PEAK", "LRA", "ФОРМАТ ФАЙЛА"]
    for col, header in enumerate(headers):
        table.rows[0].cells[col].text = header
    data_row = table.rows[1]
    data_row.cells[0].text = "5.1 cens"
    data_row.cells[3].text = "-14.7 LUFS"
    data_row.cells[4].text = "+0.60 dBFS"
    data_row.cells[5].text = "20.8 LU"

    summary = _summarize_document(doc)

    assert summary.parameters["5.1 cens"]["SAMPLE PEAK"] == "+0.60 dBFS"
    assert summary.parameters["5.1 cens"]["LOUDNESS"] == "-14.7 LUFS"


def test_compare_two_versions_compares_two_remote_reports_without_local_draft(tmp_path):
    # Сравнение v1 и v4, обе версии уже на Диске — локальный черновик не участвует.
    v1_bytes_io = io.BytesIO()
    _make_report_docx(v1_bytes_io, markers=2, track_params={})

    v4_bytes_io = io.BytesIO()
    _make_report_docx(v4_bytes_io, markers=5, track_params={})

    client = MagicMock()

    def fake_list_folder(path):
        if path == "disk:/e02/v1":
            return [{"name": "отчет_Show_s01_e02_2025_01_01_rus.docx", "type": "file",
                     "path": "disk:/e02/v1/отчет_Show_s01_e02_2025_01_01_rus.docx"}]
        if path == "disk:/e02/v4":
            return [{"name": "отчет_Show_s01_e02_2025_04_01_rus.docx", "type": "file",
                     "path": "disk:/e02/v4/отчет_Show_s01_e02_2025_04_01_rus.docx"}]
        raise AssertionError(f"unexpected list_folder call: {path}")

    def fake_download_bytes(path):
        if path == "disk:/e02/v1/отчет_Show_s01_e02_2025_01_01_rus.docx":
            return v1_bytes_io.getvalue()
        if path == "disk:/e02/v4/отчет_Show_s01_e02_2025_04_01_rus.docx":
            return v4_bytes_io.getvalue()
        raise AssertionError(f"unexpected download_bytes call: {path}")

    client.list_folder.side_effect = fake_list_folder
    client.download_bytes.side_effect = fake_download_bytes

    comparison = compare_two_versions(client, "disk:/e02/v1", "disk:/e02/v4")

    assert (comparison.marker_count_old, comparison.marker_count_new) == (2, 5)


def test_compare_two_versions_returns_none_when_one_side_missing(tmp_path):
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "source.pdf", "type": "file", "path": "disk:/e02/v1/source.pdf"},
    ]

    assert compare_two_versions(client, "disk:/e02/v1", "disk:/e02/v4.docx") is None


def test_compare_with_previous_flags_out_of_spec_parameter_as_bad(tmp_path):
    # Ячейка нового значения залита красным (как BAD_BG в самом отчёте) —
    # значит громкость вышла за допустимые пределы, помечаем "bad".
    new_docx_path = tmp_path / "отчет_Show.docx"
    _make_report_docx(
        new_docx_path, markers=0,
        track_params={"2.0 cens": {"LOUDNESS": "-19.0 LUFS"}},
        track_cell_bg={"2.0 cens": {"LOUDNESS": "E73322"}},
    )

    old_bytes_io = io.BytesIO()
    _make_report_docx(
        old_bytes_io, markers=0,
        track_params={"2.0 cens": {"LOUDNESS": "-23.0 LUFS"}},
    )

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Show_s01_e02_2025_06_23_rus.docx", "type": "file",
         "path": "disk:/e02/v1/отчет_Show_s01_e02_2025_06_23_rus.docx"},
    ]
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e02/v1", new_docx_path)

    change = comparison.parameter_changes[0]["changes"][0]
    assert change["field"] == "Громкость"
    assert change["new"] == "-19.0 LUFS"
    assert change["status"] == "bad"


def test_compare_with_previous_parameter_status_none_when_no_highlight(tmp_path):
    new_docx_path = tmp_path / "отчет_Show.docx"
    _make_report_docx(
        new_docx_path, markers=0,
        track_params={"2.0 cens": {"LOUDNESS": "-23.5 LUFS"}},
    )

    old_bytes_io = io.BytesIO()
    _make_report_docx(
        old_bytes_io, markers=0,
        track_params={"2.0 cens": {"LOUDNESS": "-23.0 LUFS"}},
    )

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Show_s01_e02_2025_06_23_rus.docx", "type": "file",
         "path": "disk:/e02/v1/отчет_Show_s01_e02_2025_06_23_rus.docx"},
    ]
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e02/v1", new_docx_path)

    assert comparison.parameter_changes[0]["changes"][0]["status"] is None


def test_upload_folder_uploads_all_files_non_recursively(tmp_path):
    (tmp_path / "отчет_Show.docx").write_text("doc")
    (tmp_path / "Show.pdf").write_text("pdf")
    (tmp_path / "Show.csv").write_text("csv")
    subdir = tmp_path / "subdir"
    subdir.mkdir()
    (subdir / "nested.txt").write_text("nested")
    (tmp_path / ".DS_Store").write_text("junk")  # дотфайлы не загружаются
    (tmp_path / "~$отчет_Show.docx").write_text("lock")  # lock-файл Word не загружается

    client = MagicMock()

    uploaded = upload_folder(client, tmp_path, "отчеты/Show/e48")

    assert client.upload.call_count == 3
    assert sorted(uploaded) == [
        "отчеты/Show/e48/Show.csv",
        "отчеты/Show/e48/Show.pdf",
        "отчеты/Show/e48/отчет_Show.docx",
    ]


def test_upload_paths_recursive_handles_mixed_files_and_nested_folder(tmp_path):
    # Как Finder отдаёт drop: обычный файл рядом с папкой (в т.ч. с
    # вложенной подпапкой) одним списком local_paths.
    single_file = tmp_path / "extra.wav"
    single_file.write_text("audio")

    folder = tmp_path / "Project"
    folder.mkdir()
    (folder / "top.txt").write_text("top")
    (folder / ".DS_Store").write_text("junk")  # дотфайлы не загружаются
    nested = folder / "sub"
    nested.mkdir()
    (nested / "deep.txt").write_text("deep")

    client = MagicMock()
    progress_calls = []

    upload_paths_recursive(
        client, [single_file, folder], "отчеты/Show/e01",
        progress_callback=lambda done, total: progress_calls.append((done, total)),
    )

    uploaded_remote_paths = sorted(call.args[1] for call in client.upload.call_args_list)
    assert uploaded_remote_paths == [
        "отчеты/Show/e01/Project/sub/deep.txt",
        "отчеты/Show/e01/Project/top.txt",
        "отчеты/Show/e01/extra.wav",
    ]
    # Родительская папка "Project" должна была быть создана до подпапки "sub"
    mkdir_paths = [call.args[0] for call in client.mkdir.call_args_list]
    assert mkdir_paths.index("отчеты/Show/e01/Project") < mkdir_paths.index("отчеты/Show/e01/Project/sub")
    assert progress_calls[-1] == (3, 3)


def test_upload_paths_recursive_does_not_recreate_same_folder_twice(tmp_path):
    folder = tmp_path / "Project"
    folder.mkdir()
    (folder / "a.txt").write_text("a")
    (folder / "b.txt").write_text("b")

    client = MagicMock()
    upload_paths_recursive(client, [folder], "отчеты/Show/e01")

    mkdir_calls = [call.args[0] for call in client.mkdir.call_args_list]
    assert mkdir_calls.count("отчеты/Show/e01/Project") == 1


def test_diff_markers_detects_added_removed_and_changed():
    old = [
        {"tc_in": "01:00:05", "tc_out": "01:00:07", "description": "провал громкости", "blocker": False, "comments": ""},
        {"tc_in": "01:02:00", "tc_out": "01:02:03", "description": "щелчок", "blocker": False, "comments": ""},
    ]
    new = [
        {"tc_in": "01:00:05", "tc_out": "01:00:07", "description": "провал громкости", "blocker": True, "comments": ""},
        {"tc_in": "01:05:00", "tc_out": "01:05:02", "description": "рассинхрон", "blocker": False, "comments": ""},
    ]

    diff = diff_markers(old, new)

    assert [m["tc_in"] for m in diff["added"]] == ["01:05:00"]
    assert [m["tc_in"] for m in diff["removed"]] == ["01:02:00"]
    assert len(diff["changed"]) == 1
    change = diff["changed"][0]
    assert change["tc_in"] == "01:00:05"
    assert change["changes"] == [{"field": "Блокер", "old": "нет", "new": "да"}]


def test_diff_markers_handles_duplicate_timecodes_positionally():
    # Два маркера на одном таймкоде: удаление одного из них — это
    # «удалён», а не ложное «изменён» для оставшегося.
    old = [
        {"tc_in": "01:00:00", "tc_out": "", "description": "первый", "blocker": False, "comments": ""},
        {"tc_in": "01:00:00", "tc_out": "", "description": "второй", "blocker": False, "comments": ""},
    ]
    new = [
        {"tc_in": "01:00:00", "tc_out": "", "description": "первый", "blocker": False, "comments": ""},
    ]

    diff = diff_markers(old, new)

    assert diff["changed"] == []
    assert [m["description"] for m in diff["removed"]] == ["второй"]
    assert diff["added"] == []


def test_diff_markers_no_changes_returns_empty_sections():
    markers = [{"tc_in": "01:00:00", "tc_out": "", "description": "шум", "blocker": False, "comments": ""}]
    assert diff_markers(markers, markers) == {"added": [], "removed": [], "changed": []}


def test_compare_with_previous_builds_marker_diff(tmp_path):
    new_docx_path = tmp_path / "отчет_Show.docx"
    _make_report_docx(
        new_docx_path,
        markers=[
            {"tc_in": "01:00:05", "description": "провал громкости стал глубже"},
            {"tc_in": "01:09:00", "description": "новый щелчок"},
        ],
        track_params={},
    )

    old_bytes_io = io.BytesIO()
    _make_report_docx(
        old_bytes_io,
        markers=[{"tc_in": "01:00:05", "description": "провал громкости"}],
        track_params={},
    )

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Show_s01_e02_2025_06_23_rus.docx", "type": "file",
         "path": "disk:/e02/v1/отчет_Show_s01_e02_2025_06_23_rus.docx"},
    ]
    client.download_bytes.return_value = old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/e02/v1", new_docx_path)

    assert [m["tc_in"] for m in comparison.marker_diff["added"]] == ["01:09:00"]
    assert comparison.marker_diff["removed"] == []
    change = comparison.marker_diff["changed"][0]
    assert change["tc_in"] == "01:00:05"
    assert change["changes"] == [
        {"field": "Описание", "old": "провал громкости", "new": "провал громкости стал глубже"},
    ]


def test_compare_with_previous_prefers_csv_persistent_id_for_moved_marker(tmp_path):
    marker_id = "M1"
    new_docx_path = tmp_path / "отчет_Show.docx"
    _make_report_docx(
        new_docx_path,
        markers=[{"tc_in": "01:00:10", "description": "Щелчок исправлен"}],
        track_params={},
    )
    (tmp_path / "Show.csv").write_text(
        f"ID\tTimecode In\tTimecode Out\tDescription\n"
        f"{marker_id}\t01:00:10:00\t\tЩелчок исправлен\n",
        encoding="utf-8",
    )

    old_bytes_io = io.BytesIO()
    _make_report_docx(
        old_bytes_io,
        markers=[{"tc_in": "01:00:01", "description": "Щелчок"}],
        track_params={},
    )
    old_csv = (
        f"ID\tTimecode In\tTimecode Out\tDescription\n"
        f"{marker_id}\t01:00:01:00\t\tЩелчок\n"
    ).encode("utf-8")

    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "отчет_Show.docx", "type": "file", "path": "disk:/v1/отчет_Show.docx"},
        {"name": "Show.csv", "type": "file", "path": "disk:/v1/Show.csv"},
    ]
    client.download_bytes.side_effect = lambda path: old_csv if path.endswith(".csv") else old_bytes_io.getvalue()

    comparison = compare_with_previous(client, "disk:/v1", new_docx_path)

    assert comparison.marker_diff["added"] == []
    assert comparison.marker_diff["removed"] == []
    assert comparison.marker_diff["changed"][0]["id"] == marker_id
    assert {change["field"] for change in comparison.marker_diff["changed"][0]["changes"]} == {
        "Timecode In", "Описание",
    }


def test_forget_uploaded_reports_removes_only_listed_entries(tmp_path):
    path = tmp_path / "uploaded.json"
    remember_uploaded_report("/local/a", "/отчеты/a", path)
    remember_uploaded_report("/local/b", "/отчеты/b", path)

    forget_uploaded_reports([{"remote_path": "/отчеты/a"}], path)

    remaining = load_uploaded_reports(path)
    assert [e["remote_path"] for e in remaining] == ["/отчеты/b"]


def test_forget_uploaded_reports_noop_for_unknown_paths(tmp_path):
    path = tmp_path / "uploaded.json"
    remember_uploaded_report("/local/a", "/отчеты/a", path)

    forget_uploaded_reports([{"remote_path": "/отчеты/нет_такого"}], path)

    assert len(load_uploaded_reports(path)) == 1


def test_series_aliases_round_trip(tmp_path):
    path = tmp_path / "aliases.json"
    remember_series_alias("Nepreklonniy_vozrast", "/отчеты/Непреклонный_возраст", path)

    assert load_series_aliases(path) == {"nepreklonniy_vozrast": "/отчеты/Непреклонный_возраст"}


def test_load_series_aliases_missing_file_returns_empty(tmp_path):
    assert load_series_aliases(tmp_path / "missing.json") == {}


def test_load_series_aliases_corrupt_file_returns_empty(tmp_path):
    path = tmp_path / "aliases.json"
    path.write_text("not valid json{{{", encoding="utf-8")

    assert load_series_aliases(path) == {}


def test_forget_series_alias_removes_entry(tmp_path):
    path = tmp_path / "aliases.json"
    remember_series_alias("Nepreklonniy_vozrast", "/отчеты/X", path)

    forget_series_alias("Nepreklonniy_vozrast", path)

    assert load_series_aliases(path) == {}


def test_forget_series_alias_missing_entry_is_noop(tmp_path):
    path = tmp_path / "aliases.json"

    forget_series_alias("Nepreklonniy_vozrast", path)  # не должно поднимать исключение

    assert load_series_aliases(path) == {}


def test_resolve_manual_pick_target_without_meta_returns_path_unchanged():
    assert resolve_manual_pick_target("/отчеты/Show", None) == ("/отчеты/Show", "/отчеты/Show")


def test_resolve_manual_pick_target_series_folder_appends_episode():
    episode_path, series_path = resolve_manual_pick_target("/отчеты/Show", META)
    assert episode_path == "/отчеты/Show/e02"
    assert series_path == "/отчеты/Show"


def test_resolve_manual_pick_target_existing_episode_folder_not_doubled():
    episode_path, series_path = resolve_manual_pick_target("/отчеты/Show/e02", META)
    assert episode_path == "/отчеты/Show/e02"
    assert series_path == "/отчеты/Show"


def test_resolve_manual_pick_target_wrong_episode_folder_uses_its_parent_as_series():
    # Пользователь по ошибке (например, при повторной отправке следующего
    # эпизода) выбрал в пикере уже существующую папку ДРУГОГО эпизода
    # ("e02" на диске, а текущий отчёт — за "e03"). Раньше это приводило к
    # "матрёшке" (e03 создавался ВНУТРИ e02); правильно — считать родителя
    # ("Show") папкой сериала и создать e03 рядом с e02, а не внутри него.
    from src.report_filename import parse_report_filename
    other_meta = parse_report_filename("отчет_Show_s01_e03_2025_07_14_rus.docx")

    episode_path, series_path = resolve_manual_pick_target("/отчеты/Show/e02", other_meta)

    assert episode_path == "/отчеты/Show/e03"
    assert series_path == "/отчеты/Show"


def test_resolve_manual_pick_target_does_not_record_episode_folder_as_series_alias():
    # Тот же сценарий, но явно проверяем, что series_path из результата не
    # указывает на папку эпизода — иначе алиас "серия -> папка эпизода"
    # ломает ВСЕ последующие отчёты серии (каждый следующий эпизод создаётся
    # уже внутри этого неверно запомненного алиаса).
    from src.report_filename import parse_report_filename
    meta = parse_report_filename("отчет_Show_s01_e02_2025_07_15_rus.docx")

    _episode_path, series_path = resolve_manual_pick_target("/отчеты/Show/e01", meta)

    assert series_path == "/отчеты/Show"


def test_find_series_folder_uses_alias_before_fuzzy_scan(tmp_path):
    path = tmp_path / "aliases.json"
    remember_series_alias("Nepreklonniy_vozrast", "/отчеты/Непреклонный_возраст", path)
    client = MagicMock()
    client.list_folder.return_value = []  # exists-check для алиаса — папка есть, но пустая

    result = find_series_folder(client, "Nepreklonniy_vozrast", aliases_path=path)

    assert result == "/отчеты/Непреклонный_возраст"
    client.list_folder.assert_called_once_with("/отчеты/Непреклонный_возраст")


def test_find_series_folder_falls_back_when_alias_target_missing_404(tmp_path):
    path = tmp_path / "aliases.json"
    remember_series_alias("Nepreklonniy_vozrast", "/отчеты/Deleted_folder", path)
    client = MagicMock()

    def list_folder_side_effect(p):
        if p == "/отчеты/Deleted_folder":
            raise YandexDiskError("not found", status_code=404)
        return [{"name": "Nepreklonniy_vozrast", "type": "dir", "path": "/отчеты/Nepreklonniy_vozrast"}]

    client.list_folder.side_effect = list_folder_side_effect

    result = find_series_folder(client, "Nepreklonniy_vozrast", aliases_path=path)

    assert result == "/отчеты/Nepreklonniy_vozrast"
    assert load_series_aliases(path) == {}  # протухший алиас забыт


def test_find_series_folder_propagates_non_404_error_on_alias_check(tmp_path):
    path = tmp_path / "aliases.json"
    remember_series_alias("Nepreklonniy_vozrast", "/отчеты/Непреклонный_возраст", path)
    client = MagicMock()
    client.list_folder.side_effect = YandexDiskError("Server error", status_code=500)

    with pytest.raises(YandexDiskError):
        find_series_folder(client, "Nepreklonniy_vozrast", aliases_path=path)

    # транзитная ошибка не должна расцениваться как «алиас протух»
    assert load_series_aliases(path) == {"nepreklonniy_vozrast": "/отчеты/Непреклонный_возраст"}


def test_find_series_folder_searches_multiple_roots_in_order(tmp_path):
    path = tmp_path / "aliases.json"
    client = MagicMock()

    def list_folder_side_effect(p):
        if p == "/отчеты":
            return [{"name": "Unrelated", "type": "dir", "path": "/отчеты/Unrelated"}]
        if p == "/ПРОЕКТЫ NUENDO":
            return [{"name": "Nepreklonniy_vozrast", "type": "dir", "path": "/ПРОЕКТЫ NUENDO/Nepreklonniy_vozrast"}]
        return []

    client.list_folder.side_effect = list_folder_side_effect

    result = find_series_folder(
        client, "Nepreklonniy_vozrast", roots=["/отчеты", "/ПРОЕКТЫ NUENDO"], aliases_path=path,
    )

    assert result == "/ПРОЕКТЫ NUENDO/Nepreklonniy_vozrast"


def test_find_series_folder_stops_at_first_matching_root(tmp_path):
    path = tmp_path / "aliases.json"
    client = MagicMock()
    client.list_folder.side_effect = lambda p: [
        {"name": "Show", "type": "dir", "path": f"{p}/Show"},
    ]

    result = find_series_folder(client, "Show", roots=["/отчеты", "/ПРОЕКТЫ NUENDO"], aliases_path=path)

    assert result == "/отчеты/Show"
    # второй корень не должен опрашиваться, раз совпадение уже найдено в первом
    client.list_folder.assert_called_once_with("/отчеты")


def test_resolve_target_path_creates_in_first_root_when_not_found_anywhere(tmp_path):
    path = tmp_path / "aliases.json"
    client = MagicMock()
    client.list_folder.return_value = []

    episode_path, created = resolve_target_path(
        client, META, create_if_missing=True, roots=["/отчеты", "/ПРОЕКТЫ NUENDO"], aliases_path=path,
    )

    assert created is True
    assert episode_path.startswith("/отчеты/")
    assert load_series_aliases(path) == {"nepreklonniy_vozrast": "/отчеты/Nepreklonniy_vozrast"}


def test_fallback_series_key_strips_report_prefix_and_extension():
    assert fallback_series_key("отчет_kp_orlov_perehodniy_vozrast_preroll_48.docx") == \
        "kp_orlov_perehodniy_vozrast_preroll_48"


def test_fallback_series_key_without_report_prefix():
    assert fallback_series_key("kp_orlov_preroll_48.docx") == "kp_orlov_preroll_48"


def test_fallback_series_key_handles_other_extensions():
    assert fallback_series_key("отчет_preroll.pdf") == "preroll"


def test_fallback_series_key_strips_npr_extension():
    assert fallback_series_key("ShowName_season01.npr") == "ShowName_season01"


def test_fallback_series_key_strips_npr_extension_with_date_suffix():
    assert fallback_series_key("ShowName_2026_03_27_v1.npr") == "ShowName"


def test_npr_aliases_use_separate_file_from_series_aliases(tmp_path):
    # NPR_ALIASES_FILE — отдельный файл от SERIES_ALIASES_FILE: одинаковый
    # ключ в двух местах не должен пересекаться.
    series_path = tmp_path / "series_aliases.json"
    npr_path = tmp_path / "npr_aliases.json"
    key = fallback_series_key("ShowName_season01.npr")

    remember_series_alias(key, "/отчеты/ShowName", series_path)
    remember_series_alias(key, "/ПРОЕКТЫ NUENDO/ShowName", npr_path)

    client = MagicMock()
    client.list_folder.return_value = []

    assert find_series_folder(client, key, aliases_path=series_path) == "/отчеты/ShowName"
    assert find_series_folder(client, key, aliases_path=npr_path) == "/ПРОЕКТЫ NUENDO/ShowName"
    assert NPR_ALIASES_FILE.name == "npr_aliases.json"


def test_fallback_series_key_strips_trailing_date_and_version():
    assert fallback_series_key("отчет_KP_Orlov_Perehodniy_vozrast_preroll_48_2026_03_27_v1.docx") == \
        "KP_Orlov_Perehodniy_vozrast_preroll_48"


def test_fallback_series_key_strips_trailing_date_version_and_take_number():
    # Повторные заходы одного и того же ролика (v1_1, v1_2, ...) должны
    # давать ОДИН и тот же ключ — иначе выбор папки никогда не запомнится
    # на следующую отправку.
    assert fallback_series_key("отчет_KP_Orlov_Perehodniy_vozrast_preroll_48_2026_03_27_v1_1.docx") == \
        "KP_Orlov_Perehodniy_vozrast_preroll_48"
    assert fallback_series_key("отчет_KP_Orlov_Perehodniy_vozrast_preroll_48_2026_03_27_v1_2.docx") == \
        "KP_Orlov_Perehodniy_vozrast_preroll_48"


def test_fallback_series_key_without_date_suffix_unaffected():
    # "_48" в конце не должен ошибочно приниматься за дату/версию.
    assert fallback_series_key("отчет_kp_orlov_perehodniy_vozrast_preroll_48.docx") == \
        "kp_orlov_perehodniy_vozrast_preroll_48"


def test_fallback_series_key_versions_share_alias(tmp_path):
    # Реальный сценарий: первая версия ролика запоминает алиас вручную
    # выбранной папки — вторая, по-другому названная версия того же
    # ролика должна найти её автоматически по общему ключу.
    path = tmp_path / "aliases.json"
    first_key = fallback_series_key("отчет_KP_Orlov_Perehodniy_vozrast_preroll_48_2026_03_27_v1.docx")
    remember_series_alias(first_key, "/отчеты/preroll", path)

    second_key = fallback_series_key("отчет_KP_Orlov_Perehodniy_vozrast_preroll_48_2026_03_27_v1_2.docx")
    client = MagicMock()
    client.list_folder.return_value = []

    result = find_series_folder(client, second_key, aliases_path=path)

    assert result == "/отчеты/preroll"


def test_fallback_series_key_used_as_series_alias_round_trips(tmp_path):
    # Имя файла отчёта не распознано parse_report_filename (нет sNN/eNN) —
    # fallback_series_key должен давать стабильный, пригодный для
    # remember_series_alias/find_series_folder ключ.
    path = tmp_path / "aliases.json"
    key = fallback_series_key("отчет_kp_orlov_perehodniy_vozrast_preroll_48.docx")
    remember_series_alias(key, "/отчеты/preroll", path)

    client = MagicMock()
    client.list_folder.return_value = []

    result = find_series_folder(client, key, aliases_path=path)

    assert result == "/отчеты/preroll"


def test_resolve_target_path_records_alias_on_fuzzy_match(tmp_path):
    path = tmp_path / "aliases.json"
    client = MagicMock()
    client.list_folder.return_value = [
        {"name": "Nepreklonniy_vozrst", "type": "dir", "path": "/отчеты/Nepreklonniy_vozrst"},
    ]

    resolve_target_path(client, META, create_if_missing=False, aliases_path=path)

    assert load_series_aliases(path) == {"nepreklonniy_vozrast": "/отчеты/Nepreklonniy_vozrst"}


def test_resolve_target_path_records_alias_on_create(tmp_path):
    path = tmp_path / "aliases.json"
    client = MagicMock()
    client.list_folder.return_value = []

    resolve_target_path(client, META, create_if_missing=True, aliases_path=path)

    assert load_series_aliases(path) == {"nepreklonniy_vozrast": "/отчеты/Nepreklonniy_vozrast"}


def test_resolve_target_path_second_call_uses_alias_not_fuzzy_scan(tmp_path):
    path = tmp_path / "aliases.json"
    client = MagicMock()
    client.list_folder.return_value = []

    resolve_target_path(client, META, create_if_missing=True, aliases_path=path)
    client.list_folder.reset_mock()
    client.list_folder.return_value = []  # exists-check алиаса — папка есть

    resolve_target_path(client, META, create_if_missing=False, aliases_path=path)

    client.list_folder.assert_called_once_with("/отчеты/Nepreklonniy_vozrast")


def test_list_series_aliases_sorted(tmp_path):
    path = tmp_path / "aliases.json"
    remember_series_alias("Zebra_show", "/отчеты/Zebra", path)
    remember_series_alias("Alpha_show", "/отчеты/Alpha", path)

    result = list_series_aliases(path)

    assert result == [
        ("alpha_show", "/отчеты/Alpha"),
        ("zebra_show", "/отчеты/Zebra"),
    ]


def test_list_series_aliases_empty(tmp_path):
    assert list_series_aliases(tmp_path / "missing.json") == []


def test_alias_key_for_path_finds_matching_key(tmp_path):
    path = tmp_path / "aliases.json"
    remember_series_alias("Zebra_show", "/отчеты/Zebra", path)
    remember_series_alias("Alpha_show", "/отчеты/Alpha", path)

    assert alias_key_for_path("/отчеты/Alpha", path) == "alpha_show"


def test_alias_key_for_path_returns_none_when_unmatched(tmp_path):
    path = tmp_path / "aliases.json"
    remember_series_alias("Zebra_show", "/отчеты/Zebra", path)

    assert alias_key_for_path("/отчеты/Unrelated", path) is None


def test_alias_key_for_path_missing_file_returns_none(tmp_path):
    assert alias_key_for_path("/отчеты/Alpha", tmp_path / "missing.json") is None


def test_remember_uploaded_report_round_trip(tmp_path):
    path = tmp_path / "uploads.json"
    remember_uploaded_report("/local/a", "/отчеты/Show/e01/report_a", path)

    entries = load_uploaded_reports(path)

    assert len(entries) == 1
    assert entries[0]["local_folder"] == "/local/a"
    assert entries[0]["remote_path"] == "/отчеты/Show/e01/report_a"
    assert "uploaded_at" in entries[0]


def test_remember_uploaded_report_newest_first(tmp_path):
    path = tmp_path / "uploads.json"
    remember_uploaded_report("/local/a", "/отчеты/Show/e01/a", path)
    remember_uploaded_report("/local/b", "/отчеты/Show/e02/b", path)

    entries = load_uploaded_reports(path)

    assert [e["remote_path"] for e in entries] == ["/отчеты/Show/e02/b", "/отчеты/Show/e01/a"]


def test_remember_uploaded_report_dedupes_same_remote_path(tmp_path):
    path = tmp_path / "uploads.json"
    remember_uploaded_report("/local/a", "/отчеты/Show/e01/a", path)
    remember_uploaded_report("/local/a", "/отчеты/Show/e01/a", path)  # повторная синхронизация того же отчёта

    entries = load_uploaded_reports(path)

    assert len(entries) == 1


def test_remember_uploaded_report_caps_at_max_entries(tmp_path):
    path = tmp_path / "uploads.json"
    for i in range(25):
        remember_uploaded_report(f"/local/{i}", f"/отчеты/Show/e{i:02d}/report", path)

    entries = load_uploaded_reports(path)

    assert len(entries) == 20
    assert entries[0]["remote_path"] == "/отчеты/Show/e24/report"  # самый новый первый


def test_load_uploaded_reports_missing_file_returns_empty(tmp_path):
    assert load_uploaded_reports(tmp_path / "missing.json") == []


def test_load_uploaded_reports_corrupt_file_returns_empty(tmp_path):
    path = tmp_path / "uploads.json"
    path.write_text("not valid json{{{", encoding="utf-8")

    assert load_uploaded_reports(path) == []


def test_set_variant_override_round_trip(tmp_path):
    path = tmp_path / "variant_overrides.json"
    set_variant_override("/отчеты/Show/weird_folder", "ME", path)

    assert load_variant_overrides(path) == {"/отчеты/Show/weird_folder": "ME"}


def test_set_variant_override_none_clears_existing_entry(tmp_path):
    path = tmp_path / "variant_overrides.json"
    set_variant_override("/отчеты/Show/weird_folder", "AD", path)
    set_variant_override("/отчеты/Show/weird_folder", None, path)

    assert load_variant_overrides(path) == {}


def test_set_variant_override_none_on_missing_entry_is_a_noop(tmp_path):
    path = tmp_path / "variant_overrides.json"
    set_variant_override("/отчеты/Show/weird_folder", None, path)

    assert load_variant_overrides(path) == {}


def test_load_variant_overrides_missing_file_returns_empty(tmp_path):
    assert load_variant_overrides(tmp_path / "missing.json") == {}


def test_load_variant_overrides_corrupt_file_returns_empty(tmp_path):
    path = tmp_path / "variant_overrides.json"
    path.write_text("not valid json{{{", encoding="utf-8")

    assert load_variant_overrides(path) == {}
