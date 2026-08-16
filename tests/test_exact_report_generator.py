from docx import Document

from src.csv_importer import Issue
from src.exact_report_generator import ExactReportGenerator


def _issue_with_marker_id() -> Issue:
    return Issue(
        timecode_in="01:02:03:04",
        timecode_out="01:02:05:06",
        description="Test marker.wav",
        audio_20_c=True,
        audio_20_uc=False,
        audio_51_c=True,
        audio_51_uc=False,
        blocker=False,
        fix_required=True,
        comment_required=False,
        comments="Test comment",
        marker_id="M1",
    )


def test_collect_pdf_paths_prefers_full_list_and_removes_duplicates():
    pdf_paths = [
        "/tmp/project_20_cens.pdf",
        "/tmp/project_20_uncens.pdf",
        "/tmp/project_51_cens.pdf",
        "/tmp/project_51_uncens.pdf",
        "/tmp/project_20_cens.pdf",
    ]

    result = ExactReportGenerator._collect_pdf_paths(
        pdf_paths=pdf_paths,
        pdf_20_path="/tmp/project_20_cens.pdf",
        pdf_51_path="/tmp/project_51_cens.pdf",
    )

    assert result == [
        "/tmp/project_20_cens.pdf",
        "/tmp/project_20_uncens.pdf",
        "/tmp/project_51_cens.pdf",
        "/tmp/project_51_uncens.pdf",
    ]


def test_collect_pdf_paths_falls_back_to_legacy_two_paths():
    result = ExactReportGenerator._collect_pdf_paths(
        pdf_20_path="/tmp/project_20_uncens.pdf",
        pdf_51_path="/tmp/project_51_uncens.pdf",
    )

    assert result == [
        "/tmp/project_20_uncens.pdf",
        "/tmp/project_51_uncens.pdf",
    ]


def test_marker_list_places_id_before_timecode_in_for_main_report():
    doc = Document()

    ExactReportGenerator()._add_marker_list_exact(
        doc, [_issue_with_marker_id()], report_type="main"
    )

    table = doc.tables[-1]
    assert [cell.text for cell in table.rows[0].cells[:4]] == [
        "ID",
        "Timecode In",
        "Timecode Out",
        "Description",
    ]
    assert [cell.text for cell in table.rows[1].cells[:4]] == [
        "M1",
        "01:02:03:04",
        "01:02:05:06",
        "Test marker",
    ]
    assert len(table.columns) == 12
    assert table.rows[1].cells[-1].text == "Test comment"


def test_marker_list_uses_ten_point_font_like_conclusion():
    doc = Document()

    ExactReportGenerator()._add_marker_list_exact(
        doc, [_issue_with_marker_id()], report_type="main"
    )

    table = doc.tables[-1]
    font_sizes = {
        run.font.size.pt
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text
    }
    assert font_sizes == {10.0}


def test_marker_list_places_id_before_timecode_in_for_me_report():
    doc = Document()

    ExactReportGenerator()._add_marker_list_exact(
        doc, [_issue_with_marker_id()], report_type="me"
    )

    table = doc.tables[-1]
    assert [cell.text for cell in table.rows[0].cells[:4]] == [
        "ID",
        "Timecode In",
        "Timecode Out",
        "Description",
    ]
    assert [cell.text for cell in table.rows[1].cells[:4]] == [
        "M1",
        "01:02:03:04",
        "01:02:05:06",
        "Test marker",
    ]
    assert [cell.text for cell in table.rows[0].cells[4:10]] == [
        "2.0 ME", "5.1 ME", "2.0 DX", "5.1 DX", "2.0 OPT", "5.1 OPT",
    ]
    assert len(table.columns) == 14
    assert table.rows[1].cells[-1].text == "Test comment"


def test_me_marker_list_places_dx_and_opt_marks_in_new_columns():
    doc = Document()
    issue = _issue_with_marker_id()
    issue.me_tracks = {
        "me_20": True,
        "me_51": False,
        "dx_20": True,
        "dx_51": True,
        "opt_20": False,
        "opt_51": True,
    }

    ExactReportGenerator()._add_marker_list_exact(doc, [issue], report_type="me")

    assert [cell.text for cell in doc.tables[-1].rows[1].cells[4:10]] == [
        "*", "", "*", "*", "", "*",
    ]


def test_me_technical_table_adds_all_dx_and_opt_variants():
    doc = Document()
    audio = {
        "duration": 60.0,
        "file_name": "episode.wav",
        "sample_rate": 48000,
        "bit_depth": "PCM_24",
        "channel_order": "L R",
    }
    tech_info = {
        "audio_20_c": dict(audio),
        "audio_51_c": {**audio, "channel_order": "L R C LFE Ls Rs"},
        "video": {"duration": 60.0, "file_name": "episode.mov", "format": "MOV", "fps": 25},
        "audio_me_dx_20": {**audio, "file_name": "episode_20_DX.wav"},
        "audio_me_dx_51": {**audio, "file_name": "episode_51_DX.wav", "channel_order": "L R C LFE Ls Rs"},
        "audio_me_opt_a_20": {**audio, "file_name": "episode_20_OPT_A.wav"},
        "audio_me_opt_a_51": {**audio, "file_name": "episode_51_OPT_A.wav", "channel_order": "L R C LFE Ls Rs"},
        "audio_me_opt_b_20": {**audio, "file_name": "episode_20_OPT_B.wav"},
        "audio_me_opt_b_51": {**audio, "file_name": "episode_51_OPT_B.wav", "channel_order": "L R C LFE Ls Rs"},
    }

    ExactReportGenerator()._add_me_technical_table(doc, tech_info, "", "")

    labels = [row.cells[0].text for row in doc.tables[-1].rows[3:12]]
    assert labels == [
        "20 ME", "51 ME", "20 DX", "51 DX",
        "20 OPT A", "51 OPT A", "20 OPT B", "51 OPT B", "VIDEO",
    ]


def test_dcp_technical_table_adds_mix_and_split_channels():
    doc = Document()
    mix = {
        "duration": 60.0,
        "file_name": "episode_51_DCP.wav",
        "sample_rate": 48000,
        "bit_depth": "PCM_24",
        "channel_order": "L R C LFE Ls Rs",
        "sample_peak": -0.1,
    }
    split = {
        "duration": 60.0,
        "sample_rate": 48000,
        "bit_depth": "PCM_24",
        "channel_order": "M",
        "sample_peak": -0.25,
    }
    tech_info = {
        "audio_51_c": mix,
        "video": {"duration": 60.0, "file_name": "episode.mov", "format": "MOV", "fps": 25},
    }
    for key, suffix in (("l", "L"), ("r", "R"), ("c", "C"), ("lfe", "LFE"), ("ls", "Ls"), ("rs", "Rs")):
        tech_info[f"audio_dcp_split_{key}"] = {**split, "file_name": f"episode_{suffix}.wav"}

    ExactReportGenerator()._add_technical_table_with_conclusion(
        doc, tech_info, "", "", report_type="dcp"
    )

    table = doc.tables[-1]
    assert [row.cells[0].text for row in table.rows[3:11]] == [
        "51 DCP", "Video ref", "Left", "Right", "Center", "LFE",
        "Left surround", "Right surround",
    ]
    assert table.rows[3].cells[4].text == "-0.10 dBFS"
    for row in table.rows[5:11]:
        assert row.cells[2].text == "0:01:00.000"
        assert row.cells[4].text == "-0.25 dBFS"
        assert row.cells[6].text == "PCM 48kHz 24 bit split"
